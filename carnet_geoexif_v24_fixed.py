import os
import sys
import json
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import csv
import math
import re
import base64
import collections
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
import multiprocessing
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
import customtkinter as ctk
from tkintermapview import TkinterMapView
import subprocess
import threading
import time
import shutil
import requests
import io
import socket
import http.server
import urllib.parse
import webbrowser
from PIL import Image, ImageTk, ImageDraw

COMPANION_PORT = 8765
COMPANION_INBOX = "companion_inbox.json"  # file d'attente si import différé

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("dark-blue")

CONFIG_FILE = "app_config.json"
NOTES_FILE = "observations.json"
DEVICES_FILE = "dispositifs.json"
PROJECT_FILE = "geoexif_projet.json"  # projet multi-sorties (ex. Meuse 2026)
SPECIES_FILE = "dictionnaire_especes.json"
WIKI_CACHE_FILE = "wikipedia_cache.json"  # cache local des résumés Wikipédia (fiche espèce)
NOTES_BACKUP_DIR = ".geoexif_notes_backups"  # sous-dossier dans chaque dossier de photos
NOTES_BACKUP_KEEP = 15                       # nombre max de versions conservées
NOTES_BACKUP_MIN_INTERVAL_SEC = 90           # intervalle mini entre 2 versions (évite le spam de l'autosave)

APP_VERSION = "4.0.17"
APP_AUTHOR = "Wilmet Sébastien"
# Version 4.0 — jalon figé : menus stabilisés, archive/recovery, Birda (semi-auto),
# espace de travail multi-écran, débrief hors-ligne. Pas de refonte framework.
# 4.0.1 — chemins Audacity/Birda configurables + import SQLite birda-catalog.db

# Palette UI (interface sombre moderne)
# Palettes de thème (sombre / clair / papier)
THEMES = {
    "sombre": {
        "bg": "#101612",
        "sidebar": "#141c17",
        "card": "#1c2620",
        "card_alt": "#24302a",
        "card_elevated": "#2c3a32",
        "border": "#354840",
        "border_soft": "#2a3830",
        "accent": "#4ec4b0",
        "accent_hover": "#3aab98",
        "accent_soft": "#1e3a34",
        "success": "#6ed99a",
        "success_hover": "#52c07f",
        "warning": "#f0b45c",
        "warning_hover": "#d99a3e",
        "danger": "#ef8a8a",
        "danger_hover": "#d97070",
        "purple": "#b39af5",
        "purple_hover": "#9a7ee6",
        "gold": "#e0c49a",
        "map_tint": "#243a32",
        "pink": "#f072a8",
        "pink_hover": "#d45a90",
        "text": "#f3f7f4",
        "text_dim": "#a8b8b0",
        "text_muted": "#7a8c84",
        "text_accent": "#8eebd8",
        "list_annotated": "#4fd67a",
        "list_pending": "#b0beb6",
        "list_video": "#8eb4e8",
        "ctk_mode": "Dark",
    },
    "clair": {
        "bg": "#e8eee9",
        "sidebar": "#f4f7f5",
        "card": "#ffffff",
        "card_alt": "#eef3f0",
        "card_elevated": "#f8faf9",
        "border": "#c5d4cb",
        "border_soft": "#d9e4dd",
        "accent": "#2a9d8f",
        "accent_hover": "#21867a",
        "accent_soft": "#d4efe9",
        "success": "#2f9e5f",
        "success_hover": "#25804c",
        "warning": "#d4920a",
        "warning_hover": "#b87a08",
        "danger": "#d64545",
        "danger_hover": "#b53636",
        "purple": "#7c5cbf",
        "purple_hover": "#6649a8",
        "gold": "#b0893e",
        "map_tint": "#dce8e2",
        "pink": "#c2185b",
        "pink_hover": "#a0144c",
        "text": "#1a2420",
        "text_dim": "#4a5c54",
        "text_muted": "#6a7d74",
        "text_accent": "#1d7a6c",
        "list_annotated": "#1a7a3c",
        "list_pending": "#3a4a42",
        "list_video": "#1a5fad",
        "ctk_mode": "Light",
    },
    "papier": {
        "bg": "#f3efe6",
        "sidebar": "#faf7f0",
        "card": "#fffdf8",
        "card_alt": "#efe9dc",
        "card_elevated": "#ffffff",
        "border": "#d4cbb8",
        "border_soft": "#e6dfd0",
        "accent": "#3d6b5a",
        "accent_hover": "#2f5547",
        "accent_soft": "#e2ebe6",
        "success": "#4a8c5c",
        "success_hover": "#3a7249",
        "warning": "#c4841a",
        "warning_hover": "#a36c12",
        "danger": "#b54a3c",
        "danger_hover": "#943c30",
        "purple": "#6b5b8a",
        "purple_hover": "#564870",
        "gold": "#a67c3a",
        "map_tint": "#e8e2d4",
        "pink": "#a64d6d",
        "pink_hover": "#8a3f5a",
        "text": "#2c2418",
        "text_dim": "#5c5346",
        "text_muted": "#7a7164",
        "text_accent": "#3d6b5a",
        "list_annotated": "#2d6b3f",
        "list_pending": "#4a4035",
        "list_video": "#2a5a9e",
        "ctk_mode": "Light",
    },
}

# Palette active (mutable — mise à jour au changement de thème)
UI = dict(THEMES["sombre"])


# Workers CPU pour tâches I/O parallèles (miniatures, lectures EXIF…)
CPU_WORKERS = max(2, min(8, (os.cpu_count() or 4)))


# Critères de confiance ID IA (Gemini / Grok)
IA_CONFIANCE_RUBRIC_FR = (
    "Niveaux de confiance (choisir EXACTEMENT : faible | moyenne | élevée) :\n"
    "\n"
    "élevée — TOUTES les conditions :\n"
    "  • traits diagnostiques d'espèce clairement visibles\n"
    "  • netteté et lumière correctes\n"
    "  • peu d'espèces confondables réalistes en Meuse pour cet aspect\n"
    "  • identification au niveau ESPÈCE (pas seulement genre/famille)\n"
    "\n"
    "moyenne — au moins une limitation :\n"
    "  • bons indices mais 2 espèces proches possibles\n"
    "  • animal partiel (flanc, arrière, tête seule)\n"
    "  • indice clair mais pas 100 % spécifique\n"
    "  → espèce la plus probable ; alternative possible dans commentaire\n"
    "\n"
    "faible — incertitude forte :\n"
    "  • flou, nuit, fort crop, silhouette, contre-jour\n"
    "  • indice ambigu\n"
    "  • plusieurs familles possibles\n"
    "  → espece=\"Inconnu\" ou groupe (\"Cervidé indéterminé\") si besoin\n"
    "  → ne jamais inventer une espèce rare/exotique pour remplir le champ\n"
    "\n"
    "Meuse : ne pas mettre élevée pour une espèce hors aire sans critère ultra net."
)

IA_CONFIANCE_RUBRIC_EN = (
    "Confidence (use EXACTLY one French label: faible | moyenne | élevée):\n"
    "\n"
    "élevée — ALL must hold:\n"
    "  • species-level diagnostic traits clearly visible\n"
    "  • adequate sharpness and light\n"
    "  • few realistic Meuse (NE France) look-alikes for this view\n"
    "  • SPECIES-level ID (not only genus/family)\n"
    "\n"
    "moyenne — at least one limitation:\n"
    "  • solid clues but 2 close species possible\n"
    "  • partial animal (flank, rear, head only)\n"
    "  • clear sign but not fully species-specific\n"
    "  → most probable species; optional alternative in commentaire\n"
    "\n"
    "faible — uncertain:\n"
    "  • blur, night, heavy crop, silhouette, backlight\n"
    "  • ambiguous sign; several families possible\n"
    "  → espece=\"Inconnu\" or group name if species-level ID unreasonable\n"
    "  → never invent rare/exotic species to fill the field\n"
    "\n"
    "Meuse prior: do not use élevée for out-of-range taxa without razor-sharp diagnostics."
)


# Couleurs des marqueurs sur la carte, par catégorie
CATEGORY_COLORS = {
    "Mammifère": "#e67e22",
    "Oiseau": "#3498db",
    "Insecte": "#f1c40f",
    "Autre": "#95a5a6",
    "Non classé": "#7f7f7f",
}
# Couleur distincte pour les observations sans photo (placées sur la carte ou saisie manuelle)
SANS_PHOTO_COLOR = "#e91e63"  # rose / magenta — bien visible sur fonds IGN/OSM

# Modèle d'observation unifié (fiche photo / vidéo / sans photo / indice)
COMPORTEMENTS_OBS = [
    "", "Affût", "Fuite", "Nourrissage", "Repos", "Repos / marche", "Stationnaire",
    "Toilette", "Jeu", "Vocalisation / chant", "Vol", "Nage", "Repos", "Autre",
]
CERTITUDES_OBS = ["", "Sûr", "Probable", "Possible"]
TYPES_INDICE = [
    "", "Aucun (contact direct)", "Empreinte / Trace", "Terrier", "Coulée / Passage",
    "Latrines / Crottes", "Écorçage / Frottis", "Restes de repas / Pelote",
    "Nid / Gîte", "Plumes / Poils",
    "Entendu / Chant", "Prise de son (Birda)", "Autre indice",
]


# Formes de marqueurs pour les indices / traces (différenciés des observations d'animaux)
TRACE_SHAPES = {
    "empreinte": "triangle",
    "trace": "triangle",
    "terrier": "square",
    "coulée": "diamond",
    "coulee": "diamond",
    "passage": "diamond",
    "latrine": "square",
    "écorçage": "diamond",
    "ecorcage": "diamond",
    "pelote": "triangle",
    "restes": "triangle",
    "frottis": "diamond",
    "indice": "triangle",
}

JOURS_FR = ["Lundi", "Mardi", "Mercredi", "Jeudi", "Vendredi", "Samedi", "Dimanche"]

# Suggestions naturalistes Meuse (55) — calendrier simple, non exhaustif
MEUSE_SAISON_ESPECES = {
    1:  {"habitats": ["forêt", "lisière", "plan d'eau gelé/bord"], "cibles": ["Cerf élaphe (brame terminé, observation)", "Sanglier", "Renard roux", "Buse variable", "Mésange charbonnière"], "notes": "Cœur d'hiver : traces dans la neige, affûts lisière."},
    2:  {"habitats": ["forêt", "prairie", "cours d'eau"], "cibles": ["Blaireau européen", "Chevreuil d'Europe", "Pic noir", "Héron cendré", "Traces / Empreintes"], "notes": "Février : reprise d'activité mustélidés, pics."},
    3:  {"habitats": ["lisière", "prairie", "zone humide"], "cibles": ["Crapaud / amphibiens (migration)", "Milan royal", "Grue cendrée (migration)", "Renard roux", "Lièvre d'Europe"], "notes": "Migration pré-nuptiale, retour des milans."},
    4:  {"habitats": ["bocage", "lisière", "mare"], "cibles": ["Milan noir", "Pie-grièche écorcheur (arrivée)", "Chevreuil d'Europe", "Amphibiens", "Papillons précoces"], "notes": "Avril : chant, territoires, mare."},
    5:  {"habitats": ["lisière", "prairie fleurie", "forêt"], "cibles": ["Lucane cerf-volant (fin mai)", "Chevreuil d'Europe", "Milan royal", "Engoulevent (écoute)", "Orchidées / flore"], "notes": "Mai : insectes, floraison, écoute nocturne."},
    6:  {"habitats": ["forêt", "clairière", "rivière"], "cibles": ["Lucane cerf-volant", "Cerf élaphe (bois en velours)", "Caloptéryx / libellules", "Martin-pêcheur", "Sanglier"], "notes": "Juin : lucanes, odonates, affûts longs."},
    7:  {"habitats": ["lisière", "culture", "forêt"], "cibles": ["Sanglier", "Chevreuil d'Europe", "Buse variable", "Papillons", "Coulée / Passage"], "notes": "Juillet : chaleur, points d'eau, coulées."},
    8:  {"habitats": ["lisière", "gagnage", "forêt"], "cibles": ["Cerf élaphe (approche du brame)", "Sanglier", "Renard roux", "Milan noir", "Indices (frottis, écorçage)"], "notes": "Août : préparation brame, frottis."},
    9:  {"habitats": ["forêt", "lisière", "clairière"], "cibles": ["Cerf élaphe (brame)", "Chevreuil d'Europe", "Sanglier", "Buse variable", "Grue cendrée (premiers passages)"], "notes": "Septembre : brame, migration amorce."},
    10: {"habitats": ["forêt", "plan d'eau", "plaine"], "cibles": ["Grue cendrée", "Cigogne", "Cerf élaphe", "Renard roux", "Champignon"], "notes": "Octobre : grues, champignons, chasse d'affût."},
    11: {"habitats": ["forêt", "lisière", "humide"], "cibles": ["Sanglier", "Cerf élaphe", "Canards / oiseaux d'eau", "Traces / Empreintes", "Blaireau européen"], "notes": "Novembre : traces, oiseaux d'eau."},
    12: {"habitats": ["forêt", "lisière", "village/lisière"], "cibles": ["Renard roux", "Sanglier", "Buse variable", "Mésanges", "Empreinte (détail)"], "notes": "Décembre : observation hivernale, traces."},
}




# Calendrier de chasse INDICATIF — Meuse (55)
# Verifier chaque saison : federation chasseurs Meuse + arrete prefectorel.
MEUSE_CALENDRIER_CHASSE = {
    "disclaimer": (
        "Calendrier INDICATIF pour la Meuse (55). "
        "Les dates reelles sont fixees chaque saison par arrete prefectorel "
        "et la Federation departementale des chasseurs. "
        "Verifiez toujours les textes officiels avant toute sortie."
    ),
    "sources_suggerees": [
        "Federation departementale des chasseurs de la Meuse",
        "Arretes prefectorels Meuse (chasse)",
        "Office francais de la biodiversite (OFB)",
    ],
    "periodes": [
        {
            "espece": "Sanglier",
            "mode": "approche / affut / battue (selon arrete)",
            "mois": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12],
            "note": "Souvent praticable une grande partie de l'annee selon modalites locales.",
            "photo": "Pression possible presque toute l'annee : discretion renforcee.",
        },
        {
            "espece": "Chevreuil d'Europe",
            "mode": "approche / affut (ete) puis battue",
            "mois": [6, 7, 8, 9, 10, 11, 12, 1, 2],
            "note": "Ouverture estivale puis automne-hiver selon arrete.",
            "photo": "Ete : affuts lisiere ; automne : plus de frequentation forestiere.",
        },
        {
            "espece": "Cerf elaphe",
            "mode": "approche / affut / battue",
            "mois": [9, 10, 11, 12, 1, 2],
            "note": "Saison automne-hiver ; brame souvent autour de sept.-oct.",
            "photo": "Brame : priorite observation discrete hors actions de chasse.",
        },
        {
            "espece": "Renard roux",
            "mode": "divers (selon arrete)",
            "mois": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12],
            "note": "Cadre reglementaire variable (nuisible / chassable).",
            "photo": "Possible toute l'annee selon contexte local.",
        },
        {
            "espece": "Gibier d'eau (canards, limicoles...)",
            "mode": "chasse au gibier d'eau",
            "mois": [8, 9, 10, 11, 12, 1, 2],
            "note": "Ouvertures souvent fin ete / automne-hiver ; horaires particuliers.",
            "photo": "Zones humides : verifier jours et heures autorises.",
        },
        {
            "espece": "Petit gibier sedéntaire (faisan, perdrix...)",
            "mode": "chasse devant soi",
            "mois": [9, 10, 11, 12, 1, 2],
            "note": "Classiquement automne-hiver ; variabilite selon especes.",
            "photo": "Bocage et plaines : weekends souvent plus frequentes.",
        },
        {
            "espece": "Lievre d'Europe",
            "mode": "chasse devant soi",
            "mois": [9, 10, 11, 12],
            "note": "Saison typique automnale.",
            "photo": "Plaines / cultures : anticiper la frequentation.",
        },
    ],
}


def meuse_chasse_pour_mois(mois):
    """Periodes de chasse indicatives actives pour un mois (1-12)."""
    try:
        mois = int(mois)
    except Exception:
        mois = datetime.now().month
    out = []
    for p in MEUSE_CALENDRIER_CHASSE.get("periodes") or []:
        if mois in (p.get("mois") or []):
            out.append(p)
    return out


# Base par défaut (sauvegardée si le fichier JSON n'existe pas)
DEFAULT_FAUNE = {
    "Mammifère": ["Chevreuil d'Europe", "Sanglier", "Renard roux", "Cerf élaphe", "Chat forestier (Europe)", "Blaireau européen", "Lièvre d'Europe", "Fouine", "Martre des pins", "Écureuil roux", "Hérisson d'Europe", "Autre mammifère"],
    "Oiseau": ["Milan royal", "Milan noir", "Buse variable", "Faucon créerelle", "Cigogne noire", "Cigogne blanche", "Héron cendré", "Grande Aigrette", "Grue cendrée (en migration/halte)", "Pic noir", "Pic épeiche", "Geai des chênes", "Pie-grièche écorcheur", "Mésange charbonnière", "Mésange huppée", "Autre oiseau"],
    "Insecte": ["Lucane cerf-volant", "Grand Paon du nuit", "Paon du jour (papillon)", "Aurore (papillon)", "Flambé (papillon)", "Robert-le-Diable (papillon)", "Abeille charpentière", "Bourdon terrestre", "Libellule déprime", "Caloptéryx éclatant (cours d'eau)", "Carabe doré", "Gendarme (Pyrrhocore)", "Autre insecte"],
    "Autre": [
        "Traces / Empreintes", "Empreinte (détail)", "Terrier", "Coulée / Passage",
        "Latrines", "Écorçage / Frottis", "Restes de repas / Pelote",
        "Amphibien (Grenouille/Crapaud/Triton)", "Reptile (Lézard/Couleuvre)",
        "Champignon", "Flore / Plante", "Inconnu"
    ]
}


def _exiftool_chunk_worker(payload):
    """Worker multiprocessing (picklable) : un lot de fichiers → liste de dicts JSON ExifTool.

    payload = (exe_path, file_paths, tags, timeout_seconds)
    Process séparé = vrai parallélisme CPU autour des processus ExifTool enfants.
    """
    exe_path, file_paths, tags, timeout_seconds = payload
    if not file_paths:
        return []
    cmd = [exe_path, "-n", "-fast2", "-q", "-q", "-json"]
    for t in tags:
        cmd.append(f"-{t}")
    cmd.extend([os.path.normpath(fp) for fp in file_paths])

    startupinfo = None
    if os.name == "nt":
        try:
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        except Exception:
            startupinfo = None
    try:
        proc = subprocess.run(
            cmd, startupinfo=startupinfo,
            capture_output=True, text=True, encoding="utf-8", errors="ignore",
            timeout=timeout_seconds,
        )
        if proc.returncode not in (0, 1):
            return []
        out = (proc.stdout or "").strip()
        if not out:
            return []
        data = json.loads(out)
        return data if isinstance(data, list) else [data]
    except Exception:
        return []


class GeoExifIgnApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        try:
            ctk.set_appearance_mode("dark")
            ctk.set_default_color_theme("green")
        except Exception:
            pass

        self.title(f"GeoExif Meuse 55  ·  v{APP_VERSION}")
        self.geometry("1450x850")
        self.configure(fg_color=UI["bg"])
        self._maximize_window()
        # Menu natif après init fenêtre (évite conflits CTk au démarrage)
        self.after(50, self._build_menubar)
        self.after(80, lambda: self.apply_theme(getattr(self, "_theme_name", "sombre"), save=False, silent=True))
        self.after(400, self._startup_config_check)

        self.photo_folder_path = ""
        self.gpx_file_path = ""
        self.devices_data = []
        self.device_place_mode = False
        self._device_map_markers = []
        self.photos_data = {} 
        self.selected_photo_path = ""
        self.is_synced = False  
        self.is_processing = False 
        self.weather_cache = {}          # filename -> {date, heure, temperature, humidite, ciel, ...}
        self.location_cache = {}         # filename -> nom de lieu (reverse geocoding)
        self._current_full_preview = None  # PIL.Image plein format de la photo sélectionnée (pour le slider de taille)
        self.last_backup_dir = None      # dossier de sauvegarde de la dernière synchro (pour annulation)
        self.active_markers = []         # marqueurs actuellement affichés sur la carte
        self.category_filters = {}       # catégorie -> BooleanVar (filtres carte)
        self.ia_config = self._load_ia_config()  # clé API Gemini (Lot 3 — identification IA)
        self.app_config = self._load_app_config()  # dossier rapports, préférences export
        self._theme_name = (self.app_config or {}).get("theme", "sombre")
        
        # Piste 1 : Charger le dictionnaire d'espèces évolutif
        self.load_species_dict()

        self.grid_columnconfigure(0, weight=1, minsize=300)
        self.grid_columnconfigure(1, weight=4)
        self.grid_rowconfigure(0, weight=0)  # barre d'icônes
        self.grid_rowconfigure(1, weight=1)  # contenu

        # --- BARRE D'OUTILS (style QGIS / FastStone) ---
        self._build_icon_toolbar()

        # --- PANNEAU DE CONTRÔLE (Gauche) — UI allégée 3 blocs ---
        self.sidebar = ctk.CTkFrame(self, corner_radius=18, fg_color=UI["sidebar"], border_width=1, border_color=UI["border"])
        self.sidebar.grid(row=1, column=0, sticky="nsew", padx=(10, 6), pady=(4, 10))
        self.sidebar.grid_rowconfigure(4, weight=1)  # log prend l'espace restant
        self.sidebar.grid_columnconfigure(0, weight=1)
        self._ui_mode = "sortie"  # "sortie" | "saison"

        # —— En-tête + bascule mode ——
        header = ctk.CTkFrame(
            self.sidebar, fg_color=UI["card"], corner_radius=16,
            border_width=1, border_color=UI["border"],
        )
        header.grid(row=0, column=0, padx=12, pady=(12, 8), sticky="ew")
        ctk.CTkFrame(header, fg_color=UI["accent"], height=4, corner_radius=0).pack(fill="x")
        title_row = ctk.CTkFrame(header, fg_color="transparent")
        title_row.pack(fill="x", padx=12, pady=(12, 4))
        self.title_label = ctk.CTkLabel(
            title_row, text="🦌  GeoExif",
            font=ctk.CTkFont(size=18, weight="bold"), text_color=UI["text"],
        )
        self.title_label.pack(side="left")
        ctk.CTkLabel(
            title_row, text=f" v{APP_VERSION} ",
            font=ctk.CTkFont(size=10, weight="bold"),
            text_color=UI["bg"], fg_color=UI["accent"], corner_radius=8,
        ).pack(side="right")
        ctk.CTkLabel(
            header, text="Meuse 55 · carnet naturaliste",
            font=ctk.CTkFont(size=12), text_color=UI["text_accent"],
        ).pack(anchor="w", padx=12, pady=(0, 8))

        self.mode_switch = ctk.CTkSegmentedButton(
            header,
            values=["Sortie", "Saison"],
            command=self._on_ui_mode_change,
            font=ctk.CTkFont(size=12, weight="bold"),
            selected_color=UI["accent"],
            selected_hover_color=UI["accent_hover"],
            unselected_color=UI["card_alt"],
            unselected_hover_color=UI["border"],
            height=32,
        )
        self.mode_switch.pack(fill="x", padx=12, pady=(0, 12))
        self.mode_switch.set("Sortie")

        # —— Bloc 1 : Dossier ——
        card_dossier = ctk.CTkFrame(
            self.sidebar, fg_color=UI["card"], corner_radius=14,
            border_width=1, border_color=UI["border"],
        )
        card_dossier.grid(row=1, column=0, padx=12, pady=6, sticky="ew")
        ctk.CTkLabel(
            card_dossier, text="1 · Sortie",
            font=ctk.CTkFont(size=12, weight="bold"), text_color=UI["text_dim"],
        ).pack(anchor="w", padx=12, pady=(10, 4))
        self.folder_btns = ctk.CTkFrame(card_dossier, fg_color="transparent")
        self.folder_btns.pack(fill="x", padx=10, pady=(0, 10))
        self.folder_btns.grid_columnconfigure(0, weight=1)
        self.btn_browse_photos = ctk.CTkButton(
            self.folder_btns, text="📂  Ouvrir un dossier",
            command=self.select_photo_folder,
            fg_color=UI["card_alt"], hover_color=UI["border"],
            border_width=1, border_color=UI["border"],
            height=34, corner_radius=10,
        )
        self.btn_browse_photos.grid(row=0, column=0, sticky="ew", pady=(0, 4))
        self.btn_carnet_only = ctk.CTkButton(
            self.folder_btns, text="👁️  Carnet sans photos",
            command=self.create_or_open_carnet_folder,
            fg_color=UI.get("warning", "#f0b45c"),
            hover_color=UI.get("warning_hover", "#d99a3e"),
            height=32, corner_radius=10,
            font=ctk.CTkFont(size=12, weight="bold"),
        )
        self.btn_carnet_only.grid(row=1, column=0, sticky="ew")
        # labels conservés pour compatibilité
        self.lbl_step1 = ctk.CTkLabel(self.sidebar, text="")

        # —— Bloc 2 : GPS (essentiel) ——
        self.card_gps = ctk.CTkFrame(
            self.sidebar, fg_color=UI["card"], corner_radius=14,
            border_width=1, border_color=UI["border"],
        )
        self.card_gps.grid(row=2, column=0, padx=12, pady=6, sticky="ew")
        ctk.CTkLabel(
            self.card_gps, text="2 · GPS",
            font=ctk.CTkFont(size=12, weight="bold"), text_color=UI["text_dim"],
        ).pack(anchor="w", padx=12, pady=(10, 4))

        self.btn_browse_gpx = ctk.CTkButton(
            self.card_gps, text="🛰️  Trace GPX",
            command=self.select_gpx_file,
            fg_color=UI["card_alt"], hover_color=UI["border"],
            border_width=1, border_color=UI["border"],
            height=32, corner_radius=10,
        )
        self.btn_browse_gpx.pack(fill="x", padx=10, pady=(0, 6))
        self.lbl_step2 = ctk.CTkLabel(self.sidebar, text="")

        # Fond de carte : conservé mais discret (aussi accessible sur l'onglet carte)
        self.lbl_step3 = ctk.CTkLabel(self.sidebar, text="")
        self.options_cartes = ["Plan IGN (Moderne)", "Photos Aériennes", "OpenStreetMap"]
        map_row = ctk.CTkFrame(self.card_gps, fg_color="transparent")
        map_row.pack(fill="x", padx=10, pady=(0, 6))
        ctk.CTkLabel(map_row, text="Fond", font=ctk.CTkFont(size=11), text_color=UI["text_muted"]).pack(side="left")
        self.selecteur_carte = ctk.CTkOptionMenu(
            map_row, values=self.options_cartes, command=self.changer_fond_carte,
            fg_color=UI["card_alt"], button_color=UI["accent"], button_hover_color=UI["accent_hover"],
            height=28, width=160, corner_radius=8, font=ctk.CTkFont(size=11),
        )
        self.selecteur_carte.pack(side="right", fill="x", expand=True, padx=(8, 0))

        self.geosync_frame = ctk.CTkFrame(self.card_gps, fg_color="transparent")
        self.geosync_frame.pack(fill="x", padx=10, pady=(0, 4))
        self.geosync_frame.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(
            self.geosync_frame, text="Décalage photo → GPX",
            font=ctk.CTkFont(size=10), text_color=UI["text_muted"], anchor="w",
        ).grid(row=0, column=0, columnspan=2, sticky="ew")
        self.entry_geosync = ctk.CTkEntry(
            self.geosync_frame, placeholder_text="-1:00:00 ou -2:00:00", height=28,
        )
        self.entry_geosync.grid(row=1, column=0, sticky="ew", padx=(0, 4))
        ctk.CTkButton(
            self.geosync_frame, text="?", width=28, height=28,
            fg_color=UI.get("card_alt", "#3a3a3a"),
            command=self.show_geosync_help,
        ).grid(row=1, column=1, sticky="e")

        tz_row = ctk.CTkFrame(self.geosync_frame, fg_color="transparent")
        tz_row.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(6, 2))
        tz_row.grid_columnconfigure((0, 1), weight=1)
        ctk.CTkButton(
            tz_row, text="☀ Été −1 h", height=28,
            font=ctk.CTkFont(size=12, weight="bold"),
            fg_color=UI.get("warning", "#f0b45c"),
            hover_color=UI.get("warning_hover", "#d99a3e"),
            command=lambda: self.apply_geosync_preset("-1:00:00", "été"),
        ).grid(row=0, column=0, sticky="ew", padx=(0, 3))
        ctk.CTkButton(
            tz_row, text="❄ Hiver −2 h", height=28,
            font=ctk.CTkFont(size=12, weight="bold"),
            fg_color=UI.get("accent", "#4ec4b0"),
            hover_color=UI.get("accent_hover", "#3aab98"),
            command=lambda: self.apply_geosync_preset("-2:00:00", "hiver"),
        ).grid(row=0, column=1, sticky="ew", padx=(3, 0))

        # Options avancées repliées
        self.force_resync = ctk.BooleanVar(value=False)
        self.no_backup = ctk.BooleanVar(value=False)
        self._gps_advanced_visible = False
        self.btn_gps_advanced = ctk.CTkButton(
            self.card_gps, text="▸ Options avancées",
            height=24, font=ctk.CTkFont(size=11),
            fg_color="transparent", text_color=UI["text_muted"],
            hover_color=UI["card_alt"],
            command=self._toggle_gps_advanced,
        )
        self.btn_gps_advanced.pack(anchor="w", padx=8, pady=(2, 0))
        self.gps_advanced_frame = ctk.CTkFrame(self.card_gps, fg_color="transparent")
        self.check_force_resync = ctk.CTkCheckBox(
            self.gps_advanced_frame,
            text="Forcer la réécriture GPS",
            variable=self.force_resync, font=ctk.CTkFont(size=10),
        )
        self.check_force_resync.pack(anchor="w", padx=4, pady=2)
        self.check_no_backup = ctk.CTkCheckBox(
            self.gps_advanced_frame,
            text="Sans sauvegarde (irréversible)",
            variable=self.no_backup, font=ctk.CTkFont(size=10),
        )
        self.check_no_backup.pack(anchor="w", padx=4, pady=2)

        self.btn_sync = ctk.CTkButton(
            self.card_gps, text="⚡  Synchroniser GPS",
            fg_color=UI["success"], hover_color=UI["success_hover"],
            command=self.start_sync_thread,
            font=ctk.CTkFont(size=14, weight="bold"),
            height=40, corner_radius=12,
        )
        self.btn_sync.pack(fill="x", padx=10, pady=(8, 4))

        self.progress_bar = ctk.CTkProgressBar(
            self.card_gps, orientation="horizontal",
            progress_color=UI["accent"], fg_color=UI["card_alt"],
            height=6, corner_radius=3,
        )
        self.progress_bar.pack(fill="x", padx=10, pady=(2, 2))
        self.progress_bar.set(0)
        self.lbl_progress = ctk.CTkLabel(
            self.card_gps, text="Prêt",
            font=ctk.CTkFont(size=11), text_color=UI["text_muted"],
        )
        self.lbl_progress.pack(anchor="w", padx=12, pady=(0, 4))

        self.tools_frame = ctk.CTkFrame(self.card_gps, fg_color="transparent")
        self.tools_frame.pack(fill="x", padx=10, pady=(0, 10))
        self.btn_undo_sync = ctk.CTkButton(
            self.tools_frame, text="↩️ Annuler synchro",
            fg_color=UI["card_alt"], hover_color=UI["border"],
            command=self.undo_last_sync, state="disabled",
            font=ctk.CTkFont(size=11), height=28, corner_radius=8,
        )
        self.btn_undo_sync.pack(fill="x")

        # —— Bloc 3 saison (masqué en mode Sortie) ——
        self.card_saison = ctk.CTkFrame(
            self.sidebar, fg_color=UI["card"], corner_radius=14,
            border_width=1, border_color=UI["border"],
        )
        # pas de grid tant que mode Sortie
        ctk.CTkLabel(
            self.card_saison, text="Saison · analyse",
            font=ctk.CTkFont(size=12, weight="bold"), text_color=UI["text_dim"],
        ).pack(anchor="w", padx=12, pady=(10, 6))
        for lab, cmd, col in (
            ("📁  Projet Meuse", "open_project_hub", UI["accent"]),
            ("🗂️  Carte cumulée", "open_multi_sorties_map", UI["success"]),
            ("📊  Analyse / richesse", "open_analysis_panel", UI.get("warning", "#f0b45c")),
            ("🌿  Cibles de saison", "open_season_suggestions", UI.get("purple", "#b39af5")),
            ("📅  Brief multi-sorties", "open_aggregated_report", UI.get("purple", "#b39af5")),
        ):
            ctk.CTkButton(
                self.card_saison, text=lab, height=32, corner_radius=10,
                fg_color=col, hover_color=UI.get("accent_hover", col),
                command=getattr(self, cmd),
                font=ctk.CTkFont(size=12, weight="bold"),
            ).pack(fill="x", padx=10, pady=3)
        ctk.CTkFrame(self.card_saison, height=8, fg_color="transparent").pack()

        # Hub outils (accès discret)
        self.btn_outils_hub = ctk.CTkButton(
            self.sidebar, text="✨  Plus d'outils…",
            fg_color=UI["card_alt"], hover_color=UI["border"],
            command=self.open_outils_hub,
            font=ctk.CTkFont(size=12), height=30, corner_radius=10,
        )
        self.btn_outils_hub.grid(row=3, column=0, padx=12, pady=(4, 4), sticky="ew")

        # —— Journal ——
        self.log_box = ctk.CTkTextbox(
            self.sidebar, activate_scrollbars=True, fg_color=UI["card"],
            border_width=1, border_color=UI["border"], corner_radius=12,
            text_color=UI["text_dim"], font=ctk.CTkFont(size=11),
        )
        self.log_box.grid(row=4, column=0, padx=12, pady=(4, 12), sticky="nsew")
        self.log_box.insert(
            "0.0",
            "Bienvenue 🌿\n"
            "Mode Sortie : ouvrir un dossier → GPX → synchroniser → annoter.\n"
            "Mode Saison : projet, carte cumulée, analyses.\n\n",
        )
        self.log_box.configure(state="disabled")

        # --- ZONE PRINCIPALE (Droite) ---
        self.main_frame = ctk.CTkFrame(
            self, corner_radius=16, fg_color=UI["sidebar"],
            border_width=1, border_color=UI["border"]
        )
        self.main_frame.grid(row=1, column=1, sticky="nsew", padx=(6, 10), pady=(4, 10))

        # Bouton afficher/masquer la barre d'outils (bord gauche de la zone principale)
        self._sidebar_visible = True
        self.btn_toggle_sidebar = ctk.CTkButton(
            self.main_frame, text="◀", width=26, height=52,
            font=ctk.CTkFont(size=13, weight="bold"),
            fg_color=UI.get("card_elevated", "#2c3a32"),
            hover_color=UI.get("accent", "#4ec4b0"),
            border_width=1, border_color=UI.get("border", "#354840"),
            corner_radius=8,
            command=self.toggle_sidebar,
        )
        self.btn_toggle_sidebar.place(x=6, rely=0.42, anchor="w")
        try:
            self.btn_toggle_sidebar.lift()
        except Exception:
            pass
        self.main_frame.grid_rowconfigure(0, weight=1)
        self.main_frame.grid_columnconfigure(0, weight=1)

        self.tab_view = ctk.CTkTabview(
            self.main_frame, fg_color=UI["card"], segmented_button_fg_color=UI["card_alt"],
            segmented_button_selected_color=UI["accent"],
            segmented_button_selected_hover_color=UI["accent_hover"],
            segmented_button_unselected_color=UI["card_alt"],
            segmented_button_unselected_hover_color=UI["border"],
            text_color=UI["text"],
            corner_radius=14, border_width=0
        )
        self.tab_view.grid(row=0, column=0, sticky="nsew", padx=12, pady=12)
        
        self.tab_map = self.tab_view.add("🗺️  Carte")
        self.tab_notebook = self.tab_view.add("📝  Carnet")

        # Configuration Carte
        self.tab_map.grid_rowconfigure(0, weight=0)
        self.tab_map.grid_rowconfigure(1, weight=1)
        self.tab_map.grid_columnconfigure(0, weight=1)

        self.map_filter_frame = ctk.CTkFrame(self.tab_map, fg_color="transparent")
        self.map_filter_frame.grid(row=0, column=0, sticky="ew", padx=5, pady=(5, 0))
        ctk.CTkLabel(self.map_filter_frame, text="🔎 Filtrer par catégorie :", font=ctk.CTkFont(size=12, weight="bold")).pack(side="left", padx=(5, 12))

        for cat in list(self.faune_meuse.keys()) + ["Non classé"]:
            var = ctk.BooleanVar(value=True)
            couleur = CATEGORY_COLORS.get(cat, "#7f7f7f")
            cb = ctk.CTkCheckBox(
                self.map_filter_frame, text=cat, variable=var, command=self.refresh_map_markers,
                checkbox_width=18, checkbox_height=18, fg_color=couleur, hover_color=couleur
            )
            cb.pack(side="left", padx=6)
            self.category_filters[cat] = var

        self.var_afficher_trace = ctk.BooleanVar(value=False)
        self.trace_gpx_path_obj = None
        ctk.CTkCheckBox(
            self.map_filter_frame, text="🗺️ Trace GPX", variable=self.var_afficher_trace, command=self.toggle_trace_gpx,
            checkbox_width=18, checkbox_height=18, fg_color="#e74c3c", hover_color="#c0392b"
        ).pack(side="left", padx=(18, 6))

        ctk.CTkButton(
            self.map_filter_frame, text="🔄 Rafraîchir", width=90, height=24, font=ctk.CTkFont(size=11),
            fg_color="#2b2b2b", hover_color="#3a3a3a", command=self.recharger_donnees_gps
        ).pack(side="left", padx=(12, 6))

        ctk.CTkLabel(
            self.map_filter_frame,
            text="  ● animal  ▲ empreinte  ■ terrier  ◆ coulée  |  ✦ sans photo (rose)",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim", "#666666")
        ).pack(side="left", padx=(10, 4))

        self.btn_place_obs = ctk.CTkButton(
            self.map_filter_frame, text="➕ Placer observation", width=150, height=24,
            font=ctk.CTkFont(size=11), fg_color=UI.get("pink", SANS_PHOTO_COLOR), hover_color=UI.get("pink_hover", "#c2185b"),
            command=self.toggle_place_observation_mode
        )
        self.btn_place_obs.pack(side="left", padx=(12, 4))
        self.btn_point_ecoute = ctk.CTkButton(
            self.map_filter_frame, text="🎧 Point d'écoute GPS", width=150, height=24,
            font=ctk.CTkFont(size=11), fg_color=UI.get("accent"),
            command=self.toggle_point_ecoute_gps,
        )
        self.btn_point_ecoute.pack(side="left", padx=(4, 4))
        self._place_obs_mode = False
        self._point_ecoute_mode = False

        self.map_widget = TkinterMapView(self.tab_map, corner_radius=10)
        self.map_widget.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
        self.map_widget.set_position(49.1627, 5.3854) 
        self.map_widget.set_zoom(9)
        self.changer_fond_carte("Plan IGN (Moderne)")
        # Clic gauche sur la carte → placer une observation (si le mode est activé)
        try:
            self.map_widget.add_left_click_map_command(self._on_map_left_click)
        except Exception:
            pass

        # Configuration Carnet de Terrain
        self.tab_notebook.grid_columnconfigure(0, weight=1, minsize=250) 
        self.tab_notebook.grid_columnconfigure(1, weight=2)              
        self.tab_notebook.grid_rowconfigure(0, weight=0)
        self.tab_notebook.grid_rowconfigure(1, weight=0)
        self.tab_notebook.grid_rowconfigure(2, weight=1)

        self.carnet_actions_frame = ctk.CTkFrame(self.tab_notebook, fg_color="transparent")
        self.carnet_actions_frame.grid(row=0, column=0, sticky="ew", padx=5, pady=(5, 2))
        self.carnet_actions_frame.grid_columnconfigure((0, 1, 2, 3), weight=1)

        self.btn_ai_summary = ctk.CTkButton(
            self.carnet_actions_frame, text="🤖  Brief IA", fg_color=UI["purple"], hover_color=UI["purple_hover"],
            command=self.generate_ai_summary, font=ctk.CTkFont(size=12, weight="bold"),
            height=34, corner_radius=12
        )
        self.btn_ai_summary.grid(row=0, column=0, sticky="ew", padx=(0, 2))

        self.btn_manual_obs = ctk.CTkButton(
            self.carnet_actions_frame, text="👁️  Sans photo", fg_color=UI["warning"], hover_color=UI["warning_hover"],
            command=self.open_manual_observation_dialog, font=ctk.CTkFont(size=12, weight="bold"),
            height=34, corner_radius=12
        )
        self.btn_manual_obs.grid(row=0, column=1, sticky="ew", padx=2)

        self.btn_import_son = ctk.CTkButton(
            self.carnet_actions_frame, text="🎧  Import son",
            fg_color=UI.get("success", "#2f9e5f"), hover_color=UI.get("accent_hover", "#3da892"),
            command=self.open_import_son_menu, font=ctk.CTkFont(size=12, weight="bold"),
            height=34, corner_radius=12
        )
        self.btn_import_son.grid(row=0, column=2, sticky="ew", padx=2)

        self.btn_edit_gps = ctk.CTkButton(
            self.carnet_actions_frame, text="📍  GPS", fg_color=UI["accent"], hover_color=UI["accent_hover"],
            command=self.open_edit_gps_dialog, font=ctk.CTkFont(size=12, weight="bold"),
            height=34, corner_radius=12
        )
        self.btn_edit_gps.grid(row=0, column=3, sticky="ew", padx=(2, 0))

        self.entry_search_photo = ctk.CTkEntry(
            self.tab_notebook,
            placeholder_text="🔎 Rechercher (nom, espèce, catégorie, lieu...)",
            height=32, corner_radius=10,
        )
        self.entry_search_photo.grid(row=1, column=0, sticky="ew", padx=5, pady=(0, 2))
        self.entry_search_photo.bind("<KeyRelease>", self.filter_photo_list)

        # Explorateur média : barre de mode + liste ou miniatures
        self.media_explorer = ctk.CTkFrame(
            self.tab_notebook, fg_color=UI["card"], corner_radius=12,
            border_width=1, border_color=UI["border"],
        )
        self.media_explorer.grid(row=2, column=0, sticky="nsew", padx=5, pady=(0, 5))
        self.media_explorer.grid_rowconfigure(1, weight=1)
        self.media_explorer.grid_columnconfigure(0, weight=1)

        mode_bar = ctk.CTkFrame(self.media_explorer, fg_color="transparent")
        mode_bar.grid(row=0, column=0, sticky="ew", padx=6, pady=(6, 2))
        ctk.CTkLabel(
            mode_bar, text="Explorateur",
            font=ctk.CTkFont(size=11, weight="bold"), text_color=UI["text_dim"]
        ).pack(side="left", padx=(2, 8))
        self.btn_view_list = ctk.CTkButton(
            mode_bar, text="☰ Liste", width=72, height=26, font=ctk.CTkFont(size=11),
            fg_color=UI["accent"], hover_color=UI["accent_hover"],
            command=lambda: self.set_carnet_view_mode("list"),
        )
        self.btn_view_list.pack(side="left", padx=2)
        self.btn_view_thumbs = ctk.CTkButton(
            mode_bar, text="▢ Miniatures", width=96, height=26, font=ctk.CTkFont(size=11),
            fg_color=UI["card_alt"], hover_color=UI["border"],
            command=lambda: self.set_carnet_view_mode("thumbs"),
        )
        self.btn_view_thumbs.pack(side="left", padx=2)
        self.lbl_explorer_count = ctk.CTkLabel(
            mode_bar, text="", font=ctk.CTkFont(size=10), text_color=UI["text_muted"]
        )
        self.lbl_explorer_count.pack(side="right", padx=4)

        # Liste (rapide, léger)
        self.photo_listbox = tk.Listbox(
            self.media_explorer,
            bg=UI["card_alt"], fg=UI.get("list_pending", UI["text"]),
            selectbackground=UI["accent"], selectforeground=UI.get("bg", "#0a1612"),
            borderwidth=0, highlightthickness=0,
            font=("Segoe UI", 10), selectmode=tk.EXTENDED, exportselection=False,
            activestyle="none",
        )
        self.photo_listbox.grid(row=1, column=0, sticky="nsew", padx=6, pady=(2, 6))
        self.photo_listbox.bind("<<ListboxSelect>>", self.on_photo_select)
        self.photo_listbox.bind("<Double-Button-1>", self._on_media_double_click)

        # Grille miniatures (lazy) — masquée par défaut
        self.thumb_scroll = ctk.CTkScrollableFrame(
            self.media_explorer, fg_color=UI["card_alt"], corner_radius=8
        )
        # pas de grid tant que mode list
        self.carnet_view_mode = "list"
        self._thumb_widgets = {}          # filename -> bouton
        self._thumb_photo_refs = {}       # garder refs ImageTk
        self._thumb_load_token = 0        # annule charges obsolètes
        self._filtered_media_names = []   # noms fichiers sans préfixe 🎬

        self.edit_panel = ctk.CTkScrollableFrame(self.tab_notebook, fg_color="transparent")
        self.edit_panel.grid(row=0, column=1, rowspan=3, sticky="nsew", padx=5, pady=5)

        self.compteur_frame = ctk.CTkFrame(self.edit_panel, fg_color=UI.get("card", "#1a2332"), corner_radius=12, border_width=1, border_color=UI.get("border", "#2a3544"))
        self.compteur_frame.pack(fill="x", pady=(0, 8))

        compteur_header = ctk.CTkFrame(self.compteur_frame, fg_color="transparent")
        compteur_header.pack(fill="x", padx=10, pady=(8, 2))
        ctk.CTkLabel(compteur_header, text="📊 Compteur du jour (manuel)", font=ctk.CTkFont(size=12, weight="bold")).pack(side="left")
        ctk.CTkButton(
            compteur_header, text="🔄 Depuis les photos", width=120, height=22, font=ctk.CTkFont(size=10),
            fg_color="#2b2b2b", hover_color="#3a3a3a", command=self.initialiser_compteur_depuis_photos
        ).pack(side="right", padx=(4, 0))
        ctk.CTkButton(
            compteur_header, text="+ Espèce", width=80, height=22, font=ctk.CTkFont(size=10),
            fg_color="#2b2b2b", hover_color="#3a3a3a", command=self.ajouter_espece_au_compteur
        ).pack(side="right")

        self.compteur_lignes_frame = ctk.CTkFrame(self.compteur_frame, fg_color="transparent")
        self.compteur_lignes_frame.pack(fill="x", padx=10, pady=(0, 8))

        self.lbl_compteur_hint = ctk.CTkLabel(
            self.compteur_frame, text="", font=ctk.CTkFont(size=10), text_color=UI.get("text_dim", "#666666"),
            justify="left", anchor="w", wraplength=420
        )
        self.lbl_compteur_hint.pack(anchor="w", padx=10, pady=(0, 8), fill="x")

        self.upper_panel = ctk.CTkFrame(self.edit_panel, fg_color="transparent")
        self.upper_panel.pack(fill="x", pady=5)

        self.preview_container = ctk.CTkFrame(self.upper_panel, fg_color="transparent")
        self.preview_container.pack(side="left", padx=5, anchor="n")

        self.lbl_preview = ctk.CTkLabel(
            self.preview_container,
            text="Sélectionnez une photo\nou une vidéo",
            width=260, height=195, fg_color=UI["card_alt"], corner_radius=10,
            text_color=UI["text_dim"],
        )
        self.lbl_preview.pack()
        self.lbl_preview.bind("<Double-Button-1>", lambda e: self.open_media_large())
        self.lbl_preview.bind("<Enter>", self._on_preview_hover_enter)
        self.lbl_preview.bind("<Leave>", self._on_preview_hover_leave)
        try:
            self.lbl_preview.configure(cursor="hand2")
        except Exception:
            pass

        self.preview_btn_row = ctk.CTkFrame(self.preview_container, fg_color="transparent")
        self.preview_btn_row.pack(fill="x", padx=0, pady=(4, 0))
        self.btn_play_video = ctk.CTkButton(
            self.preview_btn_row, text="▶  Lire la vidéo",
            fg_color=UI.get("purple", "#9b59b6"), hover_color=UI.get("purple_hover", "#8e44ad"),
            height=30, corner_radius=10, command=self.open_video_external,
        )
        self.btn_open_large = ctk.CTkButton(
            self.preview_btn_row, text="🔍 Agrandir",
            fg_color=UI["accent"], hover_color=UI["accent_hover"],
            height=30, corner_radius=10, command=self.open_media_large,
        )
        self.btn_open_large.pack(side="left", fill="x", expand=True, padx=(0, 2))
        # play vidéo empilé à la demande

        self.lbl_thumb_size_value = ctk.CTkLabel(self.preview_container, text="Taille de l'aperçu : 260 px", font=ctk.CTkFont(size=10))
        self.lbl_thumb_size_value.pack(pady=(6, 0))

        self.slider_thumb_size = ctk.CTkSlider(
            self.preview_container, from_=100, to=600, number_of_steps=25,
            command=self.on_thumb_size_change, width=260
        )
        self.slider_thumb_size.set(260)
        self.slider_thumb_size.pack(pady=(2, 0))

        self.weather_box = ctk.CTkTextbox(self.upper_panel, height=150, width=300, activate_scrollbars=False)
        self.weather_box.pack(side="left", fill="x", expand=True, padx=5)
        self.weather_box.insert("0.0", "Météo : En attente du géotagging...")
        self.weather_box.configure(state="disabled")

        # --- FORMULAIRE DE SAISIE ---
        self.form_frame = ctk.CTkFrame(
            self.edit_panel, corner_radius=10, border_width=1,
            border_color=UI.get("border", "#333333"), fg_color=UI.get("card", "#1c2620"),
        )
        self.form_frame.pack(fill="x", padx=5, pady=10)

        ctk.CTkLabel(self.form_frame, text="Catégorie :", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, padx=15, pady=10, sticky="w")
        self.choice_category = ctk.CTkOptionMenu(self.form_frame, values=list(self.faune_meuse.keys()), command=self.on_category_change)
        self.choice_category.grid(row=0, column=1, padx=15, pady=10, sticky="ew")

        ctk.CTkLabel(self.form_frame, text="Espèce (Meuse 55) :", font=ctk.CTkFont(weight="bold")).grid(row=1, column=0, padx=15, pady=10, sticky="w")
        
        # Piste 1 : Conteneur pour aligner menu déroulant espèce + bouton d'ajout "+"
        self.species_frame = ctk.CTkFrame(self.form_frame, fg_color="transparent")
        self.species_frame.grid(row=1, column=1, padx=15, pady=10, sticky="ew")
        self.species_frame.grid_columnconfigure(0, weight=1)

        self.choice_species = ctk.CTkOptionMenu(self.species_frame, values=["Sélectionnez..."])
        self.choice_species.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        
        self.btn_add_custom_species = ctk.CTkButton(self.species_frame, text="+", width=30, fg_color="#2b2b2b", hover_color="#3a3a3a", command=self.add_custom_species)
        self.btn_add_custom_species.grid(row=0, column=1, sticky="e", padx=(0, 5))

        self.btn_remove_species = ctk.CTkButton(self.species_frame, text="✕", width=30, fg_color="#3a2020", hover_color="#4a2828", text_color="#dd8888", command=self.remove_current_species)
        self.btn_remove_species.grid(row=0, column=2, sticky="e", padx=(0, 5))

        self.btn_suggest_ai = ctk.CTkButton(self.species_frame, text="🤖", width=30, fg_color="#59339d", hover_color="#452878", command=self.suggest_species_ai)
        self.btn_suggest_ai.grid(row=0, column=3, sticky="e")

        ctk.CTkLabel(self.form_frame, text="Nombre :", font=ctk.CTkFont(weight="bold")).grid(row=2, column=0, padx=15, pady=10, sticky="w")
        self.choice_count = ctk.CTkComboBox(
            self.form_frame,
            values=["1", "2", "3", "4", "5", "6-10", "10+", "Nombreux", "Vol important", "Indéterminé"]
        )
        self.choice_count.grid(row=2, column=1, padx=15, pady=10, sticky="ew")
        self.choice_count.set("1")

        ctk.CTkLabel(self.form_frame, text="Heure :", font=ctk.CTkFont(weight="bold")).grid(row=3, column=0, padx=15, pady=10, sticky="w")
        self.entry_time = ctk.CTkEntry(self.form_frame, placeholder_text="HH:MM")
        self.entry_time.grid(row=3, column=1, padx=15, pady=10, sticky="ew")

        ctk.CTkLabel(self.form_frame, text="Lieu :", font=ctk.CTkFont(weight="bold")).grid(row=4, column=0, padx=15, pady=10, sticky="w")
        self.lieu_frame = ctk.CTkFrame(self.form_frame, fg_color="transparent")
        self.lieu_frame.grid(row=4, column=1, padx=15, pady=10, sticky="ew")
        self.lieu_frame.grid_columnconfigure(0, weight=1)

        self.entry_lieu = ctk.CTkEntry(self.lieu_frame, placeholder_text="Localisation automatique...")
        self.entry_lieu.grid(row=0, column=0, sticky="ew", padx=(0, 5))

        self.btn_relocaliser = ctk.CTkButton(self.lieu_frame, text="📍", width=30, fg_color="#2b2b2b", hover_color="#3a3a3a", command=self.relocaliser_photo)
        self.btn_relocaliser.grid(row=0, column=1, sticky="e")

        ctk.CTkLabel(self.form_frame, text="Comportement :", font=ctk.CTkFont(weight="bold")).grid(
            row=5, column=0, padx=15, pady=6, sticky="w"
        )
        self.choice_comportement = ctk.CTkComboBox(self.form_frame, values=COMPORTEMENTS_OBS)
        self.choice_comportement.grid(row=5, column=1, padx=15, pady=6, sticky="ew")
        self.choice_comportement.set("")

        ctk.CTkLabel(self.form_frame, text="Certitude :", font=ctk.CTkFont(weight="bold")).grid(
            row=6, column=0, padx=15, pady=6, sticky="w"
        )
        self.choice_certitude = ctk.CTkComboBox(self.form_frame, values=CERTITUDES_OBS)
        self.choice_certitude.grid(row=6, column=1, padx=15, pady=6, sticky="ew")
        self.choice_certitude.set("")

        ctk.CTkLabel(self.form_frame, text="Type d'indice :", font=ctk.CTkFont(weight="bold")).grid(
            row=7, column=0, padx=15, pady=6, sticky="w"
        )
        self.choice_type_indice = ctk.CTkComboBox(self.form_frame, values=TYPES_INDICE)
        self.choice_type_indice.grid(row=7, column=1, padx=15, pady=6, sticky="ew")
        self.choice_type_indice.set("")

        ctk.CTkLabel(self.form_frame, text="Nom scientifique :", font=ctk.CTkFont(weight="bold")).grid(
            row=8, column=0, padx=15, pady=6, sticky="w"
        )
        sci_row = ctk.CTkFrame(self.form_frame, fg_color="transparent")
        sci_row.grid(row=8, column=1, padx=15, pady=6, sticky="ew")
        sci_row.grid_columnconfigure(0, weight=1)
        self.entry_nom_sci = ctk.CTkEntry(sci_row, placeholder_text="optionnel — ex. Capreolus capreolus")
        self.entry_nom_sci.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        self.btn_inpn = ctk.CTkButton(
            sci_row, text="INPN", width=52, height=28,
            fg_color=UI.get("card_alt", "#333"), command=self.open_inpn_search,
        )
        self.btn_inpn.grid(row=0, column=1)

        # Liens Amazon Photos (affichage + ouverture navigateur)
        ctk.CTkLabel(self.form_frame, text="Lien Amazon (photo) :", font=ctk.CTkFont(weight="bold")).grid(
            row=9, column=0, padx=15, pady=6, sticky="w"
        )
        amz_row = ctk.CTkFrame(self.form_frame, fg_color="transparent")
        amz_row.grid(row=9, column=1, padx=15, pady=6, sticky="ew")
        amz_row.grid_columnconfigure(0, weight=1)
        self.entry_amazon_photo = ctk.CTkEntry(
            amz_row, placeholder_text="https://www.amazon.fr/photos/... (cette image)",
        )
        self.entry_amazon_photo.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        self.btn_open_amazon_photo = ctk.CTkButton(
            amz_row, text="Ouvrir", width=64, height=28,
            fg_color=UI.get("accent", "#3d9cf0"), command=self.open_amazon_photo_link,
        )
        self.btn_open_amazon_photo.grid(row=0, column=1)

        ctk.CTkLabel(self.form_frame, text="Repère Amazon (jour) :", font=ctk.CTkFont(weight="bold")).grid(
            row=10, column=0, padx=15, pady=6, sticky="w"
        )
        alb_row = ctk.CTkFrame(self.form_frame, fg_color="transparent")
        alb_row.grid(row=10, column=1, padx=15, pady=6, sticky="ew")
        alb_row.grid_columnconfigure(0, weight=1)
        self.entry_amazon_album = ctk.CTkEntry(
            alb_row,
            placeholder_text="ex. 19 juillet 2026 — jour dans Amazon Photos (pas d'album)",
        )
        self.entry_amazon_album.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        self.btn_fill_amazon_day = ctk.CTkButton(
            alb_row, text="Jour auto", width=72, height=28,
            fg_color=UI.get("card_alt", "#333"), command=self.fill_amazon_day_from_sortie,
        )
        self.btn_fill_amazon_day.grid(row=0, column=1, padx=(0, 4))
        self.btn_save_amazon_album = ctk.CTkButton(
            alb_row, text="Sauver", width=64, height=28,
            fg_color=UI.get("success", "#2f9e5f"), command=self.save_amazon_album_link,
        )
        self.btn_save_amazon_album.grid(row=0, column=2)

        self.burst_frame = ctk.CTkFrame(self.edit_panel, fg_color="#3a2f0f", corner_radius=8)
        self.lbl_burst_info = ctk.CTkLabel(self.burst_frame, text="", font=ctk.CTkFont(size=11), text_color="#f1c40f")
        self.lbl_burst_info.pack(side="left", padx=10, pady=8)
        self.btn_apply_burst = ctk.CTkButton(
            self.burst_frame, text="Appliquer à toute la rafale", width=190, fg_color="#8a4b08", hover_color="#6b3a06",
            command=self.apply_to_burst, font=ctk.CTkFont(size=11)
        )
        self.btn_apply_burst.pack(side="right", padx=10, pady=6)
        self.current_burst_files = []
        # (le burst_frame n'est affiché que lorsqu'une rafale est détectée, voir on_photo_select)

        self.lbl_details_header = ctk.CTkLabel(self.edit_panel, text="Détails additionnels (Comportement, habitat...) :", font=ctk.CTkFont(size=12, weight="bold"))
        self.lbl_details_header.pack(anchor="w", padx=5, pady=(10, 2))
        self.note_text = ctk.CTkTextbox(self.edit_panel, height=80, activate_scrollbars=True)
        self.note_text.pack(fill="x", padx=5, pady=5)

        self.btn_save_note = ctk.CTkButton(self.edit_panel, text="💾 Enregistrer l'Observation", fg_color="#1f7d37", hover_color="#145a27", command=self.save_current_note, font=ctk.CTkFont(weight="bold"))
        self.btn_save_note.pack(anchor="e", padx=5, pady=10)

        self.on_category_change("Mammifère")

        # Raccourcis clavier (configurables — voir Affichage → Raccourcis)
        self._shortcut_bindings = []  # list of sequences currently bound
        self.after(200, self._setup_keyboard_shortcuts)

        # Lot 1 : Sauvegarde automatique périodique
        self.after(60000, self._autosave_tick)
        # Espace de travail (géométries / détachements) si déjà enregistré
        self.after(900, lambda: self.restore_workspace_layout(silent=True))

    # —— Fenêtres outils au premier plan + raccourcis ——
    DEFAULT_SHORTCUTS = {
        "save_note": "<Control-s>",
        "prev_photo": "<Control-Left>",
        "next_photo": "<Control-Right>",
        "open_folder": "<Control-o>",
        "open_gpx": "<Control-g>",
        "sync_gps": "<Control-y>",
        "manual_obs": "<Control-n>",
        "brief_ia": "<F2>",
        "import_son": "<Control-i>",
        "edit_gps": "<Control-e>",
        "multi_map": "<Control-m>",
        "sidebar": "<Control-b>",
        "tools_hub": "<Control-t>",
        "fullscreen": "<F11>",
        "fullscreen2": "<F12>",
        "shortcuts_cfg": "<F1>",
    }

    SHORTCUT_LABELS = {
        "save_note": "Enregistrer l'observation",
        "prev_photo": "Photo / obs. précédente",
        "next_photo": "Photo / obs. suivante",
        "open_folder": "Ouvrir un dossier sortie",
        "open_gpx": "Charger une trace GPX",
        "sync_gps": "Synchroniser photos & GPX",
        "manual_obs": "Observation sans photo",
        "brief_ia": "Brief IA de la sortie",
        "import_son": "Import son (menu)",
        "edit_gps": "Éditer le GPS",
        "multi_map": "Carte cumulée",
        "sidebar": "Afficher / masquer barre latérale",
        "tools_hub": "Plus d'outils…",
        "fullscreen": "Plein écran",
        "fullscreen2": "Plein écran (alt.)",
        "shortcuts_cfg": "Configurer les raccourcis",
    }

    def _prepare_tool_window(self, win, *, stay_on_top=None):
        """Place une fenêtre outil devant le carnet (transient + focus + option topmost)."""
        if win is None:
            return
        cfg = {}
        try:
            cfg = self._load_app_config() or {}
        except Exception:
            pass
        if stay_on_top is None:
            stay_on_top = bool(cfg.get("tool_windows_on_top", True))
        try:
            win.transient(self)
        except Exception:
            pass
        try:
            win.lift()
        except Exception:
            pass
        if stay_on_top:
            try:
                win.attributes("-topmost", True)
            except Exception:
                pass

        def _front():
            try:
                if not win.winfo_exists():
                    return
                win.lift()
                win.focus_force()
            except Exception:
                pass

        try:
            win.after(40, _front)
            win.after(180, _front)
        except Exception:
            _front()
        # Laisser l'utilisateur basculer vers la carte si besoin : topmost soft
        if stay_on_top:
            def _release_topmost():
                # garder topmost si option active — ne pas relâcher
                # (sinon la fenêtre passe derrière dès un clic ailleurs)
                pass
            try:
                win.after(300, _release_topmost)
            except Exception:
                pass

    def _default_shortcuts(self):
        return dict(self.DEFAULT_SHORTCUTS)

    def _load_shortcuts(self):
        cfg = self._load_app_config() or {}
        sc = dict(self._default_shortcuts())
        user = cfg.get("shortcuts") or {}
        if isinstance(user, dict):
            for k, v in user.items():
                if k in sc and isinstance(v, str) and v.strip():
                    sc[k] = v.strip()
        return sc

    def _save_shortcuts(self, shortcuts):
        cfg = dict(self._load_app_config() or {})
        cfg["shortcuts"] = dict(shortcuts or {})
        self._save_app_config(cfg)

    def _shortcut_actions(self):
        """Map action_id → callable."""
        return {
            "save_note": lambda: self.save_current_note(silent=False) if hasattr(self, "save_current_note") else None,
            "prev_photo": lambda: self.navigate_photo(-1),
            "next_photo": lambda: self.navigate_photo(1),
            "open_folder": self.select_photo_folder,
            "open_gpx": self.select_gpx_file,
            "sync_gps": self.start_sync_thread,
            "manual_obs": self.open_manual_observation_dialog,
            "brief_ia": self.generate_ai_summary,
            "import_son": self.open_import_son_menu,
            "edit_gps": self.open_edit_gps_dialog,
            "multi_map": self.open_multi_sorties_map,
            "sidebar": self.toggle_sidebar,
            "tools_hub": self.open_outils_hub,
            "fullscreen": self.toggle_fullscreen,
            "fullscreen2": self.toggle_fullscreen,
            "shortcuts_cfg": self.open_shortcuts_settings,
        }

    def _unbind_keyboard_shortcuts(self):
        for seq in list(getattr(self, "_shortcut_bindings", []) or []):
            try:
                self.unbind_all(seq)
            except Exception:
                pass
        self._shortcut_bindings = []

    def _setup_keyboard_shortcuts(self):
        """(Ré)applique les raccourcis depuis la config."""
        self._unbind_keyboard_shortcuts()
        sc = self._load_shortcuts()
        actions = self._shortcut_actions()
        bound = []
        for action, seq in sc.items():
            fn = actions.get(action)
            if not fn or not seq:
                continue
            # Éviter double bind sur la même séquence
            if seq in bound:
                continue

            def _make(handler):
                def _cb(event=None):
                    try:
                        handler()
                    except Exception as e:
                        try:
                            self.log("Raccourci : %s" % e)
                        except Exception:
                            pass
                    return "break"
                return _cb

            try:
                self.bind_all(seq, _make(fn))
                bound.append(seq)
            except Exception as e:
                try:
                    self.log("Raccourci invalide %s → %s : %s" % (action, seq, e))
                except Exception:
                    pass
        # Échap quitte toujours le plein écran
        try:
            self.bind_all("<Escape>", self._quit_fullscreen)
            bound.append("<Escape>")
        except Exception:
            pass
        self._shortcut_bindings = bound

    def open_shortcuts_settings(self):
        """Configurer les raccourcis clavier et l'option fenêtres au premier plan."""
        win = ctk.CTkToplevel(self)
        win.title("Raccourcis clavier")
        win.geometry("640x720")
        try:
            win.minsize(520, 480)
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Raccourcis clavier",
            font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text="Format Tk : <Control-s>, <Control-Shift-S>, <F2>, <Control-Left>…\n"
                 "Les modifications s'appliquent après « Enregistrer ».",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            justify="left",
        ).pack(anchor="w", padx=16, pady=(0, 8))

        cfg0 = self._load_app_config() or {}
        topmost_var = tk.BooleanVar(value=bool(cfg0.get("tool_windows_on_top", True)))
        top_fr = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=10)
        top_fr.pack(fill="x", padx=16, pady=(0, 8))
        ctk.CTkCheckBox(
            top_fr,
            text="Garder les fenêtres outils au premier plan (recommandé)",
            variable=topmost_var,
            text_color=UI.get("text"),
        ).pack(anchor="w", padx=12, pady=10)

        sc = self._load_shortcuts()
        scroll = ctk.CTkScrollableFrame(win, fg_color=UI.get("card"), corner_radius=10)
        scroll.pack(fill="both", expand=True, padx=16, pady=4)
        entries = {}
        for key in self.DEFAULT_SHORTCUTS:
            row = ctk.CTkFrame(scroll, fg_color="transparent")
            row.pack(fill="x", padx=8, pady=3)
            ctk.CTkLabel(
                row, text=self.SHORTCUT_LABELS.get(key, key),
                width=240, anchor="w", text_color=UI.get("text"),
            ).pack(side="left")
            var = tk.StringVar(value=sc.get(key, self.DEFAULT_SHORTCUTS[key]))
            ent = ctk.CTkEntry(row, textvariable=var, width=200)
            ent.pack(side="left", padx=6)
            entries[key] = var

        def save():
            new_sc = {}
            for k, var in entries.items():
                v = (var.get() or "").strip()
                if not v:
                    v = self.DEFAULT_SHORTCUTS[k]
                new_sc[k] = v
            self._save_shortcuts(new_sc)
            cfg = dict(self._load_app_config() or {})
            cfg["tool_windows_on_top"] = bool(topmost_var.get())
            self._save_app_config(cfg)
            self._setup_keyboard_shortcuts()
            messagebox.showinfo("Raccourcis", "Raccourcis enregistrés et appliqués.")
            try:
                win.destroy()
            except Exception:
                pass

        def reset():
            for k, var in entries.items():
                var.set(self.DEFAULT_SHORTCUTS[k])
            topmost_var.set(True)

        actions = ctk.CTkFrame(win, fg_color="transparent")
        actions.pack(fill="x", padx=16, pady=12)
        ctk.CTkButton(
            actions, text="Enregistrer", height=36,
            fg_color=UI.get("success", "#2f9e5f"), command=save,
        ).pack(side="left", padx=(0, 8))
        ctk.CTkButton(
            actions, text="Réinitialiser", height=36,
            fg_color=UI.get("card_alt"), command=reset,
        ).pack(side="left", padx=4)
        ctk.CTkButton(actions, text="Fermer", height=36, width=90, command=win.destroy).pack(side="right")

    def _maximize_window(self):
        """Ouvre l'application en grande fenêtre (maximisée) au démarrage."""
        try:
            self.state("zoomed")  # Windows et certains environnements Linux
        except Exception:
            try:
                self.attributes("-zoomed", True)  # repli Linux (certains WM)
            except Exception:
                # Repli ultime : occuper la quasi-totalité de l'écran
                sw = self.winfo_screenwidth()
                sh = self.winfo_screenheight()
                self.geometry(f"{sw - 60}x{sh - 80}+30+20")

    # --- PISTE 1 : GESTION DES ESPÈCES PERSONNALISÉES ---
    def open_species_manager(self):
        """Fenêtre de gestion complète du dictionnaire d'espèces : ajouter, renommer, supprimer,
        pour chacune des 4 catégories (Mammifère, Oiseau, Insecte, Autre)."""
        win = ctk.CTkToplevel(self)
        win.title("🗂️ Gérer les espèces")
        win.geometry("520x680")
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Renommer/supprimer ne modifie que la liste de choix : les observations déjà "
                      "enregistrées gardent le nom tel qu'il était au moment de la saisie.",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim", "#666666"), wraplength=470, justify="left"
        ).pack(anchor="w", padx=15, pady=(15, 10))

        top = ctk.CTkFrame(win, fg_color="transparent")
        top.pack(fill="x", padx=15)
        ctk.CTkLabel(top, text="Catégorie :", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=(0, 8))
        choix_cat = ctk.CTkOptionMenu(top, values=list(self.faune_meuse.keys()))
        choix_cat.pack(side="left", fill="x", expand=True)

        liste_frame = ctk.CTkScrollableFrame(win)
        liste_frame.pack(fill="both", expand=True, padx=15, pady=12)

        def rafraichir_liste():
            for w in liste_frame.winfo_children():
                w.destroy()
            cat = choix_cat.get()
            especes = sorted(self.faune_meuse.get(cat, []), key=lambda s: s.lower())
            if not especes:
                ctk.CTkLabel(liste_frame, text="Aucune espèce dans cette catégorie.", text_color=UI.get("text_dim", "#666666")).pack(anchor="w", pady=6)
                return
            for espece in especes:
                ligne = ctk.CTkFrame(liste_frame, fg_color="#1e2c22", corner_radius=6)
                ligne.pack(fill="x", pady=2)
                ctk.CTkLabel(ligne, text=espece, anchor="w", font=ctk.CTkFont(size=12)).pack(side="left", padx=(10, 4), pady=8, fill="x", expand=True)
                ctk.CTkButton(
                    ligne, text="✎", width=30, height=26, fg_color="#2b2b2b", hover_color="#3a3a3a",
                    command=lambda e=espece: renommer(e)
                ).pack(side="left", padx=2)
                ctk.CTkButton(
                    ligne, text="✕", width=30, height=26, fg_color="transparent", hover_color="#3a2020", text_color="#aa6666",
                    command=lambda e=espece: supprimer(e)
                ).pack(side="left", padx=(2, 8))

        def ajouter():
            cat = choix_cat.get()
            nom = simpledialog.askstring("Ajouter une espèce", f"Nom de la nouvelle espèce dans « {cat} » :", parent=win)
            if not nom or not nom.strip():
                return
            nom = nom.strip()
            self.faune_meuse.setdefault(cat, [])
            if nom in self.faune_meuse[cat]:
                messagebox.showinfo("Déjà présente", f"« {nom} » existe déjà dans cette catégorie.")
                return
            self.faune_meuse[cat].append(nom)
            self.faune_meuse[cat].sort(key=lambda s: s.lower())
            self.save_species_dict()
            rafraichir_liste()
            self._synchroniser_dropdown_especes()

        def renommer(ancien_nom):
            cat = choix_cat.get()
            nouveau = simpledialog.askstring("Renommer l'espèce", f"Nouveau nom pour « {ancien_nom} » :", initialvalue=ancien_nom, parent=win)
            if not nouveau or not nouveau.strip() or nouveau.strip() == ancien_nom:
                return
            nouveau = nouveau.strip()
            liste = self.faune_meuse.get(cat, [])
            if nouveau in liste:
                messagebox.showinfo("Déjà présente", f"« {nouveau} » existe déjà dans cette catégorie.")
                return
            idx = liste.index(ancien_nom)
            liste[idx] = nouveau
            liste.sort(key=lambda s: s.lower())
            self.save_species_dict()
            rafraichir_liste()
            self._synchroniser_dropdown_especes()
            self.log(f"✎ Espèce renommée dans le dictionnaire : « {ancien_nom} » → « {nouveau} » (les observations déjà enregistrées ne sont pas modifiées).")

        def supprimer(espece):
            cat = choix_cat.get()
            if not messagebox.askyesno("Confirmer la suppression", f"Retirer « {espece} » de la liste « {cat} » ?\n(les observations déjà enregistrées avec ce nom ne sont pas touchées)"):
                return
            if espece in self.faune_meuse.get(cat, []):
                self.faune_meuse[cat].remove(espece)
                self.save_species_dict()
            rafraichir_liste()
            self._synchroniser_dropdown_especes()

        choix_cat.configure(command=lambda c: rafraichir_liste())
        rafraichir_liste()

        ctk.CTkButton(win, text="➕ Ajouter une espèce", command=ajouter, fg_color="#2ba14b", hover_color="#1f7d37", font=ctk.CTkFont(weight="bold")).pack(fill="x", padx=15, pady=(0, 15))

    def _synchroniser_dropdown_especes(self):
        """Rafraîchit le menu déroulant Espèce du formulaire principal après une modification du dictionnaire."""
        cat_active = self.choice_category.get()
        valeurs = self.faune_meuse.get(cat_active, [])
        try:
            self.choice_species.configure(values=valeurs)
            if valeurs and self.choice_species.get() not in valeurs:
                self.choice_species.set(valeurs[0])
        except Exception:
            pass

    def load_species_dict(self):
        if os.path.exists(SPECIES_FILE):
            try:
                with open(SPECIES_FILE, "r", encoding="utf-8") as f:
                    self.faune_meuse = json.load(f)
            except Exception:
                self.faune_meuse = DEFAULT_FAUNE.copy()
        else:
            self.faune_meuse = DEFAULT_FAUNE.copy()
            self.save_species_dict()

    def save_species_dict(self):
        try:
            with open(SPECIES_FILE, "w", encoding="utf-8") as f:
                json.dump(self.faune_meuse, f, indent=4, ensure_ascii=False)
        except Exception as e:
            self.log(f"⚠️ Impossible de sauvegarder la liste des espèces : {str(e)}")

    def remove_current_species(self):
        """Retire du dictionnaire l'espèce actuellement sélectionnée dans le menu déroulant (avec confirmation)."""
        cat_active = self.choice_category.get()
        espece_active = self.choice_species.get()

        if not espece_active or espece_active == "Sélectionnez...":
            messagebox.showinfo("Aucune espèce", "Sélectionnez d'abord une espèce à retirer.")
            return

        if espece_active not in self.faune_meuse.get(cat_active, []):
            messagebox.showinfo("Introuvable", f"« {espece_active} » ne fait pas partie du dictionnaire « {cat_active} ».")
            return

        if not messagebox.askyesno(
            "Confirmer la suppression",
            f"Retirer « {espece_active} » de la liste « {cat_active} » ?\n\n"
            "Les observations déjà enregistrées avec ce nom ne sont pas modifiées : "
            "seule la liste de choix change."
        ):
            return

        self.faune_meuse[cat_active].remove(espece_active)
        self.save_species_dict()
        self.on_category_change(cat_active)
        self.log(f"🗑️ Espèce retirée du dictionnaire : {espece_active} ({cat_active})")

    def add_custom_species(self):
        cat_active = self.choice_category.get()
        nouvelle_espece = simpledialog.askstring("Nouvelle espèce", f"Entrez le nom de la nouvelle espèce à ajouter dans '{cat_active}' :")
        
        if nouvelle_espece and nouvelle_espece.strip():
            nouvelle_espece = nouvelle_espece.strip()
            if nouvelle_espece not in self.faune_meuse[cat_active]:
                self.faune_meuse[cat_active].append(nouvelle_espece)
                self.faune_meuse[cat_active].sort()  # Tri alphabétique
                self.save_species_dict()
                
                # Rafraîchir l'interface
                self.on_category_change(cat_active)
                self.choice_species.set(nouvelle_espece)
                self.log(f"➕ Espèce ajoutée au dictionnaire : {nouvelle_espece} ({cat_active})")
            else:
                messagebox.showinfo("Déjà présente", "Cette espèce figure déjà dans la liste.")

    # --- PISTE 2 : PRÉPARATION SYNTHÈSE IA ---
    def _stats_from_notes(self, all_notes):
        """Stats simples + liste d'espèces pour brief IA (sans carte)."""
        from collections import Counter
        especes = Counter()
        categories = Counter()
        indices = 0
        avec_photo = 0
        sans_photo = 0
        lieux = Counter()
        for key, data in (all_notes or {}).items():
            esp = (data.get("espece") or "").strip() or "Inconnu"
            cat = (data.get("categorie") or "Autre").strip()
            especes[esp] += 1
            categories[cat] += 1
            if data.get("sans_photo") or str(key).startswith("_manuel_"):
                sans_photo += 1
            else:
                avec_photo += 1
            typ = (data.get("type_observation") or "") + " " + esp.lower()
            if any(x in typ.lower() for x in ("empreinte", "terrier", "coulée", "coulee", "trace", "latrine", "indice")):
                indices += 1
            lieu = (data.get("lieu") or "").strip()
            if lieu:
                lieux[lieu] += 1
        return {
            "n_obs": len(all_notes or {}),
            "n_especes": len(especes),
            "especes": especes,
            "categories": categories,
            "indices": indices,
            "avec_photo": avec_photo,
            "sans_photo": sans_photo,
            "lieux": lieux,
        }

    def generate_ai_summary(self):
        """Brief de sortie : texte + stats + liste d'espèces (pas de carte décorative)."""
        if not self.photo_folder_path:
            messagebox.showwarning(
                "Dossier manquant",
                "Ouvrez un dossier photos ou un carnet sans photos."
            )
            return

        chemin_notes = os.path.join(self.photo_folder_path, NOTES_FILE)
        if not os.path.exists(chemin_notes):
            messagebox.showinfo("Aucune donnée", "Aucune observation enregistrée dans ce carnet.")
            return

        try:
            with open(chemin_notes, "r", encoding="utf-8") as f:
                all_notes = json.load(f) or {}
        except Exception as e:
            messagebox.showerror("Erreur", f"Lecture du carnet impossible : {e}")
            return

        if not all_notes:
            messagebox.showinfo("Aucune donnée", "Le carnet d'observations est vide.")
            return

        stats = self._stats_from_notes(all_notes)
        sortie = os.path.basename(self.photo_folder_path)

        # Bloc stats pour l'humain + l'IA
        stats_lines = [
            f"Sortie : {sortie}",
            f"Observations : {stats['n_obs']} (photos {stats['avec_photo']} · sans photo {stats['sans_photo']})",
            f"Espèces distinctes : {stats['n_especes']}",
            f"Indices / traces notés : {stats['indices']}",
            "Par catégorie : " + ", ".join(f"{k} ({v})" for k, v in stats["categories"].most_common()),
            "Espèces : " + ", ".join(f"{k} ({v})" for k, v in stats["especes"].most_common()),
        ]
        if stats["lieux"]:
            stats_lines.append(
                "Lieux : " + ", ".join(f"{k} ({v})" for k, v in stats["lieux"].most_common(8))
            )

        prompt_lines = [
            "Tu es un naturaliste de terrain (Grand Est, Meuse 55).",
            "Rédige un compte-rendu de sortie clair, professionnel et vivant.",
            "Structure : 1) Contexte & conditions 2) Espèces et effectifs 3) Indices 4) Points remarquables 5) Pistes pour la prochaine sortie.",
            "Distingue contacts visuels et indices (empreintes, terriers, coulées…).",
            "N'invente aucune espèce absente des données. Pas de carte ni de schéma ASCII.",
            "",
            "=== SYNTHÈSE CHIFFRÉE ===",
            *stats_lines,
            "",
            "=== DONNÉES BRUTES ===",
        ]

        for img, data in all_notes.items():
            ligne = (
                f"- {img} | {data.get('heure', '--:--')} | {data.get('categorie', '?')} | "
                f"{data.get('espece', '?')} ×{data.get('nombre', '?')}"
            )
            if data.get("lieu"):
                ligne += f" | {data.get('lieu')}"
            if data.get("type_observation"):
                ligne += f" | {data.get('type_observation')}"
            info = self.photos_data.get(img, {})
            lat = info.get("lat") or data.get("lat")
            lon = info.get("lon") or data.get("lon")
            if lat and lon:
                try:
                    ligne += f" | GPS {float(lat):.5f},{float(lon):.5f}"
                except Exception:
                    pass
            meteo = data.get("meteo") or {}
            if meteo.get("temperature") is not None:
                ligne += f" | {meteo.get('temperature')}°C {meteo.get('ciel', '')}"
            if data.get("notes_libres"):
                ligne += f" | {data.get('notes_libres')[:120]}"
            prompt_lines.append(ligne)

        prompt_lines.append("")
        prompt_lines.append("=== FIN ===")
        prompt_text = "\n".join(prompt_lines)

        ai_window = ctk.CTkToplevel(self)
        ai_window.title(f"Brief IA — {sortie}")
        ai_window.geometry("720x640")
        ai_window.configure(fg_color=UI.get("bg", "#0c1210"))
        self._prepare_tool_window(ai_window)

        ctk.CTkLabel(
            ai_window, text="Brief de sortie (texte + stats — sans carte)",
            font=ctk.CTkFont(size=15, weight="bold")
        ).pack(anchor="w", padx=14, pady=(12, 4))

        stats_box = ctk.CTkTextbox(ai_window, height=110, font=ctk.CTkFont(size=12))
        stats_box.pack(fill="x", padx=14, pady=4)
        stats_box.insert("0.0", "\n".join(stats_lines))
        stats_box.configure(state="disabled")

        ctk.CTkLabel(
            ai_window,
            text="Prompt à coller dans Gemini / Grok / ChatGPT :",
            font=ctk.CTkFont(size=12), text_color=UI.get("text_dim", "#9db0a6")
        ).pack(anchor="w", padx=14, pady=(8, 2))

        text_box = ctk.CTkTextbox(ai_window, font=ctk.CTkFont(size=12))
        text_box.pack(fill="both", expand=True, padx=14, pady=4)
        text_box.insert("0.0", prompt_text)

        def copy_prompt():
            self.clipboard_clear()
            self.clipboard_append(prompt_text)
            self.log("Brief IA copié dans le presse-papiers")
            messagebox.showinfo("Copié", "Prompt copié — collez-le dans votre IA.")

        row = ctk.CTkFrame(ai_window, fg_color="transparent")
        row.pack(fill="x", padx=14, pady=10)
        ctk.CTkButton(
            row, text="Copier le prompt", command=copy_prompt,
            fg_color=UI.get("purple", "#a78bfa"), height=34
        ).pack(side="left")
        ctk.CTkButton(row, text="Fermer", command=ai_window.destroy, width=100).pack(side="right")
        self.log(f"Brief IA : {stats['n_obs']} obs, {stats['n_especes']} espèces — {sortie}")

    # --- LE RESTE DU MOTEUR (IDENTIQUE ET STABLE) ---
    def log(self, message):
        """Ajoute une ligne au journal. Thread-safe : si appelée depuis un thread d'arrière-plan
        (ExifTool, météo, IA...), la mise à jour du widget est automatiquement reportée sur le thread
        principal via after() — écrire un widget Tkinter depuis un autre thread n'est pas supporté et
        peut provoquer des blocages ou fermetures aléatoires."""
        if threading.current_thread() is not threading.main_thread():
            self.after(0, lambda: self._log_direct(message))
            return
        self._log_direct(message)

    def _log_direct(self, message):
        self.log_box.configure(state="normal")
        self.log_box.insert("end", message + "\n")
        self.log_box.see("end")
        self.log_box.configure(state="disabled")
        # Miroir si journal détaché
        box = getattr(self, "_detached_log_box", None)
        if box is not None:
            try:
                box.configure(state="normal")
                box.insert("end", message + "\n")
                box.see("end")
                box.configure(state="disabled")
            except Exception:
                pass

    def get_exiftool_path(self):
        """Cherche ExifTool : dossier app, tools/, PyInstaller, PATH."""
        candidates = []
        # Dossier de l'exe / du script
        try:
            dossier_application = os.path.dirname(os.path.abspath(sys.argv[0]))
            candidates.append(os.path.join(dossier_application, "exiftool.exe"))
            candidates.append(os.path.join(dossier_application, "tools", "exiftool.exe"))
            candidates.append(os.path.join(dossier_application, "exiftool", "exiftool.exe"))
        except Exception:
            pass
        # PyInstaller onefile (bundle temporaire)
        try:
            if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
                base = sys._MEIPASS
                candidates.append(os.path.join(base, "exiftool.exe"))
                candidates.append(os.path.join(base, "tools", "exiftool.exe"))
        except Exception:
            pass
        # Dossier du fichier source (dev)
        try:
            here = os.path.dirname(os.path.abspath(__file__))
            candidates.append(os.path.join(here, "exiftool.exe"))
            candidates.append(os.path.join(here, "tools", "exiftool.exe"))
        except Exception:
            pass
        for c in candidates:
            if c and os.path.isfile(c):
                return c
        return shutil.which("exiftool") or shutil.which("exiftool.exe")

    def get_supported_images(self):
        """Photos / RAW uniquement (géotag, miniatures, synchro GPX)."""
        if not self.photo_folder_path:
            return []
        valid_extensions = ('.jpg', '.jpeg', '.png', '.tif', '.tiff', '.cr2', '.nef', '.arw', '.dng', '.orf', '.rw2')
        return [f for f in os.listdir(self.photo_folder_path) if f.lower().endswith(valid_extensions)]

    def get_supported_videos(self):
        """Vidéos d'affût / billebaude dans le dossier de sortie."""
        if not self.photo_folder_path:
            return []
        valid_extensions = ('.mp4', '.mov', '.avi', '.mkv', '.mts', '.m2ts', '.wmv', '.m4v', '.mpg', '.mpeg')
        return sorted(
            f for f in os.listdir(self.photo_folder_path)
            if f.lower().endswith(valid_extensions)
        )

    def get_supported_media(self):
        """Photos + vidéos pour la liste du carnet."""
        photos = self.get_supported_images()
        videos = self.get_supported_videos()
        # Photos d'abord, puis vidéos (préfixe visuel géré à l'affichage)
        return photos + videos

    @staticmethod
    def _is_video_file(filename):
        return os.path.splitext(filename or "")[1].lower() in (
            '.mp4', '.mov', '.avi', '.mkv', '.mts', '.m2ts', '.wmv', '.m4v', '.mpg', '.mpeg'
        )

    def open_media_large(self, path=None):
        """Agrandir : ouvre dans la visionneuse Windows (zoom, plein ecran natif).
        Pour les RAW, genere un JPEG temporaire haute resolution pour eviter Photoshop."""
        path = path or self.selected_photo_path
        if not path or not os.path.isfile(path):
            messagebox.showinfo("Aperçu", "Aucun média sélectionné.")
            return

        if self._is_video_file(path):
            self.open_video_external(path)
            return

        # JPEG/PNG : ouverture directe Windows
        ext = os.path.splitext(path)[1].lower()
        if ext in (".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tif", ".tiff", ".webp"):
            self._open_with_windows_viewer(path)
            return

        # RAW : extraire un apercu JPEG HD puis l'ouvrir (evite l'asso Photoshop)
        try:
            self._open_raw_in_windows_viewer(path)
        except Exception as e:
            self.log("Aperçu Windows RAW: %s" % e)
            # Repli : visionneuse interne ou startfile
            try:
                self._open_internal_image_viewer(path)
            except Exception:
                try:
                    os.startfile(path)
                except Exception as e2:
                    messagebox.showerror("Aperçu", str(e2))

    def _open_with_windows_viewer(self, path):
        """Ouvre un fichier image avec l'app Photos Windows si possible, sinon association par defaut."""
        path = os.path.normpath(path)
        # 1) Essayer Windows Photos (URI)
        try:
            uri = "ms-photos:viewer?file=" + path.replace("\\", "/")
            # Start-Process plus fiable pour ms-photos:
            subprocess.Popen(
                ["cmd", "/c", "start", "", path],
                shell=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            return
        except Exception:
            pass
        try:
            os.startfile(path)
        except Exception as e:
            messagebox.showerror("Aperçu", "Impossible d'ouvrir :\n%s" % e)

    def _open_raw_in_windows_viewer(self, path):
        """Extrait un JPEG HD du RAW (cache disque) puis ouvre avec Windows Photos."""
        exe = self.get_exiftool_path()
        if not exe:
            os.startfile(path)
            return

        out_jpg = self._disk_cache_path(path, kind="hd", max_dim=2400)
        need = True
        try:
            if os.path.isfile(out_jpg) and os.path.getsize(out_jpg) > 2000:
                if os.path.getmtime(out_jpg) >= os.path.getmtime(path):
                    need = False
        except Exception:
            need = True

        if need:
            startupinfo = self._exiftool_startupinfo()
            written = False
            for tag in ("-PreviewImage", "-JpgFromRaw"):
                try:
                    # -W ecrase le chemin cible
                    cmd = [exe, "-b", tag, "-W", out_jpg, path]
                    subprocess.run(
                        cmd, capture_output=True, timeout=60, startupinfo=startupinfo,
                    )
                    if os.path.isfile(out_jpg) and os.path.getsize(out_jpg) > 2000:
                        written = True
                        break
                except Exception:
                    pass
            if not written:
                data = self._get_thumbnail_bytes(path, exe, max_dim=2400)
                if data:
                    self._write_disk_cache(out_jpg, data)
                else:
                    raise RuntimeError("Impossible d'extraire un apercu JPEG du RAW")

        self._open_with_windows_viewer(out_jpg)
        self.log("Aperçu Windows (cache) : %s" % out_jpg)

    def _open_internal_image_viewer(self, path):
        """Repli interne minimal si Windows Photos indisponible."""
        try:
            exe = self.get_exiftool_path()
            data = self._get_thumbnail_bytes(path, exe, max_dim=1600)
            if data:
                img = Image.open(io.BytesIO(data)).convert("RGB")
            else:
                img = Image.open(path).convert("RGB")
            img.thumbnail((1200, 800), Image.Resampling.LANCZOS if hasattr(Image, "Resampling") else Image.LANCZOS)
            win = ctk.CTkToplevel(self)
            win.title("Aperçu — %s" % os.path.basename(path))
            win.geometry("%dx%d" % (img.width + 40, img.height + 80))
            ctk_img = ctk.CTkImage(light_image=img, dark_image=img, size=img.size)
            win._img = ctk_img
            ctk.CTkLabel(win, text="", image=ctk_img).pack(padx=10, pady=10)
            ctk.CTkButton(win, text="Fermer", command=win.destroy).pack(pady=6)
        except Exception as e:
            messagebox.showerror("Aperçu", str(e))

    def _open_with_windows_dialog(self, path):
        """Laisse Windows proposer l'application (sans imposer Photoshop)."""
        try:
            if os.name == "nt":
                # Boîte « Ouvrir avec »
                subprocess.Popen(["rundll32", "shell32.dll,OpenAs_RunDLL", path])
            else:
                subprocess.Popen(["xdg-open", path])
        except Exception as e:
            messagebox.showerror("Ouvrir avec", str(e))

    def _on_preview_hover_enter(self, event=None):
        """Survol de l'aperçu : loupe temporaire après un court délai (évite les déclenchements parasites)."""
        self._preview_hover_job = self.after(450, self._show_hover_loupe)

    def _on_preview_hover_leave(self, event=None):
        job = getattr(self, "_preview_hover_job", None)
        if job:
            try:
                self.after_cancel(job)
            except Exception:
                pass
            self._preview_hover_job = None
        self._hide_hover_loupe()

    def _show_hover_loupe(self):
        """Petite fenêtre flottante avec un aperçu plus grand (photo uniquement)."""
        path = self.selected_photo_path
        if not path or not os.path.isfile(path) or self._is_video_file(path):
            return
        if getattr(self, "_hover_loupe", None) and self._hover_loupe.winfo_exists():
            return
        try:
            exe = self.get_exiftool_path()
            data = self._get_thumbnail_bytes(path, exe, max_dim=480)
            if not data and self._current_full_preview is not None:
                img = self._current_full_preview.copy()
                img.thumbnail((480, 360))
            elif data:
                img = Image.open(io.BytesIO(data)).convert("RGB")
                img.thumbnail((480, 360))
            else:
                return
            top = ctk.CTkToplevel(self)
            top.overrideredirect(True)
            top.attributes("-topmost", True)
            x = self.winfo_pointerx() + 16
            y = self.winfo_pointery() + 16
            top.geometry(f"+{x}+{y}")
            ctk_img = ctk.CTkImage(light_image=img, dark_image=img, size=(img.width, img.height))
            self._hover_loupe_img = ctk_img
            lbl = ctk.CTkLabel(top, image=ctk_img, text="", corner_radius=8)
            lbl.pack()
            top.bind("<Leave>", lambda e: self._hide_hover_loupe())
            lbl.bind("<Double-Button-1>", lambda e: (self._hide_hover_loupe(), self.open_media_large()))
            self._hover_loupe = top
        except Exception:
            pass

    def _hide_hover_loupe(self):
        win = getattr(self, "_hover_loupe", None)
        if win:
            try:
                win.destroy()
            except Exception:
                pass
        self._hover_loupe = None

    def open_video_external(self, path=None):
        """Ouvre la vidéo avec le lecteur par défaut de Windows (un clic)."""
        path = path or self.selected_photo_path
        if not path or not os.path.isfile(path):
            messagebox.showinfo("Vidéo", "Aucune vidéo sélectionnée.")
            return
        try:
            if os.name == "nt":
                os.startfile(path)  # lecteur par défaut Windows
            elif sys.platform == "darwin":
                subprocess.Popen(["open", path])
            else:
                subprocess.Popen(["xdg-open", path])
            self.log(f"🎬 Lecture : {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Vidéo", f"Impossible d'ouvrir la vidéo :\n{e}")


    def _on_media_double_click(self, event=None):
        """Double-clic listbox : lit la vidéo immédiatement."""
        sel = self.photo_listbox.curselection()
        if not sel:
            return
        filename = self._listbox_filename(self.photo_listbox.get(sel[0]))
        if self._is_video_file(filename):
            path = os.path.join(self.photo_folder_path or "", filename)
            self.open_video_external(path)

    @staticmethod
    def _exiftool_startupinfo():
        """Masque la console sous Windows ; None ailleurs."""
        if os.name != "nt":
            return None
        try:
            info = subprocess.STARTUPINFO()
            info.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            return info
        except Exception:
            return None

    def _set_progress(self, fraction, text=None):
        """Mise à jour thread-safe de la barre de progression et du libellé."""
        fraction = max(0.0, min(1.0, float(fraction or 0)))

        def _ui():
            try:
                self.progress_bar.set(fraction)
                if text is not None:
                    self.lbl_progress.configure(text=text)
            except Exception:
                pass

        if threading.current_thread() is threading.main_thread():
            _ui()
        else:
            self.after(0, _ui)

    def _start_progress_pulse(self, label="Traitement…"):
        """Animation de barre indéterminée pendant une tâche longue."""
        self._pulse_active = True
        self._pulse_phase = 0

        def tick():
            if not getattr(self, "_pulse_active", False):
                return
            self._pulse_phase = (getattr(self, "_pulse_phase", 0) + 1) % 20
            phase = self._pulse_phase
            val = 0.15 + 0.7 * (phase / 19.0 if phase <= 10 else (20 - phase) / 10.0)
            try:
                self.progress_bar.set(val)
                if label:
                    self.lbl_progress.configure(text=label)
            except Exception:
                pass
            self.after(80, tick)

        self.after(0, tick)

    def _stop_progress_pulse(self, final=0.0, text="Prêt"):
        self._pulse_active = False
        self._set_progress(final, text)

    def _exiftool_read_json(self, exe_path, targets, tags, timeout=90):
        """Lecture optimisée ExifTool → liste de dicts JSON.

        targets : dossier ou liste de chemins fichiers
        tags : liste de tags sans tiret initial, ex. ['GPSLatitude', 'DateTimeOriginal']
        Flags : -n (numérique), -fast2 (lecture accélérée), -q -q (silencieux)
        """
        if not targets:
            return []
        cmd = [exe_path, "-n", "-fast2", "-q", "-q", "-json"]
        for t in tags:
            cmd.append(f"-{t}")
        # Limite aux extensions supportées si on vise un dossier
        if isinstance(targets, str):
            for ext in ("jpg", "jpeg", "png", "tif", "tiff", "cr2", "nef", "arw", "dng", "orf", "rw2"):
                cmd.extend(["-ext", ext])
            cmd.append(os.path.normpath(targets))
        else:
            cmd.extend([os.path.normpath(t) for t in targets])

        try:
            proc = subprocess.run(
                cmd,
                startupinfo=self._exiftool_startupinfo(),
                capture_output=True, text=True, encoding="utf-8", errors="ignore",
                timeout=timeout,
            )
            if proc.returncode not in (0, 1):  # 1 = warnings mineurs parfois
                return []
            out = (proc.stdout or "").strip()
            if not out:
                return []
            data = json.loads(out)
            return data if isinstance(data, list) else [data]
        except subprocess.TimeoutExpired:
            raise
        except Exception:
            return []

    def _exiftool_read_json_chunked(self, exe_path, file_paths, tags, chunk_size=40, timeout_per_chunk=60, progress_label="Lecture EXIF"):
        """Lit les métadonnées par lots en parallèle (multiprocessing) avec barre de progression.

        - Petit volume : 1 appel ExifTool dans le process courant
        - Gros volume : ProcessPoolExecutor (vrais process OS) avec repli ThreadPool si échec
        L'écriture géotag reste séquentielle (intégrité des fichiers).
        """
        file_paths = list(file_paths)
        if not file_paths:
            return []

        if len(file_paths) <= chunk_size:
            self._set_progress(0.15, f"{progress_label}…")
            try:
                rows = self._exiftool_read_json(exe_path, file_paths, tags, timeout=timeout_per_chunk)
            except subprocess.TimeoutExpired:
                self.log(f"⚠️ {progress_label} : délai dépassé")
                rows = []
            self._set_progress(0.9, f"{progress_label} terminée")
            return rows

        chunks = [file_paths[i:i + chunk_size] for i in range(0, len(file_paths), chunk_size)]
        payloads = [(exe_path, chunk, list(tags), timeout_per_chunk) for chunk in chunks]
        # 2–4 process max : au-delà le disque sature souvent avant le CPU
        workers = max(1, min(4, CPU_WORKERS, len(chunks)))
        results = []
        done = 0

        def _consume(futures_map):
            nonlocal done, results
            for fut in as_completed(futures_map):
                try:
                    rows = fut.result()
                except Exception:
                    rows = []
                if rows:
                    results.extend(rows)
                done += 1
                frac = 0.1 + 0.8 * (done / len(chunks))
                self._set_progress(frac, f"{progress_label} : lot {done}/{len(chunks)} (×{workers} process)")

        # --- Multiprocessing (processus séparés) ---
        used_mp = False
        # Sous EXE Windows, privilégier les threads : plus stable (pas de re-spawn de l'appli).
        # freeze_support() est déjà appelé dans __main__ ; les process restent dispo hors EXE.
        use_processes = not getattr(sys, "frozen", False)
        try:
            if use_processes:
                ctx = multiprocessing.get_context("spawn" if os.name == "nt" else "fork")
                with ProcessPoolExecutor(max_workers=workers, mp_context=ctx) as pool:
                    fmap = {pool.submit(_exiftool_chunk_worker, pl): i for i, pl in enumerate(payloads)}
                    _consume(fmap)
                used_mp = True
            else:
                raise RuntimeError("mode EXE → threads")
        except Exception as e:
            if use_processes:
                self.log(f"⚠️ Multiprocessing indisponible ({e}) — repli sur threads.")

        if not used_mp:
            results = []
            done = 0
            with ThreadPoolExecutor(max_workers=workers) as pool:
                fmap = {pool.submit(_exiftool_chunk_worker, pl): i for i, pl in enumerate(payloads)}
                _consume(fmap)

        mode = "process" if used_mp else "threads"
        self.log(f"📡 {progress_label} : {len(results)} fiche(s) via {len(chunks)} lot(s), mode {mode} ×{workers}")
        return results

    def select_photo_folder(self):
        self.photo_folder_path = filedialog.askdirectory()
        if self.photo_folder_path:
            self.btn_browse_photos.configure(text=f"📁 {os.path.basename(self.photo_folder_path)}")
            self.log(f"Dossier chargé : {self.photo_folder_path}")
            
            self.entry_search_photo.delete(0, tk.END)
            images = self.get_supported_images()
            videos = self.get_supported_videos()
            self._filtered_media_names = images + videos
            if getattr(self, "carnet_view_mode", "list") == "thumbs":
                self._rebuild_thumb_explorer(lazy=True)
            else:
                self._populate_listbox_from_filtered()
            self.log(
                f"📋 {len(images)} photo(s)"
                + (f" · {len(videos)} vidéo(s)" if videos else "")
                + ". Prêt pour le géotagging / lecture vidéo."
            )
            self._register_known_folder(self.photo_folder_path)
            # Observations sans photo déjà notées → carte + compteur
            try:
                notes_path = os.path.join(self.photo_folder_path, NOTES_FILE)
                if os.path.exists(notes_path):
                    with open(notes_path, "r", encoding="utf-8") as f:
                        all_notes = json.load(f) or {}
                    for key, data in all_notes.items():
                        if not isinstance(data, dict):
                            continue
                        if not self._is_virtual_observation(key, data):
                            continue
                        if data.get("lat") is None or data.get("lon") is None:
                            continue
                        self.photos_data[key] = {
                            "path": "",
                            "lat": data.get("lat"),
                            "lon": data.get("lon"),
                            "date": data.get("heure") or "",
                            "sans_photo": True,
                            "source": data.get("source") or "",
                        }
            except Exception:
                pass
            self.refresh_daily_counter()
            self._refresh_listbox_annotation_status()
            try:
                self.refresh_map_markers()
            except Exception:
                pass

            # Le carnet (annotation d'espèces) n'a pas besoin du GPS : on pré-remplit tout de suite une
            # entrée par photo (sans coordonnées) pour que la saisie soit possible immédiatement, sans
            # attendre ni exiger de synchronisation. Le GPS viendra enrichir ces entrées dès qu'il sera
            # disponible (détection automatique ci-dessous, ou synchro manuelle).
            self.photos_data = {
                img: {"path": os.path.join(self.photo_folder_path, img), "lat": None, "lon": None, "date": None}
                for img in images
            }
            for vid in videos:
                self.photos_data[vid] = {
                    "path": os.path.join(self.photo_folder_path, vid),
                    "lat": None, "lon": None, "date": None, "is_video": True,
                }
            self.is_synced = True
            self.last_backup_dir = None
            self.btn_undo_sync.configure(state="disabled")

            # Réhydrate GPS depuis le carnet (photos déjà taguées + sans photo)
            # puis ExifTool relira les balises EXIF en arrière-plan pour confirmer / compléter.
            notes_existantes = self._load_notes_dict()
            n_from_notes = 0
            for cle, data in notes_existantes.items():
                if not isinstance(data, dict):
                    continue
                lat, lon = data.get("lat"), data.get("lon")
                if lat is None or lon is None:
                    continue
                try:
                    lat, lon = float(lat), float(lon)
                except Exception:
                    continue
                if cle in self.photos_data:
                    self.photos_data[cle]["lat"] = lat
                    self.photos_data[cle]["lon"] = lon
                    if data.get("heure") and not self.photos_data[cle].get("date"):
                        self.photos_data[cle]["date"] = data.get("heure")
                    n_from_notes += 1
                elif self._is_virtual_observation(cle, data):
                    self.photos_data[cle] = {
                        "path": "",
                        "lat": lat,
                        "lon": lon,
                        "date": data.get("heure") or "",
                        "sans_photo": True,
                        "source": data.get("source") or "",
                    }
                    n_from_notes += 1
            if n_from_notes:
                self.log("📍 %d point(s) GPS repris depuis le carnet (observations.json)." % n_from_notes)

            self.refresh_map_markers()
            exe_path = self.get_exiftool_path()
            if exe_path:
                self.log("🔎 Lecture des balises GPS dans les fichiers (ExifTool, arrière-plan)...")
                threading.Thread(target=self._check_existing_gps_worker, args=(exe_path,), daemon=True).start()
            else:
                self.log("⚠️ ExifTool introuvable : GPS des fichiers photo non relus (carnet uniquement).")

    def select_gpx_file(self):
        self.gpx_file_path = filedialog.askopenfilename(filetypes=[("GPS Track Files", "*.gpx")])
        if self.gpx_file_path:
            self.btn_browse_gpx.configure(text=f"🗺️ {os.path.basename(self.gpx_file_path)}")
            self.log(f"GPX sélectionné : {os.path.basename(self.gpx_file_path)}")
            if self.var_afficher_trace.get():
                self._afficher_trace_gpx()

    def changer_fond_carte(self, nom_carte: str):
        if nom_carte == "Plan IGN (Moderne)":
            url = "https://data.geopf.fr/wmts?SERVICE=WMTS&REQUEST=GetTile&VERSION=1.0.0&LAYER=GEOGRAPHICALGRIDSYSTEMS.PLANIGNV2&STYLE=normal&FORMAT=image/png&TILEMATRIXSET=PM&TILEMATRIX={z}&TILEROW={y}&TILECOL={x}"
            self.map_widget.set_tile_server(url, max_zoom=19)
        elif nom_carte == "Photos Aériennes":
            url = "https://data.geopf.fr/wmts?SERVICE=WMTS&REQUEST=GetTile&VERSION=1.0.0&LAYER=ORTHOIMAGERY.ORTHOPHOTOS&STYLE=normal&FORMAT=image/jpeg&TILEMATRIXSET=PM&TILEMATRIXSET=PM&TILEMATRIX={z}&TILEROW={y}&TILECOL={x}"
            self.map_widget.set_tile_server(url, max_zoom=19)
        elif nom_carte == "OpenStreetMap":
            url = "https://a.tile.openstreetmap.org/{z}/{x}/{y}.png"
            self.map_widget.set_tile_server(url)

    def _parse_nombre_approx(self, valeur):
        """Extrait un nombre d'un champ 'Nombre' exact ('3') ou approximatif/libre ('vol d'une trentaine',
        'Nombreux', 'Vol important'...). Retourne (nombre_estimé, est_approximatif)."""
        txt = str(valeur or "1").strip()
        m = re.search(r"\d+", txt)
        if m:
            nombre = int(m.group())
            return nombre, not txt.isdigit()
        return 1, True  # texte purement qualitatif : compté pour au moins 1, marqué approximatif

    def _compteur_manuel_path(self):
        return os.path.join(self.photo_folder_path, "compteur_manuel.json")

    def _load_compteur_manuel(self):
        if not self.photo_folder_path:
            return {}
        path = self._compteur_manuel_path()
        if not os.path.exists(path):
            return {}
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}

    def _save_compteur_manuel(self, dict_compteur):
        if not self.photo_folder_path:
            return
        try:
            with open(self._compteur_manuel_path(), "w", encoding="utf-8") as f:
                json.dump(dict_compteur, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _compute_auto_species_sum(self):
        """Somme automatique par espèce à partir des photos/observations — purement indicative :
        peut surcompter si le même individu a été photographié plusieurs fois."""
        notes = self._load_notes_dict()
        images = self.get_supported_images()
        compteur = collections.Counter()
        annotees = 0
        for img in images:
            data = notes.get(img)
            if not data or not data.get("espece"):
                continue
            annotees += 1
            n, _ = self._parse_nombre_approx(data.get("nombre", "1"))
            compteur[data["espece"]] += n
        for cle, data in notes.items():
            if cle in images or not isinstance(data, dict) or not data.get("espece"):
                continue
            n, _ = self._parse_nombre_approx(data.get("nombre", "1"))
            compteur[data["espece"]] += n
        non_annotees = max(0, len(images) - annotees)
        return compteur, non_annotees, len(images)

    def _parser_gpx_points(self, chemin_gpx):
        """Extrait la liste ordonnée des points (lat, lon) d'une trace GPX (trkpt, ou wpt à défaut)."""
        points = []
        try:
            import xml.etree.ElementTree as ET
            tree = ET.parse(chemin_gpx)
            root = tree.getroot()
            ns = root.tag.split("}")[0] + "}" if root.tag.startswith("{") else ""
            for trkpt in root.iter(f"{ns}trkpt"):
                lat, lon = trkpt.get("lat"), trkpt.get("lon")
                if lat and lon:
                    points.append((float(lat), float(lon)))
            if not points:
                for wpt in root.iter(f"{ns}wpt"):
                    lat, lon = wpt.get("lat"), wpt.get("lon")
                    if lat and lon:
                        points.append((float(lat), float(lon)))
        except Exception:
            pass
        return points

    def toggle_trace_gpx(self):
        if self.var_afficher_trace.get():
            self._afficher_trace_gpx()
        else:
            self._masquer_trace_gpx()

    def _afficher_trace_gpx(self):
        self._masquer_trace_gpx()
        if not self.gpx_file_path or not os.path.exists(self.gpx_file_path):
            messagebox.showinfo("Aucune trace", "Sélectionnez d'abord un fichier GPX (étape 2 de la sidebar).")
            self.var_afficher_trace.set(False)
            return
        points = self._parser_gpx_points(self.gpx_file_path)
        if not points:
            messagebox.showinfo("Trace vide", "Aucun point exploitable trouvé dans ce fichier GPX.")
            self.var_afficher_trace.set(False)
            return
        try:
            self.trace_gpx_path_obj = self.map_widget.set_path(points, color="#e74c3c", width=4)
            self.log(f"🗺️ Trace GPX affichée ({len(points)} points).")
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible d'afficher la trace : {e}")
            self.var_afficher_trace.set(False)

    def _masquer_trace_gpx(self):
        if self.trace_gpx_path_obj is not None:
            try:
                self.trace_gpx_path_obj.delete()
            except Exception:
                pass
            self.trace_gpx_path_obj = None

    def refresh_daily_counter(self):
        """Reconstruit le panneau 'Compteur du jour' — désormais MANUEL : c'est vous qui validez le nombre
        d'individus distincts par espèce. La somme calculée depuis les photos n'est affichée qu'à titre indicatif
        (elle peut surcompter le même animal photographié plusieurs fois)."""
        if not hasattr(self, "compteur_lignes_frame"):
            return
        for w in self.compteur_lignes_frame.winfo_children():
            w.destroy()

        if not self.photo_folder_path:
            ctk.CTkLabel(self.compteur_lignes_frame, text="Aucun dossier chargé.", font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#666666")).pack(anchor="w")
            self.lbl_compteur_hint.configure(text="")
            return

        manuel = self._load_compteur_manuel()
        auto_sum, non_annotees, total_photos = self._compute_auto_species_sum()

        if not manuel:
            ctk.CTkLabel(
                self.compteur_lignes_frame,
                text="Aucune espèce comptée pour l'instant. « 🔄 Depuis les photos » pour démarrer, ou « + Espèce » pour saisir à la main.",
                font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#666666"), wraplength=400, justify="left"
            ).pack(anchor="w")
        else:
            for espece in sorted(manuel.keys(), key=lambda e: -manuel[e]):
                valeur = manuel[espece]
                ligne = ctk.CTkFrame(self.compteur_lignes_frame, fg_color=UI.get("card_alt", "#1e2a3a"), corner_radius=8, border_width=1, border_color=UI.get("border", "#2a3544"))
                ligne.pack(fill="x", pady=2)

                ctk.CTkLabel(ligne, text=espece, font=ctk.CTkFont(size=12), anchor="w").pack(side="left", padx=(8, 4), pady=6, fill="x", expand=True)

                auto_n = auto_sum.get(espece)
                if auto_n is not None:
                    ctk.CTkLabel(ligne, text=f"(photos : {auto_n})", font=ctk.CTkFont(size=9), text_color="#777777").pack(side="left", padx=(0, 6))

                ctk.CTkButton(
                    ligne, text="–", width=26, height=24, font=ctk.CTkFont(size=13),
                    fg_color="#2b2b2b", hover_color="#3a3a3a", command=lambda e=espece: self._ajuster_compteur_manuel(e, -1)
                ).pack(side="left", padx=1)

                entry_val = ctk.CTkEntry(ligne, width=44, height=24, justify="center", font=ctk.CTkFont(size=12))
                entry_val.insert(0, str(valeur))
                entry_val.pack(side="left", padx=1)
                entry_val.bind("<Return>", lambda ev, e=espece, w=entry_val: self._definir_compteur_manuel(e, w.get()))
                entry_val.bind("<FocusOut>", lambda ev, e=espece, w=entry_val: self._definir_compteur_manuel(e, w.get()))

                ctk.CTkButton(
                    ligne, text="+", width=26, height=24, font=ctk.CTkFont(size=13),
                    fg_color="#2b2b2b", hover_color="#3a3a3a", command=lambda e=espece: self._ajuster_compteur_manuel(e, 1)
                ).pack(side="left", padx=1)

                ctk.CTkButton(
                    ligne, text="✕", width=26, height=24, font=ctk.CTkFont(size=11),
                    fg_color="transparent", hover_color="#3a2020", text_color="#aa6666",
                    command=lambda e=espece: self._retirer_espece_du_compteur(e)
                ).pack(side="left", padx=(1, 6))

        morceaux = []
        if non_annotees > 0:
            morceaux.append(f"⚠️ {non_annotees}/{total_photos} photo(s) pas encore annotée(s).")
        elif total_photos:
            morceaux.append(f"✅ Les {total_photos} photos du dossier sont annotées.")
        morceaux.append("Le nombre affiché est celui que VOUS validez ; « photos : N » n'est qu'une indication automatique (peut surcompter un même individu photographié plusieurs fois).")
        self.lbl_compteur_hint.configure(text=" ".join(morceaux))

    def _ajuster_compteur_manuel(self, espece, delta):
        manuel = self._load_compteur_manuel()
        manuel[espece] = max(0, manuel.get(espece, 0) + delta)
        self._save_compteur_manuel(manuel)
        self.refresh_daily_counter()

    def _definir_compteur_manuel(self, espece, valeur_texte):
        try:
            valeur = max(0, int(str(valeur_texte).strip()))
        except Exception:
            valeur = 0
        manuel = self._load_compteur_manuel()
        manuel[espece] = valeur
        self._save_compteur_manuel(manuel)
        self.refresh_daily_counter()

    def _retirer_espece_du_compteur(self, espece):
        manuel = self._load_compteur_manuel()
        if espece in manuel:
            del manuel[espece]
            self._save_compteur_manuel(manuel)
        self.refresh_daily_counter()

    def ajouter_espece_au_compteur(self):
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Chargez d'abord un dossier de photos.")
            return
        nom = simpledialog.askstring("Ajouter une espèce", "Nom de l'espèce à compter :", parent=self)
        if not nom:
            return
        nom = nom.strip()
        if not nom:
            return
        manuel = self._load_compteur_manuel()
        if nom not in manuel:
            manuel[nom] = 1
            self._save_compteur_manuel(manuel)
        self.refresh_daily_counter()

    def initialiser_compteur_depuis_photos(self):
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Chargez d'abord un dossier de photos.")
            return
        auto_sum, _, _ = self._compute_auto_species_sum()
        if not auto_sum:
            messagebox.showinfo("Rien à initialiser", "Aucune espèce détectée dans les photos annotées.")
            return
        manuel = self._load_compteur_manuel()
        ajoutes = 0
        for espece, n in auto_sum.items():
            if espece not in manuel:
                manuel[espece] = n
                ajoutes += 1
        self._save_compteur_manuel(manuel)
        self.refresh_daily_counter()
        self.log(f"🔄 Compteur initialisé depuis les photos ({ajoutes} nouvelle(s) espèce(s) ajoutée(s), valeurs existantes conservées).")
        messagebox.showinfo("Initialisé", f"{ajoutes} espèce(s) ajoutée(s) depuis les photos.\nAjustez ensuite les valeurs si plusieurs photos montrent le même individu.")

    def _refresh_listbox_annotation_status(self):
        """Colore les entrées selon annotation / type, adapté au thème actif."""
        if not hasattr(self, "photo_listbox"):
            return
        notes = self._load_notes_dict()
        items = self.photo_listbox.get(0, tk.END)
        col_ok = UI.get("list_annotated", "#4fd67a")
        col_wait = UI.get("list_pending", UI.get("text_dim", "#a8b8b0"))
        col_vid = UI.get("list_video", "#5b9bd5")
        video_ext = (".mp4", ".mov", ".avi", ".mkv", ".m4v", ".mpg", ".mpeg")
        for idx, img in enumerate(items):
            name = str(img)
            # retirer éventuels préfixes icônes
            pure = name.strip()
            for prefix in ("🎬 ", "🎥 ", "🎞 ", "📷 ", "🖼 "):
                if pure.startswith(prefix):
                    pure = pure[len(prefix):].strip()
            annotee = bool(notes.get(pure, {}).get("espece") or notes.get(name, {}).get("espece"))
            is_video = pure.lower().endswith(video_ext) or name.strip().startswith(("🎬", "🎥", "🎞"))
            if annotee:
                couleur = col_ok
            elif is_video:
                couleur = col_vid
            else:
                couleur = col_wait
            try:
                self.photo_listbox.itemconfig(idx, fg=couleur)
            except Exception:
                pass

    def _load_notes_dict(self, folder_path=None):
        """Charge le fichier observations.json d'un dossier, avec cache mémoire basé sur la date de
        modification du fichier (invalidation automatique dès que le fichier change, y compris de l'extérieur)."""
        folder_path = folder_path or self.photo_folder_path
        if not folder_path:
            return {}
        if not hasattr(self, "_notes_cache"):
            self._notes_cache = {}
        chemin_notes = os.path.join(folder_path, NOTES_FILE)
        if not os.path.exists(chemin_notes):
            self._notes_cache.pop(folder_path, None)
            return {}
        try:
            mtime = os.path.getmtime(chemin_notes)
        except Exception:
            mtime = None
        entree_cache = self._notes_cache.get(folder_path)
        if entree_cache and entree_cache[0] == mtime:
            return entree_cache[1]
        try:
            with open(chemin_notes, "r", encoding="utf-8") as f:
                data = json.load(f)
            self._notes_cache[folder_path] = (mtime, data)
            return data
        except Exception:
            return {}

    def _notes_backup_dir(self, folder_path=None):
        folder_path = folder_path or self.photo_folder_path
        if not folder_path:
            return None
        return os.path.join(folder_path, NOTES_BACKUP_DIR)

    def _list_notes_backups(self, folder_path=None):
        """Liste les sauvegardes versionnées (plus récentes d'abord) : (chemin, datetime, taille)."""
        bdir = self._notes_backup_dir(folder_path)
        if not bdir or not os.path.isdir(bdir):
            return []
        result = []
        for name in os.listdir(bdir):
            if not (name.startswith("observations_") and name.endswith(".json")):
                continue
            full = os.path.join(bdir, name)
            try:
                # observations_20260726_163045.json
                stamp = name[len("observations_"):-len(".json")]
                dt = datetime.strptime(stamp, "%Y%m%d_%H%M%S")
            except Exception:
                try:
                    dt = datetime.fromtimestamp(os.path.getmtime(full))
                except Exception:
                    dt = datetime.min
            try:
                size = os.path.getsize(full)
            except Exception:
                size = 0
            result.append((full, dt, size))
        result.sort(key=lambda x: x[1], reverse=True)
        return result

    def _rotate_notes_backups(self, folder_path=None):
        """Supprime les anciennes versions au-delà de NOTES_BACKUP_KEEP."""
        backups = self._list_notes_backups(folder_path)
        for path, _, _ in backups[NOTES_BACKUP_KEEP:]:
            try:
                os.remove(path)
            except Exception:
                pass

    def _should_create_notes_backup(self, folder_path, force=False):
        """Évite de créer une version à chaque autosave (60 s) : intervalle minimum entre versions."""
        if force:
            return True
        backups = self._list_notes_backups(folder_path)
        if not backups:
            return True
        last_dt = backups[0][1]
        try:
            age = (datetime.now() - last_dt).total_seconds()
        except Exception:
            return True
        return age >= NOTES_BACKUP_MIN_INTERVAL_SEC

    def _save_notes_dict(self, notes_dict, folder_path=None, *, create_backup=True, force_backup=False, silent=False):
        """Écrit observations.json de façon atomique, avec sauvegarde versionnée de l'ancienne version.

        - Copie l'ancien fichier vers .geoexif_notes_backups/observations_YYYYMMDD_HHMMSS.json
        - Écrit via fichier temporaire + replace (évite un JSON tronqué si crash)
        - Rotation : conserve NOTES_BACKUP_KEEP versions max
        """
        folder_path = folder_path or self.photo_folder_path
        if not folder_path:
            raise RuntimeError("Aucun dossier de photos chargé.")

        chemin_notes = os.path.join(folder_path, NOTES_FILE)
        bdir = self._notes_backup_dir(folder_path)

        # Contenu sérialisé une seule fois
        payload = json.dumps(notes_dict, indent=4, ensure_ascii=False)

        # Sauvegarde versionnée de l'état précédent (si le fichier existe et change)
        if create_backup and os.path.exists(chemin_notes) and self._should_create_notes_backup(folder_path, force=force_backup):
            try:
                with open(chemin_notes, "r", encoding="utf-8") as f:
                    ancien = f.read()
                if ancien.strip() != payload.strip():
                    os.makedirs(bdir, exist_ok=True)
                    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    dest = os.path.join(bdir, f"observations_{stamp}.json")
                    # Évite collision si 2 sauvegardes dans la même seconde
                    if os.path.exists(dest):
                        stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                        dest = os.path.join(bdir, f"observations_{stamp}.json")
                    with open(dest, "w", encoding="utf-8") as f:
                        f.write(ancien)
                    self._rotate_notes_backups(folder_path)
                    if not silent:
                        self.log(f"💾 Version notes : {os.path.basename(dest)}")
            except Exception as e:
                # Une erreur de backup ne doit pas bloquer la sauvegarde principale
                self.log(f"⚠️ Sauvegarde versionnée impossible : {e}")

        # Écriture atomique
        tmp_path = chemin_notes + ".tmp"
        try:
            with open(tmp_path, "w", encoding="utf-8") as f:
                f.write(payload)
                f.flush()
                os.fsync(f.fileno())
            os.replace(tmp_path, chemin_notes)

            try:
                self._run_intervention_backup(folder_path)
            except Exception:
                pass
            # Miroir local (PC) : recuperable meme si le disque externe est absent
            try:
                self._mirror_notes_to_local_archive(folder_path, payload)
            except Exception:
                pass
        except Exception:
            try:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception:
                pass
            raise

        # Invalide le cache mémoire
        if hasattr(self, "_notes_cache"):
            self._notes_cache.pop(folder_path, None)
        return chemin_notes

    def _local_archive_root(self):
        """Archive centrale des carnets sur le PC (hors disque externe)."""
        root = os.path.join(self._user_data_dir(), "archive_carnets")
        try:
            os.makedirs(root, exist_ok=True)
        except Exception:
            pass
        return root

    def _archive_key_for_folder(self, folder_path):
        """Cle stable pour un dossier de sortie (nom + hash court du chemin)."""
        import hashlib
        abs_path = os.path.normpath(os.path.abspath(folder_path))
        base = os.path.basename(abs_path) or "sortie"
        base = re.sub(r'[<>:"/\\|?*]', "_", base)[:60]
        h = hashlib.md5(abs_path.encode("utf-8", errors="replace")).hexdigest()[:8]
        return "%s__%s" % (base, h)

    def _mirror_notes_to_local_archive(self, folder_path, payload_json_str):
        """Copie observations.json (+ meta) dans AppData pour survie hors disque externe."""
        if not folder_path or not payload_json_str:
            return
        key = self._archive_key_for_folder(folder_path)
        dest_dir = os.path.join(self._local_archive_root(), key)
        os.makedirs(dest_dir, exist_ok=True)
        dest_notes = os.path.join(dest_dir, NOTES_FILE)
        with open(dest_notes, "w", encoding="utf-8") as f:
            f.write(payload_json_str)
        meta = {
            "path_origine": os.path.normpath(os.path.abspath(folder_path)),
            "label": os.path.basename(folder_path),
            "last_mirror": datetime.now().isoformat(timespec="seconds"),
            "nb_obs": 0,
        }
        try:
            data = json.loads(payload_json_str)
            meta["nb_obs"] = len(data) if isinstance(data, dict) else 0
        except Exception:
            pass
        try:
            alb = self._load_sortie_amazon_album(folder_path)
            if alb:
                meta["lien_amazon_album"] = alb
        except Exception:
            pass
        # conserver ancien lien album si present
        try:
            old_meta_p = os.path.join(dest_dir, "meta.json")
            if os.path.isfile(old_meta_p) and not meta.get("lien_amazon_album"):
                with open(old_meta_p, "r", encoding="utf-8") as f:
                    old = json.load(f) or {}
                if old.get("lien_amazon_album"):
                    meta["lien_amazon_album"] = old["lien_amazon_album"]
        except Exception:
            pass
        with open(os.path.join(dest_dir, "meta.json"), "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
        # dispositifs + meta sortie si presents
        try:
            dev_src = os.path.join(folder_path, DEVICES_FILE)
            if os.path.isfile(dev_src):
                shutil.copy2(dev_src, os.path.join(dest_dir, DEVICES_FILE))
        except Exception:
            pass
        try:
            sm = self._sortie_meta_path(folder_path)
            if os.path.isfile(sm):
                shutil.copy2(sm, os.path.join(dest_dir, "sortie_meta.json"))
        except Exception:
            pass

    def open_local_archive_browser(self):
        """Consulter / exporter les carnets archives sur le PC (disque externe non requis)."""
        root = self._local_archive_root()
        entries = []
        try:
            for name in sorted(os.listdir(root)):
                d = os.path.join(root, name)
                if not os.path.isdir(d):
                    continue
                notes_p = os.path.join(d, NOTES_FILE)
                if not os.path.isfile(notes_p):
                    continue
                meta = {}
                meta_p = os.path.join(d, "meta.json")
                if os.path.isfile(meta_p):
                    try:
                        with open(meta_p, "r", encoding="utf-8") as f:
                            meta = json.load(f) or {}
                    except Exception:
                        pass
                entries.append({
                    "dir": d,
                    "notes": notes_p,
                    "label": meta.get("label") or name,
                    "path_origine": meta.get("path_origine") or "",
                    "last_mirror": meta.get("last_mirror") or "",
                    "nb_obs": meta.get("nb_obs") or "?",
                })
        except Exception as e:
            messagebox.showerror("Archive locale", str(e))
            return

        win = ctk.CTkToplevel(self)
        win.title("Archive locale des carnets (PC)")
        win.geometry("700x520")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Archive locale — recovery sans disque externe",
            font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text="A chaque enregistrement, GeoExif copie le carnet ici :\n%s\n"
                 "Les photos ne sont pas dupliquees — seulement observations.json (+ dispositifs)."
                 % root,
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            wraplength=660, justify="left",
        ).pack(anchor="w", padx=16, pady=(0, 8))

        if not entries:
            ctk.CTkLabel(
                win,
                text="Aucune archive pour l'instant.\n"
                     "Enregistrez au moins une fiche carnet (sortie ouverte) pour creer le miroir.",
                text_color=UI.get("text_dim"),
            ).pack(padx=16, pady=30)
            ctk.CTkButton(win, text="Fermer", command=win.destroy).pack(pady=10)
            return

        scroll = ctk.CTkScrollableFrame(win, fg_color=UI.get("card"))
        scroll.pack(fill="both", expand=True, padx=16, pady=6)

        for e in entries:
            row = ctk.CTkFrame(scroll, fg_color="transparent")
            row.pack(fill="x", pady=4, padx=4)
            txt = "%s  ·  %s obs.  ·  miroir %s" % (
                e["label"], e["nb_obs"], (e["last_mirror"] or "?")[:16],
            )
            ctk.CTkLabel(row, text=txt, text_color=UI.get("text"), anchor="w").pack(
                side="left", fill="x", expand=True
            )

            def exporter(src=e["notes"], lab=e["label"]):
                dest = filedialog.asksaveasfilename(
                    title="Exporter le carnet archive",
                    initialfile="%s_observations.json" % re.sub(r"[^\w\-]+", "_", lab)[:40],
                    defaultextension=".json",
                    filetypes=[("JSON", "*.json")],
                )
                if not dest:
                    return
                try:
                    shutil.copy2(src, dest)
                    messagebox.showinfo("Archive", "Copie enregistree :\n%s" % dest)
                except Exception as ex:
                    messagebox.showerror("Archive", str(ex))

            def restaurer_vers_dossier(src=e["notes"], lab=e["label"]):
                target = filedialog.askdirectory(
                    title="Dossier de sortie ou restaurer « %s »" % lab
                )
                if not target:
                    return
                dest = os.path.join(target, NOTES_FILE)
                if os.path.isfile(dest):
                    if not messagebox.askyesno(
                        "Restaurer",
                        "Remplacer observations.json dans :\n%s ?" % target,
                    ):
                        return
                    try:
                        shutil.copy2(dest, dest + ".avant_restore")
                    except Exception:
                        pass
                try:
                    shutil.copy2(src, dest)
                    messagebox.showinfo(
                        "Restaurer",
                        "Carnet restaure.\nOuvrez ce dossier dans GeoExif pour le revoir.",
                    )
                    self._register_known_folder(target)
                except Exception as ex:
                    messagebox.showerror("Restaurer", str(ex))

            ctk.CTkButton(row, text="Exporter…", width=90, height=28, command=exporter).pack(
                side="right", padx=2
            )
            ctk.CTkButton(
                row, text="Restaurer…", width=90, height=28,
                fg_color=UI.get("success", "#2f9e5f"), command=restaurer_vers_dossier,
            ).pack(side="right", padx=2)

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=16, pady=10)
        ctk.CTkButton(
            bf, text="Ouvrir le dossier archive",
            command=lambda: os.startfile(root) if os.name == "nt" else None,
        ).pack(side="left", padx=3)
        ctk.CTkButton(bf, text="Fermer", command=win.destroy).pack(side="right", padx=3)

    def open_notes_backup_manager(self):
        """Fenêtre pour consulter et restaurer une version antérieure du carnet."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Chargez d'abord un dossier de photos.")
            return

        win = ctk.CTkToplevel(self)
        win.title("💾 Sauvegardes du carnet")
        win.geometry("560x520")
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win,
            text="Versions horodatées de observations.json (rotation automatique, "
                 f"{NOTES_BACKUP_KEEP} max). La restauration remplace le carnet actuel "
                 "(une version de sécurité est créée avant).",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#888888"),
            wraplength=520, justify="left"
        ).pack(anchor="w", padx=15, pady=(15, 8))

        scroll = ctk.CTkScrollableFrame(win, height=340)
        scroll.pack(fill="both", expand=True, padx=15, pady=5)

        def rafraichir():
            for w in scroll.winfo_children():
                w.destroy()
            backups = self._list_notes_backups()
            if not backups:
                ctk.CTkLabel(
                    scroll, text="Aucune sauvegarde versionnée pour l'instant.\n"
                                 "Elles apparaissent après la première modification du carnet.",
                    text_color=UI.get("text_dim", "#666666"), justify="left"
                ).pack(anchor="w", pady=12)
                return
            for path, dt, size in backups:
                ligne = ctk.CTkFrame(scroll, fg_color=UI.get("card", "#1e2c22"), corner_radius=8)
                ligne.pack(fill="x", pady=3)
                n_obs = "?"
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    n_obs = str(len(data)) if isinstance(data, dict) else "?"
                except Exception:
                    pass
                txt = f"{dt.strftime('%d/%m/%Y  %H:%M:%S')}   ·   {n_obs} obs.   ·   {size // 1024} Ko"
                ctk.CTkLabel(ligne, text=txt, font=ctk.CTkFont(size=12), anchor="w").pack(
                    side="left", padx=10, pady=8, fill="x", expand=True
                )
                ctk.CTkButton(
                    ligne, text="Restaurer", width=90, height=28,
                    fg_color=UI.get("accent", "#2b6cb0"), hover_color=UI.get("accent_hover", "#1f4f80"),
                    command=lambda p=path, d=dt: restaurer(p, d)
                ).pack(side="right", padx=8)

        def restaurer(path, dt):
            if not messagebox.askyesno(
                "Confirmer la restauration",
                f"Remplacer le carnet actuel par la version du\n{dt.strftime('%d/%m/%Y à %H:%M:%S')} ?\n\n"
                "L'état actuel sera sauvegardé avant restauration."
            ):
                return
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if not isinstance(data, dict):
                    raise ValueError("Fichier de sauvegarde invalide.")
                # force_backup=True pour garder l'état courant avant écrasement
                self._save_notes_dict(data, create_backup=True, force_backup=True)
                self.refresh_map_markers()
                self.refresh_daily_counter()
                self._refresh_listbox_annotation_status()
                # Réhydrate photos_data pour les observations manuelles
                for cle, obs in data.items():
                    if isinstance(obs, dict) and obs.get("sans_photo") and obs.get("lat") and obs.get("lon"):
                        self.photos_data[cle] = {"lat": obs["lat"], "lon": obs["lon"], "date": ""}
                self.log(f"↩️ Carnet restauré depuis {os.path.basename(path)}")
                messagebox.showinfo("Restauré", f"Carnet restauré ({len(data)} observation(s)).")
                rafraichir()
            except Exception as e:
                messagebox.showerror("Erreur", str(e))

        rafraichir()
        btn_row = ctk.CTkFrame(win, fg_color="transparent")
        btn_row.pack(fill="x", padx=15, pady=12)
        ctk.CTkButton(btn_row, text="🔄 Actualiser", command=rafraichir, width=120,
                      fg_color="#3a3a3a", hover_color="#4a4a4a").pack(side="left")
        ctk.CTkButton(btn_row, text="Fermer", command=win.destroy).pack(side="right")

    @staticmethod
    def _detect_marker_shape(espece):
        """Détermine la forme du marqueur selon le nom d'espèce (traces/indices → formes distinctes)."""
        if not espece:
            return "circle"
        txt = espece.lower()
        for mot_cle, forme in TRACE_SHAPES.items():
            if mot_cle in txt:
                return forme
        return "circle"

    def _get_marker_icon(self, categorie, count=1, shape="circle", sans_photo=False):
        """Génère (et met en cache) une icône par catégorie + forme.
        Formes : circle (animal vu), triangle (empreinte), square (terrier), diamond (coulée).
        sans_photo=True → couleur rose distincte + contour noir pour se démarquer des photos."""
        if not hasattr(self, "_marker_icon_cache"):
            self._marker_icon_cache = {}
        cle = (categorie, shape, count if count > 1 else 1, bool(sans_photo))
        if cle in self._marker_icon_cache:
            return self._marker_icon_cache[cle]

        if sans_photo:
            couleur_hex = SANS_PHOTO_COLOR
            outline = (20, 20, 20, 255)
            outline_w = 3
        else:
            couleur_hex = CATEGORY_COLORS.get(categorie, "#7f7f7f")
            outline = (255, 255, 255, 255)
            outline_w = 2

        couleur_rgb = tuple(int(couleur_hex[i:i + 2], 16) for i in (1, 3, 5))

        size = 32 if count <= 1 else 38
        if sans_photo:
            size += 2  # légèrement plus grand pour rester visible
        img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)
        fill = couleur_rgb + (240,)
        m = 2

        if shape == "triangle":
            pts = [(size // 2, m), (size - m, size - m), (m, size - m)]
            draw.polygon(pts, fill=fill, outline=outline)
        elif shape == "square":
            draw.rectangle([m + 2, m + 2, size - m - 2, size - m - 2], fill=fill, outline=outline, width=outline_w)
        elif shape == "diamond":
            cx, cy = size // 2, size // 2
            pts = [(cx, m), (size - m, cy), (cx, size - m), (m, cy)]
            draw.polygon(pts, fill=fill, outline=outline)
        else:
            draw.ellipse([m, m, size - m, size - m], fill=fill, outline=outline, width=outline_w)
            # Anneau intérieur blanc pour les obs. sans photo (effet « cible »)
            if sans_photo:
                inn = m + 5
                draw.ellipse([inn, inn, size - inn, size - inn], outline=(255, 255, 255, 220), width=2)

        if count > 1:
            texte = str(count) if count < 100 else "99+"
            try:
                bbox = draw.textbbox((0, 0), texte)
                tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
                draw.text(((size - tw) / 2 - bbox[0], (size - th) / 2 - bbox[1]), texte, fill=(255, 255, 255, 255))
            except Exception:
                pass

        photo = ImageTk.PhotoImage(img)
        self._marker_icon_cache[cle] = photo
        return photo

    def _on_marker_click(self, fichiers):
        """Clic sur un point de la carte : ouvre la photo, ou affiche le détail d'une observation manuelle."""
        notes = self._load_notes_dict()
        # Sépare observations manuelles (sans photo) et fichiers photo
        manuels = [f for f in fichiers if str(f).startswith("_manuel_") or notes.get(f, {}).get("sans_photo")]
        photos = [f for f in fichiers if f not in manuels]

        if len(fichiers) == 1 and photos:
            self.select_photo_by_filename(photos[0])
            self.tab_view.set("📝  Carnet")
            return
        if len(fichiers) == 1 and manuels:
            self._show_manual_observation_detail(manuels[0])
            return

        win = ctk.CTkToplevel(self)
        win.title("%d observation(s) a ce point" % len(fichiers))
        win.geometry("560x600")
        try:
            win.configure(fg_color=UI.get("bg", "#f5f7f6"))
        except Exception:
            pass
        self._prepare_tool_window(win)
        ctk.CTkLabel(
            win,
            text="%d element(s) au meme endroit :" % len(fichiers),
            font=ctk.CTkFont(size=15, weight="bold"),
            text_color=UI.get("text", "#111111"),
        ).pack(anchor="w", padx=14, pady=12)
        scroll = ctk.CTkScrollableFrame(win, fg_color=UI.get("card", "#ffffff"))
        scroll.pack(fill="both", expand=True, padx=12, pady=(0, 12))

        # Contraste force : texte clair sur fond sombre OU texte sombre sur fond clair
        is_light = UI.get("ctk_mode") == "Light"
        if is_light:
            btn_fg, btn_hover, btn_txt = "#2c3e50", "#1a252f", "#ffffff"
            man_fg, man_hover = "#b7791f", "#975a16"
        else:
            btn_fg, btn_hover, btn_txt = "#4a5568", "#2d3748", "#ffffff"
            man_fg, man_hover = "#c4841a", "#9a6510"

        for fn in sorted(fichiers):
            if fn in manuels:
                data = notes.get(fn, {})
                label = "Sans photo · %s — %s (%s)" % (
                    data.get("espece", "?"),
                    data.get("type_observation", "sans photo"),
                    data.get("heure", "?"),
                )
                def voir_manuel(f=fn):
                    self._show_manual_observation_detail(f)
                    win.destroy()
                ctk.CTkButton(
                    scroll, text=label, command=voir_manuel, anchor="w",
                    fg_color=man_fg, hover_color=man_hover,
                    text_color="#ffffff",
                    font=ctk.CTkFont(size=13),
                    height=34,
                ).pack(fill="x", pady=3, padx=4)
            else:
                def choisir(f=fn):
                    self.select_photo_by_filename(f)
                    try:
                        self.tab_view.set("📝  Carnet")
                    except Exception:
                        pass
                    win.destroy()
                ctk.CTkButton(
                    scroll, text=fn, command=choisir, anchor="w",
                    fg_color=btn_fg, hover_color=btn_hover,
                    text_color=btn_txt,
                    font=ctk.CTkFont(size=13),
                    height=34,
                ).pack(fill="x", pady=3, padx=4)

    def _show_manual_observation_detail(self, cle):
        """Affiche le détail d'une observation sans photo (et propose la suppression)."""
        notes = self._load_notes_dict()
        data = notes.get(cle, {})
        win = ctk.CTkToplevel(self)
        win.title("👁️ Observation sans photo")
        win.geometry("420x360")
        self._prepare_tool_window(win)

        lignes = [
            f"Type : {data.get('type_observation', '—')}",
            f"Catégorie : {data.get('categorie', '—')}",
            f"Espèce : {data.get('espece', '—')}",
            f"Nombre : {data.get('nombre', '—')}",
            f"Heure : {data.get('heure', '—')}",
            f"Lieu : {data.get('lieu', '—')}",
        ]
        if data.get("lat") and data.get("lon"):
            lignes.append(f"GPS : {data['lat']:.5f}, {data['lon']:.5f}")
        if data.get("notes_libres"):
            lignes.append(f"Notes : {data.get('notes_libres')}")

        ctk.CTkLabel(win, text="\n".join(lignes), justify="left", anchor="w",
                     font=ctk.CTkFont(size=12)).pack(anchor="w", padx=15, pady=15)

        def supprimer():
            if not messagebox.askyesno("Confirmer", "Supprimer cette observation sans photo ?"):
                return
            notes2 = self._load_notes_dict()
            if cle in notes2:
                del notes2[cle]
                try:
                    self._save_notes_dict(notes2, force_backup=True)
                except Exception as e:
                    messagebox.showerror("Erreur", str(e))
                    return
            self.photos_data.pop(cle, None)
            self.refresh_map_markers()
            self.refresh_daily_counter()
            self.log(f"🗑️ Observation manuelle supprimée : {data.get('espece', cle)}")
            win.destroy()

        btn_row = ctk.CTkFrame(win, fg_color="transparent")
        btn_row.pack(fill="x", padx=15, pady=15)
        ctk.CTkButton(btn_row, text="🗑️ Supprimer", fg_color="#8a2d24", hover_color="#6b221c",
                      command=supprimer).pack(side="left")
        ctk.CTkButton(btn_row, text="Fermer", command=win.destroy).pack(side="right")

    def toggle_place_observation_mode(self):
        """Active/désactive le mode « clic sur la carte pour placer une observation »."""
        if not self.photo_folder_path:
            messagebox.showwarning(
                "Dossier manquant",
                "Chargez d'abord un dossier de photos : l'observation sera enregistrée dans ce dossier."
            )
            return
        self._place_obs_mode = not self._place_obs_mode
        if self._place_obs_mode:
            self.btn_place_obs.configure(
                text="✕ Cliquez sur la carte…", fg_color="#c0392b", hover_color="#a93226"
            )
            self.tab_view.set("🗺️  Carte")
            self.log("📍 Mode placement actif : cliquez sur la carte → formulaire (icône rose).")
        else:
            self.btn_place_obs.configure(
                text="➕ Placer observation", fg_color=SANS_PHOTO_COLOR, hover_color="#c2185b"
            )
            self.log("Mode placement désactivé.")

    def _on_map_left_click(self, coords):
        """Callback carte : point d'écoute / observation sans photo / dispositif."""
        try:
            lat, lon = float(coords[0]), float(coords[1])
        except Exception:
            return
        # Pick GPS pour la fenêtre Birda CLI (remplit Lat/Lon)
        pick = getattr(self, "_birda_cli_gps_pick", None)
        if pick:
            self._birda_cli_gps_pick = None
            try:
                lv = pick.get("lat_var")
                ov = pick.get("lon_var")
                if lv is not None:
                    lv.set("%.6f" % lat)
                if ov is not None:
                    ov.set("%.6f" % lon)
                hint = pick.get("hint")
                if hint is not None:
                    try:
                        hint.configure(text="Point carte → %.5f , %.5f" % (lat, lon))
                    except Exception:
                        pass
                self.log("Birda CLI GPS (clic carte) : %.5f, %.5f" % (lat, lon))
                try:
                    w = pick.get("win")
                    if w is not None and w.winfo_exists():
                        w.lift()
                        w.focus_force()
                except Exception:
                    pass
            except Exception as e:
                self.log("GPS Birda CLI : %s" % e)
            return
        if getattr(self, "device_place_mode", False):
            self.device_place_mode = False
            self.log("Dispositif — point : %.5f, %.5f" % (lat, lon))
            self.open_device_dialog(lat=lat, lon=lon)
            return
        # Point d'écoute pour obs. 🎧 / 📍 déjà sélectionnée
        if getattr(self, "_point_ecoute_mode", False):
            self._point_ecoute_mode = False
            key = getattr(self, "_current_note_key", None)
            if not key or not self.photo_folder_path:
                return
            notes = self._load_notes_dict(self.photo_folder_path) or {}
            data = notes.get(key)
            if not isinstance(data, dict):
                data = {}
                notes[key] = data
            data["lat"] = lat
            data["lon"] = lon
            notes[key] = data
            try:
                self._save_notes_dict(notes, force_backup=True)
            except Exception as e:
                messagebox.showerror("GPS", str(e))
                return
            self.photos_data[key] = dict(self.photos_data.get(key) or {})
            self.photos_data[key].update({"lat": lat, "lon": lon, "sans_photo": True})
            try:
                self.refresh_map_markers()
            except Exception:
                pass
            self.log("Point d'écoute GPS pour %s : %.5f, %.5f" % (key, lat, lon))
            messagebox.showinfo(
                "Point d'écoute",
                "GPS enregistré sur « %s »\n%.5f, %.5f"
                % (data.get("espece") or key, lat, lon),
            )
            return
        if not getattr(self, "_place_obs_mode", False):
            return
        self._place_obs_mode = False
        try:
            self.btn_place_obs.configure(
                text="➕ Placer observation", fg_color=SANS_PHOTO_COLOR, hover_color="#c2185b"
            )
        except Exception:
            pass
        self.log("Point choisi : %.5f, %.5f" % (lat, lon))
        self.open_manual_observation_dialog(lat=lat, lon=lon)

    def recharger_donnees_gps(self):
        """Recharge les coordonnées GPS depuis le disque (contrairement à refresh_map_markers qui ne
        fait que redessiner ce qui est déjà en mémoire). Utile si le chargement initial a échoué
        (délai dépassé, dossier très volumineux...)."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Sélectionnez d'abord un dossier de photos.")
            return
        exe_path = self.get_exiftool_path()
        if not exe_path:
            messagebox.showerror("ExifTool introuvable", "ExifTool n'est pas accessible.")
            return
        self.log("🔄 Rechargement des coordonnées GPS depuis le disque...")
        threading.Thread(target=self._check_existing_gps_worker, args=(exe_path,), daemon=True).start()

    def refresh_map_markers(self):
        """Redessine les marqueurs de la carte selon les filtres de catégorie actifs, en regroupant les points très proches."""
        if not hasattr(self, "map_widget"):
            return
        for m in self.active_markers:
            try:
                m.delete()
            except Exception:
                pass
        self.active_markers = []

        notes = self._load_notes_dict()
        actives = {cat for cat, var in self.category_filters.items() if var.get()}

        # Regroupe les photos prises quasiment au même endroit (~10 m) pour éviter la superposition illisible de marqueurs
        clusters = {}
        for filename, info in self.photos_data.items():
            lat, lon = info.get("lat"), info.get("lon")
            if not (lat and lon):
                continue
            note = notes.get(filename, {})
            categorie = note.get("categorie") or "Non classé"
            if categorie not in actives:
                continue
            espece = note.get("espece") or ""
            forme = self._detect_marker_shape(espece)
            is_manuel = self._is_virtual_observation(filename, note)
            # Forme d'après le type d'observation si présent
            type_obs = (note.get("type_observation") or note.get("type_indice") or "").lower()
            if "empreinte" in type_obs or "trace" in type_obs:
                forme = "triangle"
            elif "terrier" in type_obs:
                forme = "square"
            elif "coulée" in type_obs or "coulee" in type_obs or "passage" in type_obs:
                forme = "diamond"
            elif "prise de son" in type_obs or "birda" in type_obs or "chirpity" in type_obs or "écoute" in type_obs:
                forme = "diamond"  # point d'écoute distinct
            cle = (round(lat, 4), round(lon, 4))
            c = clusters.setdefault(cle, {
                "lat": lat, "lon": lon, "fichiers": [],
                "categories": collections.Counter(), "shapes": collections.Counter(),
                "sans_photo": 0,
            })
            c["fichiers"].append(filename)
            c["categories"][categorie] += 1
            c["shapes"][forme] += 1
            if is_manuel:
                c["sans_photo"] += 1

        coords_shown = []
        for c in clusters.values():
            categorie_dominante = c["categories"].most_common(1)[0][0]
            forme_dominante = "circle"
            for forme_candidat in ("triangle", "square", "diamond"):
                if c["shapes"].get(forme_candidat, 0) > 0:
                    forme_dominante = forme_candidat
                    break
            fichiers = c["fichiers"]
            # Si le cluster est majoritairement « sans photo », icône rose distincte
            est_sans_photo = c.get("sans_photo", 0) >= max(1, (len(fichiers) + 1) // 2)
            icon = self._get_marker_icon(
                categorie_dominante, len(fichiers),
                shape=forme_dominante, sans_photo=est_sans_photo
            )
            marker = self.map_widget.set_marker(
                c["lat"], c["lon"], text=None, icon=icon,
                command=lambda m, fl=fichiers: self._on_marker_click(fl)
            )
            self.active_markers.append(marker)
            coords_shown.append((c["lat"], c["lon"]))

        if coords_shown:
            lats = [c[0] for c in coords_shown]
            lons = [c[1] for c in coords_shown]
            if len(coords_shown) == 1:
                self.map_widget.set_position(*coords_shown[0])
                self.map_widget.set_zoom(14)
            else:
                self.map_widget.fit_bounding_box((max(lats), min(lons)), (min(lats), max(lons)))

    def filter_photo_list(self, event=None):
        """Filtre la liste / grilles selon recherche (nom, espèce, catégorie, lieu)."""
        query = self.entry_search_photo.get().strip().lower()
        notes = self._load_notes_dict() if query else {}
        names = []
        for img in self.get_supported_media():
            if not query:
                names.append(img)
                continue
            data = notes.get(img, {})
            haystack = " ".join([
                img.lower(), str(data.get("espece", "")).lower(),
                str(data.get("categorie", "")).lower(), str(data.get("lieu", "")).lower(),
                "video" if self._is_video_file(img) else "",
            ])
            if query in haystack:
                names.append(img)
        self._filtered_media_names = names
        if getattr(self, "carnet_view_mode", "list") == "thumbs":
            self._rebuild_thumb_explorer(lazy=True)
        else:
            self._populate_listbox_from_filtered()

    def navigate_photo(self, direction):
        """Sélectionne la photo précédente/suivante dans la liste (raccourci Ctrl+Flèches)."""
        taille = self.photo_listbox.size()
        if taille == 0:
            return
        courant = self.photo_listbox.curselection()
        idx = courant[0] if courant else -1
        nouvel_idx = idx + direction
        if 0 <= nouvel_idx < taille:
            self.photo_listbox.selection_clear(0, tk.END)
            self.photo_listbox.selection_set(nouvel_idx)
            self.photo_listbox.activate(nouvel_idx)
            self.photo_listbox.see(nouvel_idx)
            self.on_photo_select(None)

    def _autosave_tick(self):
        """Sauvegarde silencieusement l'observation en cours toutes les 60 secondes."""
        if self.selected_photo_path and self.photo_folder_path:
            try:
                self.save_current_note(silent=True)
            except Exception:
                pass
        self.after(60000, self._autosave_tick)

    @staticmethod
    def moon_phase(date_obj):
        """Calcule la phase lunaire approximative pour une date donnée (algorithme local, sans API)."""
        reference = datetime(2000, 1, 6)
        diff_days = (date_obj - reference).total_seconds() / 86400.0
        lunations = diff_days / 29.53058867
        phase = lunations % 1
        if phase < 0:
            phase += 1
        if phase < 0.03 or phase > 0.97:
            return "🌑 Nouvelle lune"
        elif phase < 0.22:
            return "🌒 Premier croissant"
        elif phase < 0.28:
            return "🌓 Premier quartier"
        elif phase < 0.47:
            return "🌔 Gibbeuse croissante"
        elif phase < 0.53:
            return "🌕 Pleine lune"
        elif phase < 0.72:
            return "🌖 Gibbeuse décroissante"
        elif phase < 0.78:
            return "🌗 Dernier quartier"
        else:
            return "🌘 Dernier croissant"

    def undo_last_sync(self):
        """Restaure les photos à leur état d'avant la dernière synchronisation GPX."""
        if not self.last_backup_dir or not os.path.isdir(self.last_backup_dir):
            messagebox.showinfo("Rien à annuler", "Aucune synchronisation récente à annuler.")
            return

        fichiers = os.listdir(self.last_backup_dir)
        if not messagebox.askyesno(
            "Confirmer l'annulation",
            f"Restaurer {len(fichiers)} photo(s) à leur état d'avant la dernière synchronisation GPX ?\n"
            "Les coordonnées ajoutées lors de cette synchro seront perdues."
        ):
            return

        dossier_propre = os.path.normpath(self.photo_folder_path)
        restaurees = 0
        for fname in fichiers:
            src = os.path.join(self.last_backup_dir, fname)
            dst = os.path.join(dossier_propre, fname)
            try:
                shutil.copy2(src, dst)
                restaurees += 1
            except Exception:
                pass

        self.log(f"↩️ {restaurees}/{len(fichiers)} photo(s) restaurée(s) à leur état précédent.")
        self.last_backup_dir = None
        self.btn_undo_sync.configure(state="disabled")

        exe_path = self.get_exiftool_path()
        if exe_path:
            threading.Thread(target=self.extract_coordinates_after_sync, args=(exe_path,), daemon=True).start()

    # --- Édition manuelle / batch des points GPS ---


    def open_affut_series_dialog(self):
        """Série d'affût : même point GPS + tranche horaire approximative pour un groupe de photos."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Chargez d'abord un dossier de photos.")
            return
        selection = self.photo_listbox.curselection()
        if not selection:
            messagebox.showinfo(
                "Affût / série",
                "Sélectionnez un groupe de photos dans la liste\n"
                "(clic, Maj+clic ou Ctrl+clic), puis relancez.",
            )
            return
        fichiers = [self.photo_listbox.get(i) for i in selection]
        # ignorer les vidéos pour l'écriture EXIF date/gps photo
        video_ext = (".mp4", ".mov", ".avi", ".mkv", ".m4v", ".mpg", ".mpeg", ".wmv", ".mts")
        fichiers = [f for f in fichiers if not str(f).lower().endswith(video_ext)]
        if not fichiers:
            messagebox.showinfo("Affût / série", "Aucune photo dans la sélection (vidéos ignorées).")
            return

        premier = self.photos_data.get(fichiers[0], {})
        lat0 = premier.get("lat")
        lon0 = premier.get("lon")

        win = ctk.CTkToplevel(self)
        win.title("Affût / série — GPS + tranche horaire")
        win.geometry("780x720")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win,
            text="Même lieu GPS pour %d photo(s) + horaire approximatif d'affût" % len(fichiers),
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color=UI.get("text"),
        ).pack(anchor="w", padx=14, pady=(12, 4))
        ctk.CTkLabel(
            win,
            text="Cas typique : série prise au même poste sans trace GPX continue. "
                 "Les heures sont reparties entre début et fin d'affût (ou identiques).",
            font=ctk.CTkFont(size=11),
            text_color=UI.get("text_dim"),
            wraplength=740, justify="left",
        ).pack(anchor="w", padx=14, pady=(0, 8))

        # Liste courte
        if len(fichiers) <= 8:
            ctk.CTkLabel(
                win, text="• " + "  |  ".join(fichiers),
                font=ctk.CTkFont(size=10), text_color=UI.get("text_muted"),
                wraplength=740, justify="left",
            ).pack(anchor="w", padx=14, pady=(0, 6))
        else:
            ctk.CTkLabel(
                win,
                text="• %s  …  %s  (%d fichiers)" % (fichiers[0], fichiers[-1], len(fichiers)),
                font=ctk.CTkFont(size=10), text_color=UI.get("text_muted"),
            ).pack(anchor="w", padx=14, pady=(0, 6))

        # GPS
        gps_fr = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=10)
        gps_fr.pack(fill="x", padx=14, pady=6)
        ctk.CTkLabel(
            gps_fr, text="Point GPS commun",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=10, pady=(8, 4))

        coord_fr = ctk.CTkFrame(gps_fr, fg_color="transparent")
        coord_fr.pack(fill="x", padx=10, pady=4)
        lat_var = tk.StringVar(value=("" if lat0 is None else "%.6f" % float(lat0)))
        lon_var = tk.StringVar(value=("" if lon0 is None else "%.6f" % float(lon0)))
        ctk.CTkLabel(coord_fr, text="Lat", text_color=UI.get("text")).pack(side="left")
        ctk.CTkEntry(coord_fr, textvariable=lat_var, width=120).pack(side="left", padx=4)
        ctk.CTkLabel(coord_fr, text="Lon", text_color=UI.get("text")).pack(side="left", padx=(10, 0))
        ctk.CTkEntry(coord_fr, textvariable=lon_var, width=120).pack(side="left", padx=4)

        map_host = tk.Frame(gps_fr, bg=UI.get("card", "#1c2620"), highlightthickness=0)
        map_host.pack(fill="both", expand=True, padx=10, pady=(4, 10))
        map_host.configure(height=220)
        try:
            map_w = TkinterMapView(map_host, corner_radius=0)
            map_w.pack(fill="both", expand=True)
            if lat0 and lon0:
                map_w.set_position(float(lat0), float(lon0))
                map_w.set_zoom(14)
            else:
                map_w.set_position(49.16, 5.38)
                map_w.set_zoom(10)
            marker = {"m": None}

            def on_map_click(coords):
                try:
                    la, lo = float(coords[0]), float(coords[1])
                except Exception:
                    return
                lat_var.set("%.6f" % la)
                lon_var.set("%.6f" % lo)
                try:
                    if marker["m"] is not None:
                        marker["m"].delete()
                except Exception:
                    pass
                try:
                    marker["m"] = map_w.set_marker(la, lo, text="Affût")
                except Exception:
                    pass

            map_w.add_left_click_map_command(on_map_click)
            if lat0 and lon0:
                try:
                    marker["m"] = map_w.set_marker(float(lat0), float(lon0), text="Affût")
                except Exception:
                    pass
        except Exception:
            ctk.CTkLabel(
                map_host, text="Carte indisponible — saisissez lat/lon à la main.",
                text_color=UI.get("text_dim"),
            ).pack(pady=20)

        # Horaires
        time_fr = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=10)
        time_fr.pack(fill="x", padx=14, pady=6)
        ctk.CTkLabel(
            time_fr, text="Tranche horaire approximative",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=10, pady=(8, 4))

        # date du jour par défaut = aujourd'hui ou date dossier
        default_date = datetime.now().strftime("%Y-%m-%d")
        # tenter date depuis nom dossier
        try:
            base = os.path.basename(self.photo_folder_path or "")
            if len(base) >= 10 and base[4] == "-" and base[7] == "-":
                default_date = base[:10]
        except Exception:
            pass

        row1 = ctk.CTkFrame(time_fr, fg_color="transparent")
        row1.pack(fill="x", padx=10, pady=4)
        date_var = tk.StringVar(value=default_date)
        start_var = tk.StringVar(value="06:30")
        end_var = tk.StringVar(value="09:00")
        ctk.CTkLabel(row1, text="Date (AAAA-MM-JJ)", text_color=UI.get("text")).pack(side="left")
        ctk.CTkEntry(row1, textvariable=date_var, width=120).pack(side="left", padx=6)
        ctk.CTkLabel(row1, text="Début", text_color=UI.get("text")).pack(side="left", padx=(12, 0))
        ctk.CTkEntry(row1, textvariable=start_var, width=70).pack(side="left", padx=4)
        ctk.CTkLabel(row1, text="Fin", text_color=UI.get("text")).pack(side="left", padx=(8, 0))
        ctk.CTkEntry(row1, textvariable=end_var, width=70).pack(side="left", padx=4)

        mode_var = tk.StringVar(value="repartir")
        row2 = ctk.CTkFrame(time_fr, fg_color="transparent")
        row2.pack(fill="x", padx=10, pady=(4, 10))
        ctk.CTkRadioButton(
            row2, text="Repartir les heures entre début et fin (ordre de la selection)",
            variable=mode_var, value="repartir", text_color=UI.get("text"),
        ).pack(anchor="w", pady=2)
        ctk.CTkRadioButton(
            row2, text="Même heure pour toutes (heure de début)",
            variable=mode_var, value="identique", text_color=UI.get("text"),
        ).pack(anchor="w", pady=2)
        ctk.CTkLabel(
            time_fr,
            text="Format heure : HH:MM  ·  L'ordre des photos suit la selection dans la liste.",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_muted"),
        ).pack(anchor="w", padx=10, pady=(0, 8))

        apply_gps = tk.BooleanVar(value=True)
        apply_time = tk.BooleanVar(value=True)
        opt = ctk.CTkFrame(win, fg_color="transparent")
        opt.pack(fill="x", padx=14, pady=4)
        ctk.CTkCheckBox(opt, text="Ecrire le GPS", variable=apply_gps, text_color=UI.get("text")).pack(side="left", padx=4)
        ctk.CTkCheckBox(opt, text="Ecrire les horaires", variable=apply_time, text_color=UI.get("text")).pack(side="left", padx=12)

        def apply():
            do_gps = bool(apply_gps.get())
            do_time = bool(apply_time.get())
            if not do_gps and not do_time:
                messagebox.showwarning("Affût", "Cochez GPS et/ou horaires.")
                return
            lat = lon = None
            if do_gps:
                try:
                    lat = float(lat_var.get().replace(",", "."))
                    lon = float(lon_var.get().replace(",", "."))
                except Exception:
                    messagebox.showwarning("GPS", "Latitude / longitude invalides (ou cliquez la carte).")
                    return
            times = None
            if do_time:
                try:
                    d = datetime.strptime(date_var.get().strip()[:10], "%Y-%m-%d")
                    t0 = datetime.strptime(start_var.get().strip(), "%H:%M")
                    t1 = datetime.strptime(end_var.get().strip(), "%H:%M")
                    start_dt = d.replace(hour=t0.hour, minute=t0.minute, second=0)
                    end_dt = d.replace(hour=t1.hour, minute=t1.minute, second=0)
                    if end_dt < start_dt:
                        messagebox.showwarning("Horaires", "L'heure de fin doit etre >= debut.")
                        return
                    n = len(fichiers)
                    if mode_var.get() == "identique" or n == 1:
                        times = [start_dt] * n
                    else:
                        span = (end_dt - start_dt).total_seconds()
                        times = [
                            start_dt + timedelta(seconds=span * i / (n - 1))
                            for i in range(n)
                        ]
                except Exception as e:
                    messagebox.showwarning(
                        "Horaires",
                        "Date ou heures invalides.\nDate: AAAA-MM-JJ · Heure: HH:MM\n%s" % e,
                    )
                    return

            if not messagebox.askyesno(
                "Confirmer",
                "Appliquer a %d photo(s) ?\nGPS: %s\nHoraires: %s" % (
                    len(fichiers),
                    ("%.5f, %.5f" % (lat, lon)) if do_gps else "non",
                    mode_var.get() if do_time else "non",
                ),
            ):
                return
            win.destroy()
            self._write_affut_series(fichiers, lat, lon, times, do_gps=do_gps, do_time=do_time)

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=14, pady=(6, 14))
        ctk.CTkButton(
            bf, text="Appliquer a la selection", height=36,
            fg_color=UI.get("success", "#2f9e5f"), command=apply,
        ).pack(side="left", fill="x", expand=True, padx=(0, 6))
        ctk.CTkButton(
            bf, text="Annuler", height=36,
            fg_color=UI.get("card_alt"), command=win.destroy,
        ).pack(side="left", fill="x", expand=True)

    def _write_affut_series(self, filenames, lat, lon, times, do_gps=True, do_time=True):
        """Ecrit GPS et/ou DateTimeOriginal sur un groupe de photos (serie d'affut) avec barre de progression."""
        exe_path = self.get_exiftool_path()
        if not exe_path:
            messagebox.showerror("ExifTool", "exiftool.exe introuvable.")
            return
        dossier = os.path.normpath(self.photo_folder_path)
        total = len(filenames)
        self.log(
            "Serie affut : %d photo(s) — GPS=%s horaires=%s" % (
                total, "oui" if do_gps else "non", "oui" if do_time else "non",
            )
        )

        # Fenetre de progression dediee
        prog_win = ctk.CTkToplevel(self)
        prog_win.title("Affût / série — écriture en cours")
        prog_win.geometry("420x160")
        try:
            prog_win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        try:
            prog_win.transient(self)
        except Exception:
            pass
        ctk.CTkLabel(
            prog_win,
            text="Ecriture ExifTool (GPS / horaires)…",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(16, 6))
        lbl = ctk.CTkLabel(
            prog_win, text="0 / %d" % total,
            text_color=UI.get("text_dim"),
        )
        lbl.pack(anchor="w", padx=16)
        bar = ctk.CTkProgressBar(prog_win, height=16)
        bar.pack(fill="x", padx=16, pady=12)
        bar.set(0)
        detail = ctk.CTkLabel(
            prog_win, text="",
            font=ctk.CTkFont(size=11),
            text_color=UI.get("text_muted", UI.get("text_dim")),
        )
        detail.pack(anchor="w", padx=16, pady=(0, 10))
        try:
            prog_win.grab_set()
        except Exception:
            pass

        cancel = {"flag": False}

        def ask_cancel():
            cancel["flag"] = True
            detail.configure(text="Annulation apres le fichier en cours…")

        ctk.CTkButton(
            prog_win, text="Annuler", width=100, height=28,
            fg_color=UI.get("card_alt"), command=ask_cancel,
        ).pack(pady=(0, 12))

        def set_prog(i, fn, ok, err):
            def _ui():
                try:
                    frac = (i / total) if total else 1.0
                    bar.set(max(0.0, min(1.0, frac)))
                    lbl.configure(text="%d / %d  —  OK %d · erreurs %d" % (i, total, ok, err))
                    detail.configure(text=fn or "")
                    try:
                        self._set_progress(frac, "Affût %d/%d" % (i, total))
                    except Exception:
                        pass
                except Exception:
                    pass
            try:
                self.after(0, _ui)
            except Exception:
                pass

        def worker():
            ok, err = 0, 0
            startupinfo = self._exiftool_startupinfo()
            for i, fn in enumerate(filenames):
                if cancel["flag"]:
                    break
                set_prog(i, fn, ok, err)
                path = os.path.join(dossier, fn)
                if not os.path.isfile(path):
                    err += 1
                    continue
                cmd = [exe_path, "-overwrite_original"]
                if do_gps and lat is not None and lon is not None:
                    cmd += [
                        "-GPSLatitude=%s" % abs(lat),
                        "-GPSLatitudeRef=%s" % ("N" if lat >= 0 else "S"),
                        "-GPSLongitude=%s" % abs(lon),
                        "-GPSLongitudeRef=%s" % ("E" if lon >= 0 else "W"),
                    ]
                if do_time and times is not None:
                    dt = times[i]
                    stamp = dt.strftime("%Y:%m:%d %H:%M:%S")
                    cmd += [
                        "-DateTimeOriginal=%s" % stamp,
                        "-CreateDate=%s" % stamp,
                        "-ModifyDate=%s" % stamp,
                    ]
                cmd.append(path)
                try:
                    r = subprocess.run(
                        cmd, capture_output=True, text=True,
                        timeout=60, startupinfo=startupinfo,
                        encoding="utf-8", errors="replace",
                    )
                    if r.returncode == 0:
                        ok += 1
                        if fn not in self.photos_data:
                            self.photos_data[fn] = {}
                        if do_gps and lat is not None:
                            self.photos_data[fn]["lat"] = lat
                            self.photos_data[fn]["lon"] = lon
                        if do_time and times is not None:
                            self.photos_data[fn]["datetime"] = times[i].isoformat(timespec="seconds")
                    else:
                        err += 1
                        self.log("Affut EXIF echec %s: %s" % (fn, (r.stderr or r.stdout or "")[:200]))
                except Exception as e:
                    err += 1
                    self.log("Affut erreur %s: %s" % (fn, e))
                set_prog(i + 1, fn, ok, err)

            def done():
                try:
                    prog_win.grab_release()
                except Exception:
                    pass
                try:
                    prog_win.destroy()
                except Exception:
                    pass
                try:
                    self._set_progress(0, "Prêt")
                except Exception:
                    pass
                stopped = " (annule)" if cancel["flag"] else ""
                self.log("Serie affut terminee%s — OK %d · erreurs %d" % (stopped, ok, err))
                messagebox.showinfo(
                    "Affût / série",
                    "Traitement termine%s.\nReussis : %d\nErreurs : %d" % (stopped, ok, err),
                )
                try:
                    self.recharger_donnees_gps()
                except Exception:
                    pass
                try:
                    self.refresh_map_markers()
                except Exception:
                    pass
            try:
                self.after(0, done)
            except Exception:
                pass

        threading.Thread(target=worker, daemon=True).start()

    def open_edit_gps_dialog(self):
        """Permet de modifier le point GPS d'une ou plusieurs photos sélectionnées dans la liste
        (Ctrl+clic / Maj+clic pour multi-sélection) : clic sur la carte OU saisie manuelle,
        puis écriture des nouvelles coordonnées via ExifTool."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Chargez d'abord un dossier de photos.")
            return

        selection = self.photo_listbox.curselection()
        if not selection:
            messagebox.showinfo(
                "Aucune sélection",
                "Sélectionnez une ou plusieurs photos dans la liste (Ctrl+clic pour en prendre plusieurs)."
            )
            return

        fichiers = [self.photo_listbox.get(i) for i in selection]
        premier = self.photos_data.get(fichiers[0], {})
        lat0 = premier.get("lat")
        lon0 = premier.get("lon")

        win = ctk.CTkToplevel(self)
        win.title(f"📍 Modifier le GPS — {len(fichiers)} photo(s)")
        win.geometry("760x680")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)
        win.after(250, lambda: self._style_dialog(win))

        ctk.CTkLabel(
            win,
            text=f"{len(fichiers)} photo(s) sélectionnée(s). Cliquez sur la carte pour placer le point, "
                 "ou saisissez les coordonnées à la main.",
            font=ctk.CTkFont(size=12), justify="left", wraplength=730
        ).pack(anchor="w", padx=15, pady=(15, 8))

        if len(fichiers) <= 6:
            ctk.CTkLabel(
                win, text="• " + "\n• ".join(fichiers), font=ctk.CTkFont(size=10),
                text_color=UI.get("text_dim", "#666666"), justify="left", anchor="w"
            ).pack(anchor="w", padx=20, pady=(0, 8))
        else:
            ctk.CTkLabel(
                win, text=f"• {fichiers[0]}\n• …\n• {fichiers[-1]}", font=ctk.CTkFont(size=10),
                text_color=UI.get("text_dim", "#666666"), justify="left", anchor="w"
            ).pack(anchor="w", padx=20, pady=(0, 8))

        # --- Carte interactive : clic gauche = déplace/pose le marqueur ---
        # NB : un parent CTkFrame en fg_color="transparent" fait planter TkinterMapView
        # ("unknown color name 'transparent'"), car il détecte sa propre couleur de fond
        # d'après celle de son parent. On fixe donc bg_color explicitement pour l'éviter.
        map_frame = ctk.CTkFrame(win, fg_color="transparent")
        map_frame.pack(fill="both", expand=True, padx=15, pady=(0, 8))

        map_w = TkinterMapView(map_frame, corner_radius=10, bg_color="#242424")
        map_w.pack(fill="both", expand=True)
        try:
            map_w.set_tile_server(
                "https://data.geopf.fr/wmts?SERVICE=WMTS&REQUEST=GetTile&VERSION=1.0.0&LAYER=GEOGRAPHICALGRIDSYSTEMS.PLANIGNV2"
                "&STYLE=normal&FORMAT=image/png&TILEMATRIXSET=PM&TILEMATRIX={z}&TILEROW={y}&TILECOL={x}",
                max_zoom=19
            )
        except Exception:
            pass

        lat_depart = lat0 if lat0 is not None else 49.1627
        lon_depart = lon0 if lon0 is not None else 5.3854
        map_w.set_position(lat_depart, lon_depart)
        map_w.set_zoom(15 if lat0 is not None else 9)

        marker_ref = {"m": None}

        def poser_marqueur(lat, lon):
            if marker_ref["m"] is not None:
                try:
                    marker_ref["m"].delete()
                except Exception:
                    pass
            marker_ref["m"] = map_w.set_marker(lat, lon, text=f"{len(fichiers)} photo(s)")

        if lat0 is not None and lon0 is not None:
            poser_marqueur(lat0, lon0)

        def on_map_click(coords):
            lat, lon = coords
            poser_marqueur(lat, lon)
            entry_lat.delete(0, "end")
            entry_lat.insert(0, f"{lat:.6f}")
            entry_lon.delete(0, "end")
            entry_lon.insert(0, f"{lon:.6f}")

        try:
            map_w.add_left_click_map_command(on_map_click)
        except Exception:
            pass

        # --- Saisie manuelle (toujours disponible, synchronisée avec la carte) ---
        form = ctk.CTkFrame(win, fg_color="transparent")
        form.pack(fill="x", padx=15, pady=8)
        form.grid_columnconfigure((1, 3), weight=1)

        ctk.CTkLabel(form, text="Latitude :", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, sticky="w", pady=6)
        entry_lat = ctk.CTkEntry(form, placeholder_text="ex : 49.16270")
        entry_lat.grid(row=0, column=1, sticky="ew", pady=6, padx=(8, 12))
        if lat0 is not None:
            entry_lat.insert(0, f"{lat0:.6f}")

        ctk.CTkLabel(form, text="Longitude :", font=ctk.CTkFont(weight="bold")).grid(row=0, column=2, sticky="w", pady=6)
        entry_lon = ctk.CTkEntry(form, placeholder_text="ex : 5.38540")
        entry_lon.grid(row=0, column=3, sticky="ew", pady=6, padx=(8, 0))
        if lon0 is not None:
            entry_lon.insert(0, f"{lon0:.6f}")

        def centrer_depuis_saisie():
            try:
                lat = float(entry_lat.get().strip().replace(",", "."))
                lon = float(entry_lon.get().strip().replace(",", "."))
            except ValueError:
                messagebox.showerror("Coordonnées invalides", "Entrez une latitude et une longitude numériques.")
                return
            if not (-90 <= lat <= 90 and -180 <= lon <= 180):
                messagebox.showerror("Hors plage", "Latitude doit être entre -90 et 90, longitude entre -180 et 180.")
                return
            poser_marqueur(lat, lon)
            map_w.set_position(lat, lon)
            map_w.set_zoom(15)

        ctk.CTkButton(
            form, text="🎯 Centrer la carte sur ces coordonnées", command=centrer_depuis_saisie,
            fg_color="#3a3a3a", hover_color="#4a4a4a", font=ctk.CTkFont(size=11)
        ).grid(row=1, column=0, columnspan=4, sticky="ew", pady=(4, 0))

        def appliquer():
            try:
                lat = float(entry_lat.get().strip().replace(",", "."))
                lon = float(entry_lon.get().strip().replace(",", "."))
            except ValueError:
                messagebox.showerror("Coordonnées invalides", "Entrez une latitude et une longitude numériques (ex. 49.1627).")
                return
            if not (-90 <= lat <= 90 and -180 <= lon <= 180):
                messagebox.showerror("Hors plage", "Latitude doit être entre -90 et 90, longitude entre -180 et 180.")
                return
            if not messagebox.askyesno(
                "Confirmer",
                f"Écrire les coordonnées\n  lat = {lat:.6f}\n  lon = {lon:.6f}\n"
                f"dans les métadonnées de {len(fichiers)} photo(s) ?"
            ):
                return
            win.destroy()
            self._write_gps_to_photos(fichiers, lat, lon)

        btn_row = ctk.CTkFrame(win, fg_color="transparent")
        btn_row.pack(fill="x", padx=15, pady=15)
        btn_row.grid_columnconfigure((0, 1), weight=1)
        ctk.CTkButton(
            btn_row, text="Annuler", command=win.destroy,
            fg_color="#3a3a3a", hover_color="#4a4a4a"
        ).grid(row=0, column=0, sticky="ew", padx=(0, 4))
        ctk.CTkButton(
            btn_row, text="💾 Écrire le GPS", command=appliquer,
            fg_color="#2ba14b", hover_color="#1f7d37", font=ctk.CTkFont(weight="bold")
        ).grid(row=0, column=1, sticky="ew", padx=(4, 0))

    def _write_gps_to_photos(self, filenames, lat, lon):
        """Écrit lat/lon dans les EXIF via ExifTool, met à jour photos_data et rafraîchit la carte."""
        exe_path = self.get_exiftool_path()
        if not exe_path:
            messagebox.showerror("ExifTool introuvable", "Posez exiftool.exe à côté du script ou installez-le dans le PATH.")
            return

        dossier = os.path.normpath(self.photo_folder_path)
        chemins = [os.path.join(dossier, fn) for fn in filenames if os.path.isfile(os.path.join(dossier, fn))]
        if not chemins:
            messagebox.showerror("Fichiers introuvables", "Aucune des photos sélectionnées n'existe sur le disque.")
            return

        self.log(f"📍 Écriture GPS manuelle sur {len(chemins)} photo(s) → {lat:.6f}, {lon:.6f} ...")
        self.btn_edit_gps.configure(state="disabled", text="⏳ GPS…")

        def worker():
            ok, err = 0, None
            try:
                startupinfo = self._exiftool_startupinfo()
                cmd = [
                    exe_path, "-P", "-overwrite_original",
                    f"-GPSLatitude={abs(lat)}", f"-GPSLatitudeRef={'N' if lat >= 0 else 'S'}",
                    f"-GPSLongitude={abs(lon)}", f"-GPSLongitudeRef={'E' if lon >= 0 else 'W'}",
                ] + chemins
                proc = subprocess.run(
                    cmd, startupinfo=startupinfo, capture_output=True, text=True,
                    encoding="utf-8", errors="ignore", timeout=120
                )
                if proc.returncode == 0:
                    ok = len(chemins)
                else:
                    err = (proc.stderr or proc.stdout or "code retour non nul").strip()[:300]
            except Exception as e:
                err = str(e)

            def finish():
                self.btn_edit_gps.configure(state="normal", text="📍 GPS")
                if ok:
                    for fn in filenames:
                        if fn in self.photos_data:
                            self.photos_data[fn]["lat"] = lat
                            self.photos_data[fn]["lon"] = lon
                        else:
                            self.photos_data[fn] = {
                                "path": os.path.join(dossier, fn),
                                "lat": lat, "lon": lon, "date": None
                            }
                    self.refresh_map_markers()
                    self.log(f"✅ GPS écrit sur {ok} photo(s).")
                    messagebox.showinfo("GPS mis à jour", f"{ok} photo(s) géolocalisée(s) avec succès.")
                else:
                    self.log(f"❌ Échec écriture GPS : {err}")
                    messagebox.showerror("Échec", f"Impossible d'écrire le GPS :\n{err}")

            self.after(0, finish)

        threading.Thread(target=worker, daemon=True).start()

    def _generate_map_image(self, dataset, max_size=(900, 600)):
        """Génère une image PNG (bytes) d'une carte simple avec les points GPS du dataset
        (utilisée dans les rapports PDF/Word et le brief IA). Retourne None si aucun point."""
        points = [(o.get("lat"), o.get("lon"), o.get("categorie") or "Non classé", o.get("espece") or "")
                  for o in dataset if o.get("lat") and o.get("lon")]
        if not points:
            return None
        try:
            from matplotlib.figure import Figure
            import matplotlib
            matplotlib.use("Agg")
        except ImportError:
            return None

        lats = [p[0] for p in points]
        lons = [p[1] for p in points]
        dlat = max(0.008, (max(lats) - min(lats)) * 0.25 or 0.01)
        dlon = max(0.008, (max(lons) - min(lons)) * 0.25 or 0.01)

        fig = Figure(figsize=(max_size[0] / 100, max_size[1] / 100), dpi=100, facecolor="#1a1a1a")
        ax = fig.add_subplot(111)
        ax.set_facecolor("#222222")
        ax.set_xlim(min(lons) - dlon, max(lons) + dlon)
        ax.set_ylim(min(lats) - dlat, max(lats) + dlat)
        ax.set_xlabel("Longitude", color="#aaaaaa", fontsize=8)
        ax.set_ylabel("Latitude", color="#aaaaaa", fontsize=8)
        ax.tick_params(colors="#aaaaaa", labelsize=7)
        for spine in ax.spines.values():
            spine.set_color("#555555")
        ax.set_title("Carte des observations", color="white", fontsize=11, pad=8)
        ax.set_aspect("equal", adjustable="box")
        ax.grid(True, color="#333333", linewidth=0.5)

        from matplotlib.lines import Line2D
        legend_handles = []
        for label, marker, color in [
            ("Animal / observation", "o", "#e67e22"),
            ("Empreinte / trace", "^", "#95a5a6"),
            ("Terrier / latrine", "s", "#95a5a6"),
            ("Coulée / passage", "D", "#95a5a6"),
        ]:
            legend_handles.append(Line2D([0], [0], marker=marker, color="w", markerfacecolor=color,
                                         markersize=8, linestyle="None", label=label))

        for lat, lon, cat, espece in points:
            couleur = CATEGORY_COLORS.get(cat, "#7f7f7f")
            forme = self._detect_marker_shape(espece)
            marker = {"triangle": "^", "square": "s", "diamond": "D"}.get(forme, "o")
            ax.plot(lon, lat, marker=marker, markersize=10, color=couleur,
                    markeredgecolor="white", markeredgewidth=0.9, linestyle="None", zorder=5)
            # Étiquette courte de l'espèce (évite le trop-plein si beaucoup de points)
            if espece and len(points) <= 40:
                label = espece if len(espece) <= 22 else espece[:20] + "…"
                ax.annotate(
                    label, (lon, lat), textcoords="offset points", xytext=(6, 5),
                    fontsize=6.5, color="#dddddd",
                    bbox=dict(boxstyle="round,pad=0.15", facecolor="#1a1a1a", edgecolor="none", alpha=0.75),
                    zorder=6,
                )

        ax.legend(handles=legend_handles, loc="upper right", fontsize=7,
                  facecolor="#2b2b2b", edgecolor="#555555", labelcolor="white")
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=120, facecolor=fig.get_facecolor())
        import matplotlib.pyplot as plt
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()


    def export_csv(self):
        """Exporte toutes les observations enregistrées vers un fichier CSV (Excel-compatible)."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Sélectionnez d'abord un dossier de photos.")
            return
        notes = self._load_notes_dict()
        if not notes:
            messagebox.showinfo("Aucune donnée", "Aucune observation enregistrée à exporter.")
            return

        chemin = self._ask_save_report(
            ".csv", [("Fichier CSV", "*.csv")], "observations_meuse.csv"
        )
        if not chemin:
            return

        try:
            with open(chemin, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f, delimiter=";")
                writer.writerow([
                    "Fichier", "Date", "Heure", "Catégorie", "Espèce", "Nombre", "Lieu",
                    "Latitude", "Longitude", "Température (°C)", "Humidité (%)", "Ciel",
                    "Pluie 3j précédents (mm)", "Phase lunaire", "Remarques"
                ])
                for img, data in notes.items():
                    info = self.photos_data.get(img, {})
                    meteo = data.get("meteo") or {}
                    writer.writerow([
                        img, meteo.get("date", ""), data.get("heure", ""), data.get("categorie", ""),
                        data.get("espece", ""), data.get("nombre", ""), data.get("lieu", ""),
                        info.get("lat", ""), info.get("lon", ""),
                        meteo.get("temperature", ""), meteo.get("humidite", ""), meteo.get("ciel", ""),
                        meteo.get("pluie_3j_precedents", ""), meteo.get("phase_lunaire", ""),
                        data.get("notes_libres", "")
                    ])
            self.log(f"📄 Export CSV réussi : {os.path.basename(chemin)}")
            messagebox.showinfo("Export réussi", f"{len(notes)} observation(s) exportée(s) en CSV.")
        except Exception as e:
            messagebox.showerror("Erreur d'export", str(e))

    def export_gpx(self):
        """Exporte les observations géolocalisées vers un fichier GPX (points d'intérêt)."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Sélectionnez d'abord un dossier de photos.")
            return
        notes = self._load_notes_dict()
        points = []
        for img, data in notes.items():
            info = self.photos_data.get(img, {})
            if info.get("lat") and info.get("lon"):
                points.append((img, info["lat"], info["lon"], data))

        if not points:
            messagebox.showinfo("Aucune donnée", "Aucune observation géolocalisée à exporter.")
            return

        chemin = self._ask_save_report(
            ".gpx", [("Fichier GPX", "*.gpx")], "observations_meuse.gpx"
        )
        if not chemin:
            return

        def echapper(txt):
            return str(txt).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        try:
            lignes = [
                '<?xml version="1.0" encoding="UTF-8"?>',
                '<gpx version="1.1" creator="GeoExif Meuse 55" xmlns="http://www.topografix.com/GPX/1/1">'
            ]
            for img, lat, lon, data in points:
                nom = data.get("espece") or img
                morceaux = []
                if data.get("categorie"): morceaux.append(data.get("categorie"))
                if data.get("nombre"): morceaux.append(f"Nombre: {data.get('nombre')}")
                if data.get("lieu"): morceaux.append(f"Lieu: {data.get('lieu')}")
                if data.get("heure"): morceaux.append(f"Heure: {data.get('heure')}")
                if data.get("notes_libres"): morceaux.append(data.get("notes_libres"))
                desc = " | ".join(morceaux)
                lignes.append(f'  <wpt lat="{lat}" lon="{lon}"><name>{echapper(nom)}</name><desc>{echapper(desc)}</desc></wpt>')
            lignes.append("</gpx>")

            with open(chemin, "w", encoding="utf-8") as f:
                f.write("\n".join(lignes))
            self.log(f"🌍 Export GPX réussi : {os.path.basename(chemin)}")
            messagebox.showinfo("Export réussi", f"{len(points)} point(s) exporté(s) en GPX.")
        except Exception as e:
            messagebox.showerror("Erreur d'export", str(e))

    def _gather_geolocated_observations(self):
        """Retourne la liste (cle, lat, lon, data) de toutes les observations géolocalisées du dossier actif."""
        if not self.photo_folder_path:
            return []
        notes = self._load_notes_dict()
        points = []
        for cle, data in notes.items():
            if not isinstance(data, dict):
                continue
            info = self.photos_data.get(cle, {})
            lat = info.get("lat") if info.get("lat") is not None else data.get("lat")
            lon = info.get("lon") if info.get("lon") is not None else data.get("lon")
            if lat and lon:
                points.append((cle, float(lat), float(lon), data))
        return points

    @staticmethod
    def _xml_escape(txt):
        return (
            str(txt).replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;")
        )

    def _build_kml_content(self, points, doc_name="Observations Meuse 55", track_points=None):
        """Construit un KML enrichi : styles uniques, icônes distinctes, dossiers par catégorie,
        trace GPX optionnelle, sans doublons de points (même GPS + même espèce)."""

        def kml_color(hex_rgb, alpha="ff"):
            h = (hex_rgb or "#7f7f7f").lstrip("#")
            if len(h) != 6:
                return alpha + "7f7f7f"
            r, g, b = h[0:2], h[2:4], h[4:6]
            return f"{alpha}{b}{g}{r}"  # KML = aabbggrr

        # Icônes Google Maps distinctes (catégorie × type d'observation)
        # circle = animal vu | triangle ≈ empreinte | square ≈ terrier | diamond ≈ coulée
        ICONS = {
            ("Mammifère", "circle"):   "http://maps.google.com/mapfiles/kml/paddle/orange-circle.png",
            ("Mammifère", "triangle"): "http://maps.google.com/mapfiles/kml/paddle/orange-blank.png",
            ("Mammifère", "square"):    "http://maps.google.com/mapfiles/kml/paddle/orange-square.png",
            ("Mammifère", "diamond"):  "http://maps.google.com/mapfiles/kml/paddle/orange-diamond.png",
            ("Oiseau", "circle"):      "http://maps.google.com/mapfiles/kml/paddle/blu-circle.png",
            ("Oiseau", "triangle"):    "http://maps.google.com/mapfiles/kml/paddle/blu-blank.png",
            ("Oiseau", "square"):      "http://maps.google.com/mapfiles/kml/paddle/blu-square.png",
            ("Oiseau", "diamond"):     "http://maps.google.com/mapfiles/kml/paddle/blu-diamond.png",
            ("Insecte", "circle"):     "http://maps.google.com/mapfiles/kml/paddle/ylw-circle.png",
            ("Insecte", "triangle"):   "http://maps.google.com/mapfiles/kml/paddle/ylw-blank.png",
            ("Insecte", "square"):     "http://maps.google.com/mapfiles/kml/paddle/ylw-square.png",
            ("Insecte", "diamond"):    "http://maps.google.com/mapfiles/kml/paddle/ylw-diamond.png",
            ("Autre", "circle"):       "http://maps.google.com/mapfiles/kml/paddle/wht-circle.png",
            ("Autre", "triangle"):     "http://maps.google.com/mapfiles/kml/paddle/wht-blank.png",
            ("Autre", "square"):       "http://maps.google.com/mapfiles/kml/paddle/wht-square.png",
            ("Autre", "diamond"):      "http://maps.google.com/mapfiles/kml/paddle/wht-diamond.png",
            ("Non classé", "circle"):  "http://maps.google.com/mapfiles/kml/paddle/grn-circle.png",
            ("Non classé", "triangle"):"http://maps.google.com/mapfiles/kml/paddle/grn-blank.png",
            ("Non classé", "square"):  "http://maps.google.com/mapfiles/kml/paddle/grn-square.png",
            ("Non classé", "diamond"): "http://maps.google.com/mapfiles/kml/paddle/grn-diamond.png",
        }
        FALLBACK_ICON = "http://maps.google.com/mapfiles/kml/paddle/red-circle.png"

        # Déduplication : même position (~11 m) + même espèce → on garde la première, on cumule le nombre
        uniques = {}
        for cle, lat, lon, data in points:
            espece = (data.get("espece") or cle or "?").strip()
            key = (round(lat, 4), round(lon, 4), espece.lower())
            if key in uniques:
                # Cumule un compteur indicatif dans les notes si plusieurs occurrences
                prev = uniques[key]
                prev_data = prev[3]
                prev_data.setdefault("_occurrences", 1)
                prev_data["_occurrences"] += 1
                continue
            # copie légère pour ne pas muter l'original
            data_copy = dict(data)
            data_copy["_occurrences"] = 1
            uniques[key] = (cle, lat, lon, data_copy)
        points_dedup = list(uniques.values())

        par_cat = {}
        styles_needed = set()  # (cat, shape) réellement utilisés
        for cle, lat, lon, data in points_dedup:
            cat = data.get("categorie") or "Non classé"
            shape = self._detect_marker_shape(data.get("espece") or "")
            styles_needed.add((cat, shape))
            par_cat.setdefault(cat, []).append((cle, lat, lon, data, shape))

        lignes = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<kml xmlns="http://www.opengis.net/kml/2.2">',
            "  <Document>",
            f"    <name>{self._xml_escape(doc_name)}</name>",
            f"    <description>Export GeoExif Meuse 55 — {self._xml_escape(APP_AUTHOR)} — "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')} — "
            f"{len(points_dedup)} point(s)"
            f"{' — trace GPX incluse' if track_points else ''}</description>",
        ]

        # Styles avancés : Style normal + highlight (StyleMap), uniquement combinaisons utilisées
        for cat, shape in sorted(styles_needed):
            base = f"s_{cat.replace(' ', '_')}_{shape}".replace("é", "e").replace("è", "e")
            icon = ICONS.get((cat, shape)) or ICONS.get((cat, "circle")) or FALLBACK_ICON
            couleur = CATEGORY_COLORS.get(cat, "#7f7f7f")
            lignes += [
                f'    <Style id="{base}_n">',
                "      <IconStyle>",
                f"        <color>{kml_color(couleur)}</color>",
                "        <scale>1.1</scale>",
                f"        <Icon><href>{icon}</href></Icon>",
                '        <hotSpot x="0.5" y="0" xunits="fraction" yunits="fraction"/>',
                "      </IconStyle>",
                "      <LabelStyle>",
                f"        <color>{kml_color('#ffffff', 'cc')}</color>",
                "        <scale>0.8</scale>",
                "      </LabelStyle>",
                "      <BalloonStyle>",
                "        <bgColor>ff1e1e1e</bgColor>",
                "        <textColor>ffffffff</textColor>",
                "        <text><![CDATA[<b><font size='+1'>$[name]</font></b><br/><br/>$[description]]]></text>",
                "      </BalloonStyle>",
                "    </Style>",
            ]
            lignes += [
                f'    <Style id="{base}_h">',
                "      <IconStyle>",
                f"        <color>{kml_color(couleur)}</color>",
                "        <scale>1.35</scale>",
                f"        <Icon><href>{icon}</href></Icon>",
                '        <hotSpot x="0.5" y="0" xunits="fraction" yunits="fraction"/>',
                "      </IconStyle>",
                "      <LabelStyle>",
                f"        <color>{kml_color('#ffffff')}</color>",
                "        <scale>1.0</scale>",
                "      </LabelStyle>",
                "      <BalloonStyle>",
                "        <bgColor>ff1e1e1e</bgColor>",
                "        <textColor>ffffffff</textColor>",
                "        <text><![CDATA[<b><font size='+1'>$[name]</font></b><br/><br/>$[description]]]></text>",
                "      </BalloonStyle>",
                "    </Style>",
            ]
            lignes += [
                f'    <StyleMap id="{base}">',
                f"      <Pair><key>normal</key><styleUrl>#{base}_n</styleUrl></Pair>",
                f"      <Pair><key>highlight</key><styleUrl>#{base}_h</styleUrl></Pair>",
                "    </StyleMap>",
            ]

        # Style de la trace GPX
        if track_points and len(track_points) >= 2:
            lignes += [
                '    <Style id="style_trace_gpx">',
                "      <LineStyle>",
                "        <color>ff2222ff</color>",
                "        <width>4</width>",
                "      </LineStyle>",
                "      <PolyStyle><color>00000000</color></PolyStyle>",
                "    </Style>",
            ]

        # Dossier Trace GPX (en premier)
        if track_points and len(track_points) >= 2:
            coords_txt = " ".join(f"{lon:.7f},{lat:.7f},0" for lat, lon in track_points)
            lignes += [
                "    <Folder>",
                "      <name>🗺️ Trace GPX</name>",
                "      <open>1</open>",
                "      <Placemark>",
                f"        <name>Trace ({len(track_points)} points)</name>",
                "        <styleUrl>#style_trace_gpx</styleUrl>",
                "        <LineString>",
                "          <tessellate>1</tessellate>",
                f"          <coordinates>{coords_txt}</coordinates>",
                "        </LineString>",
                "      </Placemark>",
                "    </Folder>",
            ]

        # Dossiers observations par catégorie
        for cat in sorted(par_cat.keys()):
            items = par_cat[cat]
            lignes.append("    <Folder>")
            lignes.append(f"      <name>{self._xml_escape(cat)} ({len(items)})</name>")
            lignes.append("      <open>1</open>")

            for cle, lat, lon, data, shape in sorted(items, key=lambda x: x[3].get("heure") or ""):
                nom = data.get("espece") or cle
                base = f"s_{cat.replace(' ', '_')}_{shape}".replace("é", "e").replace("è", "e")

                desc_parts = []
                if data.get("type_observation"):
                    desc_parts.append(f"<b>Type</b> : {self._xml_escape(data.get('type_observation'))}")
                if data.get("categorie"):
                    desc_parts.append(f"<b>Catégorie</b> : {self._xml_escape(data.get('categorie'))}")
                if data.get("nombre"):
                    n_txt = self._xml_escape(data.get("nombre"))
                    if data.get("_occurrences", 1) > 1:
                        n_txt += f" (×{data['_occurrences']} obs. regroupées)"
                    desc_parts.append(f"<b>Nombre</b> : {n_txt}")
                if data.get("heure"):
                    desc_parts.append(f"<b>Heure</b> : {self._xml_escape(data.get('heure'))}")
                if data.get("lieu"):
                    desc_parts.append(f"<b>Lieu</b> : {self._xml_escape(data.get('lieu'))}")
                if data.get("sans_photo") or str(cle).startswith("_manuel_"):
                    desc_parts.append("<i>Observation sans photo</i>")
                elif not str(cle).startswith("_manuel_"):
                    desc_parts.append(f"<b>Fichier</b> : {self._xml_escape(cle)}")
                meteo = data.get("meteo") or {}
                if meteo.get("temperature") is not None:
                    desc_parts.append(
                        f"<b>Météo</b> : {meteo.get('temperature')}°C, "
                        f"{meteo.get('humidite')}% hum., {self._xml_escape(meteo.get('ciel') or '')}"
                    )
                if meteo.get("phase_lunaire"):
                    desc_parts.append(f"<b>Lune</b> : {self._xml_escape(meteo.get('phase_lunaire'))}")
                if data.get("notes_libres"):
                    desc_parts.append(f"<b>Notes</b> : {self._xml_escape(data.get('notes_libres'))}")
                desc_parts.append(f"<b>GPS</b> : {lat:.6f}, {lon:.6f}")
                # Légende forme
                forme_label = {
                    "circle": "● Animal / observation",
                    "triangle": "▲ Empreinte / trace",
                    "square": "■ Terrier / latrine",
                    "diamond": "◆ Coulée / passage",
                }.get(shape, "●")
                desc_parts.append(f"<b>Marqueur</b> : {forme_label}")
                desc_html = "<br/>".join(desc_parts)

                lignes += [
                    "      <Placemark>",
                    f"        <name>{self._xml_escape(nom)}</name>",
                    f"        <description><![CDATA[{desc_html}]]></description>",
                    f"        <styleUrl>#{base}</styleUrl>",
                    "        <Point>",
                    f"          <coordinates>{lon:.7f},{lat:.7f},0</coordinates>",
                    "        </Point>",
                    "      </Placemark>",
                ]

            lignes.append("    </Folder>")

        # Légende visuelle : un placemark par style utilisé, avec l'icône réelle
        # Positionnés légèrement hors de la zone des données pour ne pas masquer les obs.
        if styles_needed and points_dedup:
            lats_d = [p[1] for p in points_dedup]
            lons_d = [p[2] for p in points_dedup]
            lat_leg = max(lats_d) + max(0.01, (max(lats_d) - min(lats_d)) * 0.08 or 0.01)
            lon0 = min(lons_d)
            span = max(0.01, (max(lons_d) - min(lons_d)) or 0.02)
            step = span / max(len(styles_needed), 1)

            shape_labels = {
                "circle": "● Animal vu",
                "triangle": "▲ Empreinte / trace",
                "square": "■ Terrier / latrine",
                "diamond": "◆ Coulée / passage",
            }
            lignes.append("    <Folder>")
            lignes.append("      <name>ℹ️ Légende des icônes</name>")
            lignes.append("      <open>0</open>")
            lignes.append(
                "      <description><![CDATA["
                "<b>Couleur</b> = catégorie taxonomique<br/>"
                "<b>Forme</b> = type d'observation (vu / empreinte / terrier / coulée)<br/>"
                "Survolez une icône pour l'agrandir. Les doublons GPS+espèce sont regroupés."
                "]]></description>"
            )
            for i, (cat, shape) in enumerate(sorted(styles_needed)):
                base = f"s_{cat.replace(' ', '_')}_{shape}".replace("é", "e").replace("è", "e")
                label = f"{cat} — {shape_labels.get(shape, shape)}"
                lon_i = lon0 + i * step
                lignes += [
                    "      <Placemark>",
                    f"        <name>{self._xml_escape(label)}</name>",
                    f"        <description><![CDATA[Légende : <b>{self._xml_escape(cat)}</b> — {shape_labels.get(shape, shape)}]]></description>",
                    f"        <styleUrl>#{base}</styleUrl>",
                    "        <Point>",
                    f"          <coordinates>{lon_i:.7f},{lat_leg:.7f},0</coordinates>",
                    "        </Point>",
                    "      </Placemark>",
                ]
            lignes.append("    </Folder>")

        lignes += ["  </Document>", "</kml>"]
        return "\n".join(lignes)

    def _get_track_points_for_kml(self):
        """Retourne la liste (lat, lon) de la trace GPX chargée, ou [] si absente."""
        if not getattr(self, "gpx_file_path", None) or not os.path.exists(self.gpx_file_path):
            return []
        try:
            return self._parser_gpx_points(self.gpx_file_path) or []
        except Exception:
            return []

    def _write_kml_file(self, filepath, points, doc_name=None, track_points=None):
        """Écrit un fichier .kml (observations + trace GPX optionnelle)."""
        track_points = track_points if track_points is not None else self._get_track_points_for_kml()
        if not points and not track_points:
            return False
        name = doc_name or f"Observations — {os.path.basename(self.photo_folder_path or 'Meuse')}"
        content = self._build_kml_content(points or [], name, track_points=track_points or None)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return True

    def _write_kmz_file(self, filepath, points, doc_name=None, track_points=None):
        """Écrit un fichier .kmz (zip contenant doc.kml)."""
        import zipfile
        track_points = track_points if track_points is not None else self._get_track_points_for_kml()
        if not points and not track_points:
            return False
        name = doc_name or f"Observations — {os.path.basename(self.photo_folder_path or 'Meuse')}"
        content = self._build_kml_content(points or [], name, track_points=track_points or None)
        with zipfile.ZipFile(filepath, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("doc.kml", content.encode("utf-8"))
        return True

    def export_kml(self):
        """Exporte les observations géolocalisées (+ trace GPX si chargée) vers KML/KMZ."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Sélectionnez d'abord un dossier de photos.")
            return
        points = self._gather_geolocated_observations()
        track = self._get_track_points_for_kml()
        if not points and not track:
            messagebox.showinfo(
                "Aucune donnée",
                "Aucune observation géolocalisée ni trace GPX à exporter.\n"
                "Chargez un GPX (étape 2) et/ou annotez des photos géotaguées."
            )
            return
        chemin = self._ask_save_report(
            ".kml",
            [("Google Earth KML", "*.kml"), ("Google Earth KMZ", "*.kmz")],
            "observations_meuse.kml",
        )
        if not chemin:
            return
        try:
            if chemin.lower().endswith(".kmz"):
                self._write_kmz_file(chemin, points, track_points=track)
            else:
                self._write_kml_file(chemin, points, track_points=track)
            msg_parts = []
            if points:
                msg_parts.append(f"{len(points)} observation(s)")
            if track:
                msg_parts.append(f"trace GPX ({len(track)} pts)")
            self.log(f"🌍 Export KML/KMZ réussi : {os.path.basename(chemin)} — {', '.join(msg_parts)}")
            messagebox.showinfo(
                "Export réussi",
                f"Exporté : {', '.join(msg_parts)}.\n\n"
                "Ouvrez le fichier avec Google Earth.\n"
                "Les points en double (même lieu + espèce) ont été regroupés."
            )
        except Exception as e:
            messagebox.showerror("Erreur d'export", str(e))

    # --- Configuration application (dossier rapports, etc.) ---

    def _user_data_dir(self):
        """Dossier stable entre versions (independant de l'emplacement du .py / .exe)."""
        if getattr(self, "_user_data_dir_cache", None):
            return self._user_data_dir_cache
        candidates = []
        if os.name == "nt":
            base = os.environ.get("LOCALAPPDATA") or os.environ.get("APPDATA")
            if base:
                candidates.append(os.path.join(base, "GeoExif"))
        candidates.append(os.path.join(os.path.expanduser("~"), ".geoexif"))
        candidates.append(os.path.join(os.path.dirname(os.path.abspath(sys.argv[0])), "geoexif_data"))
        chosen = candidates[-1]
        for c in candidates:
            try:
                os.makedirs(c, exist_ok=True)
                probe = os.path.join(c, ".write_test")
                with open(probe, "w", encoding="utf-8") as f:
                    f.write("ok")
                os.remove(probe)
                chosen = c
                break
            except Exception:
                continue
        self._user_data_dir_cache = chosen
        return chosen

    def _migrate_legacy_config_files(self):
        """Recupere sorties_connues / config depuis l'ancien emplacement (a cote du script)."""
        data_dir = self._user_data_dir()
        search_dirs = [
            os.path.dirname(os.path.abspath(sys.argv[0])),
            os.getcwd(),
        ]
        names = [CONFIG_FILE, "sorties_connues.json", "ia_config.json", PROJECT_FILE]
        import shutil
        for name in names:
            dst = os.path.join(data_dir, name)
            if os.path.isfile(dst):
                continue
            for d in search_dirs:
                src = os.path.join(d, name)
                if os.path.isfile(src):
                    try:
                        shutil.copy2(src, dst)
                        try:
                            self.log("Migration : %s" % name)
                        except Exception:
                            pass
                    except Exception:
                        pass
                    break

    def _app_config_path(self):
        try:
            self._migrate_legacy_config_files()
        except Exception:
            pass
        return os.path.join(self._user_data_dir(), CONFIG_FILE)


    def _load_app_config(self):
        path = self._app_config_path()
        if not os.path.exists(path):
            return {}
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}

    def _save_app_config(self, config=None):
        cfg = config if config is not None else getattr(self, "app_config", {})
        self.app_config = cfg
        try:
            with open(self._app_config_path(), "w", encoding="utf-8") as f:
                json.dump(cfg, f, ensure_ascii=False, indent=2)
        except Exception as e:
            self.log(f"⚠️ Impossible de sauvegarder la config : {e}")


    def _startup_config_check(self):
        """Premier demarrage : assistant config. Sinon backups periodiques / chemins par defaut."""
        try:
            self._migrate_legacy_config_files()
        except Exception:
            pass
        cfg = self._load_app_config() or {}
        self.app_config = cfg
        try:
            self.log("Dossier donnees (stable) : %s" % self._user_data_dir())
            n = len(self._load_known_folders() or [])
            if n:
                self.log("%d sortie(s) connue(s) rechargee(s) depuis sorties_connues.json" % n)
        except Exception:
            pass
        # Appliquer chemins par defaut si dossiers vides
        try:
            if not self.photo_folder_path and cfg.get("default_photo_folder"):
                dp = cfg["default_photo_folder"]
                if os.path.isdir(dp):
                    # ne charge pas auto le carnet (lourd) — juste memorise
                    pass
            if not self.gpx_file_path and cfg.get("default_gpx_folder"):
                pass
        except Exception:
            pass
        # Backup mensuel global si active
        try:
            self._maybe_monthly_backup()
        except Exception:
            pass
        if not cfg.get("setup_done"):
            try:
                self.open_setup_wizard(first_run=True)
            except Exception as e:
                self.log("Assistant config : %s" % e)

    def scan_parent_for_sorties(self):
        """Parcourt un dossier parent et enregistre tous les sous-dossiers contenant observations.json."""
        parent = filedialog.askdirectory(title="Dossier parent contenant vos sorties")
        if not parent:
            return
        found = []
        try:
            for root, dirs, files in os.walk(parent):
                if NOTES_FILE in files:
                    found.append(os.path.normpath(root))
                # limiter profondeur
                depth = root[len(parent):].count(os.sep)
                if depth >= 3:
                    dirs.clear()
        except Exception as e:
            messagebox.showerror("Scan", str(e))
            return
        if not found:
            messagebox.showinfo("Scan", "Aucun observations.json trouve sous :\n%s" % parent)
            return
        for fp in found:
            self._register_known_folder(fp)
        # Copie immediate vers archive PC (indispensable avant debrancher le disque)
        n_mir = 0
        try:
            n_mir = self.mirror_all_known_sorties_to_archive(silent=True)
        except Exception:
            pass
        messagebox.showinfo(
            "Scan",
            "%d sortie(s) enregistree(s).\n"
            "%d carnet(s) copies dans l'archive locale PC.\n\n"
            "Vous pouvez debrancher le disque : tableau de bord / especes / carte\n"
            "utiliseront l'archive pour les observations.\n\n"
            "Archive : %s"
            % (len(found), n_mir, self._local_archive_root()),
        )
        self.log("Scan sorties : %d dossiers, %d miroirs archive" % (len(found), n_mir))

    def open_backup_settings(self):
        """Panneau dedie : sauvegardes automatiques et emplacement."""
        cfg = self._load_app_config() or {}
        win = ctk.CTkToplevel(self)
        win.title("Sauvegardes automatiques")
        win.geometry("560x520")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Sauvegardes automatiques",
            font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text="GeoExif copie observations.json, dispositifs.json et la config "
                 "pour limiter la perte de donnees. Les photos elles-memes ne sont pas dupliquees.",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            wraplength=520, justify="left",
        ).pack(anchor="w", padx=16, pady=(0, 10))

        card = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=10)
        card.pack(fill="both", expand=True, padx=16, pady=6)

        ctk.CTkLabel(
            card, text="Mode de sauvegarde",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=12, pady=(12, 4))

        mode_var = tk.StringVar(value=cfg.get("backup_mode") or "each_save")
        modes = [
            ("each_save", "A chaque enregistrement du carnet (recommande)"),
            ("monthly", "Une fois par mois (plus leger)"),
            ("off", "Desactive"),
        ]
        for val, label in modes:
            ctk.CTkRadioButton(
                card, text=label, variable=mode_var, value=val,
                text_color=UI.get("text"),
            ).pack(anchor="w", padx=20, pady=3)

        ctk.CTkLabel(
            card, text="Dossier de destination des backups",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=12, pady=(16, 4))

        bak_var = tk.StringVar(value=cfg.get("backup_folder") or "")
        row = ctk.CTkFrame(card, fg_color="transparent")
        row.pack(fill="x", padx=12, pady=4)
        ctk.CTkEntry(
            row, textvariable=bak_var,
            placeholder_text="Vide = geoexif_backups a cote de l'application",
        ).pack(side="left", fill="x", expand=True)
        ctk.CTkButton(
            row, text="…", width=36,
            command=lambda: bak_var.set(
                filedialog.askdirectory(title="Dossier des sauvegardes") or bak_var.get()
            ),
        ).pack(side="left", padx=4)

        root_now = self._backup_root_dir()
        ctk.CTkLabel(
            card, text="Emplacement actuel :\n%s" % root_now,
            font=ctk.CTkFont(size=11), text_color=UI.get("text_accent", UI.get("accent")),
            wraplength=500, justify="left",
        ).pack(anchor="w", padx=12, pady=(8, 4))
        ctk.CTkLabel(
            card,
            text="Sous-dossiers : auto\\ (chaque save) · monthly\\AAAA-MM\\ (mensuel)",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim"),
        ).pack(anchor="w", padx=12, pady=(0, 8))

        last = cfg.get("last_monthly_backup") or "jamais"
        ctk.CTkLabel(
            card, text="Dernier backup mensuel enregistre : %s" % last,
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
        ).pack(anchor="w", padx=12, pady=(0, 12))

        def save_cfg():
            new_cfg = dict(cfg)
            new_cfg["backup_mode"] = mode_var.get()
            new_cfg["backup_folder"] = bak_var.get().strip()
            self._save_app_config(new_cfg)
            self.app_config = new_cfg
            messagebox.showinfo(
                "Sauvegardes",
                "Reglages enregistres.\n\nMode : %s\nDossier : %s" % (
                    new_cfg["backup_mode"], self._backup_root_dir(),
                ),
            )
            win.destroy()

        def run_now():
            try:
                # force copies
                prev = (self.app_config or {}).get("backup_mode")
                if not self.app_config:
                    self.app_config = self._load_app_config() or {}
                self.app_config["backup_mode"] = "each_save"
                self._run_intervention_backup(self.photo_folder_path)
                # monthly snapshot too
                self.app_config["last_monthly_backup"] = None
                self._maybe_monthly_backup()
                if prev:
                    self.app_config["backup_mode"] = prev
                    self._save_app_config(self.app_config)
                messagebox.showinfo("Sauvegardes", "Backup manuel effectue dans :\n%s" % self._backup_root_dir())
            except Exception as e:
                messagebox.showerror("Sauvegardes", str(e))

        def restore_notes():
            if not self.photo_folder_path:
                messagebox.showwarning("Restaurer", "Ouvrez d abord le dossier de sortie cible.")
                return
            src = filedialog.askopenfilename(
                title="Choisir un backup observations.json",
                initialdir=self._backup_root_dir(),
                filetypes=[("JSON", "*.json"), ("Tous", "*.*")],
            )
            if not src:
                return
            dest = os.path.join(self.photo_folder_path, NOTES_FILE)
            if os.path.isfile(dest) and not messagebox.askyesno(
                "Restaurer", "Remplacer le carnet actuel par ce backup ?"
            ):
                return
            try:
                import shutil
                if os.path.isfile(dest):
                    shutil.copy2(dest, dest + ".avant_restore")
                shutil.copy2(src, dest)
                messagebox.showinfo("Restaurer", "Carnet restaure. Rouvrez le dossier si besoin.")
                self.log("Restore depuis %s" % src)
            except Exception as e:
                messagebox.showerror("Restaurer", str(e))

        def open_folder():
            d = self._backup_root_dir()
            try:
                os.makedirs(d, exist_ok=True)
                os.startfile(d)
            except Exception as e:
                messagebox.showinfo("Sauvegardes", "%s\n%s" % (d, e))

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=16, pady=12)
        ctk.CTkButton(bf, text="Sauvegarder maintenant", command=run_now).pack(side="left", padx=3)
        ctk.CTkButton(bf, text="Ouvrir le dossier", command=open_folder).pack(side="left", padx=3)
        ctk.CTkButton(bf, text="Restaurer un carnet…", command=restore_notes).pack(side="left", padx=3)
        ctk.CTkButton(
            bf, text="Enregistrer", width=120,
            fg_color=UI.get("success", "#2f9e5f"), command=save_cfg,
        ).pack(side="right", padx=3)
        ctk.CTkButton(
            bf, text="Fermer", width=90, fg_color=UI.get("card_alt"), command=win.destroy,
        ).pack(side="right", padx=3)
        win.after(200, lambda: self._style_dialog(win) if hasattr(self, "_style_dialog") else None)

    def open_setup_wizard(self, first_run=False):
        """Assistant de configuration (1er demarrage ou menu Configuration)."""
        cfg = self._load_app_config() or {}
        ia = {}
        try:
            ia = self._load_ia_config() or {}
        except Exception:
            ia = {}

        win = ctk.CTkToplevel(self)
        win.title("Configuration GeoExif" + (" — premier démarrage" if first_run else ""))
        win.geometry("640x720")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        try:
            win.transient(self)
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win,
            text="Configuration" + (" — bienvenue !" if first_run else ""),
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text="Reglez les chemins par defaut, le theme, le cloud, les sauvegardes et les cles IA.\n"
                 "Vous pourrez rouvrir cet ecran a tout moment : menu Plus → Configuration…",
            font=ctk.CTkFont(size=11),
            text_color=UI.get("text_dim"),
            wraplength=600, justify="left",
        ).pack(anchor="w", padx=16, pady=(0, 8))

        # Emplacement fichier config
        cfg_path = self._app_config_path()
        data_dir = self._user_data_dir()
        path_fr = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=8)
        path_fr.pack(fill="x", padx=16, pady=6)
        ctk.CTkLabel(
            path_fr, text="Dossier donnees (conserve entre versions) :",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=10, pady=(8, 2))
        ctk.CTkLabel(
            path_fr, text=data_dir,
            font=ctk.CTkFont(size=11), text_color=UI.get("text_accent", UI.get("accent")),
            wraplength=580, justify="left",
        ).pack(anchor="w", padx=10, pady=(0, 2))
        ctk.CTkLabel(
            path_fr, text="Config : %s" % cfg_path,
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim"),
            wraplength=580, justify="left",
        ).pack(anchor="w", padx=10, pady=(0, 4))
        ctk.CTkLabel(
            path_fr,
            text="Cles IA : %s  ·  A cote de l'executable / du script Python." % (
                os.path.join(os.path.dirname(cfg_path), "ia_config.json")
            ),
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim"),
            wraplength=580, justify="left",
        ).pack(anchor="w", padx=10, pady=(0, 8))

        scroll = ctk.CTkScrollableFrame(win, fg_color=UI.get("card"))
        scroll.pack(fill="both", expand=True, padx=16, pady=8)
        scroll.grid_columnconfigure(1, weight=1)

        def add_label(r, text):
            ctk.CTkLabel(scroll, text=text, text_color=UI.get("text"), anchor="w").grid(
                row=r, column=0, sticky="w", padx=8, pady=6
            )

        # Theme
        add_label(0, "Theme")
        theme_var = tk.StringVar(value=cfg.get("theme") or getattr(self, "_theme_name", "sombre") or "sombre")
        theme_menu = ctk.CTkOptionMenu(
            scroll, values=["sombre", "clair", "papier"], variable=theme_var, width=200,
        )
        theme_menu.grid(row=0, column=1, sticky="w", padx=8, pady=6)

        # Dossier photos par defaut
        add_label(1, "Dossier photos (defaut)")
        photo_var = tk.StringVar(value=cfg.get("default_photo_folder") or "")
        row1 = ctk.CTkFrame(scroll, fg_color="transparent")
        row1.grid(row=1, column=1, sticky="ew", padx=8, pady=6)
        ctk.CTkEntry(row1, textvariable=photo_var).pack(side="left", fill="x", expand=True)
        ctk.CTkButton(
            row1, text="…", width=36,
            command=lambda: photo_var.set(filedialog.askdirectory(title="Dossier photos par defaut") or photo_var.get()),
        ).pack(side="left", padx=4)

        # Dossier / dernier GPX
        add_label(2, "Dossier traces GPX")
        gpx_var = tk.StringVar(value=cfg.get("default_gpx_folder") or "")
        row2 = ctk.CTkFrame(scroll, fg_color="transparent")
        row2.grid(row=2, column=1, sticky="ew", padx=8, pady=6)
        ctk.CTkEntry(row2, textvariable=gpx_var).pack(side="left", fill="x", expand=True)
        ctk.CTkButton(
            row2, text="…", width=36,
            command=lambda: gpx_var.set(filedialog.askdirectory(title="Dossier GPX par defaut") or gpx_var.get()),
        ).pack(side="left", padx=4)

        # Cloud
        add_label(3, "Dossier cloud (Drive/Dropbox…)")
        cloud_var = tk.StringVar(value=cfg.get("cloud_folder") or "")
        row3 = ctk.CTkFrame(scroll, fg_color="transparent")
        row3.grid(row=3, column=1, sticky="ew", padx=8, pady=6)
        ctk.CTkEntry(row3, textvariable=cloud_var).pack(side="left", fill="x", expand=True)
        ctk.CTkButton(
            row3, text="…", width=36,
            command=lambda: cloud_var.set(
                filedialog.askdirectory(title="Dossier cloud local synchronise") or cloud_var.get()
            ),
        ).pack(side="left", padx=4)

        # Dossier rapports
        add_label(4, "Dossier rapports / exports")
        rap_var = tk.StringVar(value=cfg.get("reports_folder") or "")
        row4 = ctk.CTkFrame(scroll, fg_color="transparent")
        row4.grid(row=4, column=1, sticky="ew", padx=8, pady=6)
        ctk.CTkEntry(row4, textvariable=rap_var).pack(side="left", fill="x", expand=True)
        ctk.CTkButton(
            row4, text="…", width=36,
            command=lambda: rap_var.set(filedialog.askdirectory(title="Dossier rapports") or rap_var.get()),
        ).pack(side="left", padx=4)

        # Backups
        add_label(5, "Sauvegarde auto")
        backup_var = tk.StringVar(value=cfg.get("backup_mode") or "each_save")
        ctk.CTkOptionMenu(
            scroll,
            values=["each_save", "monthly", "off"],
            variable=backup_var,
            width=220,
        ).grid(row=5, column=1, sticky="w", padx=8, pady=6)
        ctk.CTkLabel(
            scroll,
            text="each_save = a chaque save carnet · monthly = 1x/mois · off = off  (detail : menu Sauvegardes automatiques)",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim"),
            wraplength=400, justify="left",
        ).grid(row=6, column=1, sticky="w", padx=8)

        add_label(7, "Dossier backup global")
        bak_var = tk.StringVar(value=cfg.get("backup_folder") or "")
        row7 = ctk.CTkFrame(scroll, fg_color="transparent")
        row7.grid(row=7, column=1, sticky="ew", padx=8, pady=6)
        ctk.CTkEntry(row7, textvariable=bak_var, placeholder_text="(optionnel — sinon sous le dossier app)").pack(
            side="left", fill="x", expand=True
        )
        ctk.CTkButton(
            row7, text="…", width=36,
            command=lambda: bak_var.set(filedialog.askdirectory(title="Dossier backup") or bak_var.get()),
        ).pack(side="left", padx=4)

        # IA keys
        ctk.CTkLabel(
            scroll, text="Cles IA (optionnel)",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).grid(row=8, column=0, columnspan=2, sticky="w", padx=8, pady=(14, 4))

        add_label(9, "Gemini API key")
        gem_var = tk.StringVar(value=ia.get("gemini_api_key") or ia.get("google_api_key") or "")
        ctk.CTkEntry(scroll, textvariable=gem_var, show="*").grid(row=9, column=1, sticky="ew", padx=8, pady=6)

        add_label(10, "xAI / Grok API key")
        xai_var = tk.StringVar(value=ia.get("xai_api_key") or ia.get("grok_api_key") or "")
        ctk.CTkEntry(scroll, textvariable=xai_var, show="*").grid(row=10, column=1, sticky="ew", padx=8, pady=6)

        def open_cfg_folder():
            d = data_dir
            try:
                os.makedirs(d, exist_ok=True)
                os.startfile(d)
            except Exception as e:
                messagebox.showinfo("Config", "Dossier :\n%s\n%s" % (d, e))

        def save_and_close(mark_done=True):
            new_cfg = dict(cfg)
            new_cfg["theme"] = theme_var.get().strip() or "sombre"
            new_cfg["default_photo_folder"] = photo_var.get().strip()
            new_cfg["default_gpx_folder"] = gpx_var.get().strip()
            new_cfg["cloud_folder"] = cloud_var.get().strip()
            new_cfg["reports_folder"] = rap_var.get().strip()
            new_cfg["backup_mode"] = backup_var.get().strip() or "each_save"
            new_cfg["backup_folder"] = bak_var.get().strip()
            if mark_done:
                new_cfg["setup_done"] = True
            self._save_app_config(new_cfg)
            self.app_config = new_cfg
            # IA
            new_ia = dict(ia)
            if gem_var.get().strip():
                new_ia["gemini_api_key"] = gem_var.get().strip()
            if xai_var.get().strip():
                new_ia["xai_api_key"] = xai_var.get().strip()
            try:
                self._save_ia_config(new_ia)
            except Exception:
                pass
            # Theme
            try:
                self.apply_theme(new_cfg["theme"], save=True, silent=True)
            except Exception:
                pass
            self.log("Configuration enregistree → %s" % cfg_path)
            messagebox.showinfo(
                "Configuration",
                "Parametres enregistres.\n\nFichier :\n%s" % cfg_path,
            )
            win.destroy()

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=16, pady=12)
        ctk.CTkButton(bf, text="Ouvrir le dossier config", width=160, command=open_cfg_folder).pack(side="left", padx=3)
        if first_run:
            ctk.CTkButton(
                bf, text="Passer pour l'instant", width=140,
                fg_color=UI.get("card_alt"),
                command=lambda: (self._save_app_config({**cfg, "setup_done": True}), win.destroy()),
            ).pack(side="left", padx=3)
        ctk.CTkButton(
            bf, text="Enregistrer", width=140,
            fg_color=UI.get("success", "#2f9e5f"),
            command=lambda: save_and_close(True),
        ).pack(side="right", padx=3)

        win.after(300, lambda: self._style_dialog(win) if hasattr(self, "_style_dialog") else None)

    def _backup_root_dir(self):
        cfg = getattr(self, "app_config", None) or self._load_app_config() or {}
        custom = (cfg.get("backup_folder") or "").strip()
        if custom:
            try:
                os.makedirs(custom, exist_ok=True)
                return custom
            except Exception:
                pass
        root = os.path.join(os.path.dirname(self._app_config_path()), "geoexif_backups")
        try:
            os.makedirs(root, exist_ok=True)
        except Exception:
            pass
        return root

    def _copy_file_backup(self, src, dest_dir, label):
        if not src or not os.path.isfile(src):
            return
        try:
            os.makedirs(dest_dir, exist_ok=True)
            base = os.path.basename(src)
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            dest = os.path.join(dest_dir, "%s_%s_%s" % (label, stamp, base))
            import shutil
            shutil.copy2(src, dest)
        except Exception as e:
            self.log("Backup : %s" % e)

    def _run_intervention_backup(self, folder_path=None):
        """Backup declenche a chaque intervention (si mode each_save)."""
        cfg = getattr(self, "app_config", None) or self._load_app_config() or {}
        if (cfg.get("backup_mode") or "each_save") != "each_save":
            return
        folder_path = folder_path or self.photo_folder_path
        if not folder_path:
            return
        dest = os.path.join(self._backup_root_dir(), "auto", os.path.basename(folder_path) or "sortie")
        notes = os.path.join(folder_path, NOTES_FILE)
        devices = os.path.join(folder_path, DEVICES_FILE)
        self._copy_file_backup(notes, dest, "notes")
        self._copy_file_backup(devices, dest, "devices")
        # config globale
        self._copy_file_backup(self._app_config_path(), self._backup_root_dir(), "config")

    def _maybe_monthly_backup(self):
        """Une fois par mois : snapshot config + projet + notes dossier courant."""
        cfg = getattr(self, "app_config", None) or self._load_app_config() or {}
        mode = cfg.get("backup_mode") or "each_save"
        if mode == "off":
            return
        # monthly flag even if each_save (extra safety once/month)
        month_key = datetime.now().strftime("%Y-%m")
        if cfg.get("last_monthly_backup") == month_key and mode != "monthly":
            return
        if mode == "monthly" and cfg.get("last_monthly_backup") == month_key:
            return
        dest = os.path.join(self._backup_root_dir(), "monthly", month_key)
        try:
            os.makedirs(dest, exist_ok=True)
        except Exception:
            return
        self._copy_file_backup(self._app_config_path(), dest, "config")
        try:
            proj = os.path.join(os.path.dirname(self._app_config_path()), PROJECT_FILE)
            self._copy_file_backup(proj, dest, "projet")
        except Exception:
            pass
        if self.photo_folder_path:
            self._copy_file_backup(os.path.join(self.photo_folder_path, NOTES_FILE), dest, "notes")
            self._copy_file_backup(os.path.join(self.photo_folder_path, DEVICES_FILE), dest, "devices")
        cfg["last_monthly_backup"] = month_key
        self._save_app_config(cfg)
        self.log("Backup mensuel : %s" % dest)

    # --- Fiche espèce : résumé Wikipédia (recherche + cache local disque) ---

    def _wiki_cache_path(self):
        return os.path.join(os.path.dirname(os.path.abspath(sys.argv[0])), WIKI_CACHE_FILE)

    def _load_wiki_cache(self):
        path = self._wiki_cache_path()
        if not os.path.exists(path):
            return {}
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}

    def _save_wiki_cache(self, cache):
        try:
            with open(self._wiki_cache_path(), "w", encoding="utf-8") as f:
                json.dump(cache, f, ensure_ascii=False, indent=2)
        except Exception:
            pass  # cache = confort, pas de donnée critique : un échec d'écriture n'est pas bloquant

    def _fetch_wikipedia_info(self, terme):
        """Cherche 'terme' sur Wikipédia FR (recherche puis résumé), avec cache local disque.
        Renvoie un dict {title, extract, thumbnail, url}, ou {"not_found": True}, ou None en cas
        d'erreur réseau (dans ce dernier cas, rien n'est mis en cache pour permettre un nouvel essai)."""
        cle = (terme or "").strip().lower()
        if not cle:
            return None

        cache = self._load_wiki_cache()
        if cle in cache:
            return cache[cle]

        resultat = None
        try:
            r = requests.get(
                "https://fr.wikipedia.org/w/api.php",
                params={"action": "query", "list": "search", "srsearch": terme, "format": "json", "srlimit": 1},
                headers={"User-Agent": "GeoExifSync-NaturalistNotebook/1.0"},
                timeout=6,
            )
            r.raise_for_status()
            hits = r.json().get("query", {}).get("search", [])
            if not hits:
                resultat = {"not_found": True}
            else:
                titre = hits[0]["title"]
                r2 = requests.get(
                    f"https://fr.wikipedia.org/api/rest_v1/page/summary/{urllib.parse.quote(titre)}",
                    headers={"User-Agent": "GeoExifSync-NaturalistNotebook/1.0"},
                    timeout=6,
                )
                if r2.status_code == 200:
                    d = r2.json()
                    if d.get("type") == "disambiguation":
                        resultat = {"not_found": True}
                    else:
                        resultat = {
                            "title": d.get("title") or titre,
                            "extract": d.get("extract") or "",
                            "thumbnail": (d.get("thumbnail") or {}).get("source"),
                            "url": ((d.get("content_urls") or {}).get("desktop") or {}).get("page")
                                   or f"https://fr.wikipedia.org/wiki/{urllib.parse.quote(titre)}",
                        }
                else:
                    resultat = {"not_found": True}
        except Exception:
            return None

        cache[cle] = resultat
        self._save_wiki_cache(cache)
        return resultat

    def _reports_dir(self):
        """Dossier préféré pour les exports (KML, PDF, GeoJSON…). Fallback = dossier photos."""
        d = (getattr(self, "app_config", {}) or {}).get("reports_dir") or ""
        if d and os.path.isdir(d):
            return d
        return self.photo_folder_path or None

    def _ask_save_report(self, defaultextension, filetypes, initialfile):
        """filedialog.asksaveasfilename prérempli avec le dossier de rapports configuré."""
        kwargs = {
            "defaultextension": defaultextension,
            "filetypes": filetypes,
            "initialfile": initialfile,
        }
        init_dir = self._reports_dir()
        if init_dir:
            kwargs["initialdir"] = init_dir
        return filedialog.asksaveasfilename(**kwargs)

    def open_reports_dir_settings(self):
        """Choisir le dossier par défaut où sauvegarder les rapports (KML, PDF, GeoJSON…)."""
        win = ctk.CTkToplevel(self)
        win.title("📁 Dossier des rapports")
        win.geometry("520x220")
        self._prepare_tool_window(win)

        actuel = (self.app_config or {}).get("reports_dir") or "(non défini — dossier photos utilisé)"
        ctk.CTkLabel(
            win,
            text="Dossier par défaut pour les exports (KML, KMZ, GeoJSON, PDF, Word, CSV, GPX) :",
            font=ctk.CTkFont(size=12), wraplength=480, justify="left"
        ).pack(anchor="w", padx=15, pady=(15, 6))

        lbl = ctk.CTkLabel(win, text=actuel, font=ctk.CTkFont(size=11), text_color="#8ab4ff", wraplength=480, justify="left")
        lbl.pack(anchor="w", padx=15, pady=(0, 12))

        def choisir():
            d = filedialog.askdirectory(initialdir=self._reports_dir() or None)
            if not d:
                return
            cfg = dict(self.app_config or {})
            cfg["reports_dir"] = d
            self._save_app_config(cfg)
            lbl.configure(text=d)
            self.log(f"📁 Dossier des rapports : {d}")

        def reinit():
            cfg = dict(self.app_config or {})
            cfg.pop("reports_dir", None)
            self._save_app_config(cfg)
            lbl.configure(text="(non défini — dossier photos utilisé)")
            self.log("📁 Dossier des rapports réinitialisé (dossier photos).")

        row = ctk.CTkFrame(win, fg_color="transparent")
        row.pack(fill="x", padx=15, pady=10)
        ctk.CTkButton(row, text="📂 Choisir un dossier", command=choisir, fg_color="#1f7d5a", hover_color="#175f45").pack(side="left", padx=(0, 8))
        ctk.CTkButton(row, text="Réinitialiser", command=reinit, fg_color="#3a3a3a", hover_color="#4a4a4a").pack(side="left")
        ctk.CTkButton(row, text="Fermer", command=win.destroy).pack(side="right")

    def export_geojson(self):
        """Exporte les observations (+ trace GPX) en GeoJSON (QGIS, Leaflet, Mapbox…)."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Sélectionnez d'abord un dossier de photos.")
            return
        points = self._gather_geolocated_observations()
        track = self._get_track_points_for_kml()
        if not points and not track:
            messagebox.showinfo("Aucune donnée", "Aucune observation géolocalisée ni trace GPX à exporter.")
            return

        chemin = self._ask_save_report(
            ".geojson",
            [("GeoJSON", "*.geojson"), ("JSON", "*.json")],
            "observations_meuse.geojson",
        )
        if not chemin:
            return
        try:
            features = []
            # Trace GPX
            if track and len(track) >= 2:
                features.append({
                    "type": "Feature",
                    "geometry": {
                        "type": "LineString",
                        "coordinates": [[lon, lat] for lat, lon in track],
                    },
                    "properties": {
                        "name": "Trace GPX",
                        "type": "track",
                        "points": len(track),
                        "source": os.path.basename(self.gpx_file_path or ""),
                    },
                })
            # Observations (dédupliquées comme le KML)
            seen = set()
            for cle, lat, lon, data in points:
                espece = (data.get("espece") or cle or "?").strip()
                key = (round(lat, 4), round(lon, 4), espece.lower())
                if key in seen:
                    continue
                seen.add(key)
                shape = self._detect_marker_shape(espece)
                props = {
                    "name": espece,
                    "categorie": data.get("categorie"),
                    "espece": espece,
                    "nombre": data.get("nombre"),
                    "heure": data.get("heure"),
                    "lieu": data.get("lieu"),
                    "type_observation": data.get("type_observation"),
                    "sans_photo": bool(data.get("sans_photo") or str(cle).startswith("_manuel_")),
                    "notes": data.get("notes_libres"),
                    "fichier": None if str(cle).startswith("_manuel_") else cle,
                    "marqueur": shape,
                    "marker-symbol": shape,
                    "marker-color": CATEGORY_COLORS.get(data.get("categorie") or "Non classé", "#7f7f7f"),
                    "marker-size": 4,
                    "legende_icone": {
                        "circle": "Animal / observation",
                        "triangle": "Empreinte / trace",
                        "square": "Terrier / latrine",
                        "diamond": "Coulée / passage",
                    }.get(shape, "Observation"),
                    "couleur": CATEGORY_COLORS.get(data.get("categorie") or "Non classé", "#7f7f7f"),
                    "icon": {"circle": "●", "triangle": "▲", "square": "■", "diamond": "◆"}.get(shape, "●"),
                }
                meteo = data.get("meteo") or {}
                if meteo:
                    props["meteo"] = {
                        k: meteo.get(k) for k in (
                            "date", "heure", "temperature", "humidite", "ciel",
                            "phase_lunaire", "pluie_3j_precedents"
                        ) if meteo.get(k) is not None
                    }
                features.append({
                    "type": "Feature",
                    "geometry": {"type": "Point", "coordinates": [lon, lat]},
                    "properties": props,
                })

            geojson = {
                "type": "FeatureCollection",
                "name": f"Observations — {os.path.basename(self.photo_folder_path or 'Meuse')}",
                "crs": {"type": "name", "properties": {"name": "urn:ogc:def:crs:OGC:1.3:CRS84"}},
                "features": features,
                "legend": {
                    "categories": {k: v for k, v in CATEGORY_COLORS.items()},
                    "shapes": {
                        "circle": "Animal / observation",
                        "triangle": "Empreinte / trace",
                        "square": "Terrier / latrine",
                        "diamond": "Coulée / passage",
                    },
                },
            }
            with open(chemin, "w", encoding="utf-8") as f:
                json.dump(geojson, f, ensure_ascii=False, indent=2)
            self.log(f"🌍 Export GeoJSON réussi : {os.path.basename(chemin)} ({len(features)} entités)")
            messagebox.showinfo("Export réussi", f"{len(features)} entité(s) exportée(s) en GeoJSON.")
        except Exception as e:
            messagebox.showerror("Erreur d'export", str(e))

    # --- LOT 2 : Registre des sorties connues (pour la carte cumulée et les rapports agrégés) ---

    def _known_folders_path(self):
        try:
            self._migrate_legacy_config_files()
        except Exception:
            pass
        return os.path.join(self._user_data_dir(), "sorties_connues.json")

    def _load_known_folders(self):
        path = self._known_folders_path()
        if not os.path.exists(path):
            return []
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []

    def _save_known_folders(self, folders):
        try:
            with open(self._known_folders_path(), "w", encoding="utf-8") as f:
                json.dump(folders, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _register_known_folder(self, folder_path):
        """Ajoute (ou met à jour) le dossier courant dans le registre des sorties, utilisé par la carte
        cumulée et les rapports agrégés multi-sorties."""
        if not folder_path:
            return
        abs_path = os.path.normpath(folder_path)
        folders = self._load_known_folders()
        label = os.path.basename(abs_path) or abs_path
        sortie_date = None
        try:
            sortie_date = self._parse_sortie_date({"path": abs_path, "label": label})
        except Exception:
            sortie_date = None
        for entry in folders:
            if os.path.normpath(entry.get("path", "")) == abs_path:
                entry["last_opened"] = datetime.now().isoformat()
                entry["label"] = entry.get("label") or label
                if sortie_date and not entry.get("sortie_date"):
                    entry["sortie_date"] = sortie_date.isoformat()
                self._save_known_folders(folders)
                return
        folders.append({
            "path": abs_path,
            "label": label,
            "last_opened": datetime.now().isoformat(),
            "last_sync": datetime.now().isoformat(),
            "sortie_date": sortie_date.isoformat() if sortie_date else "",
        })
        self._save_known_folders(folders)

    def _unregister_known_folder(self, folder_path):
        """Retire une sortie du registre (sans supprimer les fichiers disque)."""
        if not folder_path:
            return False
        abs_path = os.path.normpath(folder_path)
        folders = self._load_known_folders()
        new_list = [
            e for e in folders
            if os.path.normpath(e.get("path", "")) != abs_path
        ]
        if len(new_list) == len(folders):
            return False
        self._save_known_folders(new_list)
        return True

    def open_manage_sorties(self):
        """Liste les sorties connues : retirer (ex. séance drone sans faune)."""
        folders = self._load_known_folders()
        win = ctk.CTkToplevel(self)
        win.title("Gérer les sorties")
        win.geometry("640x520")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Gérer les sorties enregistrées",
            font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text="Retirer une sortie la enlève du tableau de bord, de la carte cumulée et des briefs.\n"
                 "Les photos et le fichier observations.json sur le disque ne sont PAS effacés "
                 "(sauf si vous cochez l'option ci-dessous).",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            wraplength=600, justify="left",
        ).pack(anchor="w", padx=16, pady=(0, 8))

        delete_disk_var = tk.BooleanVar(value=False)
        ctk.CTkCheckBox(
            win,
            text="Aussi supprimer le dossier sur le disque (irrémédiable — photos comprises)",
            variable=delete_disk_var,
            text_color=UI.get("danger", "#c0392b"),
        ).pack(anchor="w", padx=16, pady=(0, 8))

        scroll = ctk.CTkScrollableFrame(win, fg_color=UI.get("card"))
        scroll.pack(fill="both", expand=True, padx=16, pady=6)

        state = {"folders": list(folders)}

        def refresh_list():
            for w in scroll.winfo_children():
                w.destroy()
            if not state["folders"]:
                ctk.CTkLabel(
                    scroll, text="Aucune sortie enregistrée.",
                    text_color=UI.get("text_dim"),
                ).pack(anchor="w", padx=8, pady=12)
                return

            def sort_key(e):
                d = self._parse_sortie_date(e)
                return d.toordinal() if d else 0

            for entry in sorted(state["folders"], key=sort_key, reverse=True):
                path = entry.get("path") or ""
                label = entry.get("label") or os.path.basename(path) or path
                d = self._parse_sortie_date(entry)
                date_txt = d.strftime("%d/%m/%Y") if d else "?"
                exists = os.path.isdir(path)
                row = ctk.CTkFrame(scroll, fg_color="transparent")
                row.pack(fill="x", pady=3, padx=4)
                ctk.CTkLabel(
                    row,
                    text="%s  ·  %s%s" % (label, date_txt, "" if exists else "  (dossier introuvable)"),
                    text_color=UI.get("text") if exists else UI.get("text_dim"),
                    anchor="w",
                ).pack(side="left", fill="x", expand=True)

                def retirer(p=path, lab=label):
                    msg = (
                        "Retirer cette sortie de GeoExif ?\n\n%s\n%s\n\n"
                        "Elle disparaîtra du tableau de bord et des cartes cumulées."
                        % (lab, p)
                    )
                    if delete_disk_var.get():
                        msg += (
                            "\n\n⚠ Vous avez coché la suppression disque :\n"
                            "le dossier et toutes les photos seront EFFACÉS."
                        )
                    if not messagebox.askyesno("Retirer la sortie", msg):
                        return
                    if delete_disk_var.get() and os.path.isdir(p):
                        try:
                            shutil.rmtree(p)
                            self.log("Dossier sortie supprimé : %s" % p)
                        except Exception as e:
                            messagebox.showerror("Suppression", "Impossible de supprimer le dossier :\n%s" % e)
                            return
                    self._unregister_known_folder(p)
                    # Si c'était le dossier courant
                    try:
                        if self.photo_folder_path and os.path.normpath(self.photo_folder_path) == os.path.normpath(p):
                            self.photo_folder_path = ""
                            self.photos_data = {}
                            try:
                                self.btn_browse_photos.configure(text="📁 Dossier photos")
                            except Exception:
                                pass
                    except Exception:
                        pass
                    state["folders"] = self._load_known_folders()
                    refresh_list()
                    self.log("Sortie retirée du registre : %s" % lab)

                ctk.CTkButton(
                    row, text="Retirer", width=90, height=28,
                    fg_color=UI.get("danger", "#c0392b"),
                    hover_color=UI.get("danger_hover", "#a93226"),
                    command=retirer,
                ).pack(side="right", padx=4)

        refresh_list()

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=16, pady=10)
        ctk.CTkButton(
            bf, text="Rescanner un dossier parent…",
            command=lambda: (win.destroy(), self.scan_parent_for_sorties()),
        ).pack(side="left", padx=3)
        ctk.CTkButton(
            bf, text="Fermer", fg_color=UI.get("card_alt"), command=win.destroy,
        ).pack(side="right", padx=3)
        win.after(200, lambda: self._style_dialog(win) if hasattr(self, "_style_dialog") else None)

    def _load_notes_from_archive(self, folder_path):
        """Charge observations.json depuis archive_carnets si le dossier photo est hors-ligne."""
        if not folder_path:
            return None
        try:
            root = self._local_archive_root()
            if not os.path.isdir(root):
                return None
            abs_path = os.path.normpath(os.path.abspath(folder_path))
            key = self._archive_key_for_folder(folder_path)
            cand = os.path.join(root, key, NOTES_FILE)
            if os.path.isfile(cand):
                with open(cand, "r", encoding="utf-8") as f:
                    data = json.load(f)
                return data if isinstance(data, dict) else None
            for name in os.listdir(root):
                d = os.path.join(root, name)
                meta_p = os.path.join(d, "meta.json")
                notes_p = os.path.join(d, NOTES_FILE)
                if not os.path.isfile(notes_p):
                    continue
                meta = {}
                if os.path.isfile(meta_p):
                    try:
                        with open(meta_p, "r", encoding="utf-8") as f:
                            meta = json.load(f) or {}
                    except Exception:
                        pass
                if os.path.normpath(meta.get("path_origine") or "") == abs_path:
                    with open(notes_p, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    return data if isinstance(data, dict) else None
                # fallback : même nom de dossier
                if (meta.get("label") or name.split("__")[0]) == os.path.basename(abs_path):
                    with open(notes_p, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    if isinstance(data, dict) and data:
                        return data
        except Exception:
            pass
        return None

    def _sortie_meta_path(self, folder_path):
        return os.path.join(folder_path or "", "sortie_meta.json")

    def _load_sortie_amazon_album(self, folder_path):
        if not folder_path:
            return ""
        # 1) meta locale
        try:
            p = self._sortie_meta_path(folder_path)
            if os.path.isfile(p):
                with open(p, "r", encoding="utf-8") as f:
                    meta = json.load(f) or {}
                return (meta.get("lien_amazon_album") or meta.get("amazon_album") or "").strip()
        except Exception:
            pass
        # 2) archive
        try:
            root = self._local_archive_root()
            key = self._archive_key_for_folder(folder_path)
            meta_p = os.path.join(root, key, "meta.json")
            if os.path.isfile(meta_p):
                with open(meta_p, "r", encoding="utf-8") as f:
                    meta = json.load(f) or {}
                return (meta.get("lien_amazon_album") or "").strip()
        except Exception:
            pass
        return ""

    def _save_sortie_amazon_album(self, folder_path, url):
        if not folder_path or not url:
            return
        meta = {}
        p = self._sortie_meta_path(folder_path)
        if os.path.isfile(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    meta = json.load(f) or {}
            except Exception:
                meta = {}
        meta["lien_amazon_album"] = url.strip()
        meta["updated"] = datetime.now().isoformat(timespec="seconds")
        try:
            if os.path.isdir(folder_path):
                with open(p, "w", encoding="utf-8") as f:
                    json.dump(meta, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
        # maj archive meta
        try:
            root = self._local_archive_root()
            key = self._archive_key_for_folder(folder_path)
            dest_dir = os.path.join(root, key)
            os.makedirs(dest_dir, exist_ok=True)
            meta_p = os.path.join(dest_dir, "meta.json")
            arch = {}
            if os.path.isfile(meta_p):
                with open(meta_p, "r", encoding="utf-8") as f:
                    arch = json.load(f) or {}
            arch["lien_amazon_album"] = url.strip()
            arch["path_origine"] = arch.get("path_origine") or os.path.normpath(os.path.abspath(folder_path))
            arch["label"] = arch.get("label") or os.path.basename(folder_path)
            with open(meta_p, "w", encoding="utf-8") as f:
                json.dump(arch, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def open_amazon_photo_link(self):
        url = ""
        try:
            url = (self.entry_amazon_photo.get() or "").strip()
        except Exception:
            pass
        if not url:
            messagebox.showinfo(
                "Amazon Photos",
                "Collez d'abord le lien de partage Amazon Photos pour cette image,\n"
                "puis cliquez Ouvrir.\n\n"
                "(Amazon Photos → partage / lien de l'élément)",
            )
            return
        try:
            webbrowser.open(url)
        except Exception as e:
            messagebox.showerror("Amazon Photos", str(e))

    def _format_french_day(self, d):
        """Date → libellé type « 19 juillet 2026 » (comme dans Amazon Photos)."""
        mois = (
            "janvier", "février", "mars", "avril", "mai", "juin",
            "juillet", "août", "septembre", "octobre", "novembre", "décembre",
        )
        try:
            return "%d %s %d" % (d.day, mois[d.month - 1], d.year)
        except Exception:
            return d.strftime("%d/%m/%Y")

    def fill_amazon_day_from_sortie(self):
        """Remplit le repère Amazon à partir de la date de la sortie / photo."""
        d = None
        # 1) date dossier
        if self.photo_folder_path:
            try:
                d = self._parse_sortie_date({
                    "path": self.photo_folder_path,
                    "label": os.path.basename(self.photo_folder_path),
                })
            except Exception:
                d = None
        # 2) date photo selectionnee
        if d is None and self.selected_photo_path:
            try:
                fn = os.path.basename(self.selected_photo_path)
                info = (self.photos_data or {}).get(fn) or {}
                raw = info.get("date") or ""
                if raw:
                    d = datetime.strptime(str(raw)[:10].replace(":", "-").replace("/", "-")[:10], "%Y-%m-%d").date()
            except Exception:
                try:
                    mtime = os.path.getmtime(self.selected_photo_path)
                    d = datetime.fromtimestamp(mtime).date()
                except Exception:
                    d = None
        if d is None:
            d = datetime.now().date()
        label = self._format_french_day(d)
        try:
            self.entry_amazon_album.delete(0, tk.END)
            self.entry_amazon_album.insert(0, label)
        except Exception:
            pass
        self.log("Repère Amazon : %s" % label)

    def open_amazon_album_link(self):
        """Si c'est une URL → navigateur ; sinon rappel du jour dans Amazon Photos."""
        txt = ""
        try:
            txt = (self.entry_amazon_album.get() or "").strip()
        except Exception:
            pass
        if not txt and self.photo_folder_path:
            txt = self._load_sortie_amazon_album(self.photo_folder_path)
        if not txt:
            messagebox.showinfo(
                "Amazon Photos",
                "Aucun repère pour cette sortie.\n"
                "Utilisez « Jour auto » (ex. 19 juillet 2026), puis Sauver.\n"
                "Sur Amazon Photos, ouvrez ce jour dans la vue calendrier / bibliothèque.",
            )
            return
        if txt.lower().startswith("http://") or txt.lower().startswith("https://"):
            try:
                webbrowser.open(txt)
            except Exception as e:
                messagebox.showerror("Amazon Photos", str(e))
            return
        # Texte libre = jour Amazon
        messagebox.showinfo(
            "Amazon Photos",
            "Repère enregistré :\n\n« %s »\n\n"
            "Ouvrez Amazon Photos → bibliothèque / calendrier\n"
            "et allez à cette date (vos photos y sont classées au jour le jour).\n\n"
            "Astuce : un lien de partage direct reste possible dans « Lien Amazon (photo) »."
            % txt,
        )
        try:
            webbrowser.open("https://www.amazon.fr/photos")
        except Exception:
            pass

    def save_amazon_album_link(self):
        if not self.photo_folder_path:
            messagebox.showwarning("Amazon Photos", "Ouvrez d'abord une sortie.")
            return
        txt = ""
        try:
            txt = (self.entry_amazon_album.get() or "").strip()
        except Exception:
            pass
        if not txt:
            messagebox.showwarning(
                "Amazon Photos",
                "Indiquez le jour (ex. 19 juillet 2026) via « Jour auto », puis Sauver.",
            )
            return
        self._save_sortie_amazon_album(self.photo_folder_path, txt)
        messagebox.showinfo(
            "Amazon Photos",
            "Repère enregistré pour cette sortie :\n« %s »\n\n"
            "(Visible dans le carnet et l'archive, même sans disque photo.)"
            % txt,
        )
        self.log("Repère Amazon (jour) : %s" % txt)

    def create_recovery_pack_zip(self):
        """ZIP unique : archive_carnets + sorties_connues + config (+ ia) pour recovery hors machine."""
        try:
            import zipfile
        except Exception:
            messagebox.showerror("Recovery", "Module zipfile indisponible.")
            return
        # S'assurer que l'archive est a jour si disques accessibles
        try:
            self.mirror_all_known_sorties_to_archive(silent=True)
        except Exception:
            pass

        default_name = "GeoExif_recovery_%s.zip" % datetime.now().strftime("%Y%m%d_%H%M")
        dest = filedialog.asksaveasfilename(
            title="Enregistrer le pack recovery",
            defaultextension=".zip",
            initialfile=default_name,
            filetypes=[("ZIP", "*.zip")],
        )
        if not dest:
            return

        data_dir = self._user_data_dir()
        archive_root = self._local_archive_root()
        files_flat = []
        for name in (
            "sorties_connues.json",
            CONFIG_FILE,
            "ia_config.json",
            PROJECT_FILE,
            SPECIES_FILE,
        ):
            p = os.path.join(data_dir, name)
            if os.path.isfile(p):
                files_flat.append((p, name))
        # aussi backup_folder app_config path variants
        try:
            cfg_p = self._app_config_path()
            if os.path.isfile(cfg_p) and not any(a[1] == CONFIG_FILE for a in files_flat):
                files_flat.append((cfg_p, CONFIG_FILE))
        except Exception:
            pass

        n_files = 0
        try:
            with zipfile.ZipFile(dest, "w", compression=zipfile.ZIP_DEFLATED) as zf:
                for src, arcname in files_flat:
                    zf.write(src, arcname=arcname)
                    n_files += 1
                if os.path.isdir(archive_root):
                    for root, _dirs, files in os.walk(archive_root):
                        for fn in files:
                            full = os.path.join(root, fn)
                            rel = os.path.relpath(full, data_dir)
                            zf.write(full, arcname=rel.replace("\\", "/"))
                            n_files += 1
                # readme
                readme = (
                    "GeoExif — pack recovery\n"
                    "Contient : sorties_connues, config, archive_carnets (observations JSON).\n"
                    "Les photos ne sont PAS incluses (stockage Amazon / disque photo).\n"
                    "Restauration : extraire dans %%LOCALAPPDATA%%\\GeoExif\\ "
                    "ou utiliser GeoExif → Archive locale / Restaurer.\n"
                    "Créé : %s\nVersion : %s\n"
                    % (datetime.now().isoformat(timespec="seconds"), APP_VERSION)
                )
                zf.writestr("LISEZMOI_recovery.txt", readme)
            messagebox.showinfo(
                "Pack recovery",
                "Pack créé :\n%s\n\n%d fichier(s) inclus.\n"
                "À copier sur clé USB / cloud (sans les photos)."
                % (dest, n_files),
            )
            self.log("Pack recovery ZIP : %s (%d fichiers)" % (dest, n_files))
        except Exception as e:
            messagebox.showerror("Pack recovery", str(e))

    def mirror_all_known_sorties_to_archive(self, *, silent=False):
        """Copie tous les carnets accessibles vers l'archive PC (a faire disque branche)."""
        folders = self._load_known_folders()
        ok, skip, err = 0, 0, 0
        for entry in folders:
            path = entry.get("path") or ""
            if not path or not os.path.isdir(path):
                skip += 1
                continue
            notes_path = os.path.join(path, NOTES_FILE)
            if not os.path.isfile(notes_path):
                skip += 1
                continue
            try:
                with open(notes_path, "r", encoding="utf-8") as f:
                    payload = f.read()
                self._mirror_notes_to_local_archive(path, payload)
                ok += 1
            except Exception:
                err += 1
        msg = (
            "Archive locale mise a jour.\n\n"
            "%d sortie(s) copiee(s) sur le PC\n"
            "%d ignoree(s) (hors-ligne ou sans carnet)\n"
            "%d erreur(s)\n\n"
            "Dossier :\n%s"
            % (ok, skip, err, self._local_archive_root())
        )
        if not silent:
            messagebox.showinfo("Archive locale", msg)
        try:
            self.log("Miroir archive : %d ok, %d skip, %d err" % (ok, skip, err))
        except Exception:
            pass
        return ok

    def _gather_observations_for_folder(self, folder_path, exe_path=None, *, read_exif=False):
        """Observations d'un dossier (notes JSON + GPS).

        Par défaut : rapide, GPS depuis le carnet ou le cache mémoire (pas d'ExifTool).
        read_exif=True : complète les GPS manquants via ExifTool (peut être lent sur gros dossiers).
        Si le dossier est hors-ligne : bascule sur l'archive locale PC.
        """
        notes = {}
        offline = False
        if folder_path and os.path.isdir(folder_path):
            notes = self._load_notes_dict(folder_path) or {}
            # Miroir opportuniste si le carnet existe (pour survivre au debranchement)
            if notes:
                try:
                    payload = json.dumps(notes, indent=4, ensure_ascii=False)
                    self._mirror_notes_to_local_archive(folder_path, payload)
                except Exception:
                    pass
        else:
            offline = True
            notes = self._load_notes_from_archive(folder_path) or {}

        if not notes and folder_path:
            # dernier recours archive meme si dossier existe mais sans notes
            notes = self._load_notes_from_archive(folder_path) or {}

        if not notes:
            return []

        # Hors-ligne : pas d'ExifTool possible
        if offline:
            read_exif = False
            exe_path = None

        coords_source = {}
        if os.path.normpath(folder_path) == os.path.normpath(self.photo_folder_path or ""):
            coords_source = dict(self.photos_data or {})

        # Compléter via ExifTool seulement si demandé et si des photos n'ont pas encore de GPS
        need_exif = False
        if read_exif and exe_path:
            for filename, data in notes.items():
                if not isinstance(data, dict):
                    continue
                if data.get("lat") and data.get("lon"):
                    continue
                if coords_source.get(filename, {}).get("lat") and coords_source.get(filename, {}).get("lon"):
                    continue
                if str(filename).startswith("_"):
                    continue  # obs manuelle / terrain
                need_exif = True
                break

        if need_exif and exe_path:
            try:
                startupinfo = None
                if os.name == "nt":
                    startupinfo = self._exiftool_startupinfo()
                cmd = [
                    exe_path, "-n", "-GPSLatitude", "-GPSLongitude", "-json",
                    "-q", "-q", os.path.normpath(folder_path),
                ]
                proc = subprocess.run(
                    cmd, startupinfo=startupinfo, capture_output=True,
                    text=True, encoding="utf-8", errors="ignore", timeout=20,
                )
                if proc.returncode == 0 and proc.stdout.strip():
                    for item in json.loads(proc.stdout):
                        fn = os.path.basename(item.get("SourceFile", "").replace("/", os.sep))
                        lat = self.parse_coord(item.get("GPSLatitude"))
                        lon = self.parse_coord(item.get("GPSLongitude"))
                        if lat and lon:
                            coords_source[fn] = {"lat": lat, "lon": lon}
            except Exception:
                pass

        resultats = []
        for filename, data in notes.items():
            if not isinstance(data, dict):
                continue
            info = coords_source.get(filename, {})
            enrichi = dict(data)
            enrichi["fichier"] = filename
            enrichi["dossier"] = os.path.basename(folder_path) or folder_path
            enrichi["dossier_path"] = folder_path
            # Chemin photo seulement si fichier réel
            full = os.path.join(folder_path, filename)
            enrichi["chemin_complet"] = full if os.path.isfile(full) else ""
            # Priorité : GPS déjà dans le carnet > cache / ExifTool
            lat = data.get("lat") if data.get("lat") is not None else info.get("lat")
            lon = data.get("lon") if data.get("lon") is not None else info.get("lon")
            enrichi["lat"] = lat
            enrichi["lon"] = lon
            resultats.append(enrichi)
        return resultats

    def _build_report_dataset(self, folder_paths, *, read_exif=False, progress_cb=None):
        """Agrège les observations de plusieurs dossiers.

        read_exif=False (défaut) : rapide, idéal brief IA — lit uniquement observations.json.
        """
        exe_path = self.get_exiftool_path() if read_exif else None
        dataset = []
        total = max(1, len(folder_paths))
        for i, fp in enumerate(folder_paths):
            if progress_cb:
                try:
                    progress_cb(i + 1, total, os.path.basename(fp))
                except Exception:
                    pass
            dataset.extend(
                self._gather_observations_for_folder(fp, exe_path, read_exif=read_exif)
            )
        return dataset


    def _disk_cache_dir(self):
        """Dossier cache apercus (temp utilisateur) — inoffensif pour le disque."""
        import tempfile
        d = os.path.join(tempfile.gettempdir(), "geoexif_cache")
        try:
            os.makedirs(d, exist_ok=True)
        except Exception:
            d = tempfile.gettempdir()
        return d

    def _disk_cache_path(self, chemin_complet, kind="thumb", max_dim=500):
        """Chemin de fichier cache stable pour un RAW/photo + taille demandee."""
        import hashlib
        try:
            mtime = os.path.getmtime(chemin_complet)
            size = os.path.getsize(chemin_complet)
        except Exception:
            mtime, size = 0, 0
        key = "%s|%s|%s|%s" % (os.path.normcase(os.path.abspath(chemin_complet)), mtime, size, max_dim)
        h = hashlib.md5(key.encode("utf-8", errors="replace")).hexdigest()[:16]
        base = os.path.splitext(os.path.basename(chemin_complet))[0][:40]
        safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in base)
        return os.path.join(self._disk_cache_dir(), "%s_%s_%s.jpg" % (kind, safe, h))

    def _read_disk_cache(self, cache_path):
        try:
            if cache_path and os.path.isfile(cache_path) and os.path.getsize(cache_path) > 200:
                with open(cache_path, "rb") as f:
                    return f.read()
        except Exception:
            pass
        return None

    def _write_disk_cache(self, cache_path, data):
        if not cache_path or not data:
            return
        try:
            tmp = cache_path + ".tmp"
            with open(tmp, "wb") as f:
                f.write(data)
            os.replace(tmp, cache_path)
        except Exception:
            try:
                with open(cache_path, "wb") as f:
                    f.write(data)
            except Exception:
                pass

    def clear_preview_cache(self):
        """Vide le cache disque des apercus (menu / outils)."""
        d = self._disk_cache_dir()
        n = 0
        try:
            for name in os.listdir(d):
                if name.endswith(".jpg") or name.endswith(".tmp"):
                    try:
                        os.remove(os.path.join(d, name))
                        n += 1
                    except Exception:
                        pass
        except Exception as e:
            messagebox.showerror("Cache", str(e))
            return
        # cache memoire aussi
        try:
            self._thumb_cache = {}
        except Exception:
            pass
        messagebox.showinfo("Cache", "Cache apercus vide (%d fichier(s))." % n)
        self.log("Cache apercus vide : %d fichier(s)" % n)

    def _get_thumbnail_bytes(self, chemin_complet, exe_path, max_dim=500):
        """Miniature JPEG (bytes), RAW inclus.
        Cache memoire + cache disque pour eviter de relancer ExifTool."""
        if not hasattr(self, "_thumb_cache"):
            self._thumb_cache = {}
        try:
            mtime = os.path.getmtime(chemin_complet) if os.path.exists(chemin_complet) else 0
        except Exception:
            mtime = 0
        mem_key = (chemin_complet, max_dim, mtime)
        if mem_key in self._thumb_cache:
            return self._thumb_cache[mem_key]

        # Cache disque
        disk_path = self._disk_cache_path(chemin_complet, kind="thumb", max_dim=max_dim)
        data_disk = self._read_disk_cache(disk_path)
        if data_disk:
            self._thumb_cache[mem_key] = data_disk
            if len(self._thumb_cache) > 250:
                for k in list(self._thumb_cache.keys())[:80]:
                    del self._thumb_cache[k]
            return data_disk

        try:
            ext = os.path.splitext(chemin_complet)[1].lower()
            data = None
            img = None
            if ext in (".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"):
                img = Image.open(chemin_complet)
            else:
                if not exe_path:
                    return None
                startupinfo = self._exiftool_startupinfo()
                # Preférer PreviewImage puis JpgFromRaw puis ThumbnailImage
                for tag in ("-PreviewImage", "-JpgFromRaw", "-ThumbnailImage"):
                    try:
                        proc = subprocess.run(
                            [exe_path, "-b", tag, chemin_complet],
                            startupinfo=startupinfo, capture_output=True, timeout=25,
                        )
                        data = proc.stdout
                        if data and len(data) >= 100:
                            break
                    except Exception:
                        data = None
                if not data or len(data) < 100:
                    return None
                img = Image.open(io.BytesIO(data))

            img = img.convert("RGB")
            img.thumbnail((max_dim, max_dim), Image.Resampling.LANCZOS if hasattr(Image, "Resampling") else Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            result = buf.getvalue()
            self._thumb_cache[mem_key] = result
            self._write_disk_cache(disk_path, result)
            if len(self._thumb_cache) > 250:
                for k in list(self._thumb_cache.keys())[:80]:
                    del self._thumb_cache[k]
            return result
        except Exception:
            return None

    def _prefetch_thumbnails_parallel(self, chemins, exe_path, max_dim=500):
        """Précharge des miniatures en parallèle (ThreadPool) pour accélérer les rapports PDF/Word."""
        chemins = [c for c in chemins if c and os.path.exists(c)]
        if not chemins:
            return
        def _one(c):
            return c, self._get_thumbnail_bytes(c, exe_path, max_dim)
        with ThreadPoolExecutor(max_workers=CPU_WORKERS) as pool:
            list(pool.map(_one, chemins))

    def export_report_pdf(self, dataset, title, filepath):
        """Génère un rapport PDF (miniatures + détails) à partir d'une liste d'observations enrichies."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
            from reportlab.lib import colors as rl_colors
        except ImportError:
            messagebox.showerror(
                "Bibliothèque manquante",
                "L'export PDF nécessite le module 'reportlab'.\nInstallez-le avec :\n\npip install reportlab"
            )
            return False

        exe_path = self.get_exiftool_path()
        # Précharge les miniatures en parallèle (CPU multi-cœur / I/O concurrent)
        chemins_thumbs = [
            obs.get("chemin_complet", "") for obs in dataset
            if obs.get("chemin_complet") and os.path.exists(obs.get("chemin_complet", ""))
        ]
        if chemins_thumbs:
            self._prefetch_thumbnails_parallel(chemins_thumbs, exe_path)
        styles = getSampleStyleSheet()
        style_titre = ParagraphStyle("TitreCustom", parent=styles["Title"], textColor=rl_colors.HexColor("#1f7d37"))
        style_h2 = ParagraphStyle("H2Custom", parent=styles["Heading2"], textColor=rl_colors.HexColor("#2b6cb0"), spaceBefore=14)
        style_normal = styles["Normal"]
        style_meta = ParagraphStyle("Meta", parent=styles["Normal"], textColor=rl_colors.HexColor("#555555"), fontSize=9)

        try:
            doc = SimpleDocTemplate(filepath, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
            story = [
                Paragraph(title, style_titre),
                Paragraph(f"Généré par GeoExif Meuse 55 — {APP_AUTHOR} — {datetime.now().strftime('%d/%m/%Y %H:%M')}", style_meta),
                Spacer(1, 0.5 * cm),
            ]

            # Carte des points GPS en tête de rapport
            map_bytes = self._generate_map_image(dataset)
            if map_bytes:
                try:
                    story.append(Paragraph("Carte des observations", style_h2))
                    story.append(RLImage(io.BytesIO(map_bytes), width=16 * cm, height=10.5 * cm))
                    story.append(Spacer(1, 0.4 * cm))
                    story.append(Paragraph(
                        "Légende : ● animal vu · ▲ empreinte/trace · ■ terrier/latrine · ◆ coulée/passage — couleurs = catégorie taxonomique.",
                        style_meta
                    ))
                    story.append(Spacer(1, 0.6 * cm))
                except Exception:
                    pass

            if not dataset:
                story.append(Paragraph("Aucune observation à inclure dans ce rapport.", style_normal))

            par_categorie = {}
            for obs in dataset:
                cat = obs.get("categorie") or "Non classé"
                par_categorie.setdefault(cat, []).append(obs)

            for cat in sorted(par_categorie.keys()):
                obs_list = par_categorie[cat]
                story.append(Paragraph(f"{cat} ({len(obs_list)} observation{'s' if len(obs_list) > 1 else ''})", style_h2))

                for obs in sorted(obs_list, key=lambda o: (o.get("dossier", ""), o.get("heure") or "")):
                    story.append(Paragraph(f"<b>{obs.get('espece', '?')}</b> — {obs.get('nombre', '?')} ind.", style_normal))

                    meta_parts = []
                    if obs.get("heure"): meta_parts.append(f"Heure : {obs['heure']}")
                    if obs.get("lieu"): meta_parts.append(f"Lieu : {obs['lieu']}")
                    if obs.get("dossier"): meta_parts.append(f"Sortie : {os.path.basename(obs['dossier'])}")
                    meteo = obs.get("meteo") or {}
                    if meteo.get("temperature") is not None:
                        meta_parts.append(f"Météo : {meteo.get('temperature')}°C, {meteo.get('humidite')}% hum., {meteo.get('ciel')}")
                    if meteo.get("phase_lunaire"):
                        meta_parts.append(meteo.get("phase_lunaire"))
                    if meta_parts:
                        story.append(Paragraph(" · ".join(meta_parts), style_meta))
                    if obs.get("notes_libres"):
                        story.append(Paragraph(obs["notes_libres"], style_normal))

                    chemin_complet = obs.get("chemin_complet", "")
                    thumb = self._get_thumbnail_bytes(chemin_complet, exe_path) if os.path.exists(chemin_complet) else None
                    if thumb:
                        try:
                            pil_img = Image.open(io.BytesIO(thumb))
                            w, h = pil_img.size
                            target_w = 6 * cm
                            target_h = target_w * h / w
                            if target_h > 8 * cm:
                                target_h = 8 * cm
                                target_w = target_h * w / h
                            story.append(RLImage(io.BytesIO(thumb), width=target_w, height=target_h))
                        except Exception:
                            pass
                    story.append(Spacer(1, 0.5 * cm))

            doc.build(story)
            return True
        except Exception as e:
            messagebox.showerror("Erreur d'export PDF", str(e))
            return False

    def export_report_docx(self, dataset, title, filepath):
        """Génère un rapport Word (miniatures + détails) à partir d'une liste d'observations enrichies."""
        try:
            from docx import Document
            from docx.shared import Cm, Pt, RGBColor
        except ImportError:
            messagebox.showerror(
                "Bibliothèque manquante",
                "L'export Word nécessite le module 'python-docx'.\nInstallez-le avec :\n\npip install python-docx"
            )
            return False

        exe_path = self.get_exiftool_path()
        try:
            doc = Document()
            doc.add_heading(title, level=0)
            p_meta = doc.add_paragraph(f"Généré par GeoExif Meuse 55 — {APP_AUTHOR} — {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            p_meta.runs[0].font.size = Pt(9)
            p_meta.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Carte des points GPS en tête de rapport
            map_bytes = self._generate_map_image(dataset)
            if map_bytes:
                try:
                    doc.add_heading("Carte des observations", level=1)
                    doc.add_picture(io.BytesIO(map_bytes), width=Cm(16))
                    p_leg = doc.add_paragraph(
                        "Légende : ● animal vu · ▲ empreinte/trace · ■ terrier/latrine · ◆ coulée/passage — couleurs = catégorie taxonomique."
                    )
                    p_leg.runs[0].font.size = Pt(8)
                    p_leg.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                except Exception:
                    pass

            if not dataset:
                doc.add_paragraph("Aucune observation à inclure dans ce rapport.")

            par_categorie = {}
            for obs in dataset:
                cat = obs.get("categorie") or "Non classé"
                par_categorie.setdefault(cat, []).append(obs)

            for cat in sorted(par_categorie.keys()):
                obs_list = par_categorie[cat]
                doc.add_heading(f"{cat} ({len(obs_list)} observation{'s' if len(obs_list) > 1 else ''})", level=1)

                for obs in sorted(obs_list, key=lambda o: (o.get("dossier", ""), o.get("heure") or "")):
                    doc.add_heading(f"{obs.get('espece', '?')} — {obs.get('nombre', '?')} ind.", level=2)

                    meta_parts = []
                    if obs.get("heure"): meta_parts.append(f"Heure : {obs['heure']}")
                    if obs.get("lieu"): meta_parts.append(f"Lieu : {obs['lieu']}")
                    if obs.get("dossier"): meta_parts.append(f"Sortie : {os.path.basename(obs['dossier'])}")
                    meteo = obs.get("meteo") or {}
                    if meteo.get("temperature") is not None:
                        meta_parts.append(f"Météo : {meteo.get('temperature')}°C, {meteo.get('humidite')}% hum., {meteo.get('ciel')}")
                    if meteo.get("phase_lunaire"):
                        meta_parts.append(meteo.get("phase_lunaire"))
                    if meta_parts:
                        p = doc.add_paragraph(" | ".join(meta_parts))
                        p.runs[0].font.size = Pt(9)
                        p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    if obs.get("notes_libres"):
                        doc.add_paragraph(obs["notes_libres"])

                    chemin_complet = obs.get("chemin_complet", "")
                    thumb = self._get_thumbnail_bytes(chemin_complet, exe_path) if os.path.exists(chemin_complet) else None
                    if thumb:
                        try:
                            doc.add_picture(io.BytesIO(thumb), width=Cm(7))
                        except Exception:
                            pass
                    doc.add_paragraph("")

            doc.save(filepath)
            return True
        except Exception as e:
            messagebox.showerror("Erreur d'export Word", str(e))
            return False

    def export_daily_pdf(self):
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Sélectionnez d'abord un dossier de photos.")
            return
        dataset = self._gather_observations_for_folder(self.photo_folder_path, self.get_exiftool_path())
        if not dataset:
            messagebox.showinfo("Aucune donnée", "Aucune observation enregistrée à exporter.")
            return
        chemin = self._ask_save_report(".pdf", [("Fichier PDF", "*.pdf")], "rapport_meuse.pdf")
        if not chemin:
            return
        titre = f"Carnet d'observation naturaliste — {os.path.basename(self.photo_folder_path)}"
        if self.export_report_pdf(dataset, titre, chemin):
            self.log(f"📕 Export PDF réussi : {os.path.basename(chemin)}")
            messagebox.showinfo("Export réussi", "Le rapport PDF a été généré.")

    def export_daily_docx(self):
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Sélectionnez d'abord un dossier de photos.")
            return
        dataset = self._gather_observations_for_folder(self.photo_folder_path, self.get_exiftool_path())
        if not dataset:
            messagebox.showinfo("Aucune donnée", "Aucune observation enregistrée à exporter.")
            return
        chemin = self._ask_save_report(".docx", [("Document Word", "*.docx")], "rapport_meuse.docx")
        if not chemin:
            return
        titre = f"Carnet d'observation naturaliste — {os.path.basename(self.photo_folder_path)}"
        if self.export_report_docx(dataset, titre, chemin):
            self.log(f"📘 Export Word réussi : {os.path.basename(chemin)}")
            messagebox.showinfo("Export réussi", "Le rapport Word a été généré.")

    def _show_cluster_info(self, obs_list):
        """Popup listant les observations regroupées sous un même point sur la carte cumulée."""
        win = ctk.CTkToplevel(self)
        win.title("%d observation(s) a ce point" % len(obs_list))
        win.geometry("560x560")
        try:
            win.configure(fg_color=UI.get("bg", "#f5f7f6"))
        except Exception:
            pass
        self._prepare_tool_window(win)
        ctk.CTkLabel(
            win, text="%d observation(s) :" % len(obs_list),
            font=ctk.CTkFont(size=15, weight="bold"),
            text_color=UI.get("text", "#111111"),
        ).pack(anchor="w", padx=14, pady=12)
        scroll = ctk.CTkScrollableFrame(win, fg_color=UI.get("card", "#fff"))
        scroll.pack(fill="both", expand=True, padx=12, pady=(0, 12))
        for o in obs_list:
            texte = "%s  ·  %s  ·  %s" % (
                o.get("espece", "?"),
                os.path.basename(str(o.get("dossier", "") or o.get("_folder_path", "") or "")),
                o.get("heure", "?"),
            )
            ctk.CTkLabel(
                scroll, text=texte, anchor="w",
                font=ctk.CTkFont(size=13),
                text_color=UI.get("text", "#111111"),
            ).pack(fill="x", pady=4, padx=6)

    def open_existing_sortie(self, folder_path, *, switch_tab=True, close_window=None):
        """Ouvre une sortie connue pour consultation / édition dans le carnet principal."""
        path = os.path.normpath(folder_path or "")
        if not path or not os.path.isdir(path):
            messagebox.showwarning(
                "Sortie introuvable",
                f"Le dossier n'existe plus ou est inaccessible :\n{folder_path}"
            )
            return False
        try:
            if close_window is not None:
                try:
                    close_window.destroy()
                except Exception:
                    pass
        except Exception:
            pass

        # Réutilise le chargeur carnet (photos optionnelles + observations)
        if hasattr(self, "_activate_carnet_folder"):
            self._activate_carnet_folder(path)
        else:
            self.photo_folder_path = path
            try:
                self.btn_browse_photos.configure(text=f"📁 {os.path.basename(path)}")
            except Exception:
                pass

        self._register_known_folder(path)

        # Précharge GPS depuis le JSON + lance lecture EXIF en arrière-plan si besoin
        try:
            exe = self.get_exiftool_path()
            if exe and hasattr(self, "extract_coordinates_after_sync"):
                threading.Thread(
                    target=self.extract_coordinates_after_sync, args=(exe,), daemon=True
                ).start()
        except Exception:
            pass

        if switch_tab:
            try:
                # Noms d'onglets possibles selon version UI
                for name in ("📒   Carnet", "📝  Carnet", "📝 Carnet", "Carnet"):
                    try:
                        self.tab_view.set(name)
                        break
                    except Exception:
                        continue
            except Exception:
                pass

        try:
            self.reload_devices_for_folder()
        except Exception:
            pass
        self.log(f"Sortie ouverte pour édition : {path}")
        return True

    def open_multi_sorties_map(self):
        """Carte cumulée : filtres toutes / par sortie / par espèce / masquer + liste cliquable."""
        folders = self._load_known_folders()
        if not folders:
            messagebox.showinfo(
                "Aucune sortie",
                "Aucun dossier de sortie connu.\nChargez au moins un dossier de photos ou un carnet."
            )
            return

        win = ctk.CTkToplevel(self)
        win.title("Carte cumulée — Toutes les sorties")
        win.geometry("1240x800")
        self._prepare_tool_window(win)

        # État des filtres
        state = {
            "points": [],
            "markers": [],
            "mode": "toutes",       # toutes | sortie | espece | masquer
            "sortie": "Toutes",
            "espece": "Toutes",
        }

        try:
            win.configure(fg_color=UI.get("bg", "#101612"))
        except Exception:
            pass

        top = ctk.CTkFrame(win, fg_color="transparent")
        top.pack(fill="x", padx=10, pady=(10, 4))
        ctk.CTkLabel(
            top, text=f"{len(folders)} sortie(s) enregistrée(s)",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color=UI.get("text", "#f3f7f4"),
        ).pack(side="left")
        lbl_status = ctk.CTkLabel(
            top, text="Chargement…",
            text_color=UI.get("text_dim", UI.get("text_muted", "#666")),
        )
        lbl_status.pack(side="right")

        # Barre de filtres (couleurs explicites pour thème clair/papier)
        filt = ctk.CTkFrame(
            win, fg_color=UI.get("card", "#1c2620"), corner_radius=10,
            border_width=1, border_color=UI.get("border", "#ccc"),
        )
        filt.pack(fill="x", padx=10, pady=(0, 6))
        ctk.CTkLabel(
            filt, text="Affichage :",
            font=ctk.CTkFont(size=12, weight="bold"),
            text_color=UI.get("text", "#f3f7f4"),
        ).pack(side="left", padx=(10, 6), pady=8)

        mode_var = tk.StringVar(value="toutes")
        sortie_var = tk.StringVar(value="Toutes")
        espece_var = tk.StringVar(value="Toutes")

        def on_mode(value=None):
            state["mode"] = (value or mode_var.get() or "toutes").lower()
            # activer/désactiver menus selon mode
            try:
                if state["mode"] == "sortie":
                    menu_sortie.configure(state="normal")
                    menu_espece.configure(state="disabled")
                elif state["mode"] == "espece":
                    menu_sortie.configure(state="disabled")
                    menu_espece.configure(state="normal")
                else:
                    menu_sortie.configure(state="disabled")
                    menu_espece.configure(state="disabled")
            except Exception:
                pass
            redraw_markers()

        seg = ctk.CTkSegmentedButton(
            filt,
            values=["Toutes", "Par sortie", "Par espèce", "Masquer"],
            command=lambda v: on_mode({
                "Toutes": "toutes", "Par sortie": "sortie",
                "Par espèce": "espece", "Masquer": "masquer",
            }.get(v, "toutes")),
            font=ctk.CTkFont(size=12),
            selected_color=UI.get("accent", "#4ec4b0"),
            selected_hover_color=UI.get("accent_hover", "#3aab98"),
            unselected_color=UI.get("card_alt", "#24302a"),
            unselected_hover_color=UI.get("border", "#ccc"),
            text_color=UI.get("text", "#111"),
            text_color_disabled=UI.get("text_muted", "#888"),
            height=30,
        )
        seg.pack(side="left", padx=4, pady=8)
        seg.set("Toutes")

        ctk.CTkLabel(
            filt, text="Sortie", font=ctk.CTkFont(size=11),
            text_color=UI.get("text", "#111"),
        ).pack(side="left", padx=(12, 4))
        menu_sortie = ctk.CTkOptionMenu(
            filt, variable=sortie_var, values=["Toutes"], width=160, height=28,
            command=lambda _v: redraw_markers(),
            fg_color=UI.get("card_alt", "#24302a"),
            button_color=UI.get("accent", "#4ec4b0"),
            button_hover_color=UI.get("accent_hover", "#3aab98"),
            text_color=UI.get("text", "#111"),
            dropdown_fg_color=UI.get("card", "#fff"),
            dropdown_text_color=UI.get("text", "#111"),
            dropdown_hover_color=UI.get("accent_soft", UI.get("card_alt")),
        )
        menu_sortie.pack(side="left", padx=2)
        menu_sortie.configure(state="disabled")

        ctk.CTkLabel(
            filt, text="Espèce", font=ctk.CTkFont(size=11),
            text_color=UI.get("text", "#111"),
        ).pack(side="left", padx=(12, 4))
        menu_espece = ctk.CTkOptionMenu(
            filt, variable=espece_var, values=["Toutes"], width=160, height=28,
            command=lambda _v: redraw_markers(),
            fg_color=UI.get("card_alt", "#24302a"),
            button_color=UI.get("accent", "#4ec4b0"),
            button_hover_color=UI.get("accent_hover", "#3aab98"),
            text_color=UI.get("text", "#111"),
            dropdown_fg_color=UI.get("card", "#fff"),
            dropdown_text_color=UI.get("text", "#111"),
            dropdown_hover_color=UI.get("accent_soft", UI.get("card_alt")),
        )
        menu_espece.pack(side="left", padx=2)
        menu_espece.configure(state="disabled")

        body = ctk.CTkFrame(win, fg_color=UI.get("bg", "#101612"))
        body.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        body.grid_columnconfigure(1, weight=1)
        body.grid_rowconfigure(0, weight=1)

        side = ctk.CTkFrame(
            body, fg_color=UI.get("card", "#1c2620"), corner_radius=12,
            border_width=1, border_color=UI.get("border", "#354840"), width=280,
        )
        side.grid(row=0, column=0, sticky="nsew", padx=(0, 8))
        side.grid_propagate(False)
        ctk.CTkLabel(
            side, text="Sorties — ouvrir / filtrer",
            font=ctk.CTkFont(size=12, weight="bold"),
            text_color=UI.get("text", "#111"),
        ).pack(anchor="w", padx=10, pady=(10, 4))
        list_scroll = ctk.CTkScrollableFrame(side, fg_color="transparent")
        list_scroll.pack(fill="both", expand=True, padx=6, pady=(0, 8))

        def open_one(path):
            ok = self.open_existing_sortie(path, switch_tab=True, close_window=win)
            if ok:
                messagebox.showinfo(
                    "Sortie ouverte",
                    f"« {os.path.basename(path)} » est chargée dans le carnet.\n"
                    "Vous pouvez annoter, corriger le GPS ou ajouter des observations."
                )

        def filter_on_sortie(label):
            mode_var.set("sortie")
            seg.set("Par sortie")
            sortie_var.set(label)
            state["mode"] = "sortie"
            try:
                menu_sortie.configure(state="normal")
                menu_espece.configure(state="disabled")
            except Exception:
                pass
            redraw_markers()

        for entry in folders:
            path = entry.get("path") or ""
            label = entry.get("label") or os.path.basename(path) or path
            last = (entry.get("last_sync") or "")[:10]
            row = ctk.CTkFrame(list_scroll, fg_color=UI.get("card_alt", "#24302a"), corner_radius=8)
            row.pack(fill="x", pady=3, padx=2)
            ctk.CTkLabel(
                row, text=label, font=ctk.CTkFont(size=12, weight="bold"), anchor="w",
                text_color=UI.get("text", "#111"),
            ).pack(fill="x", padx=8, pady=(6, 0))
            sub = last if last else (path[-40:] if len(path) > 40 else path)
            ctk.CTkLabel(
                row, text=sub, font=ctk.CTkFont(size=10),
                text_color=UI.get("text_dim", UI.get("text_muted", "#555")), anchor="w",
            ).pack(fill="x", padx=8, pady=(0, 4))
            btn_row = ctk.CTkFrame(row, fg_color="transparent")
            btn_row.pack(fill="x", padx=6, pady=(0, 6))
            ctk.CTkButton(
                btn_row, text="Carte", width=60, height=26, font=ctk.CTkFont(size=11),
                fg_color=UI.get("card", "#1c2620"), hover_color=UI.get("accent", "#4ec4b0"),
                command=lambda lab=label: filter_on_sortie(lab),
            ).pack(side="left", padx=(0, 4))
            ctk.CTkButton(
                btn_row, text="Ouvrir", height=26, font=ctk.CTkFont(size=11),
                fg_color=UI.get("accent", "#4ec4b0"), hover_color=UI.get("accent_hover", "#3aab98"),
                command=lambda p=path: open_one(p),
            ).pack(side="left", fill="x", expand=True)

        map_host = tk.Frame(body, bg=UI.get("card", "#1c2620"), highlightthickness=0)
        map_host.grid(row=0, column=1, sticky="nsew")
        map_w = TkinterMapView(map_host, corner_radius=0)
        map_w.pack(fill="both", expand=True)
        try:
            map_w.set_tile_server(
                "https://data.geopf.fr/wmts?SERVICE=WMTS&REQUEST=GetTile&VERSION=1.0.0"
                "&LAYER=GEOGRAPHICALGRIDSYSTEMS.PLANIGNV2&STYLE=normal&FORMAT=image/png"
                "&TILEMATRIXSET=PM&TILEMATRIX={z}&TILEROW={y}&TILECOL={x}",
                max_zoom=19,
            )
        except Exception:
            pass
        map_w.set_position(49.1627, 5.3854)
        map_w.set_zoom(9)

        def clear_markers():
            for m in state["markers"]:
                try:
                    m.delete()
                except Exception:
                    pass
            state["markers"] = []

        def filtered_points():
            pts = state["points"]
            mode = state["mode"]
            if mode == "masquer":
                return []
            if mode == "sortie":
                lab = sortie_var.get()
                if lab and lab != "Toutes":
                    pts = [p for p in pts if (p.get("_folder_label") or "") == lab]
            elif mode == "espece":
                esp = espece_var.get()
                if esp and esp != "Toutes":
                    pts = [
                        p for p in pts
                        if (p.get("espece") or "?").strip().lower() == esp.strip().lower()
                    ]
            return pts

        def redraw_markers():
            if not win.winfo_exists():
                return
            clear_markers()
            points = filtered_points()
            if state["mode"] == "masquer":
                lbl_status.configure(text="Points masqués")
                return

            clusters = {}
            for p in points:
                cle = (round(p["lat"], 4), round(p["lon"], 4))
                c = clusters.setdefault(cle, {
                    "lat": p["lat"], "lon": p["lon"], "obs": [],
                    "categories": collections.Counter(), "shapes": collections.Counter(),
                })
                c["obs"].append(p)
                c["categories"][p.get("categorie") or "Non classé"] += 1
                c["shapes"][self._detect_marker_shape(p.get("espece"))] += 1

            for c in clusters.values():
                categorie_dominante = c["categories"].most_common(1)[0][0]
                forme_dominante = "circle"
                for forme_candidat in ("triangle", "square", "diamond"):
                    if c["shapes"].get(forme_candidat, 0) > 0:
                        forme_dominante = forme_candidat
                        break
                obs_list = c["obs"]
                est_sp = sum(
                    1 for o in obs_list
                    if o.get("sans_photo") or str(o.get("fichier", "")).startswith("_")
                ) >= max(1, (len(obs_list) + 1) // 2)
                icon = self._get_marker_icon(
                    categorie_dominante, len(obs_list),
                    shape=forme_dominante, sans_photo=est_sp,
                )

                def on_marker(m, ol=obs_list):
                    self._show_cluster_info(ol)
                    folders_in = []
                    for o in ol:
                        fp = o.get("_folder_path")
                        if fp and fp not in folders_in:
                            folders_in.append(fp)
                    if len(folders_in) == 1:
                        if messagebox.askyesno(
                            "Ouvrir la sortie ?",
                            f"Ouvrir « {os.path.basename(folders_in[0])} » dans le carnet pour éditer ?"
                        ):
                            open_one(folders_in[0])
                    elif len(folders_in) > 1:
                        pick = ctk.CTkToplevel(win)
                        pick.title("Choisir la sortie")
                        pick.geometry("360x220")
                        ctk.CTkLabel(pick, text="Plusieurs sorties ici :").pack(pady=8)
                        for fp in folders_in:
                            ctk.CTkButton(
                                pick, text=os.path.basename(fp),
                                command=lambda p=fp: (pick.destroy(), open_one(p)),
                            ).pack(fill="x", padx=12, pady=3)

                mk = map_w.set_marker(
                    c["lat"], c["lon"], text=None, icon=icon, command=on_marker,
                )
                state["markers"].append(mk)

            lbl_status.configure(
                text=f"{len(points)} point(s) affiché(s) · {len(clusters)} marqueur(s)"
            )
            if points and len(clusters) > 1:
                try:
                    lats = [p["lat"] for p in points]
                    lons = [p["lon"] for p in points]
                    map_w.fit_bounding_box((max(lats), min(lons)), (min(lats), max(lons)))
                except Exception:
                    pass

        def _notes_from_archive_for_path(path):
            """Si le dossier photo est hors-ligne, tente le miroir archive_carnets."""
            try:
                root = self._local_archive_root()
                key = self._archive_key_for_folder(path)
                cand = os.path.join(root, key, NOTES_FILE)
                if os.path.isfile(cand):
                    with open(cand, "r", encoding="utf-8") as f:
                        return json.load(f) or {}
                # recherche par chemin d'origine dans meta.json
                for name in os.listdir(root):
                    d = os.path.join(root, name)
                    meta_p = os.path.join(d, "meta.json")
                    notes_p = os.path.join(d, NOTES_FILE)
                    if not os.path.isfile(notes_p):
                        continue
                    meta = {}
                    if os.path.isfile(meta_p):
                        try:
                            with open(meta_p, "r", encoding="utf-8") as f:
                                meta = json.load(f) or {}
                        except Exception:
                            pass
                    if os.path.normpath(meta.get("path_origine") or "") == os.path.normpath(path):
                        with open(notes_p, "r", encoding="utf-8") as f:
                            return json.load(f) or {}
            except Exception:
                pass
            return None

        def worker():
            exe_path = self.get_exiftool_path()
            all_points = []
            n_folders = len(folders)
            skipped = 0
            from_archive = 0
            for idx, entry in enumerate(folders):
                path = entry.get("path", "")
                label = entry.get("label") or os.path.basename(path) or path
                self.after(
                    0,
                    lambda i=idx, n=n_folders, lab=label:
                    lbl_status.configure(text="Chargement %d/%d — %s…" % (i + 1, n, lab)),
                )
                pts = []
                if os.path.isdir(path):
                    # 1) carnet JSON rapide
                    pts = self._gather_observations_for_folder(path, None, read_exif=False)
                    n_gps = sum(1 for p in pts if p.get("lat") is not None and p.get("lon") is not None)
                    # 2) si peu/pas de GPS → relire les fichiers photo (ExifTool)
                    if exe_path and (not pts or n_gps < max(1, len(pts) // 2)):
                        try:
                            pts = self._gather_observations_for_folder(
                                path, exe_path, read_exif=True
                            )
                        except Exception:
                            pass
                    # 3) enrichment global (meme logique que fiche espece)
                    try:
                        if hasattr(self, "_enrich_dataset_gps"):
                            pts = self._enrich_dataset_gps(pts)
                    except Exception:
                        pass
                else:
                    # Dossier inaccessible : archive locale PC
                    notes = _notes_from_archive_for_path(path)
                    if notes:
                        from_archive += 1
                        for fn, data in notes.items():
                            if not isinstance(data, dict):
                                continue
                            lat, lon = data.get("lat"), data.get("lon")
                            if lat is None or lon is None:
                                continue
                            try:
                                lat, lon = float(lat), float(lon)
                            except Exception:
                                continue
                            row = dict(data)
                            row["fichier"] = fn
                            row["lat"] = lat
                            row["lon"] = lon
                            row["dossier"] = label
                            row["dossier_path"] = path
                            pts.append(row)
                    else:
                        skipped += 1
                        continue

                for p in pts:
                    p["_folder_path"] = path
                    p["_folder_label"] = label
                all_points.extend(pts)

            all_points = [
                p for p in all_points
                if p.get("lat") is not None and p.get("lon") is not None
            ]
            win.after(
                0,
                lambda: finish_load(all_points, skipped=skipped, from_archive=from_archive),
            )

        def finish_load(points, skipped=0, from_archive=0):
            if not win.winfo_exists():
                return
            state["points"] = points
            # Remplir listes de filtres
            labels = sorted({
                (p.get("_folder_label") or "?").strip() for p in points
            } or {"Toutes"})
            especes = sorted({
                (p.get("espece") or "?").strip() for p in points if (p.get("espece") or "").strip()
            })
            menu_sortie.configure(values=["Toutes"] + labels)
            menu_espece.configure(values=["Toutes"] + especes if especes else ["Toutes"])
            sortie_var.set("Toutes")
            espece_var.set("Toutes")
            extra = []
            if from_archive:
                extra.append("%d via archive PC" % from_archive)
            if skipped:
                extra.append("%d dossier(s) inaccessible(s)" % skipped)
            msg = "%d observation(s) géolocalisée(s)" % len(points)
            if extra:
                msg += " · " + " · ".join(extra)
            if not points:
                msg = (
                    "0 point GPS — carnets sans coordonnées ou dossiers hors-ligne. "
                    "Ouvrez une sortie et laissez relire les GPS, ou branchez le disque."
                )
            lbl_status.configure(text=msg)
            redraw_markers()
            if points:
                try:
                    lats = [p["lat"] for p in points]
                    lons = [p["lon"] for p in points]
                    map_w.fit_bounding_box((max(lats), min(lons)), (min(lats), max(lons)))
                except Exception:
                    map_w.set_position(points[0]["lat"], points[0]["lon"])
                    map_w.set_zoom(11)

        threading.Thread(target=worker, daemon=True).start()

    def _parse_sortie_date(self, entry):
        """Date de sortie : nom du dossier en priorite (pas last_sync = date d'ouverture)."""
        if not isinstance(entry, dict):
            entry = {"path": str(entry or "")}

        for key in ("sortie_date", "date_sortie", "date"):
            raw = entry.get(key)
            if raw is None or raw == "":
                continue
            try:
                if hasattr(raw, "year"):
                    return raw
                s = str(raw).replace("Z", "").strip()
                if "T" in s:
                    return datetime.fromisoformat(s).date()
                return datetime.strptime(s[:10], "%Y-%m-%d").date()
            except Exception:
                pass

        label = entry.get("label") or os.path.basename(entry.get("path") or "")
        path_base = os.path.basename(entry.get("path") or "")
        for text_src in (label, path_base, entry.get("path") or ""):
            for pat, fmt in (
                (r"(20\d{2}-\d{2}-\d{2})", "%Y-%m-%d"),
                (r"(20\d{2}_\d{2}_\d{2})", "%Y_%m_%d"),
                (r"(20\d{6})", "%Y%m%d"),
                (r"(\d{2}-\d{2}-20\d{2})", "%d-%m-%Y"),
                (r"(\d{2}\.\d{2}\.20\d{2})", "%d.%m.%Y"),
                (r"(\d{2}/\d{2}/20\d{2})", "%d/%m/%Y"),
            ):
                mm = re.search(pat, str(text_src))
                if mm:
                    try:
                        return datetime.strptime(mm.group(1), fmt).date()
                    except Exception:
                        pass

        try:
            notes = os.path.join(entry.get("path") or "", NOTES_FILE)
            if os.path.isfile(notes):
                return datetime.fromtimestamp(os.path.getmtime(notes)).date()
        except Exception:
            pass

        ls = (entry.get("last_sync") or entry.get("last_opened") or "").strip()
        if ls:
            try:
                return datetime.fromisoformat(ls.replace("Z", "")).date()
            except Exception:
                try:
                    return datetime.strptime(ls[:10], "%Y-%m-%d").date()
                except Exception:
                    pass
        return None

    def _parse_user_date(self, text):
        """Parse une date saisie : AAAA-MM-JJ, JJ/MM/AAAA, JJ-MM-AAAA, AAAAMMJJ."""
        t = (text or "").strip()
        if not t:
            return None
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%Y%m%d"):
            try:
                return datetime.strptime(t, fmt).date()
            except Exception:
                pass
        return None

    def _filter_folders_by_period(self, folders, period, date_from=None, date_to=None):
        """Filtre les sorties : semaine | mois | annee | tout | custom (date_from/date_to)."""
        today = datetime.now().date()
        if period == "custom":
            debut = date_from
            fin = date_to or today
            if debut is None and fin is None:
                return list(folders)
            if debut is None:
                debut = datetime.min.date()
            if fin is None:
                fin = today
            if debut > fin:
                debut, fin = fin, debut
        elif period == "tout":
            return list(folders)
        elif period == "semaine":
            debut = today - timedelta(days=today.weekday())
            fin = today
        elif period == "mois":
            debut = today.replace(day=1)
            fin = today
        elif period == "annee":
            debut = today.replace(month=1, day=1)
            fin = today
        elif period == "7j":
            debut = today - timedelta(days=7)
            fin = today
        elif period == "30j":
            debut = today - timedelta(days=30)
            fin = today
        elif period == "90j":
            debut = today - timedelta(days=90)
            fin = today
        else:
            return list(folders)

        out = []
        for entry in folders:
            d = self._parse_sortie_date(entry)
            if d is not None and debut <= d <= fin:
                out.append(entry)
        return out

    def _build_archive_only_dataset(self, folder_paths=None):
        """Construit un dataset 100 % hors-ligne depuis archive_carnets (pas de photos, pas d'ExifTool)."""
        dataset = []
        root = self._local_archive_root()
        if not os.path.isdir(root):
            return dataset

        wanted = None
        if folder_paths:
            wanted = {os.path.normpath(p) for p in folder_paths}

        try:
            names = sorted(os.listdir(root))
        except Exception:
            return dataset

        for name in names:
            d = os.path.join(root, name)
            notes_p = os.path.join(d, NOTES_FILE)
            if not os.path.isfile(notes_p):
                continue
            meta = {}
            meta_p = os.path.join(d, "meta.json")
            if os.path.isfile(meta_p):
                try:
                    with open(meta_p, "r", encoding="utf-8") as f:
                        meta = json.load(f) or {}
                except Exception:
                    pass
            origin = os.path.normpath(meta.get("path_origine") or "")
            label = meta.get("label") or name.split("__")[0] or name
            if wanted is not None:
                if origin not in wanted and not any(
                    os.path.basename(os.path.normpath(p)) == label for p in (folder_paths or [])
                ):
                    # aussi accepter si le path voulu a la meme cle archive
                    match = False
                    for p in folder_paths or []:
                        try:
                            if self._archive_key_for_folder(p) == name:
                                match = True
                                break
                        except Exception:
                            pass
                    if not match:
                        continue
            try:
                with open(notes_p, "r", encoding="utf-8") as f:
                    notes = json.load(f) or {}
            except Exception:
                continue
            if not isinstance(notes, dict):
                continue
            for fn, data in notes.items():
                if not isinstance(data, dict):
                    continue
                row = dict(data)
                row["fichier"] = fn
                row["dossier"] = label
                row["dossier_path"] = origin or d
                row["chemin_complet"] = ""  # offline : pas de photo
                row["_from_archive"] = True
                dataset.append(row)
        return dataset

    def open_offline_text_debrief(self):
        """Débrief 100 % texte depuis l'archive PC — disque et photos non requis."""
        folders = self._load_known_folders()
        # Construire liste depuis archive (meme sans sorties_connues)
        archive_ds = self._build_archive_only_dataset(None)
        if not archive_ds and not folders:
            messagebox.showinfo(
                "Débrief texte",
                "Aucune archive locale.\n\n"
                "Branchez le disque une fois, puis :\n"
                "Plus… → Copier toutes les sorties accessibles → archive PC",
            )
            return

        win = ctk.CTkToplevel(self)
        win.title("Débrief texte — archive (sans photos)")
        win.geometry("620x720")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Débrief texte (archive, sans photos)",
            font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text="Mode 100 % hors-ligne : lit uniquement les carnets copiés sur le PC\n"
                 "(espèces, notes, météo, indices, GPS déjà enregistrés).\n"
                 "Aucune photo, aucun disque externe requis.",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            wraplength=580, justify="left",
        ).pack(anchor="w", padx=16, pady=(0, 8))

        # Filtres periode simples
        period_frame = ctk.CTkFrame(win, fg_color="transparent")
        period_frame.pack(fill="x", padx=16, pady=4)
        period_var = tk.StringVar(value="tout")
        period_labels = {
            "7j": "7 jours", "30j": "30 jours", "mois": "Ce mois",
            "annee": "Cette année", "tout": "Toutes",
        }

        checks = {}
        scroll = ctk.CTkScrollableFrame(win, height=360, fg_color=UI.get("card"))
        scroll.pack(fill="both", expand=True, padx=16, pady=6)

        # Index archive par path/label
        by_origin = {}
        for o in archive_ds:
            key = o.get("dossier_path") or o.get("dossier") or "?"
            by_origin.setdefault(key, []).append(o)

        # Liste sorties = union known folders + archive
        list_entries = []
        seen = set()
        for entry in folders:
            path = os.path.normpath(entry.get("path") or "")
            if path in seen:
                continue
            seen.add(path)
            list_entries.append(entry)
        for o in archive_ds:
            path = os.path.normpath(o.get("dossier_path") or "")
            lab = o.get("dossier") or ""
            if path and path not in seen:
                seen.add(path)
                list_entries.append({"path": path, "label": lab})
            elif not path and lab:
                k = "label:" + lab
                if k not in seen:
                    seen.add(k)
                    list_entries.append({"path": lab, "label": lab})

        def _sk(e):
            d = self._parse_sortie_date(e)
            return d.toordinal() if d else 0

        lbl_count = ctk.CTkLabel(win, text="", text_color=UI.get("text_dim"))
        lbl_count.pack(anchor="w", padx=16)

        def refresh_checks_for_period(p):
            period_var.set(p)
            filtered = self._filter_folders_by_period(list_entries, p) if p != "tout" else list_entries
            ok = {os.path.normpath(e.get("path") or "") for e in filtered}
            for path, var in checks.items():
                var.set(os.path.normpath(path) in ok if p != "tout" else True)
            n = sum(1 for v in checks.values() if v.get())
            lbl_count.configure(
                text="%d sortie(s) cochée(s) — filtre : %s  ·  archive = %d obs."
                % (n, period_labels.get(p, p), len(archive_ds))
            )

        for key, label in period_labels.items():
            ctk.CTkButton(
                period_frame, text=label, width=90, height=28,
                fg_color=UI.get("card_alt"),
                command=lambda k=key: refresh_checks_for_period(k),
            ).pack(side="left", padx=2)

        for entry in sorted(list_entries, key=_sk, reverse=True):
            path = entry.get("path") or ""
            label = entry.get("label") or os.path.basename(path) or path
            d = self._parse_sortie_date(entry)
            date_txt = d.strftime("%d/%m/%Y") if d else "?"
            n_arch = len(by_origin.get(os.path.normpath(path), []) or by_origin.get(path, []))
            # compter aussi par label
            if not n_arch:
                n_arch = sum(1 for o in archive_ds if (o.get("dossier") or "") == label)
            var = tk.BooleanVar(value=True)
            ctk.CTkCheckBox(
                scroll,
                text="%s  ·  %s  ·  %d obs. en archive" % (label, date_txt, n_arch),
                variable=var, text_color=UI.get("text"),
            ).pack(anchor="w", pady=2, padx=4)
            checks[path] = var

        refresh_checks_for_period("tout")

        def lancer():
            paths = [p for p, v in checks.items() if v.get()]
            if not paths:
                messagebox.showwarning("Débrief", "Cochez au moins une sortie.")
                return
            dataset = self._build_archive_only_dataset(paths)
            # si vide, essayer sans filtre path strict (labels)
            if not dataset:
                dataset = [
                    o for o in archive_ds
                    if (o.get("dossier_path") in paths)
                    or (o.get("dossier") in paths)
                    or any(
                        (o.get("dossier") or "") == (os.path.basename(p) or p)
                        for p in paths
                    )
                ]
            if not dataset:
                messagebox.showinfo(
                    "Débrief texte",
                    "Aucune observation en archive pour ces sorties.\n\n"
                    "Disque branché → Plus… → Copier toutes les sorties → archive PC",
                )
                return
            plab = period_labels.get(period_var.get(), "Archive hors-ligne")
            self._show_aggregated_prompt(
                dataset,
                period_label="%s — DÉBRIEF TEXTE ARCHIVE (sans photos)" % plab,
            )
            try:
                win.destroy()
            except Exception:
                pass

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=16, pady=12)
        ctk.CTkButton(
            bf, text="Générer le débrief texte",
            fg_color=UI.get("success", "#2f9e5f"), height=36, command=lancer,
        ).pack(side="left", padx=3)
        ctk.CTkButton(
            bf, text="Ouvrir dossier archive",
            command=lambda: os.startfile(self._local_archive_root()) if os.name == "nt" else None,
        ).pack(side="left", padx=3)
        ctk.CTkButton(bf, text="Fermer", command=win.destroy).pack(side="right", padx=3)

    def open_aggregated_report(self):
        """Choisir des sorties (semaine / mois / année) puis brief IA, PDF ou Word."""
        folders = self._load_known_folders()
        if not folders:
            messagebox.showinfo(
                "Aucune sortie",
                "Aucun dossier de sortie connu.\nChargez au moins un dossier de photos pour l'enregistrer."
            )
            return

        win = ctk.CTkToplevel(self)
        win.title("Comptes rendus multi-sorties — Brief IA")
        win.geometry("640x780")
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win,
            text="Synthèses sur plusieurs sorties : semaine, mois, année ou sélection manuelle.",
            font=ctk.CTkFont(size=12), text_color=UI.get("text_dim", "#888")
        ).pack(anchor="w", padx=15, pady=(15, 6))

        period_frame = ctk.CTkFrame(win, fg_color="transparent")
        period_frame.pack(fill="x", padx=15, pady=(0, 4))
        period_var = ctk.StringVar(value="tout")
        period_labels = {
            "7j": "7 jours",
            "30j": "30 jours",
            "semaine": "Cette semaine",
            "mois": "Ce mois",
            "annee": "Cette année",
            "tout": "Toutes",
            "custom": "Personnalisé",
        }
        checks = {}
        custom_from = [None]
        custom_to = [None]

        def apply_period(p, d_from=None, d_to=None):
            period_var.set(p)
            if p == "custom":
                d_from = d_from if d_from is not None else custom_from[0]
                d_to = d_to if d_to is not None else custom_to[0]
                custom_from[0], custom_to[0] = d_from, d_to
                filtered = self._filter_folders_by_period(folders, "custom", date_from=d_from, date_to=d_to)
                if d_from and d_to:
                    lab = f"Du {d_from.strftime('%d/%m/%Y')} au {d_to.strftime('%d/%m/%Y')}"
                elif d_from:
                    lab = f"Depuis le {d_from.strftime('%d/%m/%Y')}"
                elif d_to:
                    lab = f"Jusqu'au {d_to.strftime('%d/%m/%Y')}"
                else:
                    lab = "Personnalisé (dates invalides)"
            else:
                filtered = self._filter_folders_by_period(folders, p)
                lab = period_labels.get(p, p)
            paths_ok = {os.path.normpath(e["path"]) for e in filtered}
            for path, var in checks.items():
                var.set(True if p == "tout" else os.path.normpath(path) in paths_ok)
            n = sum(1 for v in checks.values() if v.get())
            lbl_count.configure(text=f"{n} sortie(s) cochée(s) — filtre : {lab}")

        for key, label in period_labels.items():
            if key == "custom":
                continue
            ctk.CTkButton(
                period_frame, text=label, width=100, height=28,
                fg_color=UI.get("card_alt", "#333"), hover_color=UI.get("accent", "#2b6cb0"),
                command=lambda k=key: apply_period(k)
            ).pack(side="left", padx=2)

        # Plage personnalisée
        custom_frame = ctk.CTkFrame(win, fg_color=UI.get("card", "#1a2332"), corner_radius=10)
        custom_frame.pack(fill="x", padx=15, pady=(4, 4))
        ctk.CTkLabel(
            custom_frame, text="Plage personnalisée (JJ/MM/AAAA ou AAAA-MM-JJ) :",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#888")
        ).pack(anchor="w", padx=10, pady=(8, 2))
        row_d = ctk.CTkFrame(custom_frame, fg_color="transparent")
        row_d.pack(fill="x", padx=10, pady=(0, 8))
        ctk.CTkLabel(row_d, text="Du").pack(side="left", padx=(0, 4))
        entry_from = ctk.CTkEntry(row_d, width=110, placeholder_text="01/01/2026")
        entry_from.pack(side="left", padx=(0, 10))
        ctk.CTkLabel(row_d, text="au").pack(side="left", padx=(0, 4))
        entry_to = ctk.CTkEntry(row_d, width=110, placeholder_text="30/07/2026")
        entry_to.pack(side="left", padx=(0, 10))

        def apply_custom():
            d0 = self._parse_user_date(entry_from.get())
            d1 = self._parse_user_date(entry_to.get())
            if not d0 and not d1:
                messagebox.showwarning(
                    "Dates",
                    "Indiquez au moins une date (début ou fin).\nFormats : JJ/MM/AAAA ou AAAA-MM-JJ"
                )
                return
            apply_period("custom", d_from=d0, d_to=d1)

        ctk.CTkButton(
            row_d, text="Appliquer", width=90, height=28,
            fg_color=UI.get("accent", "#2b6cb0"), hover_color=UI.get("accent_hover", "#1f4f80"),
            command=apply_custom
        ).pack(side="left", padx=4)

        lbl_count = ctk.CTkLabel(win, text="", font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#888"))
        lbl_count.pack(anchor="w", padx=15, pady=(2, 4))

        ctk.CTkLabel(win, text="Sorties à inclure :", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=15, pady=(6, 2))
        scroll = ctk.CTkScrollableFrame(win, height=340)
        scroll.pack(fill="both", expand=True, padx=15, pady=5)

        def _sort_key(e):
            d = self._parse_sortie_date(e)
            return d.toordinal() if d else 0

        for entry in sorted(folders, key=_sort_key, reverse=True):
            var = ctk.BooleanVar(value=True)
            d = self._parse_sortie_date(entry)
            date_txt = d.strftime("%d/%m/%Y") if d else "?"
            ctk.CTkCheckBox(
                scroll, text=f"{entry.get('label')}   ·  {date_txt}", variable=var,
                command=lambda: lbl_count.configure(
                    text=f"{sum(1 for v in checks.values() if v.get())} sortie(s) cochée(s) — filtre : {period_labels.get(period_var.get(), period_var.get())}"
                )
            ).pack(anchor="w", pady=3)
            checks[entry["path"]] = var

        lbl_count.configure(text=f"{len(checks)} sortie(s) cochée(s) — filtre : Toutes")

        btn_frame = ctk.CTkFrame(win, fg_color="transparent")
        btn_frame.pack(fill="x", padx=15, pady=12)

        def get_selected_paths():
            return [p for p, v in checks.items() if v.get()]

        def current_period_label():
            p = period_var.get()
            if p == "custom":
                d0, d1 = custom_from[0], custom_to[0]
                if d0 and d1:
                    return f"Du {d0.strftime('%d/%m/%Y')} au {d1.strftime('%d/%m/%Y')}"
                if d0:
                    return f"Depuis le {d0.strftime('%d/%m/%Y')}"
                if d1:
                    return f"Jusqu'au {d1.strftime('%d/%m/%Y')}"
                return "Plage personnalisée"
            return period_labels.get(p, "Période sélectionnée")

        def do_prompt():
            paths = get_selected_paths()
            if not paths:
                messagebox.showwarning("Aucune sélection", "Cochez au moins une sortie.")
                return
            try:
                win.configure(cursor="watch")
                win.update_idletasks()
                # Mode rapide : pas d'ExifTool (évite le gel de plusieurs minutes)
                dataset = self._build_report_dataset(paths, read_exif=False)
            finally:
                try:
                    win.configure(cursor="")
                except Exception:
                    pass
            if not dataset:
                messagebox.showinfo("Aucune donnée", "Aucune observation dans les sorties sélectionnées.")
                return
            self._show_aggregated_prompt(dataset, period_label=current_period_label())

        def do_pdf():
            paths = get_selected_paths()
            if not paths:
                messagebox.showwarning("Aucune sélection", "Cochez au moins une sortie.")
                return
            if hasattr(self, "_ask_save_report"):
                chemin = self._ask_save_report(".pdf", [("PDF", "*.pdf")], f"rapport_{period_var.get()}.pdf")
            else:
                chemin = filedialog.asksaveasfilename(
                    defaultextension=".pdf", filetypes=[("PDF", "*.pdf")],
                    initialfile=f"rapport_{period_var.get()}.pdf",
                    initialdir=self._reports_dir() if hasattr(self, "_reports_dir") else None,
                )
            if not chemin:
                return
            try:
                win.configure(cursor="watch")
                win.update_idletasks()
                dataset = self._build_report_dataset(paths, read_exif=False)
            finally:
                try:
                    win.configure(cursor="")
                except Exception:
                    pass
            if not dataset:
                messagebox.showinfo("Aucune donnée", "Aucune observation dans les sorties sélectionnées.")
                return
            if self.export_report_pdf(dataset, f"Rapport naturaliste — {current_period_label()}", chemin):
                messagebox.showinfo(
                    "Export réussi",
                    f"PDF enregistré :\n{chemin}\n\n"
                    "Emplacement choisi dans la boîte de dialogue "
                    "(dossier des rapports si configuré, sinon dossier photos)."
                )

        def do_docx():
            paths = get_selected_paths()
            if not paths:
                messagebox.showwarning("Aucune sélection", "Cochez au moins une sortie.")
                return
            if hasattr(self, "_ask_save_report"):
                chemin = self._ask_save_report(".docx", [("Word", "*.docx")], f"rapport_{period_var.get()}.docx")
            else:
                chemin = filedialog.asksaveasfilename(
                    defaultextension=".docx", filetypes=[("Word", "*.docx")],
                    initialfile=f"rapport_{period_var.get()}.docx",
                    initialdir=self._reports_dir() if hasattr(self, "_reports_dir") else None,
                )
            if not chemin:
                return
            try:
                win.configure(cursor="watch")
                win.update_idletasks()
                dataset = self._build_report_dataset(paths, read_exif=False)
            finally:
                try:
                    win.configure(cursor="")
                except Exception:
                    pass
            if not dataset:
                messagebox.showinfo("Aucune donnée", "Aucune observation dans les sorties sélectionnées.")
                return
            if self.export_report_docx(dataset, f"Rapport naturaliste — {current_period_label()}", chemin):
                messagebox.showinfo(
                    "Export réussi",
                    f"Word enregistré :\n{chemin}\n\n"
                    "Emplacement choisi dans la boîte de dialogue "
                    "(dossier des rapports si configuré)."
                )

        ctk.CTkButton(btn_frame, text="Brief IA", command=do_prompt, fg_color=UI.get("purple", "#6f42c1"), hover_color=UI.get("purple_hover", "#59339d"), height=36).pack(side="left", padx=(0, 6))
        ctk.CTkButton(
            btn_frame, text="Débrief texte (archive, sans photos)",
            command=lambda: (win.destroy(), self.open_offline_text_debrief()),
            fg_color=UI.get("success", "#2f9e5f"), hover_color=UI.get("success_hover", "#278f50"),
            height=36,
        ).pack(side="left", padx=3)
        ctk.CTkButton(btn_frame, text="PDF", command=do_pdf, fg_color="#b03a2e", hover_color="#8a2d24", height=36, width=80).pack(side="left", padx=3)
        ctk.CTkButton(btn_frame, text="Word", command=do_docx, fg_color=UI.get("accent", "#2b6cb0"), hover_color=UI.get("accent_hover", "#1f4f80"), height=36, width=80).pack(side="left", padx=3)
        ctk.CTkButton(btn_frame, text="Fermer", command=win.destroy, width=80).pack(side="right")

    def _show_aggregated_prompt(self, dataset, period_label="Période sélectionnée"):
        """Prompt IA de synthese multi-sorties (semaine / mois / annee / selection)."""
        if not dataset:
            messagebox.showinfo("Aucune donnée", "Aucune observation à synthétiser.")
            return

        sorties = sorted({o.get("dossier") or "?" for o in dataset})
        n_sorties = len(sorties)
        esp_counter = collections.Counter()
        cat_counter = collections.Counter()
        for o in dataset:
            esp = (o.get("espece") or "").strip()
            if esp:
                try:
                    n = int(str(o.get("nombre") or "1").split("-")[0].replace("+", "").strip() or "1")
                except Exception:
                    n = 1
                esp_counter[esp] += max(1, n)
            cat_counter[o.get("categorie") or "Non classé"] += 1

        top_especes = ", ".join(f"{e} ({c})" for e, c in esp_counter.most_common(12)) or "—"
        top_cats = ", ".join(f"{c} ({n})" for c, n in cat_counter.most_common()) or "—"

        offline = any(o.get("_from_archive") for o in dataset) or (
            "ARCHIVE" in (period_label or "").upper() or "sans photos" in (period_label or "").lower()
        )
        # Synthese meteo globale si presente dans les notes
        meteo_bits = []
        for o in dataset:
            m = o.get("meteo") or {}
            if not isinstance(m, dict):
                continue
            if m.get("temperature") is not None or m.get("ciel"):
                bit = "%s°C" % m.get("temperature") if m.get("temperature") is not None else ""
                if m.get("ciel"):
                    bit = (bit + " " + str(m.get("ciel"))).strip()
                if m.get("phase_lunaire"):
                    bit += " · " + str(m.get("phase_lunaire"))
                if m.get("date") or o.get("heure"):
                    bit = "%s : %s" % (m.get("date") or o.get("heure") or "?", bit)
                if bit.strip(" :"):
                    meteo_bits.append(bit.strip())
        meteo_bits = list(dict.fromkeys(meteo_bits))[:20]
        meteo_resume = " ; ".join(meteo_bits) if meteo_bits else "non renseignée dans les carnets"

        prompt_lines = [
            "Tu es un expert naturaliste. Voici mes notes de terrain agrégées sur plusieurs sorties dans le département de la Meuse (55).",
            f"Période demandée : {period_label}.",
            f"Nombre de sorties : {n_sorties}. Nombre d'observations : {len(dataset)}.",
            f"Espèces les plus fréquentes (effectifs cumulés indicatifs) : {top_especes}.",
            f"Répartition par catégorie : {top_cats}.",
            f"Météo / conditions notées : {meteo_resume}.",
            "",
        ]
        if offline:
            prompt_lines += [
                "MODE DÉBRIEF TEXTE HORS-LIGNE : les données viennent uniquement des carnets JSON archivés.",
                "Aucune photo n'est disponible. Ne demande pas d'images. Appuie-toi sur notes, espèces, indices, lieux, GPS et météo.",
                "",
            ]
        prompt_lines += [
            "Rédige un compte-rendu professionnel, structuré et vivant :",
            "1) Introduction (contexte Meuse, periode, effort d'observation, météo globale si dispo)",
            "2) Synthèse par sortie ou par semaine selon le volume (date / dossier, faits marquants)",
            "3) Bilan par groupe taxonomique (mammifères, oiseaux, insectes, indices de presence)",
            "4) Tendances sur la periode (espèces récurrentes, lieux privilégiés, indices vs contacts visuels)",
            "5) Conclusion courte (biodiversité perçue, pistes pour les prochaines sorties)",
            "Distingue clairement animaux vus, entendus, et indices (empreintes, terriers, coulées, latrines).",
            "Intègre les notes météo (température, ciel, humidité, phase lunaire, pluie 3j) quand elles sont présentes.",
            "",
            "=== SORTIES INCLUSES ===",
        ]
        for s in sorties:
            prompt_lines.append(f"  • {s}")
        prompt_lines.append("\n=== DONNÉES BRUTES ===")

        by_dossier = {}
        for obs in dataset:
            by_dossier.setdefault(obs.get("dossier") or "?", []).append(obs)

        for dossier in sorted(by_dossier.keys()):
            prompt_lines.append(f"\n--- Sortie : {dossier} ({len(by_dossier[dossier])} obs.) ---")
            for obs in by_dossier[dossier]:
                ligne = (
                    f"- {obs.get('fichier') or 'obs'} | Heure : {obs.get('heure', '--:--')} | "
                    f"{obs.get('categorie')} | {obs.get('espece')} | nb={obs.get('nombre')}"
                )
                if obs.get("type_observation"):
                    ligne += f" | Type : {obs.get('type_observation')}"
                if obs.get("lieu"):
                    ligne += f" | Lieu : {obs.get('lieu')}"
                if obs.get("lat") and obs.get("lon"):
                    try:
                        ligne += f" | GPS : {float(obs['lat']):.5f},{float(obs['lon']):.5f}"
                    except Exception:
                        pass
                meteo = obs.get("meteo") or {}
                if isinstance(meteo, dict) and (
                    meteo.get("temperature") is not None
                    or meteo.get("ciel")
                    or meteo.get("phase_lunaire")
                    or meteo.get("pluie_3j_precedents") is not None
                ):
                    parts_m = []
                    if meteo.get("temperature") is not None:
                        parts_m.append("%s°C" % meteo.get("temperature"))
                    if meteo.get("humidite") is not None:
                        parts_m.append("%s%% hum." % meteo.get("humidite"))
                    if meteo.get("ciel"):
                        parts_m.append(str(meteo.get("ciel")))
                    if meteo.get("phase_lunaire"):
                        parts_m.append(str(meteo.get("phase_lunaire")))
                    if meteo.get("pluie_3j_precedents") is not None:
                        parts_m.append("pluie 3j=%s mm" % meteo.get("pluie_3j_precedents"))
                    if parts_m:
                        ligne += " | Meteo : " + ", ".join(parts_m)
                if obs.get("notes_libres"):
                    ligne += f" | Notes : {obs.get('notes_libres')}"
                if obs.get("sans_photo"):
                    ligne += " | [sans photo]"
                prompt_lines.append(ligne)

        prompt_lines.append("\n=================================")
        prompt_lines.append(
            "Rédige la synthèse en français, sans inventer d'espèces absentes des donnees. "
            "Si certaines sorties sont pauvres, signale-le brièvement."
        )
        full_prompt = "\n".join(prompt_lines)

        ai_window = ctk.CTkToplevel(self)
        ai_window.title(f"Brief IA — {period_label}")
        ai_window.geometry("900x680")
        self._prepare_tool_window(ai_window)
        ctk.CTkLabel(
            ai_window,
            text=f"{period_label}  ·  {n_sorties} sortie(s)  ·  {len(dataset)} observation(s)\n"
                 "Copie le texte et colle-le dans Gemini / ChatGPT :",
            font=ctk.CTkFont(weight="bold"), justify="left"
        ).pack(padx=15, pady=10, anchor="w")
        text_area = ctk.CTkTextbox(ai_window, activate_scrollbars=True)
        text_area.pack(fill="both", expand=True, padx=15, pady=5)
        text_area.insert("0.0", full_prompt)

        def copy_to_clipboard():
            self.clipboard_clear()
            self.clipboard_append(full_prompt)
            messagebox.showinfo("Copié !", "Le prompt multi-sorties a été copié dans le presse-papiers.")

        btn_row = ctk.CTkFrame(ai_window, fg_color="transparent")
        btn_row.pack(fill="x", padx=15, pady=12)
        ctk.CTkButton(
            btn_row, text="Copier le prompt pour l'IA",
            fg_color=UI.get("purple", "#6f42c1"), hover_color=UI.get("purple_hover", "#59339d"),
            command=copy_to_clipboard, height=36
        ).pack(side="right")
        self.log(f"Brief IA multi-sorties ({period_label}) : {len(dataset)} obs., {n_sorties} sortie(s).")


    # --- LOT 3 : Identification d'espèce assistée par IA (Gemini, niveau gratuit Google) ---

    def _ia_config_path(self):
        try:
            self._migrate_legacy_config_files()
        except Exception:
            pass
        return os.path.join(self._user_data_dir(), "ia_config.json")

    def _load_ia_config(self):
        path = self._ia_config_path()
        if not os.path.exists(path):
            return {}
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}

    def _save_ia_config(self, config):
        self.ia_config = config
        try:
            with open(self._ia_config_path(), "w", encoding="utf-8") as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    # --- LOT 5 : Prévisions météo à venir et rappel des phases lunaires ---

    def _get_reference_location(self):
        """Détermine une position de référence pour les prévisions : moyenne du dossier actif, sinon Verdun par défaut."""
        coords = [(i["lat"], i["lon"]) for i in self.photos_data.values() if i.get("lat") and i.get("lon")]
        if coords:
            lat = sum(c[0] for c in coords) / len(coords)
            lon = sum(c[1] for c in coords) / len(coords)
            label = os.path.basename(self.photo_folder_path) if self.photo_folder_path else "dossier actif"
            return lat, lon, f"autour de vos dernières photos ({label})"
        return 49.1627, 5.3854, "Verdun, Meuse (position par défaut)"

    def _fetch_weather_forecast(self, lat, lon):
        """Récupère les prévisions à 7 jours (Open-Meteo, gratuit, sans clé) et calcule un score de favorabilité."""
        url = (
            f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}"
            f"&daily=temperature_2m_max,temperature_2m_min,precipitation_sum,wind_speed_10m_max,cloud_cover_mean,weather_code"
            f"&timezone=auto&forecast_days=7"
        )
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        data = resp.json().get("daily", {})
        dates = data.get("time", [])

        jours = []
        for i, date_str in enumerate(dates):
            pluie = (data.get("precipitation_sum") or [None] * len(dates))[i]
            vent = (data.get("wind_speed_10m_max") or [None] * len(dates))[i]
            nuages = (data.get("cloud_cover_mean") or [None] * len(dates))[i]
            tmax = (data.get("temperature_2m_max") or [None] * len(dates))[i]
            tmin = (data.get("temperature_2m_min") or [None] * len(dates))[i]

            # Score simple et transparent : peu de pluie + vent calme = plus favorable à un affût
            score = 0.0
            if pluie is not None:
                score += max(0.0, 3.0 - pluie)
            if vent is not None:
                score += max(0.0, 3.0 - vent / 10.0)

            date_obj = datetime.strptime(date_str, "%Y-%m-%d")
            jours.append({
                "date": date_str,
                "jour_semaine": JOURS_FR[date_obj.weekday()],
                "tmin": tmin, "tmax": tmax, "pluie": pluie, "vent": vent, "nuages": nuages,
                "phase_lune": self.moon_phase(date_obj),
                "score": round(score, 1),
            })
        return jours

    def _next_moon_dates(self, horizon_days=45):
        """Cherche la prochaine pleine lune et la prochaine nouvelle lune dans les jours à venir."""
        aujourd_hui = datetime.now()
        prochaine_pleine, prochaine_nouvelle = None, None
        for i in range(horizon_days):
            d = aujourd_hui + timedelta(days=i)
            phase = self.moon_phase(d)
            if "Pleine" in phase and prochaine_pleine is None:
                prochaine_pleine = d
            if "Nouvelle" in phase and prochaine_nouvelle is None:
                prochaine_nouvelle = d
            if prochaine_pleine and prochaine_nouvelle:
                break
        return prochaine_pleine, prochaine_nouvelle

    def open_forecast_view(self):
        """Fenêtre affichant les prévisions à 7 jours et les prochaines phases lunaires."""
        win = ctk.CTkToplevel(self)
        win.title("🔮 Prévisions & Phases lunaires")
        win.geometry("650x700")
        self._prepare_tool_window(win)

        lat, lon, lieu_label = self._get_reference_location()
        ctk.CTkLabel(win, text=f"📍 Zone de référence : {lieu_label}", font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#666666")).pack(anchor="w", padx=15, pady=(15, 0))

        lbl_status = ctk.CTkLabel(win, text="⏳ Récupération des prévisions...")
        lbl_status.pack(pady=20)

        def worker():
            try:
                jours = self._fetch_weather_forecast(lat, lon)
                erreur = None
            except Exception as e:
                jours, erreur = [], str(e)
            self.after(0, lambda: self._render_forecast(win, lbl_status, jours, erreur))

        threading.Thread(target=worker, daemon=True).start()

    def _render_forecast(self, win, lbl_status, jours, erreur):
        if not win.winfo_exists():
            return
        lbl_status.destroy()

        if erreur or not jours:
            ctk.CTkLabel(
                win, text=f"❌ Impossible de récupérer les prévisions.\n{erreur or ''}",
                text_color="#e74c3c", wraplength=580, justify="left"
            ).pack(pady=20, padx=15)
            return

        prochaine_pleine, prochaine_nouvelle = self._next_moon_dates()
        lune_frame = ctk.CTkFrame(win, fg_color="#2b2b2b", corner_radius=8)
        lune_frame.pack(fill="x", padx=15, pady=(10, 10))
        morceaux_lune = []
        if prochaine_pleine:
            morceaux_lune.append(f"🌕 Prochaine pleine lune : {prochaine_pleine.strftime('%d/%m/%Y')}")
        if prochaine_nouvelle:
            morceaux_lune.append(f"🌑 Prochaine nouvelle lune : {prochaine_nouvelle.strftime('%d/%m/%Y')}")
        ctk.CTkLabel(lune_frame, text="   ·   ".join(morceaux_lune), font=ctk.CTkFont(size=12)).pack(padx=10, pady=8)

        ctk.CTkLabel(
            win, text="⭐ = jours les plus favorables (peu de pluie, vent calme) — repère indicatif, pas une garantie d'observation.",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim", "#666666"), wraplength=600, justify="left"
        ).pack(anchor="w", padx=15, pady=(0, 8))

        meilleur_score = max((j["score"] for j in jours), default=0)

        scroll = ctk.CTkScrollableFrame(win)
        scroll.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        for j in jours:
            est_top = meilleur_score > 0 and j["score"] >= meilleur_score - 0.5
            couleur = "#1f7d37" if est_top else "#2b2b2b"
            card = ctk.CTkFrame(scroll, fg_color=couleur, corner_radius=8)
            card.pack(fill="x", pady=4)

            etoile = "⭐ " if est_top else ""
            ctk.CTkLabel(card, text=f"{etoile}{j['jour_semaine']} {j['date']}", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=12, pady=(8, 2))

            tmin = f"{j['tmin']}" if j['tmin'] is not None else "?"
            tmax = f"{j['tmax']}" if j['tmax'] is not None else "?"
            pluie = f"{j['pluie']}" if j['pluie'] is not None else "?"
            vent = f"{j['vent']}" if j['vent'] is not None else "?"
            nuages = f"{j['nuages']}" if j['nuages'] is not None else "?"
            details = f"🌡️ {tmin}–{tmax}°C   💧 {pluie}mm   💨 {vent}km/h   ☁️ {nuages}%   {j['phase_lune']}"
            ctk.CTkLabel(card, text=details, font=ctk.CTkFont(size=11), text_color="#dddddd").pack(anchor="w", padx=12, pady=(0, 8))

    def open_ai_settings(self):
        """Configuration reconnaissance d'espèces : Gemini et/ou xAI (Grok)."""
        cfg = dict(self.ia_config or {})

        win = ctk.CTkToplevel(self)
        win.title("Reconnaissance d'espèces IA")
        win.geometry("560x560")
        win.configure(fg_color=UI.get("bg", "#0f1419"))
        self._prepare_tool_window(win)

        scroll = ctk.CTkScrollableFrame(win, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=12, pady=12)

        ctk.CTkLabel(
            scroll, text="Fournisseur pour le bouton 🤖",
            font=ctk.CTkFont(size=15, weight="bold")
        ).pack(anchor="w", pady=(0, 6))

        providers = ["gemini", "xai", "auto"]
        labels = {
            "gemini": "Google Gemini (vision)",
            "xai": "xAI Grok (vision)",
            "auto": "Auto (Grok si clé xAI, sinon Gemini)",
        }
        provider = cfg.get("ai_provider") or "gemini"
        if provider not in providers:
            provider = "gemini"
        choice_provider = ctk.CTkOptionMenu(
            scroll, values=providers, height=32,
            command=lambda v: lbl_prov.configure(text=labels.get(v, v)),
        )
        choice_provider.pack(fill="x", pady=(0, 2))
        choice_provider.set(provider)
        lbl_prov = ctk.CTkLabel(scroll, text=labels.get(provider, ""), font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#888"))
        lbl_prov.pack(anchor="w", pady=(0, 12))

        # --- Gemini ---
        card_g = ctk.CTkFrame(scroll, fg_color=UI.get("card", "#1a2332"), corner_radius=12)
        card_g.pack(fill="x", pady=6)
        ctk.CTkLabel(card_g, text="Google Gemini", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=12, pady=(12, 2))
        ctk.CTkLabel(
            card_g, text="Clé gratuite : aistudio.google.com/apikey",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#888")
        ).pack(anchor="w", padx=12)
        entry_gemini = ctk.CTkEntry(card_g, placeholder_text="Clé API Gemini", show="•", height=32)
        entry_gemini.pack(fill="x", padx=12, pady=6)
        entry_gemini.insert(0, cfg.get("gemini_api_key", ""))
        models_g = ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-flash-lite-latest", "gemini-1.5-flash-latest"]
        mg = cfg.get("gemini_model") or models_g[0]
        if mg not in models_g:
            models_g = [mg] + models_g
        choice_g_model = ctk.CTkOptionMenu(card_g, values=models_g, height=30)
        choice_g_model.pack(fill="x", padx=12, pady=(0, 12))
        choice_g_model.set(mg)

        # --- xAI Grok ---
        card_x = ctk.CTkFrame(scroll, fg_color=UI.get("card", "#1a2332"), corner_radius=12)
        card_x.pack(fill="x", pady=6)
        ctk.CTkLabel(card_x, text="xAI — Grok", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=12, pady=(12, 2))
        ctk.CTkLabel(
            card_x,
            text="Clé API : console.x.ai  →  Create API key\n"
                 "Endpoint : https://api.x.ai/v1 (compatible OpenAI)",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#888"), justify="left"
        ).pack(anchor="w", padx=12)
        entry_xai = ctk.CTkEntry(card_x, placeholder_text="Clé API xAI (xai-...)", show="•", height=32)
        entry_xai.pack(fill="x", padx=12, pady=6)
        entry_xai.insert(0, cfg.get("xai_api_key", ""))
        models_x = [
            "grok-2-vision-1212",
            "grok-2-latest",
            "grok-3",
            "grok-4",
        ]
        mx = cfg.get("xai_model") or models_x[0]
        if mx not in models_x:
            models_x = [mx] + models_x
        choice_x_model = ctk.CTkOptionMenu(card_x, values=models_x, height=30)
        choice_x_model.pack(fill="x", padx=12, pady=(0, 12))
        choice_x_model.set(mx)

        auto_var = ctk.BooleanVar(value=bool(cfg.get("gemini_auto_apply", False)))
        ctk.CTkCheckBox(
            scroll, text="Appliquer automatiquement catégorie + espèce (sans dialogue)",
            variable=auto_var, font=ctk.CTkFont(size=12)
        ).pack(anchor="w", pady=10)

        lbl_status = ctk.CTkLabel(scroll, text="", font=ctk.CTkFont(size=11))
        lbl_status.pack(anchor="w")

        def refresh_status():
            parts = []
            if entry_gemini.get().strip():
                parts.append("Gemini OK")
            if entry_xai.get().strip():
                parts.append("xAI/Grok OK")
            lbl_status.configure(
                text=" · ".join(parts) if parts else "Aucune clé — bouton 🤖 inactif",
                text_color=UI.get("success", "#2ecc71") if parts else UI.get("text_dim", "#888"),
            )

        refresh_status()

        def save():
            new_cfg = dict(self.ia_config or {})
            new_cfg["ai_provider"] = choice_provider.get()
            new_cfg["gemini_api_key"] = entry_gemini.get().strip()
            new_cfg["gemini_model"] = choice_g_model.get()
            new_cfg["xai_api_key"] = entry_xai.get().strip()
            new_cfg["xai_model"] = choice_x_model.get()
            new_cfg["gemini_auto_apply"] = bool(auto_var.get())
            if not new_cfg["gemini_api_key"] and not new_cfg["xai_api_key"]:
                messagebox.showwarning("Clé manquante", "Renseignez au moins une clé (Gemini ou xAI).")
                return
            self._save_ia_config(new_cfg)
            refresh_status()
            self.log(f"IA configurée — fournisseur : {new_cfg['ai_provider']}")
            messagebox.showinfo(
                "Enregistré",
                f"Fournisseur : {new_cfg['ai_provider']}\n\n"
                "Dans le carnet : sélectionnez une photo → 🤖"
            )

        ctk.CTkButton(
            scroll, text="Enregistrer", command=save,
            fg_color=UI.get("purple", "#9b59b6"), hover_color=UI.get("purple_hover", "#8e44ad"),
            height=36
        ).pack(fill="x", pady=(8, 4))
        ctk.CTkButton(scroll, text="Fermer", command=win.destroy, height=32).pack(fill="x", pady=(0, 8))

    def _prepare_image_for_ai(self, path, max_dim=768):
        """Miniature JPEG pour l'API vision (RAW via ExifTool si besoin)."""
        if not path or not os.path.isfile(path):
            return None
        if self._is_video_file(os.path.basename(path)):
            return None
        exe = self.get_exiftool_path()
        data = self._get_thumbnail_bytes(path, exe, max_dim=max_dim)
        if data:
            return data
        # Repli : ouverture PIL directe (JPEG/PNG)
        try:
            img = Image.open(path)
            img = img.convert("RGB")
            img.thumbnail((max_dim, max_dim))
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            return buf.getvalue()
        except Exception:
            return None

    def _gemini_suggest(self, image_bytes):
        """Appel Gemini vision → dict categorie / espece / confiance / details."""
        api_key = (self.ia_config or {}).get("gemini_api_key")
        if not api_key:
            raise RuntimeError(
                "Aucune clé Gemini. Outils → Identification IA pour en ajouter une (gratuite)."
            )

        b64 = base64.b64encode(image_bytes).decode("ascii")
        liste_connues = []
        for cat, especes in (self.faune_meuse or {}).items():
            for e in especes[:20]:
                liste_connues.append(f"{e} ({cat})")
        known_txt = ", ".join(liste_connues[:80])

        prompt = (
            "Rôle : naturaliste de terrain, expert faune/flore du Grand Est (Meuse, 55, France).\n"
            "Tâche : identifier le sujet PRINCIPAL de cette photo (contact visuel, indice, ou espèce non animale).\n"
            "\n"
            "Contexte géographique : forêts, lisières, cultures et zones humides de Meuse. "
            "Privilégier les espèces réalistes pour cette région (pas d'espèces exotiques hors contexte).\n"
            "\n"
            "Règles :\n"
            "1) Un seul sujet principal. Si doute entre 2 espèces proches, choisir la plus probable et confiance moyenne/faible.\n"
            "2) Empreintes, crottes, terriers, coulées, restes → type_observation adapté, pas forcément 'Vu'.\n"
            "3) Noms français usuels (ex. Chevreuil d'Europe, Sanglier, Renard roux, Buse variable).\n"
            "4) Interdit d'inventer une espèce peu plausible. Si flou/nuit/partiel → espece='Inconnu', confiance='faible'.\n"
            "5) Sortie STRICTEMENT JSON, sans markdown, sans texte avant/après.\n"
            "\n"
            "CONFIDENCE (strict — lire entièrement) :\n"
            f"{IA_CONFIANCE_RUBRIC_FR}\n"
            "\n"
            "Schéma JSON obligatoire :\n"
            '{"categorie":"Mammifère|Oiseau|Insecte|Autre",'
            '"espece":"nom français",'
            '"nom_scientifique":"Genre species ou \"\"",'
            '"type_observation":"Vu|Empreinte / Trace|Terrier|Coulée / Passage|Autre indice",'
            '"confiance":"faible|moyenne|élevée",'
            '"commentaire":"1 phrase factuelle (indices visibles)"}\n'
            "\n"
            f"Liste d'espèces déjà notées dans ce carnet (aide, non exclusive) : {known_txt}."
        )

        model = (self.ia_config or {}).get("gemini_model") or "gemini-2.0-flash"
        models_try = [model]
        for m in ("gemini-2.0-flash", "gemini-1.5-flash", "gemini-flash-lite-latest", "gemini-1.5-flash-latest"):
            if m not in models_try:
                models_try.append(m)

        last_err = None
        for model_name in models_try:
            url = (
                f"https://generativelanguage.googleapis.com/v1beta/models/"
                f"{model_name}:generateContent?key={api_key}"
            )
            body = {
                "contents": [{
                    "parts": [
                        {"text": prompt},
                        {"inline_data": {"mime_type": "image/jpeg", "data": b64}},
                    ]
                }],
                "generationConfig": {
                    "temperature": 0.2,
                    "maxOutputTokens": 512,
                },
            }
            try:
                resp = requests.post(url, json=body, timeout=40)
                if resp.status_code == 404:
                    last_err = f"Modèle {model_name} indisponible"
                    continue
                resp.raise_for_status()
                payload = resp.json()
                texte = payload["candidates"][0]["content"]["parts"][0]["text"].strip()
                texte = texte.replace("```json", "").replace("```", "").strip()
                # Extraire l'objet JSON même s'il y a du texte autour
                m = re.search(r"\{[\s\S]*\}", texte)
                if m:
                    texte = m.group(0)
                data = json.loads(texte)
                data["_model"] = model_name
                return data
            except Exception as e:
                last_err = str(e)
                continue
        raise RuntimeError(last_err or "Échec de l'appel Gemini")

    def _resolve_ai_provider(self):
        """Retourne 'xai' ou 'gemini' selon config et clés disponibles."""
        cfg = self.ia_config or {}
        pref = (cfg.get("ai_provider") or "gemini").lower()
        has_xai = bool(cfg.get("xai_api_key"))
        has_gem = bool(cfg.get("gemini_api_key"))
        if pref == "xai":
            if has_xai:
                return "xai"
            if has_gem:
                return "gemini"
            return None
        if pref == "auto":
            if has_xai:
                return "xai"
            if has_gem:
                return "gemini"
            return None
        # gemini par défaut
        if has_gem:
            return "gemini"
        if has_xai:
            return "xai"
        return None

    def _xai_suggest(self, image_bytes):
        """Appel xAI Grok vision (API compatible OpenAI) → dict identification."""
        api_key = (self.ia_config or {}).get("xai_api_key")
        if not api_key:
            raise RuntimeError(
                "Aucune clé xAI. Outils → Identification IA pour en ajouter une (console.x.ai)."
            )

        b64 = base64.b64encode(image_bytes).decode("ascii")
        liste_connues = []
        for cat, especes in (self.faune_meuse or {}).items():
            for e in especes[:20]:
                liste_connues.append(f"{e} ({cat})")
        known_txt = ", ".join(liste_connues[:80])

        # Prompt calibré pour Grok vision : consignes courtes, JSON strict, prior régional
        prompt = (
            "You are assisting a field naturalist in Meuse (department 55), Grand Est, France.\n"
            "Identify the MAIN subject in this wildlife/trail-camera/field photo.\n"
            "\n"
            "GEOGRAPHIC PRIOR (important):\n"
            "- Temperate forests, field edges, farmland, ponds typical of NE France.\n"
            "- Prefer species that actually occur in Meuse (e.g. Capreolus capreolus, Sus scrofa, "
            "Vulpes vulpes, Meles meles, Cervus elaphus, common woodland birds, local insects).\n"
            "- Do NOT invent exotic or Mediterranean species unless the image clearly shows them.\n"
            "\n"
            "WHAT TO IDENTIFY:\n"
            "- Live animal, bird, insect, OR field signs: tracks, scat, burrow, run/coulée, feathers, remains.\n"
            "- If only a sign is visible, set type_observation accordingly (not \"Vu\").\n"
            "\n"
            "CONFIDENCE (strict):\n"
            f"{IA_CONFIANCE_RUBRIC_EN}\n"
            "\n"
            "NAMES:\n"
            "- espece = usual French name (e.g. \"Chevreuil d'Europe\", \"Sanglier\", \"Renard roux\", \"Buse variable\").\n"
            "- nom_scientifique = binomial Latin or empty string.\n"
            "- categorie must be exactly one of: Mammifère | Oiseau | Insecte | Autre\n"
            "- type_observation must be exactly one of: "
            "Vu | Empreinte / Trace | Terrier | Coulée / Passage | Autre indice\n"
            "\n"
            "OUTPUT RULES (critical for Grok):\n"
            "- Reply with ONE JSON object only.\n"
            "- No markdown fences, no commentary before or after the JSON.\n"
            "- No trailing commas. Use double quotes for all keys and string values.\n"
            "\n"
            "JSON schema:\n"
            "{"
            "\"categorie\":\"Mammifère|Oiseau|Insecte|Autre\","
            "\"espece\":\"nom français ou Inconnu\","
            "\"nom_scientifique\":\"Genre species\","
            "\"type_observation\":\"Vu|Empreinte / Trace|Terrier|Coulée / Passage|Autre indice\","
            "\"confiance\":\"faible|moyenne|élevée\","
            "\"commentaire\":\"une phrase factuelle sur les critères visibles\""
            "}\n"
            "\n"
            f"Species already used in this notebook (hints only, not a closed list): {known_txt}."
        )

        model = (self.ia_config or {}).get("xai_model") or "grok-2-vision-1212"
        models_try = [model]
        for m in ("grok-2-vision-1212", "grok-2-latest", "grok-3", "grok-4"):
            if m not in models_try:
                models_try.append(m)

        last_err = None
        url = "https://api.x.ai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        for model_name in models_try:
            body = {
                "model": model_name,
                "temperature": 0.1,
                "max_tokens": 500,
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "You output only valid JSON for wildlife identification in Meuse, France. "
                            "No markdown. No prose outside the JSON object."
                        ),
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{b64}",
                                    "detail": "high",
                                },
                            },
                        ],
                    },
                ],
            }
            # Certains comptes xAI acceptent le mode JSON strict
            body_json_mode = dict(body)
            body_json_mode["response_format"] = {"type": "json_object"}

            try:
                resp = requests.post(url, headers=headers, json=body_json_mode, timeout=45)
                if resp.status_code in (400, 404, 422):
                    # repli sans response_format / modèle suivant
                    resp = requests.post(url, headers=headers, json=body, timeout=45)
                if resp.status_code in (404, 400, 422):
                    last_err = f"{model_name}: {resp.status_code} {(resp.text or '')[:200]}"
                    continue
                resp.raise_for_status()
                payload = resp.json()
                texte = payload["choices"][0]["message"]["content"].strip()
                texte = texte.replace("```json", "").replace("```", "").strip()
                mjson = re.search(r"\{[\s\S]*\}", texte)
                if mjson:
                    texte = mjson.group(0)
                data = json.loads(texte)
                data["_model"] = f"xAI/{model_name}"
                data["_provider"] = "xai"
                return data
            except Exception as e:
                last_err = str(e)
                continue
        raise RuntimeError(last_err or "Échec de l'appel xAI Grok")

    def suggest_species_ai(self):
        """Reconnaissance d'espèce IA (Gemini ou Grok xAI) sur la photo sélectionnée."""
        if not self.selected_photo_path:
            messagebox.showwarning("Aucune photo", "Sélectionnez d'abord une photo dans la liste.")
            return
        if self._is_video_file(os.path.basename(self.selected_photo_path)):
            messagebox.showinfo(
                "Vidéo",
                "La reconnaissance IA porte sur les photos.\n"
                "Pour une vidéo, extrayez une image fixe ou photographiez le sujet."
            )
            return

        provider = self._resolve_ai_provider()
        if not provider:
            if messagebox.askyesno(
                "Clé IA manquante",
                "Aucune clé Gemini ni xAI (Grok) configurée.\n\nOuvrir les paramètres d'identification IA ?"
            ):
                self.open_ai_settings()
            return

        self.btn_suggest_ai.configure(state="disabled", text="…")
        self._start_progress_pulse(f"Reconnaissance IA ({provider})…")
        path = self.selected_photo_path

        def worker():
            resultat, erreur = None, None
            try:
                thumb = self._prepare_image_for_ai(path, max_dim=768)
                if not thumb:
                    raise RuntimeError("Impossible de préparer une miniature de la photo.")
                if provider == "xai":
                    resultat = self._xai_suggest(thumb)
                else:
                    resultat = self._gemini_suggest(thumb)
                    if isinstance(resultat, dict):
                        resultat["_provider"] = "gemini"
                        if "_model" not in resultat:
                            resultat["_model"] = "Gemini"
            except Exception as e:
                erreur = str(e)
            self.after(0, lambda: self._finish_ai_suggestion(
                (resultat or {}).get("categorie"),
                (resultat or {}).get("espece"),
                (resultat or {}).get("confiance"),
                erreur,
                extra=resultat or {},
            ))

        threading.Thread(target=worker, daemon=True).start()

    def _normalize_confiance(self, confiance):
        """Normalise le libellé de confiance renvoyé par l'IA."""
        c = (confiance or "").strip().lower()
        mapping = {
            "elevee": "élevée", "élevée": "élevée", "eleve": "élevée", "high": "élevée",
            "élevé": "élevée", "forte": "élevée", "high confidence": "élevée",
            "moyenne": "moyenne", "moyen": "moyenne", "medium": "moyenne", "moderee": "moyenne",
            "modérée": "moyenne",
            "faible": "faible", "bas": "faible", "low": "faible", "incertaine": "faible",
        }
        for k, v in mapping.items():
            if k in c or c == k:
                return v
        if "élev" in c or "high" in c:
            return "élevée"
        if "moy" in c or "med" in c:
            return "moyenne"
        if "faib" in c or "low" in c:
            return "faible"
        return confiance or "moyenne"

    def _confiance_explication(self, niveau):
        n = self._normalize_confiance(niveau)
        return {
            "élevée": "Traits diagnostiques nets, bonne image, peu de confusions réalistes en Meuse.",
            "moyenne": "Bonne piste mais limitation (espèce proche, vue partielle, ou indice incomplet).",
            "faible": "Image ou indice insuffisant — vérifier sur le terrain ; « Inconnu » possible.",
        }.get(n, "Niveau non standard — vérifier manuellement.")

    def _finish_ai_suggestion(self, categorie, espece, confiance, erreur, extra=None):
        self.btn_suggest_ai.configure(state="normal", text="🤖")
        self._stop_progress_pulse(0, "Prêt")
        extra = extra or {}
        confiance = self._normalize_confiance(confiance)

        if not espece:
            messagebox.showerror(
                "Échec de l'identification",
                erreur or "Aucune suggestion obtenue.\nVérifiez la clé API et la connexion Internet."
            )
            self.log(f"IA : échec{f' — {erreur}' if erreur else ''}")
            return

        if categorie not in (self.faune_meuse or {}):
            categorie = "Autre"

        def apply_choice(cat, esp):
            self.choice_category.set(cat)
            self.on_category_change(cat)
            especes_liste = self.faune_meuse.get(cat, [])
            correspondance = next(
                (e for e in especes_liste if esp.lower() in e.lower() or e.lower() in esp.lower()),
                None,
            )
            if correspondance:
                self.choice_species.set(correspondance)
            else:
                self.faune_meuse.setdefault(cat, [])
                if esp not in self.faune_meuse[cat]:
                    self.faune_meuse[cat].append(esp)
                    self.save_species_dict()
                self.choice_species.configure(values=self.faune_meuse[cat])
                self.choice_species.set(esp)
            # Complète les notes avec le commentaire IA si vide
            comment = (extra.get("commentaire") or "").strip()
            sci = (extra.get("nom_scientifique") or "").strip()
            existing = self.note_text.get("0.0", "end-1c").strip()
            if comment and not existing:
                add = comment
                if sci:
                    add = f"{comment} ({sci})"
                self.note_text.insert("0.0", add)

        model = extra.get("_model", "")
        self.log(
            f"IA ({model or 'Gemini'}, confiance {confiance or '?'}) : "
            f"{categorie} → {espece}"
        )

        auto = bool((self.ia_config or {}).get("gemini_auto_apply", False))
        if auto:
            apply_choice(categorie, espece)
            messagebox.showinfo(
                "Suggestion appliquée",
                f"{categorie} → {espece}\nConfiance : {confiance or '?'}\n\n"
                "Vérifiez puis enregistrez l'observation."
            )
            return

        # Dialogue de validation
        dlg = ctk.CTkToplevel(self)
        dlg.title(f"Suggestion IA ({(extra or {}).get('_provider', 'IA')}) — vérifier")
        dlg.geometry("480x360")
        dlg.configure(fg_color=UI.get("bg", "#0f1419"))
        self._prepare_tool_window(dlg)
        frame = ctk.CTkFrame(dlg, fg_color=UI.get("card", "#1a2332"), corner_radius=12)
        frame.pack(fill="both", expand=True, padx=12, pady=12)

        ctk.CTkLabel(frame, text="Proposition Gemini", font=ctk.CTkFont(size=14, weight="bold")).pack(
            anchor="w", padx=14, pady=(12, 6)
        )
        details = (
            f"Catégorie : {categorie}\n"
            f"Espèce : {espece}\n"
            f"Nom scientifique : {extra.get('nom_scientifique') or '—'}\n"
            f"Type : {extra.get('type_observation') or '—'}\n"
            f"Confiance : {confiance or '—'}\n"
            f"  → {self._confiance_explication(confiance)}\n"
            f"Commentaire : {extra.get('commentaire') or '—'}\n"
            f"Modèle : {model or '—'}"
        )
        ctk.CTkLabel(
            frame, text=details, font=ctk.CTkFont(size=12),
            text_color=UI.get("text", "#eee"), justify="left", anchor="w"
        ).pack(anchor="w", padx=14, pady=6)

        ctk.CTkLabel(
            frame,
            text="Toujours confirmer sur le terrain : l'IA peut se tromper\n"
                 "(jeunes, femelles, indices partiels, flou…).",
            font=ctk.CTkFont(size=11), text_color=UI.get("warning", "#e67e22"),
            justify="left"
        ).pack(anchor="w", padx=14, pady=(4, 10))

        row = ctk.CTkFrame(frame, fg_color="transparent")
        row.pack(fill="x", padx=14, pady=(0, 12))

        def on_apply():
            apply_choice(categorie, espece)
            dlg.destroy()

        def on_reject():
            dlg.destroy()

        ctk.CTkButton(
            row, text="Appliquer au formulaire", command=on_apply,
            fg_color=UI.get("success", "#2ecc71"), hover_color=UI.get("success_hover", "#27ae60"),
            height=34
        ).pack(side="left", padx=(0, 6))
        ctk.CTkButton(
            row, text="Ignorer", command=on_reject,
            fg_color=UI.get("card_alt", "#333"), height=34, width=100
        ).pack(side="right")

    def _detect_burst(self, filename, max_gap_seconds=8):
        """Retourne la liste des fichiers du dossier pris à moins de max_gap_seconds du fichier donné (rafale)."""
        dated = []
        for fn, info in self.photos_data.items():
            date_str = info.get("date")
            if not date_str:
                continue
            try:
                dt = datetime.strptime(date_str, "%Y:%m:%d %H:%M:%S")
            except Exception:
                continue
            dated.append((fn, dt))
        if not dated:
            return []

        target_dt = next((dt for fn, dt in dated if fn == filename), None)
        if target_dt is None:
            return []

        burst = sorted([fn for fn, dt in dated if abs((dt - target_dt).total_seconds()) <= max_gap_seconds])
        return burst if len(burst) > 1 else []

    def apply_to_burst(self):
        """Applique catégorie/espèce/nombre/lieu/notes de la photo courante à toute la rafale détectée."""
        if not self.current_burst_files or not self.photo_folder_path:
            return
        if not messagebox.askyesno(
            "Confirmer",
            f"Appliquer cette observation aux {len(self.current_burst_files)} photos de la rafale ?\n"
            "(la géolocalisation et la météo propres à chaque photo sont conservées)"
        ):
            return

        chemin_notes = os.path.join(self.photo_folder_path, NOTES_FILE)
        all_notes = {}
        if os.path.exists(chemin_notes):
            try:
                with open(chemin_notes, "r", encoding="utf-8") as f:
                    all_notes = json.load(f)
            except Exception:
                pass

        valeurs_communes = {
            "departement": "55 - Meuse",
            "categorie": self.choice_category.get(),
            "espece": self.choice_species.get(),
            "nombre": self.choice_count.get(),
            "lieu": self.entry_lieu.get().strip(),
            "notes_libres": self.note_text.get("0.0", "end-1c"),
        }

        for fn in self.current_burst_files:
            entree = dict(all_notes.get(fn, {}))
            entree.update(valeurs_communes)
            entree.setdefault("heure", self.entry_time.get())
            entree.setdefault("meteo", self.weather_cache.get(fn))
            all_notes[fn] = entree

        try:
            self._save_notes_dict(all_notes, force_backup=True)
            self.log(f"📸 Observation appliquée à {len(self.current_burst_files)} photos de la rafale.")
            self.refresh_map_markers()
            self.refresh_daily_counter()
            self._refresh_listbox_annotation_status()
            # XMP en lot pour Lightroom
            if self._xmp_keywords_enabled():
                paths = [
                    os.path.join(self.photo_folder_path, fn)
                    for fn in self.current_burst_files
                    if os.path.isfile(os.path.join(self.photo_folder_path, fn))
                ]
                data_xmp = dict(valeurs_communes)
                threading.Thread(
                    target=self._write_xmp_keywords_to_files,
                    args=(paths, data_xmp),
                    kwargs={"silent": False},
                    daemon=True,
                ).start()
            messagebox.showinfo("Terminé", f"{len(self.current_burst_files)} photos mises à jour.")
        except Exception as e:
            messagebox.showerror("Erreur", str(e))

    def _style_chart_axes(self, ax, title=None):
        """Style matplotlib adapte au theme clair/sombre."""
        is_light = UI.get("ctk_mode") == "Light" or (UI.get("bg") or "").lower() in (
            "#f5f7f6", "#f7f5ef", "#ffffff", "#fafafa", "#f0f4f1", "#efe9dc",
        )
        # Heuristique supplementaire sur la luminance du fond
        bg = UI.get("bg", "#101612")
        try:
            h = bg.lstrip("#")
            if len(h) >= 6:
                r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
                is_light = is_light or ((r + g + b) / 3 > 160)
        except Exception:
            pass
        if is_light:
            face, tick, spine, title_c, grid = "#ffffff", "#222222", "#bbbbbb", "#111111", "#e0e0e0"
        else:
            face, tick, spine, title_c, grid = "#1a1a1a", "#f0f0f0", "#555555", "#ffffff", "#333333"
        ax.set_facecolor(face)
        ax.tick_params(colors=tick, labelsize=8)
        for sp in ax.spines.values():
            sp.set_color(spine)
        try:
            ax.yaxis.label.set_color(tick)
            ax.xaxis.label.set_color(tick)
        except Exception:
            pass
        if title:
            ax.set_title(title, color=title_c, fontsize=10)
        try:
            ax.grid(True, color=grid, alpha=0.5, linestyle="--", linewidth=0.6)
            ax.set_axisbelow(True)
        except Exception:
            pass
        return {"tick": tick, "title": title_c, "is_light": is_light}


    # --- Carnet terrain hors-ligne (téléphone, style MapMarker simplifié) ---

    def _field_html_source_path(self):
        """Chemin du modèle HTML (à côté du script, cwd, ou artifacts)."""
        bases = [
            os.path.dirname(os.path.abspath(sys.argv[0])),
            os.getcwd(),
        ]
        try:
            bases.append(os.path.dirname(os.path.abspath(__file__)))
        except NameError:
            pass
        candidates = [os.path.join(b, "GeoExif_Terrain.html") for b in bases]
        candidates.append("/home/workdir/artifacts/GeoExif_Terrain.html")
        for c in candidates:
            if c and os.path.isfile(c):
                return c
        return None

    def export_field_notebook_html(self, target_dir=None):
        """Copie le carnet terrain HTML dans un dossier (ex. Google Drive) pour usage téléphone hors-ligne."""
        src = self._field_html_source_path()
        if not src:
            messagebox.showerror(
                "Fichier manquant",
                "GeoExif_Terrain.html introuvable.\nPlacez-le à côté de l'exécutable / script."
            )
            return None
        if target_dir is None:
            target_dir = filedialog.askdirectory(
                title="Dossier de destination (ex. Google Drive « GeoExif »)",
                initialdir=self._reports_dir() or self.photo_folder_path or None,
            )
        if not target_dir:
            return None
        dest = os.path.join(target_dir, "GeoExif_Terrain.html")
        shutil.copy2(src, dest)
        self.log(f"📱 Carnet terrain copié : {dest}")
        return dest

    def open_companion_panel(self):
        """Guide carnet terrain hors-ligne + import le soir."""
        win = ctk.CTkToplevel(self)
        win.title("📱 Carnet terrain hors-ligne")
        win.geometry("580x640")
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win,
            text="Comme MapMarker, mais dédié au carnet naturaliste Meuse.\n"
                 "Aucun PC ni Wi‑Fi sur le terrain (15 km ou plus).",
            font=ctk.CTkFont(size=12), text_color=UI.get("text_dim", "#888"),
            justify="left", wraplength=540
        ).pack(anchor="w", padx=16, pady=(16, 10))

        def card(title, body, btn_text=None, btn_cmd=None, color=None):
            f = ctk.CTkFrame(win, fg_color=UI.get("card", "#1a2332"), corner_radius=12)
            f.pack(fill="x", padx=16, pady=5)
            ctk.CTkLabel(f, text=title, font=ctk.CTkFont(size=13, weight="bold")).pack(anchor="w", padx=12, pady=(10, 2))
            ctk.CTkLabel(f, text=body, font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#888"),
                         justify="left", wraplength=520).pack(anchor="w", padx=12, pady=(0, 8 if btn_text else 12))
            if btn_text:
                ctk.CTkButton(
                    f, text=btn_text, command=btn_cmd, height=36,
                    fg_color=color or UI.get("accent", "#3d9cf0"),
                    hover_color=UI.get("accent_hover", "#2b7fc4")
                ).pack(fill="x", padx=12, pady=(0, 12))

        def do_export():
            path = self.export_field_notebook_html()
            if path:
                messagebox.showinfo(
                    "Carnet prêt",
                    f"Fichier :\n{path}\n\n"
                    "1) Sur le téléphone, ouvrez ce fichier (Drive / fichiers)\n"
                    "2) Autorisez le GPS\n"
                    "3) Trace GPX optionnelle + points (note / photo / audio)\n"
                    "4) Export JSON+GPX le soir → importer ici"
                )

        card(
            "①  Préparer le téléphone",
            "Générez GeoExif_Terrain.html dans un dossier Google Drive (ou câble).\n"
            "Ouvrez-le dans Chrome/Safari. Ajoutez-le à l’écran d’accueil si possible.",
            "📄 Générer GeoExif_Terrain.html…", do_export, UI.get("success", "#2ecc71")
        )
        card(
            "②  Sur le terrain",
            "• Trace GPX : démarrer / arrêter selon besoin\n"
            "• Points : espèce, notes, photo, mémo audio\n"
            "• Données stockées dans le navigateur (sans réseau)\n"
            "• Bouton Export → JSON + GPX dans Téléchargements / Drive"
        )
        card(
            "③  Le soir sur le PC",
            "Importez le JSON (points + médias). Le GPX peut aussi être chargé\n"
            "comme trace pour géotagger les photos du boîtier.",
            "📥 Importer JSON terrain", self.import_field_package
        )

        ctk.CTkLabel(
            win,
            text="Limites vs MapMarker : pas de tuiles carto offline intégrées (trop lourd).\n"
                 "Pour la navigation offline, couplez avec OsmAnd / Organic Maps.\n"
                 "Vidéo non embarquée (fichiers trop volumineux) — photo + audio oui.",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_muted", "#666"),
            justify="left", wraplength=540
        ).pack(anchor="w", padx=16, pady=(8, 4))
        ctk.CTkButton(win, text="Fermer", command=win.destroy).pack(pady=12)

    def import_field_package(self):
        """Importe un export JSON du carnet terrain (points, médias, trace optionnelle)."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Ouvrez d'abord un dossier de photos.")
            return
        path = filedialog.askopenfilename(
            title="Export JSON GeoExif Terrain",
            filetypes=[("JSON", "*.json"), ("Tous", "*.*")],
            initialdir=self._reports_dir() or self.photo_folder_path,
        )
        if not path:
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                raw = json.load(f)
        except Exception as e:
            messagebox.showerror("Erreur", f"Lecture impossible :\n{e}")
            return

        if isinstance(raw, dict) and "observations" in raw:
            items = raw.get("observations") or []
            track = raw.get("track") or []
        elif isinstance(raw, list):
            items, track = raw, []
        elif isinstance(raw, dict) and ("lat" in raw or "espece" in raw):
            items, track = [raw], []
        else:
            messagebox.showerror("Format", "JSON GeoExif Terrain attendu ({observations, track}).")
            return

        media_dir = os.path.join(self.photo_folder_path, "medias_terrain")
        os.makedirs(media_dir, exist_ok=True)
        notes = self._load_notes_dict()
        ok_n = 0
        for item in items:
            if not isinstance(item, dict):
                continue
            try:
                lat = float(item.get("lat"))
                lon = float(item.get("lon"))
            except (TypeError, ValueError):
                continue
            espece = (item.get("espece") or item.get("type_observation") or "Point terrain").strip()
            cle = str(item.get("id") or f"terrain_{datetime.now().strftime('%Y%m%d%H%M%S%f')}")
            if not cle.startswith("_"):
                cle = f"_terrain_{cle}"

            photo_path = None
            photo = item.get("photo")
            if isinstance(photo, dict) and photo.get("data"):
                photo_path = os.path.join(media_dir, f"{cle}.jpg")
                try:
                    with open(photo_path, "wb") as mf:
                        mf.write(base64.b64decode(photo["data"]))
                except Exception:
                    photo_path = None

            audio_path = None
            audio = item.get("audio")
            if isinstance(audio, dict) and audio.get("data"):
                audio_path = os.path.join(media_dir, f"{cle}.webm")
                try:
                    with open(audio_path, "wb") as mf:
                        mf.write(base64.b64decode(audio["data"]))
                except Exception:
                    audio_path = None

            notes[cle] = {
                "departement": "55 - Meuse",
                "categorie": item.get("categorie") or "Autre",
                "espece": espece,
                "nombre": item.get("nombre") or "1",
                "heure": item.get("heure") or "",
                "lieu": item.get("lieu") or "",
                "notes_libres": item.get("notes_libres") or item.get("notes") or "",
                "type_observation": item.get("type_observation") or "Point GPS",
                "sans_photo": not bool(photo_path),
                "source": "field_notebook",
                "lat": lat,
                "lon": lon,
                "gps_accuracy_m": item.get("accuracy"),
                "media_photo": os.path.basename(photo_path) if photo_path else None,
                "media_audio": os.path.basename(audio_path) if audio_path else None,
                "meteo": None,
            }
            self.photos_data[cle] = {"lat": lat, "lon": lon, "date": ""}
            ok_n += 1

        gpx_written = None
        if track and len(track) >= 2:
            gpx_written = os.path.join(self.photo_folder_path, "trace_terrain_import.gpx")
            try:
                lines = [
                    '<?xml version="1.0" encoding="UTF-8"?>',
                    '<gpx version="1.1" creator="GeoExif Terrain" xmlns="http://www.topografix.com/GPX/1/1">',
                    "  <trk><name>Trace terrain importée</name><trkseg>",
                ]
                for t in track:
                    lines.append(f'    <trkpt lat="{t["lat"]}" lon="{t["lon"]}"></trkpt>')
                lines += ["  </trkseg></trk>", "</gpx>"]
                with open(gpx_written, "w", encoding="utf-8") as gf:
                    gf.write("\n".join(lines))
                self.gpx_file_path = gpx_written
            except Exception:
                gpx_written = None

        try:
            self._save_notes_dict(notes, force_backup=True)
        except Exception as e:
            messagebox.showerror("Erreur", str(e))
            return

        self.refresh_map_markers()
        self.refresh_daily_counter()
        msg = f"{ok_n} observation(s) importée(s)."
        if gpx_written:
            msg += f"\nTrace GPX : {os.path.basename(gpx_written)}"
        if ok_n:
            msg += "\nMédias : dossier medias_terrain/"
        messagebox.showinfo("Import terminé", msg)
        self.log(f"📥 Terrain : {ok_n} points depuis {os.path.basename(path)}")

    def import_companion_file(self):
        self.import_field_package()

    def open_exif_analysis(self):
        """Analyse EXIF de la photo sélectionnée + synthèse optionnelle du dossier."""
        exe = self.get_exiftool_path()
        if not exe:
            messagebox.showerror("ExifTool", "exiftool.exe introuvable.")
            return

        win = ctk.CTkToplevel(self)
        win.title("Analyse EXIF")
        win.geometry("720x640")
        win.configure(fg_color=UI.get("bg", "#0f1419"))
        self._prepare_tool_window(win)

        header = ctk.CTkFrame(win, fg_color=UI.get("card", "#1a2332"), corner_radius=12)
        header.pack(fill="x", padx=14, pady=(14, 8))
        ctk.CTkLabel(header, text="Métadonnées techniques (ExifTool)", font=ctk.CTkFont(size=15, weight="bold")).pack(
            side="left", padx=14, pady=12
        )
        lbl_st = ctk.CTkLabel(header, text="", text_color=UI.get("text_dim", "#888"))
        lbl_st.pack(side="right", padx=14)

        body = ctk.CTkTextbox(win, font=ctk.CTkFont(family="Consolas", size=12), fg_color=UI.get("card", "#1a2332"))
        body.pack(fill="both", expand=True, padx=14, pady=(0, 8))

        btn_row = ctk.CTkFrame(win, fg_color="transparent")
        btn_row.pack(fill="x", padx=14, pady=(0, 14))

        TAGS_PHOTO = [
            "FileName", "ImageSize", "ImageWidth", "ImageHeight", "Megapixels",
            "Make", "Model", "LensModel", "LensID", "FocalLength", "FocalLengthIn35mmFormat",
            "FNumber", "ExposureTime", "ISO", "ExposureProgram", "MeteringMode",
            "WhiteBalance", "CreateDate", "DateTimeOriginal", "OffsetTimeOriginal",
            "GPSLatitude", "GPSLongitude", "GPSAltitude", "GPSDateTime",
            "Orientation", "Flash", "ShutterSpeed", "Aperture", "LightValue",
        ]

        def format_item(item):
            lines = []
            fn = os.path.basename(item.get("SourceFile", item.get("FileName", "?")))
            lines.append(f"===  {fn}  ===")
            pairs = [
                ("Appareil", f"{item.get('Make', '')} {item.get('Model', '')}".strip()),
                ("Objectif", item.get("LensModel") or item.get("LensID")),
                ("Focale", item.get("FocalLength") or item.get("FocalLengthIn35mmFormat")),
                ("Ouverture", f"f/{item.get('FNumber')}" if item.get("FNumber") is not None else None),
                ("Vitesse", item.get("ExposureTime") or item.get("ShutterSpeed")),
                ("ISO", item.get("ISO")),
                ("Date prise de vue", item.get("DateTimeOriginal") or item.get("CreateDate")),
                ("Dimensions", item.get("ImageSize") or (
                    f"{item.get('ImageWidth')}x{item.get('ImageHeight')}" if item.get("ImageWidth") else None
                )),
                ("GPS lat", item.get("GPSLatitude")),
                ("GPS lon", item.get("GPSLongitude")),
                ("Altitude", item.get("GPSAltitude")),
                ("Flash", item.get("Flash")),
                ("Balance des blancs", item.get("WhiteBalance")),
            ]
            for lab, val in pairs:
                if val is not None and str(val).strip() and str(val).strip() != "None":
                    lines.append(f"  {lab:22} {val}")
            lines.append("")
            return "\n".join(lines)

        def show_text(txt):
            body.delete("0.0", "end")
            body.insert("0.0", txt)

        def analyze_selected():
            if not self.selected_photo_path or not os.path.isfile(self.selected_photo_path):
                messagebox.showinfo("Photo", "Sélectionnez une photo dans le carnet.")
                return
            lbl_st.configure(text="Lecture...")
            path = self.selected_photo_path

            def work():
                try:
                    rows = self._exiftool_read_json(exe, [path], TAGS_PHOTO, timeout=30)
                    if not rows:
                        self.after(0, lambda: (show_text("Aucune métadonnée lue."), lbl_st.configure(text="Vide")))
                        return
                    txt = format_item(rows[0])
                    self.after(0, lambda: (show_text(txt), lbl_st.configure(text="Photo sélectionnée")))
                except Exception as e:
                    msg = str(e)
                    self.after(0, lambda: (show_text(msg), lbl_st.configure(text="Erreur")))

            threading.Thread(target=work, daemon=True).start()

        def analyze_folder():
            if not self.photo_folder_path:
                messagebox.showinfo("Dossier", "Chargez un dossier de photos.")
                return
            images = self.get_supported_images()
            if not images:
                messagebox.showinfo("Dossier", "Aucune image dans le dossier.")
                return
            lbl_st.configure(text="Analyse dossier...")
            self._start_progress_pulse("Analyse EXIF du dossier...")
            dossier = self.photo_folder_path
            chemins = [os.path.join(dossier, f) for f in images]

            def work():
                try:
                    rows = self._exiftool_read_json_chunked(
                        exe, chemins, TAGS_PHOTO,
                        chunk_size=40, timeout_per_chunk=60,
                        progress_label="Analyse EXIF",
                    )
                    iso_c = collections.Counter()
                    model_c = collections.Counter()
                    lens_c = collections.Counter()
                    with_gps = 0
                    for it in rows:
                        if it.get("ISO") is not None:
                            iso_c[str(it.get("ISO"))] += 1
                        m = f"{it.get('Make', '')} {it.get('Model', '')}".strip()
                        if m:
                            model_c[m] += 1
                        lens = it.get("LensModel") or it.get("LensID")
                        if lens:
                            lens_c[str(lens)] += 1
                        if it.get("GPSLatitude") and it.get("GPSLongitude"):
                            with_gps += 1

                    lines = [
                        f"===  Synthèse dossier ({len(rows)} fichier(s) lus / {len(images)} image(s))  ===",
                        f"  Avec GPS EXIF     {with_gps}",
                        f"  Sans GPS EXIF     {max(0, len(rows) - with_gps)}",
                        "",
                        "--- Appareils ---",
                    ]
                    for k, n in model_c.most_common(8):
                        lines.append(f"  {k}: {n}")
                    lines.append("")
                    lines.append("--- Objectifs ---")
                    for k, n in lens_c.most_common(8):
                        lines.append(f"  {k}: {n}")
                    lines.append("")
                    lines.append("--- ISO (top) ---")
                    for k, n in iso_c.most_common(10):
                        lines.append(f"  ISO {k}: {n}")
                    lines.append("")
                    lines.append("--- Échantillon (5 premiers) ---")
                    for it in rows[:5]:
                        lines.append(format_item(it))
                    txt = "\n".join(lines)
                    self.after(0, lambda: (
                        show_text(txt),
                        lbl_st.configure(text=f"{len(rows)} fichier(s)"),
                        self._stop_progress_pulse(0, "Prêt"),
                    ))
                except Exception as e:
                    msg = str(e)
                    self.after(0, lambda: (
                        show_text(msg),
                        lbl_st.configure(text="Erreur"),
                        self._stop_progress_pulse(0, "Prêt"),
                    ))

            threading.Thread(target=work, daemon=True).start()

        ctk.CTkButton(
            btn_row, text="Photo sélectionnée", command=analyze_selected,
            fg_color=UI.get("accent", "#3d9cf0"), hover_color=UI.get("accent_hover", "#2b7fc4"), height=34
        ).pack(side="left", padx=(0, 6))
        ctk.CTkButton(
            btn_row, text="Synthèse du dossier", command=analyze_folder,
            fg_color=UI.get("success", "#2ecc71"), hover_color=UI.get("success_hover", "#27ae60"), height=34
        ).pack(side="left", padx=6)
        ctk.CTkButton(btn_row, text="Fermer", command=win.destroy, width=90).pack(side="right")

        if self.selected_photo_path:
            analyze_selected()
        else:
            show_text("Sélectionnez une photo dans le carnet, ou lancez la synthèse du dossier.")

    def export_geojson_all_sorties(self):
        """Export GeoJSON cumulé de toutes les sorties connues (QGIS)."""
        folders = self._load_known_folders()
        if not folders:
            messagebox.showinfo("Aucune sortie", "Aucune sortie connue.")
            return
        chemin = self._ask_save_report(
            ".geojson",
            [("GeoJSON", "*.geojson")],
            f"toutes_sorties_{datetime.now().strftime('%Y%m%d')}.geojson",
        )
        if not chemin:
            return
        self._start_progress_pulse("Export GeoJSON multi-sorties...")

        def work():
            try:
                dataset = self._build_report_dataset(
                    [f["path"] for f in folders if os.path.isdir(f.get("path", ""))],
                    read_exif=False,
                )
                features = []
                seen = set()
                for o in dataset:
                    lat, lon = o.get("lat"), o.get("lon")
                    if not (lat and lon):
                        continue
                    try:
                        lat, lon = float(lat), float(lon)
                    except Exception:
                        continue
                    espece = (o.get("espece") or "?").strip()
                    key = (round(lat, 4), round(lon, 4), espece.lower())
                    if key in seen:
                        continue
                    seen.add(key)
                    shape = self._detect_marker_shape(espece)
                    features.append({
                        "type": "Feature",
                        "geometry": {"type": "Point", "coordinates": [lon, lat]},
                        "properties": {
                            "name": espece,
                            "espece": espece,
                            "categorie": o.get("categorie"),
                            "nombre": o.get("nombre"),
                            "heure": o.get("heure"),
                            "lieu": o.get("lieu"),
                            "dossier": o.get("dossier"),
                            "type_observation": o.get("type_observation"),
                            "sans_photo": bool(o.get("sans_photo")),
                            "marqueur": shape,
                            "marker-symbol": shape,
                            "marker-color": CATEGORY_COLORS.get(o.get("categorie") or "Non classé", "#7f7f7f"),
                            "couleur": CATEGORY_COLORS.get(o.get("categorie") or "Non classé", "#7f7f7f"),
                            "icon": {"circle": "●", "triangle": "▲", "square": "■", "diamond": "◆"}.get(shape, "●"),
                        },
                    })
                geojson = {
                    "type": "FeatureCollection",
                    "name": "GeoExif_toutes_sorties",
                    "crs": {"type": "name", "properties": {"name": "urn:ogc:def:crs:OGC:1.3:CRS84"}},
                    "features": features,
                }
                with open(chemin, "w", encoding="utf-8") as f:
                    json.dump(geojson, f, ensure_ascii=False, indent=2)
                self.after(0, lambda: (
                    self._stop_progress_pulse(1.0, "Export GeoJSON terminé"),
                    messagebox.showinfo("GeoJSON", f"{len(features)} point(s) exportés :\n{chemin}"),
                    self.log(f"Export GeoJSON multi-sorties : {len(features)} points"),
                ))
                self.after(2000, lambda: self._stop_progress_pulse(0, "Prêt"))
            except Exception as e:
                msg = str(e)
                self.after(0, lambda: (
                    self._stop_progress_pulse(0, "Prêt"),
                    messagebox.showerror("Erreur", msg),
                ))

        threading.Thread(target=work, daemon=True).start()


    def open_faq(self):
        """FAQ intégrée : usage, IA, Lightroom, vidéos, limites."""
        win = ctk.CTkToplevel(self)
        win.title("FAQ — GeoExif Meuse 55")
        win.geometry("720x640")
        win.configure(fg_color=UI.get("bg", "#0f1419"))
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Questions fréquentes",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(anchor="w", padx=16, pady=(16, 8))

        box = ctk.CTkTextbox(win, font=ctk.CTkFont(size=13), fg_color=UI.get("card", "#1a2332"))
        box.pack(fill="both", expand=True, padx=16, pady=(0, 12))
        faq = """
Q. Dans quel ordre utiliser les outils ?
R. Terrain (Locus) → copie photos (FastStone) → GeoExif (GPS + carnet) → Lightroom ensuite.

Q. Pourquoi mes points sont décalés sur la carte ?
R. Décalage d'heure photo / GPX. Essayez +1:00:00 (heure d'été) puis « Forcer la réécriture ».

Q. ExifTool introuvable ?
R. Placez exiftool.exe à côté de GeoExif, ou installez-le dans le PATH Windows.

Q. Comment fonctionne la reconnaissance d'espèces IA ?
R. Outils → Identification IA : collez une clé Gemini gratuite (aistudio.google.com).
   Puis dans le carnet : photo → bouton 🤖 → vérifier → Appliquer → Enregistrer.
   L'IA peut se tromper : validation humaine obligatoire.

Q. Pourquoi pas Grok / ChatGPT intégrés à la place de Gemini ?
R. GeoExif appelle une API depuis votre PC. Gemini propose un accès vision simple et un niveau gratuit.
   Grok peut vous aider ici dans le chat (description, photo collée dans la conversation),
   mais n'est pas branché en local dans l'appli (pas de clé API xAI configurée dans GeoExif).

Q. Reconnaissance de chants d'oiseaux (BirdNET) ?
R. Possible en option si BirdNET Analyzer est installé sur le PC (modèle local, fichiers audio).
   Bouton « Analyse BirdNET » dans les outils si l'exécutable est détecté.
   Sinon : utilisez l'appli BirdNET séparément, puis saisissez l'espèce dans le carnet.

Q. Les vidéos d'affût ?
R. Mettez-les dans le même dossier que les photos. Elles apparaissent avec 🎬.
   Clic / double-clic / bouton « Lire la vidéo » → lecteur Windows par défaut.
   L'IA image ne traite pas la vidéo.

Q. Lightroom ne voit pas le GPS ni les mots-clés ?
R. Après GeoExif : dans LR, sélection → Métadonnées → Lire les métadonnées depuis le fichier.
   Vérifiez que « Mots-clés Lightroom (XMP) » est activé dans la configuration GeoExif.

Q. Carte cumulée vide pour d'anciennes sorties ?
R. Rouvrez chaque dossier une fois, ou laissez la carte cumulée relire les EXIF.
   Enregistrer une fiche écrit aussi le GPS dans le carnet.

Q. Brief IA multi-sorties : où est le fichier ?
R. Le brief est un texte à copier vers Gemini/ChatGPT. PDF/Word demandent un emplacement de sauvegarde.

Q. Données et vie privée (IA) ?
R. La miniature de la photo est envoyée à Google (Gemini) uniquement quand vous cliquez 🤖.
   Le carnet local (observations.json) reste sur votre disque.
""".strip()
        box.insert("0.0", faq)
        box.configure(state="disabled")
        ctk.CTkButton(win, text="Fermer", command=win.destroy, width=100).pack(pady=(0, 14))

    def _find_birdnet_cmd(self):
        """Détecte BirdNET Analyzer en ligne de commande s'il est installé."""
        candidates = [
            shutil.which("birdnet_analyzer"),
            shutil.which("birdnet-analyzer"),
            shutil.which("birdnet"),
        ]
        for c in candidates:
            if c:
                return c
        # Module Python : python -m birdnet_analyzer.analyze
        try:
            r = subprocess.run(
                [sys.executable, "-c", "import birdnet_analyzer; print('ok')"],
                capture_output=True, text=True, timeout=8,
            )
            if r.returncode == 0 and "ok" in (r.stdout or ""):
                return "python_module"
        except Exception:
            pass
        return None

    def open_birdnet_panel(self):
        """Ancien point d'entrée BirdNET → redirige vers les outils Birda."""
        self.open_birda_tools()

    def _configured_exe(self, key):
        """Chemin exe configuré par l'utilisateur (audacity_exe / birda_gui_exe)."""
        cfg = self._load_app_config() or {}
        p = (cfg.get(key) or "").strip()
        if p and os.path.isfile(p):
            return p
        return None

    def _set_configured_exe(self, key, path):
        cfg = dict(self._load_app_config() or {})
        cfg[key] = (path or "").strip()
        try:
            self._save_app_config(cfg)
            self.app_config = cfg
        except Exception:
            pass

    def _find_app_exe(self, names, config_key=None):
        """Cherche un executable : config utilisateur, puis PATH, puis emplacements courants."""
        if config_key:
            conf = self._configured_exe(config_key)
            if conf:
                return conf
        for n in names:
            p = shutil.which(n)
            if p:
                return p
        if os.name == "nt":
            roots = []
            for env in ("ProgramFiles", "ProgramFiles(x86)", "LOCALAPPDATA", "APPDATA"):
                v = os.environ.get(env)
                if v:
                    roots.append(v)
            roots += [
                os.path.expanduser("~"),
                os.path.join(os.path.expanduser("~"), "AppData", "Local", "Programs"),
                os.path.join(os.path.expanduser("~"), "Desktop"),
                os.path.join(os.path.expanduser("~"), "Downloads"),
            ]
            candidates = []
            for root in roots:
                for n in names:
                    candidates.append(os.path.join(root, n))
                    candidates.append(os.path.join(root, "Audacity", n))
                    candidates.append(os.path.join(root, "Birda", n))
                    candidates.append(os.path.join(root, "birda-gui", n))
                    candidates.append(os.path.join(root, "Birda GUI", n))
                    candidates.append(os.path.join(root, "birda", n))
            for c in candidates:
                if c and os.path.isfile(c):
                    return c
        return None

    def _birda_catalog_candidates(self):
        """Emplacements possibles de birda-catalog.db (Birda GUI / Electron userData)."""
        homes = []
        if os.name == "nt":
            for env in ("APPDATA", "LOCALAPPDATA"):
                v = os.environ.get(env)
                if v:
                    homes.append(v)
        homes.append(os.path.join(os.path.expanduser("~"), ".config"))
        homes.append(os.path.join(os.path.expanduser("~"), "Library", "Application Support"))
        names = (
            "birda-gui", "Birda GUI", "birda", "Birda",
            "birda-gui-updater",
        )
        out = []
        cfg = self._load_app_config() or {}
        custom = (cfg.get("birda_catalog_db") or "").strip()
        if custom:
            out.append(custom)
        for h in homes:
            for n in names:
                out.append(os.path.join(h, n, "birda-catalog.db"))
                out.append(os.path.join(h, n, "data", "birda-catalog.db"))
        # dédup en gardant l'ordre
        seen = set()
        uniq = []
        for p in out:
            pn = os.path.normpath(p)
            if pn not in seen:
                seen.add(pn)
                uniq.append(pn)
        return uniq

    def _find_birda_catalog_db(self):
        for p in self._birda_catalog_candidates():
            if os.path.isfile(p):
                return p
        return None

    def _launch_external(self, path_or_cmd, label="Application"):
        if not path_or_cmd:
            messagebox.showinfo(
                label,
                "%s introuvable.\n\n"
                "Dans la page Birda, utilisez « Parcourir… » pour indiquer\n"
                "le chemin du fichier .exe, puis Enregistrer les chemins."
                % label,
            )
            return
        try:
            if os.name == "nt":
                os.startfile(path_or_cmd)
            else:
                subprocess.Popen([path_or_cmd])
            self.log("Ouverture : %s (%s)" % (label, path_or_cmd))
        except Exception as e:
            messagebox.showerror(label, str(e))

    def _birda_watch_folder(self):
        """Dossier où chercher les derniers résultats Birda (config utilisateur)."""
        cfg = self._load_app_config() or {}
        folder = (cfg.get("birda_watch_folder") or "").strip()
        if folder and os.path.isdir(folder):
            return folder
        # défauts utiles
        for cand in (
            cfg.get("default_photo_folder"),
            self.photo_folder_path,
            os.path.join(os.path.expanduser("~"), "Documents"),
            os.path.expanduser("~"),
        ):
            if cand and os.path.isdir(cand):
                return cand
        return os.path.expanduser("~")

    def _set_birda_watch_folder(self, folder):
        if not folder:
            return
        cfg = dict(self._load_app_config() or {})
        cfg["birda_watch_folder"] = folder
        try:
            self._save_app_config(cfg)
            self.app_config = cfg
        except Exception:
            pass

    def _find_latest_birda_result(self, folder=None, *, max_age_hours=48):
        """Dernier JSON/CSV type Birda dans le dossier (et 1 sous-niveau)."""
        folder = folder or self._birda_watch_folder()
        if not folder or not os.path.isdir(folder):
            return None
        found = []
        try:
            entries = list(os.scandir(folder))
        except Exception:
            return None
        dirs = [folder]
        for e in entries:
            if e.is_dir() and not e.name.startswith("."):
                dirs.append(e.path)
        now = datetime.now().timestamp()
        max_age = max_age_hours * 3600
        for d in dirs:
            try:
                for e in os.scandir(d):
                    if not e.is_file():
                        continue
                    name = e.name.lower()
                    if not (name.endswith(".json") or name.endswith(".csv")):
                        continue
                    # indices Birda / BirdNET
                    if not (
                        "birdnet" in name
                        or "birda" in name
                        or name.endswith(".birdnet.json")
                        or name.endswith(".birdnet.results.csv")
                        or ".birdnet." in name
                        or name.endswith("_results.csv")
                        or name.endswith(".results.csv")
                    ):
                        # accepter tout json/csv recent si peu de candidats
                        pass
                    try:
                        mtime = e.stat().st_mtime
                    except Exception:
                        continue
                    if max_age_hours and (now - mtime) > max_age:
                        continue
                    found.append((mtime, e.path))
            except Exception:
                continue
        if not found:
            return None
        found.sort(key=lambda x: -x[0])
        # privilégier noms birda/birdnet
        for mtime, p in found:
            nl = os.path.basename(p).lower()
            if "birdnet" in nl or "birda" in nl or "results" in nl:
                return p
        return found[0][1]

    def import_latest_birda_result(self):
        """Semi-auto : propose d'importer le dernier fichier résultats du dossier surveillé."""
        folder = self._birda_watch_folder()
        latest = self._find_latest_birda_result(folder, max_age_hours=72)
        if not latest:
            # élargir : tous json/csv récents
            latest = self._find_latest_birda_result(folder, max_age_hours=24 * 14)
        if not latest:
            messagebox.showinfo(
                "Birda semi-auto",
                "Aucun fichier JSON/CSV récent trouvé dans :\n%s\n\n"
                "1) Choisissez le dossier où Birda écrit ses résultats\n"
                "2) Relancez une analyse dans Birda GUI\n"
                "3) Réessayez « Importer le dernier résultat »"
                % folder,
            )
            return
        age_min = int(max(0, (datetime.now().timestamp() - os.path.getmtime(latest)) / 60))
        if not messagebox.askyesno(
            "Birda semi-auto",
            "Dernier fichier trouvé :\n\n%s\n\n"
            "Modifié il y a ~%d min\n\n"
            "Importer ces détections dans le carnet de la sortie ouverte ?"
            % (latest, age_min),
        ):
            return
        self.import_birda_detections(path=latest)

    def open_birda_cli_analyzer(self):
        """Interface Birda CLI : audio → CSV → validation → carnet (comme ExifTool)."""
        win = ctk.CTkToplevel(self)
        win.title("Birda CLI — analyser un enregistrement")
        win.geometry("820x780")
        try:
            win.minsize(700, 640)
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        header = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=14)
        header.pack(fill="x", padx=16, pady=(14, 8))
        ctk.CTkLabel(
            header, text="🎧  Analyse audio avec Birda CLI",
            font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=14, pady=(12, 2))
        ctk.CTkLabel(
            header,
            text="Fichier son → birda -f csv → validation → carnet\n"
                 "(même logique qu'ExifTool : pas de base SQLite)",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            justify="left",
        ).pack(anchor="w", padx=14, pady=(0, 12))

        body = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=14)
        body.pack(fill="both", expand=True, padx=16, pady=4)

        cfg0 = self._load_app_config() or {}
        birda_cli = (
            cfg0.get("birda_cli_exe")
            or self._find_app_exe(
                ["birda.exe", "birda", "Birda.exe"],
                "birda_cli_exe",
            )
            or ""
        )
        audio_var = tk.StringVar()
        exe_var = tk.StringVar(value=birda_cli)
        out_var = tk.StringVar()
        conf_var = tk.StringVar(value="0.35")
        lat_var = tk.StringVar()
        lon_var = tk.StringVar()
        lieu_var = tk.StringVar()
        date_var = tk.StringVar(value=datetime.now().strftime("%Y-%m-%d"))
        heure_var = tk.StringVar(value=datetime.now().strftime("%H:%M"))
        DEFAULT_MODELS = [
            "(défaut Birda)",
            "birdnet-v24",
            "birdnet-v30",
            "perch-v2",
            "bsg-fi-v44",
        ]
        saved_model = (cfg0.get("birda_cli_model") or "").strip()
        model_choices = list(DEFAULT_MODELS)
        if saved_model and saved_model not in model_choices:
            model_choices.insert(1, saved_model)
        model_var = tk.StringVar(
            value=saved_model if saved_model in model_choices else DEFAULT_MODELS[0]
        )

        # Préremplir GPS depuis carte ou session
        try:
            pos = self.map_widget.get_position()
            lat_var.set("%.6f" % float(pos[0]))
            lon_var.set("%.6f" % float(pos[1]))
        except Exception:
            pass
        meta = getattr(self, "_ecoute_session_meta", None) or {}
        if meta.get("lat") is not None:
            lat_var.set(str(meta["lat"]))
        if meta.get("lon") is not None:
            lon_var.set(str(meta["lon"]))
        if meta.get("lieu"):
            lieu_var.set(meta["lieu"])

        def row_browse(parent, label, var, browse_fn, tip=""):
            fr = ctk.CTkFrame(parent, fg_color="transparent")
            fr.pack(fill="x", padx=12, pady=5)
            ctk.CTkLabel(fr, text=label, width=110, anchor="w", text_color=UI.get("text")).pack(side="left")
            ctk.CTkEntry(fr, textvariable=var).pack(side="left", fill="x", expand=True, padx=(0, 6))
            ctk.CTkButton(fr, text="…", width=36, command=browse_fn).pack(side="left")
            if tip:
                ctk.CTkLabel(
                    parent, text=tip, font=ctk.CTkFont(size=10),
                    text_color=UI.get("text_dim"),
                ).pack(anchor="w", padx=122, pady=(0, 2))

        def pick_audio():
            p = filedialog.askopenfilename(
                title="Fichier audio à analyser",
                filetypes=[
                    ("Audio", "*.wav *.flac *.mp3 *.ogg *.m4a *.aac"),
                    ("WAV", "*.wav"),
                    ("FLAC", "*.flac"),
                    ("Tous", "*.*"),
                ],
            )
            if p:
                audio_var.set(p)
                if not out_var.get().strip():
                    out_var.set(os.path.join(os.path.dirname(p), "birda_out"))

        def pick_exe():
            p = filedialog.askopenfilename(
                title="birda.exe (CLI)",
                filetypes=[("Executable", "*.exe"), ("Tous", "*.*")],
            )
            if p:
                exe_var.set(p)
                self._set_configured_exe("birda_cli_exe", p)

        def pick_out():
            d = filedialog.askdirectory(title="Dossier de sortie CSV Birda")
            if d:
                out_var.set(d)

        def use_map_gps():
            try:
                pos = self.map_widget.get_position()
                lat_var.set("%.6f" % float(pos[0]))
                lon_var.set("%.6f" % float(pos[1]))
                try:
                    lbl_gps_hint.configure(
                        text="Centre carte → %.5f , %.5f" % (float(pos[0]), float(pos[1]))
                    )
                except Exception:
                    pass
            except Exception:
                messagebox.showinfo("GPS", "Carte indisponible — saisissez lat/lon ou cliquez la carte.")

        def pick_gps_on_map():
            """Active un clic unique sur la carte principale pour remplir Lat/Lon."""
            try:
                self.tab_view.set("🗺️  Carte")
            except Exception:
                pass
            self._birda_cli_gps_pick = {
                "lat_var": lat_var,
                "lon_var": lon_var,
                "hint": lbl_gps_hint,
                "win": win,
            }
            try:
                lbl_gps_hint.configure(
                    text="Cliquez maintenant sur la carte principale pour fixer le point…"
                )
            except Exception:
                pass
            self.log("Birda CLI — cliquez sur la carte pour définir Lat/Lon.")
            messagebox.showinfo(
                "GPS carte",
                "La fenêtre reste ouverte.\n"
                "Cliquez une fois sur la carte (onglet Carte)\n"
                "à l’endroit de l’écoute : Lat / Lon seront remplis ici.",
            )

        row_browse(body, "Audio", audio_var, pick_audio, "Enregistrement .wav / .flac / .mp3…")
        row_browse(body, "birda CLI", exe_var, pick_exe, "Ex. birda.exe (pas Birda GUI)")
        row_browse(body, "Sortie CSV", out_var, pick_out, "Dossier où Birda écrira le .csv")

        # Modèle d'analyse
        mod_fr = ctk.CTkFrame(body, fg_color="transparent")
        mod_fr.pack(fill="x", padx=12, pady=(8, 2))
        ctk.CTkLabel(
            mod_fr, text="Modèle", width=110, anchor="w", text_color=UI.get("text"),
        ).pack(side="left")
        model_menu = ctk.CTkOptionMenu(
            mod_fr, variable=model_var, values=model_choices, width=220,
        )
        model_menu.pack(side="left", padx=(0, 8))

        def refresh_models():
            exe = exe_var.get().strip()
            if not exe:
                exe = shutil.which("birda") or shutil.which("birda.exe") or ""
            if not exe or (not os.path.isfile(exe) and not shutil.which(exe)):
                messagebox.showinfo(
                    "Modèles",
                    "Indiquez d'abord le chemin de birda.exe,\n"
                    "puis cliquez Rafraîchir.\n\n"
                    "Ou installez un modèle :\n"
                    "  birda models install birdnet-v24\n"
                    "  birda models install perch-v2",
                )
                return
            try:
                startupinfo = None
                if os.name == "nt":
                    startupinfo = subprocess.STARTUPINFO()
                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                proc = subprocess.run(
                    [exe, "--output-mode", "json", "models", "list"],
                    capture_output=True, text=True, encoding="utf-8",
                    errors="replace", timeout=30, startupinfo=startupinfo,
                )
                ids = []
                raw = (proc.stdout or "").strip()
                if raw:
                    try:
                        data = json.loads(raw)
                        models = (
                            (data.get("payload") or {}).get("models")
                            or data.get("models")
                            or []
                        )
                        for m in models:
                            if isinstance(m, dict):
                                mid = m.get("id") or m.get("registry_id") or ""
                                if mid:
                                    ids.append(str(mid))
                            elif isinstance(m, str):
                                ids.append(m)
                    except Exception:
                        # sortie texte ligne par ligne
                        for line in raw.splitlines():
                            line = line.strip()
                            if line and not line.startswith("{") and len(line) < 80:
                                ids.append(line.split()[0])
                if not ids:
                    # fallback humain
                    proc2 = subprocess.run(
                        [exe, "models", "list"],
                        capture_output=True, text=True, encoding="utf-8",
                        errors="replace", timeout=30, startupinfo=startupinfo,
                    )
                    for line in (proc2.stdout or "").splitlines():
                        line = line.strip()
                        if not line or line.lower().startswith("model"):
                            continue
                        tok = line.split()[0]
                        if tok and not tok.startswith("-"):
                            ids.append(tok)
                ids = sorted(set(ids), key=str.lower)
                if not ids:
                    messagebox.showinfo(
                        "Modèles",
                        "Aucun modèle listé.\n\n"
                        "Installez par exemple :\n"
                        "  birda models install birdnet-v24\n"
                        "  birda models install perch-v2\n"
                        "  birda models install birdnet-v30",
                    )
                    return
                choices = ["(défaut Birda)"] + ids
                model_menu.configure(values=choices)
                cur = model_var.get()
                if cur not in choices:
                    model_var.set(choices[0])
                messagebox.showinfo("Modèles", "%d modèle(s) trouvé(s)." % len(ids))
            except Exception as e:
                messagebox.showerror("Modèles", str(e))

        ctk.CTkButton(
            mod_fr, text="Rafraîchir listé", width=120, command=refresh_models,
        ).pack(side="left", padx=4)
        ctk.CTkLabel(
            body,
            text="BirdNET v2.4 · BirdNET v3 · Google Perch v2 — installer via : birda models install <id>",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim"),
        ).pack(anchor="w", padx=122, pady=(0, 4))

        # Contexte
        ctx = ctk.CTkFrame(body, fg_color=UI.get("card_alt", "#24302a"), corner_radius=10)
        ctx.pack(fill="x", padx=12, pady=(10, 6))
        ctk.CTkLabel(
            ctx, text="Contexte carnet",
            font=ctk.CTkFont(size=12, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=10, pady=(8, 4))

        def pair(lab1, v1, lab2, v2):
            r = ctk.CTkFrame(ctx, fg_color="transparent")
            r.pack(fill="x", padx=10, pady=3)
            ctk.CTkLabel(r, text=lab1, width=70, anchor="w", text_color=UI.get("text")).pack(side="left")
            ctk.CTkEntry(r, textvariable=v1, width=100).pack(side="left", padx=(0, 12))
            ctk.CTkLabel(r, text=lab2, width=50, anchor="w", text_color=UI.get("text")).pack(side="left")
            ctk.CTkEntry(r, textvariable=v2, width=100).pack(side="left")

        pair("Date", date_var, "Heure", heure_var)
        r_gps = ctk.CTkFrame(ctx, fg_color="transparent")
        r_gps.pack(fill="x", padx=10, pady=3)
        ctk.CTkLabel(r_gps, text="Lat", width=70, anchor="w", text_color=UI.get("text")).pack(side="left")
        ctk.CTkEntry(r_gps, textvariable=lat_var, width=130).pack(side="left", padx=(0, 6))
        ctk.CTkLabel(r_gps, text="Lon", width=36, anchor="w", text_color=UI.get("text")).pack(side="left")
        ctk.CTkEntry(r_gps, textvariable=lon_var, width=130).pack(side="left", padx=(0, 6))
        r_gps2 = ctk.CTkFrame(ctx, fg_color="transparent")
        r_gps2.pack(fill="x", padx=10, pady=(2, 2))
        ctk.CTkButton(
            r_gps2, text="Centre carte", width=120, height=32,
            command=use_map_gps,
        ).pack(side="left", padx=(70, 6))
        ctk.CTkButton(
            r_gps2, text="Clic sur la carte…", width=150, height=32,
            fg_color=UI.get("accent"),
            command=pick_gps_on_map,
        ).pack(side="left", padx=4)
        lbl_gps_hint = ctk.CTkLabel(
            ctx,
            text="Saisie manuelle, centre de la carte, ou un clic sur la carte.",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim"),
        )
        lbl_gps_hint.pack(anchor="w", padx=80, pady=(0, 4))
        r_lieu = ctk.CTkFrame(ctx, fg_color="transparent")
        r_lieu.pack(fill="x", padx=10, pady=(3, 4))
        ctk.CTkLabel(r_lieu, text="Lieu", width=70, anchor="w", text_color=UI.get("text")).pack(side="left")
        ctk.CTkEntry(r_lieu, textvariable=lieu_var).pack(side="left", fill="x", expand=True)
        r_conf = ctk.CTkFrame(ctx, fg_color="transparent")
        r_conf.pack(fill="x", padx=10, pady=(3, 10))
        ctk.CTkLabel(r_conf, text="Seuil min.", width=70, anchor="w", text_color=UI.get("text")).pack(side="left")
        ctk.CTkEntry(r_conf, textvariable=conf_var, width=80).pack(side="left")
        ctk.CTkLabel(
            r_conf, text="(0.0 – 1.0, ex. 0.35 = 35 %)",
            font=ctk.CTkFont(size=10), text_color=UI.get("text_dim"),
        ).pack(side="left", padx=8)

        # Progression
        prog_fr = ctk.CTkFrame(body, fg_color="transparent")
        prog_fr.pack(fill="x", padx=12, pady=(4, 4))
        prog = ctk.CTkProgressBar(prog_fr, height=8)
        prog.pack(fill="x")
        prog.set(0)
        status = ctk.CTkLabel(
            prog_fr, text="Prêt", font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
        )
        status.pack(anchor="w", pady=(4, 0))

        def set_status(txt, p=None):
            def _():
                try:
                    status.configure(text=txt)
                    if p is not None:
                        prog.set(p)
                except Exception:
                    pass
            try:
                win.after(0, _)
            except Exception:
                pass

        def find_csv_outputs(out_dir, audio_path):
            """Cherche le CSV produit par Birda près de la sortie / de l'audio."""
            found = []
            bases = []
            if out_dir and os.path.isdir(out_dir):
                bases.append(out_dir)
            adir = os.path.dirname(audio_path) if audio_path else ""
            if adir and adir not in bases:
                bases.append(adir)
            stem = os.path.splitext(os.path.basename(audio_path or ""))[0].lower()
            for base in bases:
                try:
                    for root, _dirs, files in os.walk(base):
                        for name in files:
                            low = name.lower()
                            if not low.endswith(".csv"):
                                continue
                            full = os.path.join(root, name)
                            # privilégier noms BirdNET / birda / stem
                            score = 0
                            if "birdnet" in low or "birda" in low:
                                score += 2
                            if stem and stem in low:
                                score += 3
                            found.append((score, os.path.getmtime(full), full))
                except Exception:
                    pass
            found.sort(key=lambda x: (x[0], x[1]), reverse=True)
            return [p for _s, _m, p in found]

        def run_analysis():
            audio = audio_var.get().strip()
            exe = exe_var.get().strip()
            out_dir = out_var.get().strip()
            if not audio or not os.path.isfile(audio):
                messagebox.showwarning("Audio", "Choisissez un fichier audio valide.")
                return
            if not exe or not os.path.isfile(exe):
                # tenter PATH
                which = shutil.which(exe) if exe else None
                if not which:
                    which = shutil.which("birda") or shutil.which("birda.exe")
                if which:
                    exe = which
                    exe_var.set(exe)
                else:
                    messagebox.showerror(
                        "Birda CLI",
                        "birda.exe introuvable.\n\n"
                        "Installez Birda CLI et indiquez le chemin avec « … ».",
                    )
                    return
            if not out_dir:
                out_dir = os.path.join(os.path.dirname(audio), "birda_out")
                out_var.set(out_dir)
            try:
                os.makedirs(out_dir, exist_ok=True)
            except Exception as e:
                messagebox.showerror("Sortie", str(e))
                return
            try:
                conf = float(conf_var.get().replace(",", "."))
            except Exception:
                conf = 0.35
            conf = max(0.0, min(1.0, conf))
            self._set_configured_exe("birda_cli_exe", exe)
            # Lire le modèle depuis la variable ET le menu (CTkOptionMenu)
            model_id = (model_var.get() or "").strip()
            try:
                mid_menu = (model_menu.get() or "").strip()
                if mid_menu:
                    model_id = mid_menu
            except Exception:
                pass
            if model_id.startswith("(") or model_id.lower() in ("défaut", "default", ""):
                model_id = ""
            if model_id:
                cfgm = dict(self._load_app_config() or {})
                cfgm["birda_cli_model"] = model_id
                try:
                    self._save_app_config(cfgm)
                except Exception:
                    pass

            lat = lon = None
            try:
                if lat_var.get().strip():
                    lat = float(lat_var.get().replace(",", "."))
                if lon_var.get().strip():
                    lon = float(lon_var.get().replace(",", "."))
            except Exception:
                lat = lon = None

            # Sous-dossier par modèle → ne pas confondre les CSV d'analyses précédentes
            run_out = out_dir
            if model_id:
                safe = "".join(c if c.isalnum() or c in "-_." else "_" for c in model_id)
                run_out = os.path.join(out_dir, safe)
            try:
                os.makedirs(run_out, exist_ok=True)
            except Exception:
                run_out = out_dir

            # Session meta pour import carnet
            self._ecoute_session_meta = {
                "lieu": lieu_var.get().strip(),
                "date": date_var.get().strip(),
                "heure": heure_var.get().strip(),
                "lat": lat,
                "lon": lon,
                "tool": "birda_cli",
                "model": model_id or "default",
                "folder": self.photo_folder_path or run_out,
            }

            # Ordre CLI Birda : options puis fichier
            # birda -m MODEL -f csv -o DIR -c CONF --force [--lat --lon] audio.wav
            cmd_base = [exe]
            if model_id:
                cmd_base.extend(["-m", model_id])
            cmd_base.extend([
                "-f", "csv",
                "-o", run_out,
                "-c", str(conf),
                "--force",
            ])
            cmd_base.append(audio)
            cmd_geo = list(cmd_base)
            if lat is not None and lon is not None:
                cmd_geo = cmd_base[:-1] + ["--lat", str(lat), "--lon", str(lon)] + [cmd_base[-1]]

            btn_run.configure(state="disabled", text="Analyse…")
            model_label = model_id or "(défaut Birda)"
            set_status("Modèle %s — lancement…" % model_label, 0.05)
            self.log("Birda CLI modèle=%s" % model_label)
            self.log("Birda CLI cmd: %s" % " ".join(cmd_geo))

            def worker():
                t0 = time.time()
                try:
                    startupinfo = None
                    if os.name == "nt":
                        startupinfo = subprocess.STARTUPINFO()
                        startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

                    def run_cmd(c):
                        return subprocess.run(
                            c,
                            capture_output=True,
                            text=True,
                            encoding="utf-8",
                            errors="replace",
                            timeout=3600,
                            startupinfo=startupinfo,
                            cwd=run_out,
                        )

                    # Vérifier que le modèle est installé (si choisi)
                    if model_id:
                        set_status("Vérif. modèle %s…" % model_id, 0.08)
                        try:
                            chk = run_cmd([exe, "--output-mode", "json", "models", "list"])
                            listed = (chk.stdout or "") + (chk.stderr or "")
                            if model_id not in listed and chk.returncode == 0:
                                # essayer sortie texte
                                chk2 = run_cmd([exe, "models", "list"])
                                listed2 = (chk2.stdout or "") + (chk2.stderr or "")
                                if model_id not in listed and model_id not in listed2:
                                    msg = (
                                        "Le modèle « %s » n'apparaît pas dans\n"
                                        "  birda models list\n\n"
                                        "Installez-le d'abord :\n"
                                        "  birda models install %s\n\n"
                                        "Puis cliquez « Rafraîchir listé » dans cette fenêtre."
                                        % (model_id, model_id)
                                    )
                                    win.after(0, lambda m=msg: messagebox.showerror("Modèle manquant", m))
                                    win.after(0, lambda: btn_run.configure(state="normal", text="▶  Analyser puis importer"))
                                    set_status("Modèle non installé : %s" % model_id, 0)
                                    return
                        except Exception as e:
                            self.log("Vérif modèle : %s (on continue)" % e)

                    proc = run_cmd(cmd_geo)
                    out = (proc.stdout or "") + "\n" + (proc.stderr or "")
                    # Si --lat/--lon non supportés → relancer sans (garder -m)
                    if proc.returncode != 0 and cmd_geo != cmd_base:
                        err_l = out.lower()
                        if any(x in err_l for x in ("lat", "lon", "unknown option", "unrecognized", "unexpected argument")):
                            set_status("Relance sans filtre GPS (modèle %s)…" % model_label, 0.15)
                            self.log("Birda CLI : retry sans --lat/--lon, modèle=%s" % model_label)
                            proc = run_cmd(cmd_base)
                            out = (proc.stdout or "") + "\n" + (proc.stderr or "")

                    err_l = out.lower()
                    if proc.returncode != 0:
                        set_status("Erreur Birda (code %s)" % proc.returncode, 0)
                        hint = ""
                        if model_id and any(
                            x in err_l
                            for x in ("model", "not found", "unknown", "no such", "install")
                        ):
                            hint = (
                                "\n\n→ Installez le modèle :\n"
                                "  birda models install %s\n"
                                "Puis « Rafraîchir listé »."
                                % model_id
                            )
                        msg = (
                            "Birda a échoué (code %s).\n"
                            "Modèle demandé : %s\n\n"
                            "Sortie :\n%s%s"
                            % (proc.returncode, model_label, (out[-900:] if out else "(vide)"), hint)
                        )
                        win.after(0, lambda m=msg: messagebox.showerror("Birda CLI", m))
                        win.after(0, lambda: btn_run.configure(state="normal", text="▶  Analyser puis importer"))
                        return

                    set_status("Recherche CSV (modèle %s)…" % model_label, 0.7)
                    csvs = find_csv_outputs(run_out, audio)
                    # Ne garder que les CSV produits pendant cette analyse
                    fresh = []
                    for p in csvs:
                        try:
                            if os.path.getmtime(p) >= (t0 - 5):
                                fresh.append(p)
                        except Exception:
                            pass
                    if not fresh:
                        # fallback : tout CSV du sous-dossier modèle
                        fresh = csvs
                    if not fresh:
                        set_status("Aucun CSV trouvé", 0)
                        msg = (
                            "Birda OK (code 0) mais aucun CSV trouvé.\n"
                            "Modèle : %s\nDossier : %s\n\nSortie :\n%s"
                            % (model_label, run_out, out[-600:] if out else "(vide)")
                        )
                        win.after(0, lambda m=msg: messagebox.showerror("Birda CLI", m))
                        win.after(0, lambda: btn_run.configure(state="normal", text="▶  Analyser puis importer"))
                        return
                    csv_path = fresh[0]
                    set_status("CSV (%s) : %s" % (model_label, os.path.basename(csv_path)), 0.9)
                    self.log("Birda CLI CSV [%s] : %s" % (model_label, csv_path))

                    def after_ok():
                        btn_run.configure(state="normal", text="▶  Analyser puis importer")
                        set_status("Import carnet — modèle %s" % model_label, 1.0)
                        try:
                            win.destroy()
                        except Exception:
                            pass
                        self.import_chirpity_csv(path=csv_path)
                        try:
                            self._ecoute_session_meta = dict(self._ecoute_session_meta or {})
                            self._ecoute_session_meta["tool"] = "birda_cli"
                            self._ecoute_session_meta["model"] = model_id or "default"
                        except Exception:
                            pass

                    win.after(0, after_ok)
                except subprocess.TimeoutExpired:
                    win.after(0, lambda: messagebox.showerror("Birda CLI", "Délai dépassé (1 h)."))
                    win.after(0, lambda: btn_run.configure(state="normal", text="▶  Analyser puis importer"))
                    set_status("Timeout", 0)
                except Exception as e:
                    err = str(e)
                    win.after(0, lambda: messagebox.showerror("Birda CLI", err))
                    win.after(0, lambda: btn_run.configure(state="normal", text="▶  Analyser puis importer"))
                    set_status("Erreur", 0)

            threading.Thread(target=worker, daemon=True).start()

        # Actions
        actions = ctk.CTkFrame(win, fg_color="transparent")
        actions.pack(fill="x", padx=16, pady=12)
        btn_run = ctk.CTkButton(
            actions, text="▶  Analyser puis importer",
            height=40, font=ctk.CTkFont(size=13, weight="bold"),
            fg_color=UI.get("success", "#2f9e5f"),
            command=run_analysis,
        )
        btn_run.pack(side="left", padx=(0, 8))
        ctk.CTkButton(
            actions, text="Importer un CSV déjà produit…",
            height=36, fg_color=UI.get("card_alt"),
            command=lambda: (win.destroy(), self.import_chirpity_csv()),
        ).pack(side="left")
        ctk.CTkButton(actions, text="Fermer", height=36, width=90, command=win.destroy).pack(side="right")

    def open_import_son_menu(self):
        """Menu rapide Import son : BirdNET Live / Chirpity / Birda / outils."""
        win = ctk.CTkToplevel(self)
        win.title("Import son")
        win.geometry("420x400")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)
        ctk.CTkLabel(
            win, text="Importer une session d'écoute",
            font=ctk.CTkFont(size=15, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=18, pady=(16, 6))
        ctk.CTkLabel(
            win,
            text="Choisissez la source des détections audio.",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
        ).pack(anchor="w", padx=18, pady=(0, 10))

        def go(fn):
            win.destroy()
            fn()

        for text, cmd, col in (
            ("Birda CLI — analyser un audio…", self.open_birda_cli_analyzer, UI.get("success", "#2f9e5f")),
            ("BirdNET Live (metadata + selections)", self.import_birdnet_live_session, UI.get("accent")),
            ("CSV Chirpity", self.import_chirpity_csv, UI.get("purple", "#7c5cbf")),
            ("Birda GUI (base SQLite)", self.import_birda_from_catalog, UI.get("success", "#2f9e5f")),
            ("JSON / CSV Birda", self.import_birda_detections, UI.get("card_alt")),
            ("Session d'écoute guidée…", self.open_ecoute_session_wizard, UI.get("warning", "#f0b45c")),
            ("Outils audio (chemins, Audacity…)", self.open_birda_tools, UI.get("card_alt")),
        ):
            ctk.CTkButton(
                win, text=text, height=36, anchor="w",
                fg_color=col, command=lambda c=cmd: go(c),
            ).pack(fill="x", padx=18, pady=3)
        ctk.CTkButton(win, text="Fermer", fg_color=UI.get("card_alt"), command=win.destroy).pack(
            pady=(10, 12)
        )

    def open_birda_tools(self):
        """Page outils audio : chemins Audacity/Birda/Chirpity + imports."""
        win = ctk.CTkToplevel(self)
        win.title("Outils audio — import son")
        win.geometry("920x820")
        try:
            win.minsize(780, 640)
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Prise de son → import → Carnet",
            font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text="Configurez les chemins une fois, puis importez BirdNET Live / Chirpity / Birda.",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            justify="left",
        ).pack(anchor="w", padx=16, pady=(0, 6))

        # Ligne 1 : ouvrir les applis
        top_btns = ctk.CTkFrame(win, fg_color="transparent")
        top_btns.pack(fill="x", padx=16, pady=(0, 4))
        ctk.CTkButton(
            top_btns, text="Ouvrir Audacity", height=36, width=150,
            command=lambda: self._launch_external(
                (self._load_app_config() or {}).get("audacity_exe")
                or self._find_app_exe(["audacity.exe", "Audacity.exe"], "audacity_exe"),
                "Audacity",
            ),
        ).pack(side="left", padx=(0, 6))
        ctk.CTkButton(
            top_btns, text="Ouvrir Birda GUI", height=36, width=150,
            fg_color=UI.get("accent"),
            command=lambda: self._launch_external(
                (self._load_app_config() or {}).get("birda_gui_exe")
                or self._find_app_exe(
                    ["birda-gui.exe", "Birda GUI.exe", "Birda.exe", "birda.exe"],
                    "birda_gui_exe",
                ),
                "Birda GUI",
            ),
        ).pack(side="left", padx=4)
        ctk.CTkButton(
            top_btns, text="Ouvrir Chirpity", height=36, width=150,
            fg_color=UI.get("purple", "#7c5cbf"),
            command=lambda: self._launch_external(
                (self._load_app_config() or {}).get("chirpity_exe")
                or self._find_app_exe(
                    ["Chirpity.exe", "chirpity.exe", "Chirpity"],
                    "chirpity_exe",
                ),
                "Chirpity",
            ),
        ).pack(side="left", padx=4)

        # Ligne 2 : imports (tous visibles)
        top_import = ctk.CTkFrame(win, fg_color="transparent")
        top_import.pack(fill="x", padx=16, pady=(4, 2))
        ctk.CTkButton(
            top_import, text="Birda CLI…",
            height=38, width=140,
            fg_color=UI.get("success", "#2f9e5f"),
            command=self.open_birda_cli_analyzer,
        ).pack(side="left", padx=(0, 6))
        ctk.CTkButton(
            top_import, text="BirdNET Live…",
            height=38, width=160,
            fg_color=UI.get("accent"),
            command=self.import_birdnet_live_session,
        ).pack(side="left", padx=4)
        ctk.CTkButton(
            top_import, text="CSV Chirpity…",
            height=38, width=150,
            fg_color=UI.get("purple", "#7c5cbf"),
            command=self.import_chirpity_csv,
        ).pack(side="left", padx=4)
        ctk.CTkButton(
            top_import, text="Birda SQLite → carnet",
            height=38, width=170,
            fg_color=UI.get("success", "#2f9e5f"),
            command=self.import_birda_from_catalog,
        ).pack(side="left", padx=4)
        ctk.CTkButton(
            top_import, text="JSON/CSV Birda…",
            height=38, width=150,
            command=self.import_birda_detections,
        ).pack(side="left", padx=4)

        top_import2 = ctk.CTkFrame(win, fg_color="transparent")
        top_import2.pack(fill="x", padx=16, pady=(2, 8))
        ctk.CTkButton(
            top_import2, text="Session d'écoute guidée…",
            height=36, width=200,
            fg_color=UI.get("warning", "#f0b45c"),
            command=lambda: (win.destroy(), self.open_ecoute_session_wizard()),
        ).pack(side="left", padx=(0, 6))
        ctk.CTkButton(
            top_import2, text="Obs. manuelle prise de son",
            height=36, width=200,
            fg_color=UI.get("card_alt"),
            command=lambda: (win.destroy(), self.open_manual_sound_observation()),
        ).pack(side="left", padx=4)

        scroll = ctk.CTkScrollableFrame(win, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=12, pady=4)

        # —— Chemins programmes ——
        paths_fr = ctk.CTkFrame(scroll, fg_color=UI.get("card"), corner_radius=12)
        paths_fr.pack(fill="x", padx=4, pady=6)
        ctk.CTkLabel(
            paths_fr, text="Chemins des programmes",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=12, pady=(10, 6))

        cfg0 = self._load_app_config() or {}
        aud_var = tk.StringVar(
            value=cfg0.get("audacity_exe")
            or self._find_app_exe(["audacity.exe", "Audacity.exe", "audacity"], "audacity_exe")
            or ""
        )
        bir_var = tk.StringVar(
            value=cfg0.get("birda_gui_exe")
            or self._find_app_exe(
                ["birda-gui.exe", "Birda GUI.exe", "Birda.exe", "birda.exe", "birda-gui", "birda"],
                "birda_gui_exe",
            )
            or ""
        )
        chirp_var = tk.StringVar(
            value=cfg0.get("chirpity_exe")
            or self._find_app_exe(
                ["Chirpity.exe", "chirpity.exe", "Chirpity"],
                "chirpity_exe",
            )
            or ""
        )
        db_var = tk.StringVar(
            value=cfg0.get("birda_catalog_db") or self._find_birda_catalog_db() or ""
        )

        def open_audacity_now():
            p = aud_var.get().strip()
            if p:
                self._set_configured_exe("audacity_exe", p)
            self._launch_external(
                p or self._find_app_exe(["audacity.exe", "Audacity.exe", "audacity"], "audacity_exe"),
                "Audacity",
            )

        def open_birda_now():
            p = bir_var.get().strip()
            if p:
                self._set_configured_exe("birda_gui_exe", p)
            self._launch_external(
                p or self._find_app_exe(
                    ["birda-gui.exe", "Birda GUI.exe", "Birda.exe", "birda.exe"],
                    "birda_gui_exe",
                ),
                "Birda GUI",
            )

        def open_chirpity_now():
            p = chirp_var.get().strip()
            if p:
                self._set_configured_exe("chirpity_exe", p)
            self._launch_external(
                p or self._find_app_exe(
                    ["Chirpity.exe", "chirpity.exe", "Chirpity"],
                    "chirpity_exe",
                ),
                "Chirpity",
            )

        def row_path(parent, label, var, filetypes, open_cmd=None):
            ctk.CTkLabel(parent, text=label, text_color=UI.get("text")).pack(anchor="w", padx=12)
            r = ctk.CTkFrame(parent, fg_color="transparent")
            r.pack(fill="x", padx=12, pady=(0, 8))
            ctk.CTkEntry(r, textvariable=var).pack(side="left", fill="x", expand=True)

            def browse():
                p = filedialog.askopenfilename(title=label, filetypes=filetypes)
                if p:
                    var.set(p)

            ctk.CTkButton(r, text="Parcourir…", width=90, command=browse).pack(side="left", padx=4)
            if open_cmd is not None:
                ctk.CTkButton(
                    r, text="Ouvrir", width=72,
                    fg_color=UI.get("accent"),
                    command=open_cmd,
                ).pack(side="left", padx=2)

        row_path(
            paths_fr, "Audacity (.exe)",
            aud_var,
            [("Executable", "*.exe"), ("Tous", "*.*")],
            open_cmd=open_audacity_now,
        )
        row_path(
            paths_fr, "Birda GUI (.exe)",
            bir_var,
            [("Executable", "*.exe"), ("Tous", "*.*")],
            open_cmd=open_birda_now,
        )
        row_path(
            paths_fr, "Chirpity (.exe)",
            chirp_var,
            [("Executable", "*.exe"), ("Tous", "*.*")],
            open_cmd=open_chirpity_now,
        )
        row_path(
            paths_fr, "Base Birda GUI (birda-catalog.db) — optionnel",
            db_var,
            [("SQLite", "*.db"), ("Tous", "*.*")],
            open_cmd=None,
        )

        def save_paths():
            self._set_configured_exe("audacity_exe", aud_var.get().strip())
            self._set_configured_exe("birda_gui_exe", bir_var.get().strip())
            self._set_configured_exe("chirpity_exe", chirp_var.get().strip())
            cfg = dict(self._load_app_config() or {})
            cfg["birda_catalog_db"] = db_var.get().strip()
            try:
                self._save_app_config(cfg)
                self.app_config = cfg
            except Exception:
                pass
            messagebox.showinfo("Chemins", "Chemins enregistrés.")
            self.log("Chemins Audacity / Birda enregistrés.")

        ctk.CTkButton(
            paths_fr, text="Enregistrer les chemins", height=32,
            fg_color=UI.get("accent"), command=save_paths,
        ).pack(anchor="w", padx=12, pady=(0, 10))

        # —— Info stockage Birda GUI ——
        info = ctk.CTkFrame(scroll, fg_color=UI.get("card_alt"), corner_radius=10)
        info.pack(fill="x", padx=4, pady=6)
        catalog = self._find_birda_catalog_db()
        ctk.CTkLabel(
            info,
            text=(
                "Où Birda GUI stocke les détections\n\n"
                "• Principalement dans une base SQLite : birda-catalog.db\n"
                "  Emplacement typique Windows :\n"
                "  %APPDATA%\\birda-gui\\birda-catalog.db\n"
                "  (ou %LOCALAPPDATA%\\birda-gui\\…)\n\n"
                "• Détecté maintenant : %s\n\n"
                "• Des JSON/CSV n’apparaissent que si vous exportez\n"
                "  ou utilisez la CLI birda -f json / csv."
            ) % (catalog or "non trouvé — utilisez Parcourir ci-dessus"),
            font=ctk.CTkFont(size=11), text_color=UI.get("text"),
            justify="left",
        ).pack(anchor="w", padx=12, pady=10)

        # —— Dossier JSON/CSV (optionnel) ——
        watch_fr = ctk.CTkFrame(scroll, fg_color=UI.get("card"), corner_radius=12)
        watch_fr.pack(fill="x", padx=4, pady=6)
        ctk.CTkLabel(
            watch_fr, text="Dossier JSON/CSV (export manuel ou CLI)",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=12, pady=(10, 4))
        watch_var = tk.StringVar(value=self._birda_watch_folder())
        row_w = ctk.CTkFrame(watch_fr, fg_color="transparent")
        row_w.pack(fill="x", padx=12, pady=(0, 8))
        ctk.CTkEntry(row_w, textvariable=watch_var).pack(side="left", fill="x", expand=True)

        def browse_watch():
            d = filedialog.askdirectory(title="Dossier résultats JSON/CSV", initialdir=watch_var.get() or None)
            if d:
                watch_var.set(d)
                self._set_birda_watch_folder(d)

        ctk.CTkButton(row_w, text="…", width=36, command=browse_watch).pack(side="left", padx=4)

        # —— Actions ——
        bf = ctk.CTkFrame(scroll, fg_color=UI.get("card"), corner_radius=12)
        bf.pack(fill="x", padx=4, pady=6)
        ctk.CTkLabel(
            bf, text="Actions", font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=12, pady=(10, 6))

        def open_aud():
            save_paths()
            p = aud_var.get().strip() or self._find_app_exe(
                ["audacity.exe", "Audacity.exe", "audacity"], "audacity_exe"
            )
            self._launch_external(p, "Audacity")

        def open_bir():
            save_paths()
            p = bir_var.get().strip() or self._find_app_exe(
                ["birda-gui.exe", "Birda GUI.exe", "Birda.exe", "birda.exe"], "birda_gui_exe"
            )
            self._launch_external(p, "Birda GUI")

        ctk.CTkButton(bf, text="1. Ouvrir Audacity", height=34, command=open_aud).pack(
            fill="x", padx=12, pady=3
        )
        ctk.CTkButton(
            bf, text="2. Ouvrir Birda GUI", height=34,
            fg_color=UI.get("accent"), command=open_bir,
        ).pack(fill="x", padx=12, pady=3)
        ctk.CTkButton(
            bf, text="3. Importer depuis la base Birda GUI (SQLite)", height=36,
            fg_color=UI.get("success", "#2f9e5f"),
            command=lambda: (save_paths(), self.import_birda_from_catalog()),
        ).pack(fill="x", padx=12, pady=3)
        ctk.CTkButton(
            bf, text="Importer le dernier JSON/CSV (semi-auto)", height=32,
            command=lambda: (
                self._set_birda_watch_folder(watch_var.get().strip()),
                self.import_latest_birda_result(),
            ),
        ).pack(fill="x", padx=12, pady=3)
        ctk.CTkButton(
            bf, text="Choisir un fichier JSON / CSV…", height=32,
            fg_color=UI.get("card_alt"),
            command=lambda: self.import_birda_detections(),
        ).pack(fill="x", padx=12, pady=3)
        ctk.CTkButton(
            bf, text="Observation manuelle « Prise de son »…", height=32,
            fg_color=UI.get("card_alt"),
            command=lambda: (win.destroy(), self.open_manual_sound_observation()),
        ).pack(fill="x", padx=12, pady=(3, 12))

        ctk.CTkButton(win, text="Fermer", command=win.destroy).pack(pady=8)

    def open_manual_sound_observation(self):
        """Raccourci : dialogue observation sans photo prérempli pour prise de son."""
        if not self.photo_folder_path:
            messagebox.showwarning(
                "Sortie",
                "Ouvrez d'abord une sortie (dossier photos ou carnet sans photos).",
            )
            return
        # Réutilise le dialogue manuel s'il accepte des défauts ; sinon ouvre et guide
        try:
            self.open_manual_observation_dialog()
            messagebox.showinfo(
                "Prise de son",
                "Dans le formulaire :\n"
                "• Catégorie : Oiseau (souvent)\n"
                "• Type d'indice : « Prise de son (Birda) » ou « Entendu / Chant »\n"
                "• Indiquez l'espèce après Birda, l'heure et le lieu\n"
                "• Enregistrez\n\n"
                "Ou utilisez « Importer détections Birda » pour remplir plusieurs espèces d'un coup.",
            )
        except Exception as e:
            messagebox.showerror("Prise de son", str(e))

    def import_birda_from_catalog(self):
        """Importe les détections depuis birda-catalog.db (base SQLite de Birda GUI)."""
        if not self.photo_folder_path:
            messagebox.showwarning(
                "Sortie",
                "Ouvrez d'abord la sortie dans laquelle enregistrer les détections.",
            )
            return
        cfg = self._load_app_config() or {}
        db_path = (cfg.get("birda_catalog_db") or "").strip() or self._find_birda_catalog_db()
        if not db_path or not os.path.isfile(db_path):
            messagebox.showinfo(
                "Base Birda GUI",
                "Fichier birda-catalog.db introuvable.\n\n"
                "1) Lancez au moins une analyse dans Birda GUI\n"
                "2) Dans GeoExif → Birda → Parcourir… vers :\n"
                "   %APPDATA%\\birda-gui\\birda-catalog.db\n"
                "   (Explorateur : tapez %APPDATA%\\birda-gui dans la barre d'adresse)\n"
                "3) Enregistrer les chemins, puis réessayez",
            )
            return

        try:
            import sqlite3
        except Exception:
            messagebox.showerror("Birda", "Module sqlite3 indisponible.")
            return

        rows = []
        try:
            conn = sqlite3.connect(db_path)
            conn.row_factory = sqlite3.Row
            cur = conn.cursor()
            # Schéma Birda GUI : tables detections (+ éventuels noms alternatifs)
            tables = [r[0] for r in cur.execute(
                "SELECT name FROM sqlite_master WHERE type='table'"
            ).fetchall()]
            if "detections" not in tables:
                messagebox.showerror(
                    "Birda",
                    "Table « detections » absente dans :\n%s\n\nTables : %s"
                    % (db_path, ", ".join(tables) or "(aucune)"),
                )
                conn.close()
                return
            # Colonnes souples
            cols = [r[1] for r in cur.execute("PRAGMA table_info(detections)").fetchall()]
            colset = {c.lower(): c for c in cols}

            def col(*names):
                for n in names:
                    if n.lower() in colset:
                        return colset[n.lower()]
                return None

            c_common = col("common_name", "commonName", "common")
            c_sci = col("scientific_name", "scientificName", "scientific")
            c_conf = col("confidence", "score", "probability")
            c_start = col("start_time", "start", "begin_time")
            c_end = col("end_time", "end")
            c_when = col("created_at", "timestamp", "detected_at", "date")
            c_file = col("source_file", "audio_path", "filename", "file")

            # Dernières détections (limite raisonnable)
            order = c_when or "rowid"
            sql = "SELECT * FROM detections ORDER BY %s DESC LIMIT 500" % order
            for r in cur.execute(sql).fetchall():
                d = dict(r)
                common = (d.get(c_common) if c_common else None) or ""
                sci = (d.get(c_sci) if c_sci else None) or ""
                if not common and not sci:
                    continue
                conf = d.get(c_conf) if c_conf else None
                rows.append({
                    "common": str(common).strip(),
                    "sci": str(sci).strip(),
                    "conf": conf,
                    "start": d.get(c_start) if c_start else None,
                    "end": d.get(c_end) if c_end else None,
                    "source_audio": str(d.get(c_file) or "") if c_file else "",
                })
            conn.close()
        except Exception as e:
            messagebox.showerror("Birda SQLite", "Lecture impossible :\n%s\n\nFichier : %s" % (e, db_path))
            return

        if not rows:
            messagebox.showinfo(
                "Birda",
                "Aucune détection dans la base.\n"
                "Analysez d'abord un audio dans Birda GUI.\n\n%s" % db_path,
            )
            return

        if not messagebox.askyesno(
            "Import Birda GUI",
            "%d détection(s) lues dans :\n%s\n\n"
            "Importer vers le carnet de la sortie ouverte\n"
            "(espèces agrégées, type Prise de son) ?"
            % (len(rows), db_path),
        ):
            return

        # Réutilise la même logique d'agrégation que l'import fichier
        self._birda_rows_to_notes(rows, source_label=os.path.basename(db_path))

    def _match_oiseau_fr(self, common, sci):
        """Rattache un nom EN/latin (Chirpity/Birda) à un libellé FR (liste Meuse + table EN→FR)."""
        oiseaux = list(self.faune_meuse.get("Oiseau") or [])
        common = (common or "").strip()
        sci = (sci or "").strip()
        low_c = common.lower()
        low_s = sci.lower()

        # 1) déjà FR dans la liste
        if common in oiseaux:
            return common
        for esp in oiseaux:
            el = esp.lower()
            if low_c and (low_c == el or low_c in el or el in low_c):
                return esp

        # 2) table anglais courant → français (Chirpity)
        EN_FR = {
            "eurasian blackcap": "Fauvette à tête noire",
            "blackcap": "Fauvette à tête noire",
            "common blackbird": "Merle noir",
            "eurasian blackbird": "Merle noir",
            "blackbird": "Merle noir",
            "european robin": "Rougegorge familier",
            "robin": "Rougegorge familier",
            "great tit": "Mésange charbonnière",
            "blue tit": "Mésange bleue",
            "coal tit": "Mésange noire",
            "long-tailed tit": "Mésange à longue queue",
            "common chaffinch": "Pinson des arbres",
            "chaffinch": "Pinson des arbres",
            "european greenfinch": "Verdier d'Europe",
            "european goldfinch": "Chardonneret élégant",
            "goldfinch": "Chardonneret élégant",
            "common linnet": "Linotte mélodieuse",
            "eurasian wren": "Troglodyte mignon",
            "wren": "Troglodyte mignon",
            "common chiffchaff": "Pouillot véloce",
            "chiffchaff": "Pouillot véloce",
            "willow warbler": "Pouillot fitis",
            "garden warbler": "Fauvette des jardins",
            "common whitethroat": "Fauvette grisette",
            "lesser whitethroat": "Fauvette babillarde",
            "song thrush": "Grive musicienne",
            "mistle thrush": "Grive draine",
            "redwing": "Grive mauvis",
            "fieldfare": "Grive litorne",
            "common wood pigeon": "Pigeon ramier",
            "woodpigeon": "Pigeon ramier",
            "wood pigeon": "Pigeon ramier",
            "common buzzard": "Buse variable",
            "eurasian jay": "Geai des chênes",
            "jay": "Geai des chênes",
            "carrion crow": "Corneille noire",
            "rook": "Freux",
            "eurasian magpie": "Pie bavarde",
            "magpie": "Pie bavarde",
            "common starling": "Étourneau sansonnet",
            "starling": "Étourneau sansonnet",
            "house sparrow": "Moineau domestique",
            "eurasian tree sparrow": "Moineau friquet",
            "common swift": "Martinet noir",
            "barn swallow": "Hirondelle rustique",
            "common house martin": "Hirondelle de fenêtre",
            "house martin": "Hirondelle de fenêtre",
            "sand martin": "Hirondelle de rivage",
            "common cuckoo": "Coucou gris",
            "cuckoo": "Coucou gris",
            "great spotted woodpecker": "Pic épeiche",
            "green woodpecker": "Pic vert",
            "black woodpecker": "Pic noir",
            "eurasian nuthatch": "Sittelle torchepot",
            "nuthatch": "Sittelle torchepot",
            "eurasian treecreeper": "Grimpereau des jardins",
            "short-toed treecreeper": "Grimpereau des jardins",
            "common firecrest": "Roitelet à triple bandeau",
            "goldcrest": "Roitelet huppé",
            "winter wren": "Troglodyte mignon",
            "eurasian skylark": "Alouette des champs",
            "skylark": "Alouette des champs",
            "yellowhammer": "Bruant jaune",
            "common reed bunting": "Bruant des roseaux",
            "reed bunting": "Bruant des roseaux",
            "corn bunting": "Bruant proyer",
            "eurasian collared dove": "Tourterelle turque",
            "collared dove": "Tourterelle turque",
            "stock dove": "Pigeon colombin",
            "common kestrel": "Faucon crécerelle",
            "kestrel": "Faucon crécerelle",
            "sparrowhawk": "Épervier d'Europe",
            "eurasian sparrowhawk": "Épervier d'Europe",
            "grey heron": "Héron cendré",
            "gray heron": "Héron cendré",
            "white stork": "Cigogne blanche",
            "common crane": "Grue cendrée",
            "mallard": "Canard colvert",
            "common kingfisher": "Martin-pêcheur d'Europe",
            "kingfisher": "Martin-pêcheur d'Europe",
            "eurasian penduline tit": "Rémiz penduline",
            "penduline tit": "Rémiz penduline",
            "marsh tit": "Mésange nonnette",
            "willow tit": "Mésange boréale",
            "crested tit": "Mésange huppée",
            "spotted flycatcher": "Gobemouche gris",
            "european pied flycatcher": "Gobemouche noir",
            "pied flycatcher": "Gobemouche noir",
            "common redstart": "Rougequeue à front blanc",
            "black redstart": "Rougequeue noir",
            "northern wheatear": "Traquet motteux",
            "whinchat": "Tarier des prés",
            "european stonechat": "Tarier pâtre",
            "dunnock": "Accenteur mouchet",
            "hedge accentor": "Accenteur mouchet",
            "white wagtail": "Bergeronnette grise",
            "grey wagtail": "Bergeronnette des ruisseaux",
            "gray wagtail": "Bergeronnette des ruisseaux",
            "tree pipit": "Pipit des arbres",
            "meadow pipit": "Pipit farlouse",
            "common nightingale": "Rossignol philomèle",
            "nightingale": "Rossignol philomèle",
            "icterine warbler": "Hypolaïs ictérine",
            "melodious warbler": "Hypolaïs polyglotte",
            "sedge warbler": "Phragmite des joncs",
            "reed warbler": "Rousserolle effarvatte",
            "eurasian reed warbler": "Rousserolle effarvatte",
            "marsh warbler": "Rousserolle verderolle",
            "great reed warbler": "Rousserolle turdoïde",
            "tawny owl": "Chouette hulotte",
            "long-eared owl": "Hibou moyen-duc",
            "barn owl": "Effraie des clochers",
            "common raven": "Grand corbeau",
            "raven": "Grand corbeau",
            "hooded crow": "Corneille mantelée",
            "jackdaw": "Choucas des tours",
            "western jackdaw": "Choucas des tours",
            "common nightjar": "Engoulevent d'Europe",
            "nightjar": "Engoulevent d'Europe",
            "european honey buzzard": "Bondrée apivore",
            "red kite": "Milan royal",
            "black kite": "Milan noir",
            "common quail": "Caille des blés",
            "eurasian woodcock": "Bécasse des bois",
            "woodcock": "Bécasse des bois",
        }
        if low_c in EN_FR:
            return EN_FR[low_c]
        # sous-chaîne (ex. "Eurasian Blackcap")
        for en, fr in EN_FR.items():
            if en in low_c or low_c in en:
                return fr

        # 3) scientifique connu dans libellés (rare)
        if low_s:
            for esp in oiseaux:
                if low_s in esp.lower():
                    return esp
            SCI_FR = {
                "turdus merula": "Merle noir",
                "turdus philomelos": "Grive musicienne",
                "turdus viscivorus": "Grive draine",
                "turdus iliacus": "Grive mauvis",
                "turdus pilaris": "Grive litorne",
                "erithacus rubecula": "Rougegorge familier",
                "parus major": "Mésange charbonnière",
                "cyanistes caeruleus": "Mésange bleue",
                "periparus ater": "Mésange noire",
                "poecile palustris": "Mésange nonnette",
                "poecile montanus": "Mésange boréale",
                "lophophanes cristatus": "Mésange huppée",
                "aegithalos caudatus": "Mésange à longue queue",
                "fringilla coelebs": "Pinson des arbres",
                "fringilla montifringilla": "Pinson du Nord",
                "chloris chloris": "Verdier d'Europe",
                "carduelis carduelis": "Chardonneret élégant",
                "linaria cannabina": "Linotte mélodieuse",
                "spinus spinus": "Tarin des aulnes",
                "pyrrhula pyrrhula": "Bouvreuil pivoine",
                "coccothraustes coccothraustes": "Gros-bec casse-noyaux",
                "troglodytes troglodytes": "Troglodyte mignon",
                "phylloscopus collybita": "Pouillot véloce",
                "phylloscopus trochilus": "Pouillot fitis",
                "phylloscopus sibilatrix": "Pouillot siffleur",
                "sylvia atricapilla": "Fauvette à tête noire",
                "curruca atricapilla": "Fauvette à tête noire",
                "sylvia borin": "Fauvette des jardins",
                "curruca communis": "Fauvette grisette",
                "sylvia communis": "Fauvette grisette",
                "curruca curruca": "Fauvette babillarde",
                "remiz pendulinus": "Rémiz penduline",
                "buteo buteo": "Buse variable",
                "pica pica": "Pie bavarde",
                "garrulus glandarius": "Geai des chênes",
                "corvus corone": "Corneille noire",
                "corvus frugilegus": "Freux",
                "corvus corax": "Grand corbeau",
                "coloeus monedula": "Choucas des tours",
                "columba palumbus": "Pigeon ramier",
                "columba oenas": "Pigeon colombin",
                "streptopelia decaocto": "Tourterelle turque",
                "hirundo rustica": "Hirondelle rustique",
                "delichon urbicum": "Hirondelle de fenêtre",
                "riparia riparia": "Hirondelle de rivage",
                "apus apus": "Martinet noir",
                "cuculus canorus": "Coucou gris",
                "regulus regulus": "Roitelet huppé",
                "regulus ignicapilla": "Roitelet à triple bandeau",
                "sitta europaea": "Sittelle torchepot",
                "certhia brachydactyla": "Grimpereau des jardins",
                "certhia familiaris": "Grimpereau des bois",
                "dendrocopos major": "Pic épeiche",
                "dendrocopos medius": "Pic mar",
                "dendrocopos minor": "Pic épeichette",
                "picus viridis": "Pic vert",
                "dryocopus martius": "Pic noir",
                "jyux torquilla": "Torcol fourmilier",
                "jynx torquilla": "Torcol fourmilier",
                "sturnus vulgaris": "Étourneau sansonnet",
                "passer domesticus": "Moineau domestique",
                "passer montanus": "Moineau friquet",
                "emberiza citrinella": "Bruant jaune",
                "emberiza schoeniclus": "Bruant des roseaux",
                "emberiza calandra": "Bruant proyer",
                "emberiza cirlus": "Bruant zizi",
                "alauda arvensis": "Alouette des champs",
                "galerida cristata": "Cochevis huppé",
                "lullula arborea": "Alouette lulu",
                "milvus milvus": "Milan royal",
                "milvus migrans": "Milan noir",
                "pernis apivorus": "Bondrée apivore",
                "accipiter nisus": "Épervier d'Europe",
                "accipiter gentilis": "Autour des palombes",
                "falco tinnunculus": "Faucon crécerelle",
                "falco subbuteo": "Faucon hobereau",
                "ardea cinerea": "Héron cendré",
                "ardea alba": "Grande Aigrette",
                "egretta garzetta": "Aigrette garzette",
                "ciconia ciconia": "Cigogne blanche",
                "grus grus": "Grue cendrée",
                "strix aluco": "Chouette hulotte",
                "asio otus": "Hibou moyen-duc",
                "asio flammeus": "Hibou des marais",
                "tyto alba": "Effraie des clochers",
                "athene noctua": "Chevêche d'Athéna",
                "caprimulgus europaeus": "Engoulevent d'Europe",
                "alcedo atthis": "Martin-pêcheur d'Europe",
                "upupa epops": "Huppe fasciée",
                "oriolus oriolus": "Loriot d'Europe",
                "muscicapa striata": "Gobemouche gris",
                "ficedula hypoleuca": "Gobemouche noir",
                "phoenicurus phoenicurus": "Rougequeue à front blanc",
                "phoenicurus ochruros": "Rougequeue noir",
                "saxicola rubetra": "Tarier des prés",
                "saxicola rubicola": "Tarier pâtre",
                "oenanthe oenanthe": "Traquet motteux",
                "prunella modularis": "Accenteur mouchet",
                "motacilla alba": "Bergeronnette grise",
                "motacilla cinerea": "Bergeronnette des ruisseaux",
                "motacilla flava": "Bergeronnette printanière",
                "anthus trivialis": "Pipit des arbres",
                "anthus pratensis": "Pipit farlouse",
                "luscinia megarhynchos": "Rossignol philomèle",
                "hippolais icterina": "Hypolaïs ictérine",
                "hippolais polyglotta": "Hypolaïs polyglotte",
                "acrocephalus schoenobaenus": "Phragmite des joncs",
                "acrocephalus scirpaceus": "Rousserolle effarvatte",
                "acrocephalus palustris": "Rousserolle verderolle",
                "acrocephalus arundinaceus": "Rousserolle turdoïde",
                "locustella naevia": "Locustelle tachetée",
                "coturnix coturnix": "Caille des blés",
                "scolopax rusticola": "Bécasse des bois",
                "gallinago gallinago": "Bécassine des marais",
                "vanellus vanellus": "Vanneau huppé",
                "charadrius dubius": "Petit Gravelot",
                "fulica atra": "Foulque macroule",
                "gallinula chloropus": "Gallinule poule-d'eau",
                "anas platyrhynchos": "Canard colvert",
                "anas crecca": "Sarcelle d'hiver",
                "tachybaptus ruficollis": "Grèbe castagneux",
                "podiceps cristatus": "Grèbe huppé",
                "phasianus colchicus": "Faisan de Colchide",
                "perdix perdix": "Perdrix grise",
                "lanius collurio": "Pie-grièche écorcheur",
                "lanius excubitor": "Pie-grièche grise",
                "bombycilla garrulus": "Jaseur boréal",
                "coccothraustes coccothraustes": "Gros-bec casse-noyaux",
                "serinus serinus": "Serin cini",
                "pyrrhula pyrrhula": "Bouvreuil pivoine",
            }
            if low_s in SCI_FR:
                return SCI_FR[low_s]
            # genre seulement si une seule espèce locale courante
            parts = low_s.split()
            if len(parts) >= 2:
                genus = parts[0]
                hits = [fr for s, fr in SCI_FR.items() if s.startswith(genus + " ")]
                if len(hits) == 1:
                    return hits[0]

        return common or sci or "Oiseau"

    def _birda_rows_to_notes(self, detections, source_label="birda"):
        """Agrège + dialogue de validation avant écriture dans le carnet."""
        if not self.photo_folder_path:
            return
        by_sp = {}
        for d in detections:
            key = (d.get("sci") or d.get("common") or "?").strip().lower()
            conf = d.get("conf")
            try:
                conf_f = float(conf) if conf is not None and str(conf) != "" else 0.0
            except Exception:
                conf_f = 0.0
            # conf peut être 0-1 ou 0-100
            if conf_f > 1.5:
                conf_f = conf_f / 100.0
            prev = by_sp.get(key)
            if prev is None or conf_f > prev[1]:
                by_sp[key] = (d, conf_f)

        # Dialogue compact : seuil + liste dense
        win = ctk.CTkToplevel(self)
        win.title("Valider les détections")
        win.geometry("480x420")
        try:
            win.minsize(400, 320)
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        head = ctk.CTkFrame(win, fg_color="transparent")
        head.pack(fill="x", padx=12, pady=(10, 2))
        ctk.CTkLabel(
            head, text="Espèces à importer",
            font=ctk.CTkFont(size=13, weight="bold"), text_color=UI.get("text"),
        ).pack(side="left")
        lbl_count = ctk.CTkLabel(head, text="", font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"))
        lbl_count.pack(side="right")

        seuil_var = tk.DoubleVar(value=0.35)
        row_s = ctk.CTkFrame(win, fg_color="transparent")
        row_s.pack(fill="x", padx=12, pady=2)
        ctk.CTkLabel(row_s, text="Seuil", width=42, text_color=UI.get("text")).pack(side="left")
        lbl_s = ctk.CTkLabel(row_s, text="35 %", width=40, text_color=UI.get("text_dim"))
        lbl_s.pack(side="right")
        ctk.CTkSlider(
            row_s, from_=0.1, to=0.9, variable=seuil_var, height=16,
            command=lambda v: rebuild_checks(),
        ).pack(side="left", fill="x", expand=True, padx=6)

        scroll = ctk.CTkScrollableFrame(win, fg_color=UI.get("card"), corner_radius=8)
        scroll.pack(fill="both", expand=True, padx=12, pady=6)
        check_vars = {}

        def rebuild_checks(*_a):
            for w in scroll.winfo_children():
                w.destroy()
            check_vars.clear()
            thr = float(seuil_var.get())
            lbl_s.configure(text="%d %%" % int(thr * 100))
            items = sorted(by_sp.items(), key=lambda x: -x[1][1])
            n_show = 0
            for key, (d, conf_f) in items:
                if conf_f < thr:
                    continue
                sci = (d.get("sci") or "").strip()
                common = (d.get("common") or "").strip()
                fr = self._match_oiseau_fr(common, sci)
                name = fr or common or sci or "?"
                # Ligne courte : Nom  72%  (sans nom scientifique si FR dispo)
                line = "%s   %d%%" % (name, int(round(conf_f * 100)))
                var = tk.BooleanVar(value=True)
                check_vars[key] = (var, d, conf_f, fr)
                ctk.CTkCheckBox(
                    scroll, text=line, variable=var,
                    text_color=UI.get("text"),
                    font=ctk.CTkFont(size=12),
                    height=22,
                ).pack(anchor="w", padx=6, pady=0)
                n_show += 1
            lbl_count.configure(text="%d / %d" % (n_show, len(by_sp)))

        rebuild_checks()

        def do_import():
            selected = [(d, conf_f, fr) for var, d, conf_f, fr in check_vars.values() if var.get()]
            if not selected:
                messagebox.showinfo("Import", "Aucune espèce cochée.")
                return
            win.destroy()
            self._write_birda_notes(selected, source_label)

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=12, pady=(2, 10))
        ctk.CTkButton(
            bf, text="Tout", width=56, height=28,
            command=lambda: [v[0].set(True) for v in check_vars.values()],
        ).pack(side="left", padx=1)
        ctk.CTkButton(
            bf, text="Rien", width=56, height=28,
            command=lambda: [v[0].set(False) for v in check_vars.values()],
        ).pack(side="left", padx=1)
        ctk.CTkButton(
            bf, text="Importer", height=30, width=110,
            fg_color=UI.get("success", "#2f9e5f"), command=do_import,
        ).pack(side="right", padx=2)
        ctk.CTkButton(bf, text="Annuler", height=30, width=80, command=win.destroy).pack(side="right")

    def _write_birda_notes(self, selected, source_label="birda"):
        """Écrit les détections validées (liste de (d, conf_f, fr))."""
        meta = getattr(self, "_ecoute_session_meta", None) or {}
        lat = meta.get("lat")
        lon = meta.get("lon")
        if lat is None or lon is None:
            for info in (self.photos_data or {}).values():
                if info.get("lat") is not None and info.get("lon") is not None:
                    try:
                        lat, lon = float(info["lat"]), float(info["lon"])
                        break
                    except Exception:
                        pass

        notes = self._load_notes_dict(self.photo_folder_path) or {}
        added = 0
        for d, conf_f, fr in selected:
            sci = (d.get("sci") or "").strip()
            common = self._match_oiseau_fr(fr or d.get("common") or "", sci)
            cle = "_birda_%s" % datetime.now().strftime("%Y%m%d%H%M%S%f")
            heure = (meta.get("heure") or "").strip()
            if not heure:
                try:
                    sec = float(d.get("start") or 0)
                    if sec > 0:
                        heure = "%02d:%02d" % (int(sec // 3600), int((sec % 3600) // 60))
                except Exception:
                    pass
            cert = "Possible"
            if conf_f >= 0.7:
                cert = "Probable"
            if conf_f >= 0.85:
                cert = "Sûr"
            # BirdNET Live détecte aussi mammifères (ex. Chevreuil)
            cat = "Oiseau"
            low_sci = (sci or "").lower()
            low_com = (common or "").lower()
            if any(x in low_sci for x in ("capreolus", "cervus", "sus scrofa", "vulpes", "meles", "martes", "sciurus")):
                cat = "Mammifère"
            elif any(x in low_com for x in ("chevreuil", "cerf", "sanglier", "renard", "blaireau", "martre", "écureuil", "ecureuil")):
                cat = "Mammifère"
            notes[cle] = {
                "departement": "55 - Meuse",
                "categorie": cat,
                "espece": common,
                "nombre": "1",
                "heure": heure,
                "lieu": (meta.get("lieu") or "").strip(),
                "notes_libres": (
                    "Import audio — confiance %.0f %% — source %s"
                    % (conf_f * 100, source_label)
                ),
                "type_observation": (
                    "Prise de son (BirdNET Live)" if str(source_label).startswith("birdnet_live")
                    else "Prise de son (Chirpity)" if str(source_label).startswith("chirpity")
                    else "Prise de son (Birda)"
                ),
                "type_indice": (
                    "Prise de son (BirdNET Live)" if str(source_label).startswith("birdnet_live")
                    else "Prise de son (Chirpity)" if str(source_label).startswith("chirpity")
                    else "Prise de son (Birda)"
                ),
                "comportement": "Vocalisation / chant",
                "certitude": cert,
                "nom_scientifique": sci,
                "sans_photo": True,
                "source": (
                    "birdnet_live" if str(source_label).startswith("birdnet_live")
                    else "chirpity" if str(source_label).startswith("chirpity")
                    else "birda"
                ),
                "birda_confidence": conf_f,
                "meteo": (meta.get("meteo_txt") or None),
            }
            if lat is not None and lon is not None:
                notes[cle]["lat"] = lat
                notes[cle]["lon"] = lon
            self.photos_data[cle] = {
                "path": "",
                "lat": lat,
                "lon": lon,
                "date": heure,
                "sans_photo": True,
                "is_video": False,
            }
            added += 1
            try:
                import time as _t
                _t.sleep(0.001)
            except Exception:
                pass

        try:
            self._save_notes_dict(notes, force_backup=True)
        except Exception as e:
            messagebox.showerror("Import Birda", str(e))
            return
        try:
            self.refresh_map_markers()
            self.refresh_daily_counter()
            self._filtered_media_names = None
            self._populate_listbox_from_filtered()
        except Exception:
            pass
        messagebox.showinfo(
            "Import audio",
            "%d espèce(s) dans le carnet.\n\n"
            "1) Cliquez une ligne 🎧 dans l'explorateur\n"
            "2) Corrigez le nom FR si besoin → Enregistrer\n"
            "3) Bouton « Point d'écoute GPS » ou Carte → placer le point"
            % added,
        )
        self.log("Import audio : %d espèces depuis %s" % (added, source_label))

    def toggle_point_ecoute_gps(self):
        """Mode carte : le prochain clic fixe le GPS de l'observation 🎧 / 📍 sélectionnée."""
        key = getattr(self, "_current_note_key", None)
        if not key:
            messagebox.showinfo(
                "Point d'écoute",
                "Sélectionnez d'abord une observation 🎧 ou 📍 dans le carnet.",
            )
            return
        self._point_ecoute_mode = True
        self._place_obs_mode = False
        try:
            self.btn_place_obs.configure(text="➕ Placer observation")
        except Exception:
            pass
        self.log(
            "Mode point d'écoute : cliquez sur la carte pour fixer le GPS de « %s »."
            % key
        )
        messagebox.showinfo(
            "Point d'écoute GPS",
            "Cliquez sur la carte à l'endroit de l'écoute.\n"
            "Le GPS sera enregistré sur l'observation sélectionnée.",
        )

    def open_ecoute_session_wizard(self):
        """Assistant guidé : dossier écoute du jour → lieu/date/heure → Birda/Chirpity → import carnet."""
        win = ctk.CTkToplevel(self)
        win.title("Session d'écoute — assistant")
        win.geometry("640x620")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Session d'écoute du jour",
            font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text="1) Dossier  2) Lieu / date / heure / GPS  3) Outil (Birda ou Chirpity)\n"
                 "4) Analyse  5) Retour ici → import validé → carnet",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            justify="left",
        ).pack(anchor="w", padx=16, pady=(0, 8))

        form = ctk.CTkScrollableFrame(win, fg_color=UI.get("card"), corner_radius=12)
        form.pack(fill="both", expand=True, padx=14, pady=6)

        # Dossier
        ctk.CTkLabel(form, text="Dossier session d'écoute", font=ctk.CTkFont(weight="bold"),
                     text_color=UI.get("text")).pack(anchor="w", padx=12, pady=(10, 2))
        folder_var = tk.StringVar(value=self.photo_folder_path or "")
        row_f = ctk.CTkFrame(form, fg_color="transparent")
        row_f.pack(fill="x", padx=12, pady=4)
        ctk.CTkEntry(row_f, textvariable=folder_var).pack(side="left", fill="x", expand=True)

        def browse_folder():
            d = filedialog.askdirectory(title="Dossier session d'écoute")
            if d:
                folder_var.set(d)

        def create_today_folder():
            base = filedialog.askdirectory(title="Parent du dossier (ex. disque sorties)")
            if not base:
                return
            name = "Ecoute_%s" % datetime.now().strftime("%Y-%m-%d")
            path = os.path.join(base, name)
            try:
                os.makedirs(path, exist_ok=True)
                folder_var.set(path)
            except Exception as e:
                messagebox.showerror("Dossier", str(e))

        ctk.CTkButton(row_f, text="…", width=36, command=browse_folder).pack(side="left", padx=3)
        ctk.CTkButton(row_f, text="Créer dossier du jour", width=140, command=create_today_folder).pack(side="left")

        # Lieu / date / heure
        ctk.CTkLabel(form, text="Contexte", font=ctk.CTkFont(weight="bold"),
                     text_color=UI.get("text")).pack(anchor="w", padx=12, pady=(12, 2))
        lieu_var = tk.StringVar()
        date_var = tk.StringVar(value=datetime.now().strftime("%Y-%m-%d"))
        heure_var = tk.StringVar(value=datetime.now().strftime("%H:%M"))
        lat_var = tk.StringVar()
        lon_var = tk.StringVar()

        def field(label, var, width=None):
            r = ctk.CTkFrame(form, fg_color="transparent")
            r.pack(fill="x", padx=12, pady=3)
            ctk.CTkLabel(r, text=label, width=90, anchor="w", text_color=UI.get("text")).pack(side="left")
            e = ctk.CTkEntry(r, textvariable=var, width=width or 280)
            e.pack(side="left", fill="x", expand=True)
            return e

        field("Lieu", lieu_var)
        field("Date", date_var)
        field("Heure", heure_var)
        row_gps = ctk.CTkFrame(form, fg_color="transparent")
        row_gps.pack(fill="x", padx=12, pady=3)
        ctk.CTkLabel(row_gps, text="GPS lat", width=90, anchor="w", text_color=UI.get("text")).pack(side="left")
        ctk.CTkEntry(row_gps, textvariable=lat_var, width=120).pack(side="left", padx=(0, 6))
        ctk.CTkLabel(row_gps, text="lon", text_color=UI.get("text")).pack(side="left")
        ctk.CTkEntry(row_gps, textvariable=lon_var, width=120).pack(side="left", padx=6)

        def use_map_center():
            try:
                pos = self.map_widget.get_position()
                lat_var.set("%.6f" % float(pos[0]))
                lon_var.set("%.6f" % float(pos[1]))
            except Exception:
                messagebox.showinfo("GPS", "Carte non disponible — saisissez lat/lon à la main.")

        ctk.CTkButton(row_gps, text="Centre carte", width=100, command=use_map_center).pack(side="left")

        # Outil
        ctk.CTkLabel(form, text="Outil d'analyse", font=ctk.CTkFont(weight="bold"),
                     text_color=UI.get("text")).pack(anchor="w", padx=12, pady=(12, 2))
        tool_var = tk.StringVar(value="chirpity")
        tools_fr = ctk.CTkFrame(form, fg_color="transparent")
        tools_fr.pack(fill="x", padx=12, pady=4)
        ctk.CTkRadioButton(tools_fr, text="Chirpity (CSV ensuite)", variable=tool_var, value="chirpity").pack(anchor="w", pady=2)
        ctk.CTkRadioButton(tools_fr, text="Birda GUI (base SQLite ensuite)", variable=tool_var, value="birda").pack(anchor="w", pady=2)

        session_meta = {}

        def apply_folder_and_meta():
            path = folder_var.get().strip()
            if not path:
                messagebox.showwarning("Session", "Indiquez ou créez un dossier.")
                return False
            if not os.path.isdir(path):
                try:
                    os.makedirs(path, exist_ok=True)
                except Exception as e:
                    messagebox.showerror("Dossier", str(e))
                    return False
            self.photo_folder_path = path
            try:
                self._register_known_folder(path)
            except Exception:
                pass
            try:
                if hasattr(self, "lbl_photo_folder"):
                    self.lbl_photo_folder.configure(text=path)
            except Exception:
                pass
            # Charger notes / liste carnet
            try:
                self.photos_data = {}
                notes = self._load_notes_dict(path) or {}
                for k, data in notes.items():
                    if isinstance(data, dict) and data.get("lat") is not None:
                        self.photos_data[k] = {
                            "path": "",
                            "lat": data.get("lat"),
                            "lon": data.get("lon"),
                            "date": data.get("heure") or "",
                            "sans_photo": True,
                        }
                self._filtered_media_names = None
                self._populate_listbox_from_filtered()
                self.refresh_map_markers()
                self.refresh_daily_counter()
            except Exception:
                pass
            lat = lon = None
            try:
                if lat_var.get().strip():
                    lat = float(lat_var.get().replace(",", "."))
                if lon_var.get().strip():
                    lon = float(lon_var.get().replace(",", "."))
            except Exception:
                lat = lon = None
            session_meta.clear()
            session_meta.update({
                "lieu": lieu_var.get().strip(),
                "date": date_var.get().strip(),
                "heure": heure_var.get().strip(),
                "lat": lat,
                "lon": lon,
                "tool": tool_var.get(),
                "folder": path,
            })
            self._ecoute_session_meta = dict(session_meta)
            return True

        def launch_tool_then_import():
            if not apply_folder_and_meta():
                return
            tool = tool_var.get()
            if tool == "chirpity":
                # tenter d'ouvrir Chirpity si chemin connu
                cfg = self._load_app_config() or {}
                exe = (cfg.get("chirpity_exe") or "").strip() or self._find_app_exe(
                    ["Chirpity.exe", "chirpity.exe", "Chirpity"], "chirpity_exe"
                )
                if exe:
                    self._launch_external(exe, "Chirpity")
                else:
                    messagebox.showinfo(
                        "Chirpity",
                        "Lancez Chirpity manuellement, analysez l'audio,\n"
                        "puis File → Export results to CSV.\n\n"
                        "Ensuite cliquez « Importer les résultats ».",
                    )
            else:
                exe = (self._load_app_config() or {}).get("birda_gui_exe") or self._find_app_exe(
                    ["birda-gui.exe", "Birda GUI.exe", "Birda.exe"], "birda_gui_exe"
                )
                self._launch_external(exe, "Birda GUI")

            # Bascule carnet + bouton prêt pour import
            try:
                self.tab_view.set("📝  Carnet")
            except Exception:
                pass
            self.log("Session écoute prête : %s — outil %s" % (session_meta.get("folder"), tool))
            messagebox.showinfo(
                "Session d'écoute",
                "1) Terminez l'analyse dans %s\n"
                "2) Exportez (CSV Chirpity) ou fermez Birda\n"
                "3) Cliquez OK puis « Importer les résultats »\n\n"
                "Le carnet de la session est déjà ouvert."
                % ("Chirpity" if tool == "chirpity" else "Birda GUI"),
            )

        def import_results():
            if not apply_folder_and_meta():
                return
            try:
                self.tab_view.set("📝  Carnet")
            except Exception:
                pass
            tool = (self._ecoute_session_meta or {}).get("tool") or tool_var.get()
            win.destroy()
            if tool == "chirpity":
                self.import_chirpity_csv()
            else:
                self.import_birda_from_catalog()

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=14, pady=10)
        ctk.CTkButton(
            bf, text="Ouvrir l'outil d'analyse", height=36,
            fg_color=UI.get("accent"), command=launch_tool_then_import,
        ).pack(side="left", padx=4)
        ctk.CTkButton(
            bf, text="Importer les résultats → carnet", height=36,
            fg_color=UI.get("success", "#2f9e5f"), command=import_results,
        ).pack(side="left", padx=4)
        ctk.CTkButton(bf, text="Fermer", command=win.destroy).pack(side="right")

    def import_birdnet_live_session(self, path=None):
        """Import session BirdNET Live : metadata.json + selections.txt (GPS, météo, espèces FR)."""
        # Choisir le .metadata.json ou un dossier d'export
        if not path:
            path = filedialog.askopenfilename(
                title="BirdNET Live — fichier .metadata.json (ou annuler pour un dossier)",
                filetypes=[
                    ("BirdNET Live metadata", "*.metadata.json"),
                    ("JSON", "*.json"),
                    ("Tous", "*.*"),
                ],
            )
            if not path:
                path = filedialog.askdirectory(title="Dossier export BirdNET Live")
        if not path:
            return

        meta_path = None
        sel_path = None
        if os.path.isdir(path):
            folder = path
            for name in os.listdir(folder):
                low = name.lower()
                full = os.path.join(folder, name)
                if low.endswith(".metadata.json") or (low.endswith(".json") and "metadata" in low):
                    meta_path = full
                if low.endswith(".selections.txt") or (low.endswith(".txt") and "selection" in low):
                    sel_path = full
            if not meta_path:
                for name in os.listdir(folder):
                    if name.lower().endswith(".json"):
                        meta_path = os.path.join(folder, name)
                        break
        else:
            meta_path = path
            folder = os.path.dirname(path)
            base = os.path.basename(path)
            # BirdNET_Live_....metadata.json → ....selections.txt
            stem = base
            for suf in (".metadata.json", ".json"):
                if stem.lower().endswith(suf):
                    stem = stem[: -len(suf)]
                    break
            cand = os.path.join(folder, stem + ".selections.txt")
            if os.path.isfile(cand):
                sel_path = cand
            else:
                for name in os.listdir(folder):
                    if name.lower().endswith(".selections.txt"):
                        sel_path = os.path.join(folder, name)
                        break

        if not meta_path or not os.path.isfile(meta_path):
            messagebox.showerror(
                "BirdNET Live",
                "Fichier .metadata.json introuvable.\n"
                "Exportez la session depuis BirdNET Live puis sélectionnez le JSON.",
            )
            return

        try:
            with open(meta_path, "r", encoding="utf-8") as f:
                meta = json.load(f)
        except Exception as e:
            messagebox.showerror("BirdNET Live", "Lecture metadata impossible :\n%s" % e)
            return

        session = meta.get("session") or {}
        lat = session.get("latitude")
        lon = session.get("longitude")
        try:
            lat = float(lat) if lat is not None else None
            lon = float(lon) if lon is not None else None
        except Exception:
            lat = lon = None
        lieu = (session.get("locationName") or "").strip()
        start = session.get("startTime") or ""
        # Heure locale affichage HH:MM depuis ISO UTC si besoin
        heure = ""
        date_s = ""
        try:
            if start:
                # 2026-08-15T05:18:39... → date + heure (UTC telle quelle ; utilisateur en Meuse)
                if "T" in start:
                    dpart, tpart = start.split("T", 1)
                    date_s = dpart
                    heure = tpart[:5]
        except Exception:
            pass
        if not date_s:
            date_s = datetime.now().strftime("%Y-%m-%d")
        weather = session.get("weather") or {}
        display = session.get("displayName") or os.path.basename(meta_path)

        # Dossier carnet : proposer création Ecoute_DATE si aucune sortie ouverte
        if not self.photo_folder_path:
            parent = filedialog.askdirectory(
                title="Choisir le dossier parent pour la session d'écoute",
            )
            if not parent:
                return
            folder_out = os.path.join(parent, "Ecoute_%s" % date_s)
            try:
                os.makedirs(folder_out, exist_ok=True)
            except Exception as e:
                messagebox.showerror("Dossier", str(e))
                return
            self.photo_folder_path = folder_out
            try:
                self._register_known_folder(folder_out)
            except Exception:
                pass
            try:
                if hasattr(self, "lbl_photo_folder"):
                    self.lbl_photo_folder.configure(text=folder_out)
            except Exception:
                pass
        else:
            folder_out = self.photo_folder_path

        # Lire selections.txt (TSV)
        detections = []
        if sel_path and os.path.isfile(sel_path):
            try:
                with open(sel_path, "r", encoding="utf-8-sig", errors="ignore") as f:
                    sample = f.read(2048)
                    f.seek(0)
                    delim = "\t" if "\t" in sample else (";" if sample.count(";") > sample.count(",") else ",")
                    reader = csv.DictReader(f, delimiter=delim)
                    for row in reader:
                        if not row:
                            continue
                        keys = {(k or "").strip().lower(): k for k in row.keys()}

                        def g(*names):
                            for n in names:
                                kk = keys.get(n.lower())
                                if kk is not None:
                                    return (row.get(kk) or "").strip()
                            return ""

                        common = g("Common Name", "common name", "common")
                        sci = g("Scientific Name", "scientific name", "scientific")
                        conf = g("Confidence", "confidence")
                        start_s = g("Begin Time (s)", "begin time (s)", "start", "survey time (s)")
                        if not common and not sci:
                            continue
                        detections.append({
                            "common": common,
                            "sci": sci,
                            "conf": conf,
                            "start": start_s,
                            "source_audio": g("Begin File", "begin file") or display,
                        })
            except Exception as e:
                messagebox.showerror("BirdNET Live", "Lecture selections.txt :\n%s" % e)
                return
        else:
            messagebox.showwarning(
                "BirdNET Live",
                "Aucun fichier .selections.txt trouvé à côté du metadata.\n"
                "Seules les infos de session (lieu/GPS) seront utilisables.",
            )

        if not detections:
            messagebox.showinfo("BirdNET Live", "Aucune détection dans le fichier selections.")
            return

        # Météo texte pour notes
        meteo_txt = ""
        if weather:
            try:
                meteo_txt = (
                    "BirdNET Live — %.1f°C, vent %.1f m/s, précip. %.1f mm, nuages %s%%"
                    % (
                        float(weather.get("temperatureC") or 0),
                        float(weather.get("windSpeedMs") or 0),
                        float(weather.get("precipitationMm") or 0),
                        weather.get("cloudCoverPercent") if weather.get("cloudCoverPercent") is not None else "?",
                    )
                )
            except Exception:
                meteo_txt = str(weather)

        # Session meta pour _write_birda_notes
        self._ecoute_session_meta = {
            "lieu": lieu,
            "date": date_s,
            "heure": heure,
            "lat": lat,
            "lon": lon,
            "tool": "birdnet_live",
            "folder": folder_out,
            "meteo_txt": meteo_txt,
            "session_name": display,
        }

        self.log(
            "BirdNET Live : %s — %d détections, GPS %s, %s"
            % (
                display,
                len(detections),
                ("%.5f,%.5f" % (lat, lon)) if lat is not None else "—",
                lieu or "lieu ?",
            )
        )
        try:
            self.tab_view.set("📝  Carnet")
        except Exception:
            pass
        # Même dialogue validation (seuil + cases) — noms déjà FR dans l'export
        self._birda_rows_to_notes(
            detections,
            source_label="birdnet_live:" + display,
        )

    def import_chirpity_csv(self, path=None):
        """Import dédié CSV Chirpity (File → Export results to CSV) + dialogue de validation."""
        if not self.photo_folder_path:
            messagebox.showwarning(
                "Sortie",
                "Ouvrez d'abord la sortie dans laquelle enregistrer les détections.",
            )
            return
        if not path:
            path = filedialog.askopenfilename(
                title="CSV Chirpity (export résultats)",
                filetypes=[
                    ("CSV Chirpity", "*.csv"),
                    ("Tous", "*.*"),
                ],
            )
        if not path or not os.path.isfile(path):
            return

        detections = []
        try:
            with open(path, "r", encoding="utf-8-sig", errors="ignore") as f:
                sample = f.read(4096)
                f.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
                except Exception:
                    dialect = csv.excel
                reader = csv.DictReader(f, dialect=dialect)
                for row in reader:
                    if not row:
                        continue
                    keys = {(k or "").strip().lower(): k for k in row.keys()}

                    def g(*names):
                        for n in names:
                            kk = keys.get(n.lower())
                            if kk is not None:
                                return (row.get(kk) or "").strip()
                        return ""

                    common = g(
                        "Common Name", "common name", "common_name", "common",
                        "Species", "species", "label", "English name",
                        "espèce", "espece", "nom",
                    )
                    sci = g(
                        "Scientific Name", "scientific name", "scientific_name",
                        "scientific", "sci_name", "latin", "nom scientifique",
                    )
                    conf = g(
                        "Confidence", "confidence", "conf", "score",
                        "probability", "max confidence", "max_confidence",
                    )
                    start = g(
                        "Start", "start", "start (s)", "start_s", "begin",
                        "timestamp", "time",
                    )
                    end = g("End", "end", "end (s)", "end_s")
                    src = g(
                        "File", "file", "filename", "audio", "path",
                        "source", "recording",
                    )
                    if not sci and common and "(" in common and common.endswith(")"):
                        try:
                            a, b = common.rsplit("(", 1)
                            common = a.strip()
                            sci = b.rstrip(")").strip()
                        except Exception:
                            pass
                    if not common and not sci:
                        continue
                    detections.append({
                        "common": common,
                        "sci": sci,
                        "conf": conf,
                        "start": start,
                        "end": end,
                        "source_audio": src,
                    })
        except Exception as e:
            messagebox.showerror("Chirpity", "Lecture CSV impossible :\n%s" % e)
            return

        if not detections:
            messagebox.showinfo(
                "Chirpity",
                "Aucune ligne exploitable dans ce CSV.\n\n"
                "Dans Chirpity : File → Export results to CSV\n"
                "(colonnes : espèce / scientific name / confidence).",
            )
            return

        self.log("Chirpity : %d ligne(s) lues dans %s" % (len(detections), os.path.basename(path)))
        self._birda_rows_to_notes(detections, source_label="chirpity:" + os.path.basename(path))

    def import_birda_detections(self, path=None):
        """Importe un JSON ou CSV Birda / BirdNET dans le carnet de la sortie ouverte."""
        if not self.photo_folder_path:
            messagebox.showwarning(
                "Sortie",
                "Ouvrez d'abord la sortie dans laquelle enregistrer les détections.",
            )
            return
        if not path:
            path = filedialog.askopenfilename(
                title="Fichier résultats Birda (JSON ou CSV)",
                filetypes=[
                    ("Birda / BirdNET", "*.json *.csv"),
                    ("JSON", "*.json"),
                    ("CSV", "*.csv"),
                    ("Tous", "*.*"),
                ],
            )
        if not path:
            return
        if not os.path.isfile(path):
            messagebox.showerror("Import Birda", "Fichier introuvable :\n%s" % path)
            return

        detections = []
        try:
            if path.lower().endswith(".json"):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                # Format fichier -f json Birda
                if isinstance(data, dict) and "detections" in data:
                    for d in data.get("detections") or []:
                        if not isinstance(d, dict):
                            continue
                        detections.append({
                            "common": d.get("common_name") or d.get("commonName") or "",
                            "sci": d.get("scientific_name") or d.get("scientificName") or "",
                            "conf": d.get("confidence"),
                            "start": d.get("start_time") or d.get("start"),
                            "end": d.get("end_time") or d.get("end"),
                            "source_audio": data.get("source_file") or "",
                        })
                # Envelope CLI --output-mode json (liste d'events) — rare en fichier
                elif isinstance(data, list):
                    for item in data:
                        if isinstance(item, dict) and item.get("common_name"):
                            detections.append({
                                "common": item.get("common_name") or "",
                                "sci": item.get("scientific_name") or "",
                                "conf": item.get("confidence"),
                                "start": item.get("start_time"),
                                "end": item.get("end_time"),
                                "source_audio": item.get("source_file") or "",
                            })
                elif isinstance(data, dict):
                    # parfois {species: count} — peu probable
                    for k, v in data.items():
                        if isinstance(v, dict) and (v.get("common_name") or v.get("scientific_name")):
                            detections.append({
                                "common": v.get("common_name") or k,
                                "sci": v.get("scientific_name") or "",
                                "conf": v.get("confidence"),
                                "start": v.get("start_time"),
                                "end": v.get("end_time"),
                                "source_audio": "",
                            })
            else:
                # CSV Birda : Start (s), End (s), Scientific name, Common name, Confidence, File
                with open(path, "r", encoding="utf-8-sig", errors="ignore") as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        if not row:
                            continue
                        # clés souples
                        keys = { (k or "").strip().lower(): k for k in row.keys() }

                        def g(*names):
                            for n in names:
                                kk = keys.get(n.lower())
                                if kk is not None:
                                    return (row.get(kk) or "").strip()
                            return ""

                        common = g("Common name", "common_name", "common", "espèce", "espece")
                        sci = g("Scientific name", "scientific_name", "scientific", "nom scientifique")
                        conf = g("Confidence", "confidence", "conf")
                        start = g("Start (s)", "Start", "start_time", "start")
                        end = g("End (s)", "End", "end_time", "end")
                        src = g("File", "file", "source_file")
                        if not common and not sci:
                            continue
                        detections.append({
                            "common": common, "sci": sci, "conf": conf,
                            "start": start, "end": end, "source_audio": src,
                        })
        except Exception as e:
            messagebox.showerror("Import Birda", "Lecture impossible :\n%s" % e)
            return

        if not detections:
            messagebox.showinfo(
                "Import Birda",
                "Aucune détection trouvée dans ce fichier.\n"
                "Formats attendus : JSON Birda (-f json) ou CSV résultats.\n\n"
                "Avec Birda GUI, préférez :\n"
                "« Importer depuis la base Birda GUI (SQLite) ».",
            )
            return

        self._birda_rows_to_notes(detections, source_label=os.path.basename(path))

    def open_cloud_settings(self):
        """Configure le dossier cloud local (Drive, Dropbox, IKDrive, OneDrive…)."""
        cfg = dict(self.app_config or {})
        win = ctk.CTkToplevel(self)
        win.title("☁ Sauvegarde cloud")
        win.geometry("560x420")
        win.configure(fg_color=UI.get("bg", "#0c1210"))
        self._prepare_tool_window(win)

        card = ctk.CTkFrame(win, fg_color=UI.get("card", "#1a2420"), corner_radius=12)
        card.pack(fill="both", expand=True, padx=14, pady=14)

        ctk.CTkLabel(
            card, text="Dossier cloud synchronisé",
            font=ctk.CTkFont(size=15, weight="bold")
        ).pack(anchor="w", padx=14, pady=(14, 4))
        ctk.CTkLabel(
            card,
            text=(
                "Principe simple et fiable :\n"
                "1) Installez Google Drive, Dropbox, IKDrive ou OneDrive sur le PC\n"
                "2) Choisissez ici le dossier synchronisé (ex. …\\Mon Drive\\GeoExif)\n"
                "3) GeoExif y copie le carnet (JSON, KML, GeoJSON)\n"
                "Le client cloud envoie ensuite les fichiers automatiquement.\n\n"
                "Pas besoin de clé API — fonctionne hors-ligne puis se synchronise."
            ),
            font=ctk.CTkFont(size=12), text_color=UI.get("text_dim", "#9db0a6"),
            justify="left",
        ).pack(anchor="w", padx=14, pady=(0, 10))

        path_var = tk.StringVar(value=cfg.get("cloud_folder", ""))
        row = ctk.CTkFrame(card, fg_color="transparent")
        row.pack(fill="x", padx=14, pady=4)
        entry = ctk.CTkEntry(row, textvariable=path_var, height=34)
        entry.pack(side="left", fill="x", expand=True, padx=(0, 6))

        def browse():
            d = filedialog.askdirectory(title="Dossier cloud (Drive / Dropbox / IKDrive…)")
            if d:
                path_var.set(d)

        ctk.CTkButton(row, text="Parcourir…", width=100, command=browse).pack(side="right")

        auto_var = ctk.BooleanVar(value=bool(cfg.get("cloud_auto_sync", False)))
        ctk.CTkCheckBox(
            card,
            text="Après chaque enregistrement d'observation, synchroniser vers le cloud",
            variable=auto_var, font=ctk.CTkFont(size=12),
        ).pack(anchor="w", padx=14, pady=10)

        lbl = ctk.CTkLabel(card, text="", font=ctk.CTkFont(size=11))
        lbl.pack(anchor="w", padx=14)

        def save_cfg():
            folder = path_var.get().strip()
            if folder and not os.path.isdir(folder):
                messagebox.showwarning("Dossier", "Ce chemin n'existe pas encore. Créez-le ou choisissez-en un autre.")
                return
            new_cfg = dict(self.app_config or {})
            new_cfg["cloud_folder"] = folder
            new_cfg["cloud_auto_sync"] = bool(auto_var.get())
            self._save_app_config(new_cfg)
            lbl.configure(
                text="Enregistré" + (f" → {folder}" if folder else " (aucun dossier)"),
                text_color=UI.get("success", "#5ecf8a"),
            )
            self.log(f"☁ Dossier cloud : {folder or 'désactivé'}")

        btn_row = ctk.CTkFrame(card, fg_color="transparent")
        btn_row.pack(fill="x", padx=14, pady=14)
        ctk.CTkButton(
            btn_row, text="Enregistrer", command=save_cfg,
            fg_color=UI.get("accent", "#3eb4a0"), height=34,
        ).pack(side="left", padx=(0, 6))
        ctk.CTkButton(
            btn_row, text="Synchroniser maintenant",
            command=lambda: (save_cfg(), self.sync_carnet_to_cloud()),
            fg_color=UI.get("success", "#5ecf8a"), height=34,
        ).pack(side="left", padx=6)
        ctk.CTkButton(btn_row, text="Fermer", command=win.destroy, width=90).pack(side="right")

    def sync_carnet_to_cloud(self, silent=False):
        """Copie le carnet de la sortie active vers le dossier cloud configuré."""
        folder = (self.app_config or {}).get("cloud_folder", "").strip()
        if not folder:
            if not silent:
                if messagebox.askyesno(
                    "Cloud non configuré",
                    "Aucun dossier cloud n'est défini.\n\nOuvrir les réglages cloud ?"
                ):
                    self.open_cloud_settings()
            return False
        if not os.path.isdir(folder):
            if not silent:
                messagebox.showerror(
                    "Dossier introuvable",
                    f"Le dossier cloud n'existe pas :\n{folder}\n\n"
                    "Vérifiez que Drive / Dropbox / IKDrive est connecté."
                )
            return False
        if not self.photo_folder_path:
            if not silent:
                messagebox.showwarning("Pas de dossier", "Ouvrez d'abord un dossier de photos / sortie.")
            return False

        sortie = os.path.basename(self.photo_folder_path.rstrip("\\/")) or "sortie"
        # Nettoie le nom pour le système de fichiers
        sortie_safe = re.sub(r'[<>:"/\\|?*]', "_", sortie).strip() or "sortie"
        dest_root = os.path.join(folder, "GeoExif_Carnets", sortie_safe)
        try:
            os.makedirs(dest_root, exist_ok=True)
        except Exception as e:
            if not silent:
                messagebox.showerror("Cloud", f"Impossible de créer le dossier :\n{e}")
            return False

        copies = []
        errors = []

        def _copy(src, name=None):
            if not src or not os.path.isfile(src):
                return
            target = os.path.join(dest_root, name or os.path.basename(src))
            try:
                shutil.copy2(src, target)
                copies.append(os.path.basename(target))
            except Exception as e:
                errors.append(f"{os.path.basename(src)}: {e}")

        # Fichiers essentiels du carnet
        notes = os.path.join(self.photo_folder_path, NOTES_FILE)
        _copy(notes)

        # Dictionnaire d'espèces (global, à côté de l'appli)
        try:
            sp = os.path.join(os.path.dirname(os.path.abspath(__file__)), SPECIES_FILE)
            if not os.path.isfile(sp):
                sp = SPECIES_FILE
            _copy(sp, "dictionnaire_especes.json")
        except Exception:
            pass

        # Exports géo s'ils existent déjà dans le dossier sortie
        for fname in (
            "observations.kml", "observations.kmz", "observations.geojson",
            "observations.gpx", "carte_observations.png",
        ):
            _copy(os.path.join(self.photo_folder_path, fname))

        # Snapshot horodaté du carnet (historique cloud)
        if os.path.isfile(notes):
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            try:
                hist = os.path.join(dest_root, "historique")
                os.makedirs(hist, exist_ok=True)
                shutil.copy2(notes, os.path.join(hist, f"observations_{stamp}.json"))
                copies.append(f"historique/observations_{stamp}.json")
            except Exception as e:
                errors.append(f"historique: {e}")

        # Petit manifeste
        try:
            manifest = {
                "sortie": sortie,
                "source": self.photo_folder_path,
                "sync_utc": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
                "app": f"GeoExif {APP_VERSION}",
                "fichiers": copies,
            }
            with open(os.path.join(dest_root, "sync_manifest.json"), "w", encoding="utf-8") as f:
                json.dump(manifest, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

        if copies:
            self.log(f"☁ Cloud : {len(copies)} fichier(s) → {dest_root}")
            if not silent:
                messagebox.showinfo(
                    "Synchronisé",
                    f"Carnet copié vers :\n{dest_root}\n\n"
                    f"{len(copies)} fichier(s)\n"
                    "Le client Drive / Dropbox / IKDrive terminera l'envoi."
                )
            return True
        if not silent:
            messagebox.showwarning(
                "Rien à synchroniser",
                "Aucun fichier de carnet trouvé (observations.json).\n"
                "Enregistrez au moins une observation avant."
            )
        if errors:
            self.log(f"☁ Erreurs cloud : {'; '.join(errors[:3])}")
        return False

    def _project_path(self):
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), PROJECT_FILE)

    def _load_project(self):
        path = self._project_path()
        if not os.path.exists(path):
            return {
                "name": f"Meuse {datetime.now().year}",
                "year": datetime.now().year,
                "folders": [],
                "created": datetime.now().isoformat(),
            }
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
            data.setdefault("name", f"Meuse {datetime.now().year}")
            data.setdefault("year", datetime.now().year)
            data.setdefault("folders", [])
            return data
        except Exception:
            return {"name": f"Meuse {datetime.now().year}", "year": datetime.now().year, "folders": []}

    def _save_project(self, project=None):
        proj = project if project is not None else getattr(self, "project", self._load_project())
        self.project = proj
        try:
            with open(self._project_path(), "w", encoding="utf-8") as f:
                json.dump(proj, f, ensure_ascii=False, indent=2)
        except Exception as e:
            self.log(f"Projet : impossible d'enregistrer ({e})")

    def project_add_current_folder(self):
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier", "Ouvrez d'abord une sortie ou un carnet.")
            return
        proj = self._load_project()
        abs_path = os.path.normpath(self.photo_folder_path)
        for f in proj.get("folders", []):
            if os.path.normpath(f.get("path", "")) == abs_path:
                messagebox.showinfo("Projet", "Cette sortie est déjà dans le projet.")
                return
        proj.setdefault("folders", []).append({
            "path": abs_path,
            "label": os.path.basename(abs_path),
            "added": datetime.now().isoformat(),
        })
        self._save_project(proj)
        self._register_known_folder(abs_path)
        self.log(f"Projet « {proj.get('name')} » : + {os.path.basename(abs_path)}")
        messagebox.showinfo("Projet", f"Sortie ajoutée au projet « {proj.get('name')} ».")

    def _obs_dedupe_key(self, data, folder_label=""):
        """Clé de fusion pour éviter les doublons (même espèce / lieu / heure / GPS approx)."""
        esp = (data.get("espece") or "").strip().lower()
        heure = (data.get("heure") or "").strip()
        lieu = (data.get("lieu") or "").strip().lower()
        try:
            lat = round(float(data.get("lat")), 4) if data.get("lat") is not None else ""
            lon = round(float(data.get("lon")), 4) if data.get("lon") is not None else ""
        except Exception:
            lat, lon = "", ""
        return (esp, heure, lieu, lat, lon)

    def gather_project_observations(self, project=None, *, species_filter=None, date_from=None, date_to=None):
        """Charge toutes les observations du projet avec fusion anti-doublons."""
        proj = project or self._load_project()
        folders = proj.get("folders") or []
        # Si projet vide, repli sur known_folders
        if not folders:
            folders = [{"path": e.get("path"), "label": e.get("label")} for e in self._load_known_folders()]

        merged = []
        seen = set()
        for entry in folders:
            path = entry.get("path") or ""
            label = entry.get("label") or os.path.basename(path)
            notes_path = os.path.join(path, NOTES_FILE)
            if not os.path.isfile(notes_path):
                continue
            try:
                with open(notes_path, "r", encoding="utf-8") as f:
                    notes = json.load(f) or {}
            except Exception:
                continue
            for key, data in notes.items():
                if not isinstance(data, dict):
                    continue
                # Filtre espèce
                if species_filter:
                    esp = (data.get("espece") or "").lower()
                    if species_filter.lower() not in esp:
                        continue
                # Filtre date via nom de dossier ou champ
                if date_from or date_to:
                    d = self._parse_sortie_date(label) if hasattr(self, "_parse_sortie_date") else None
                    if d:
                        if date_from and d < date_from:
                            continue
                        if date_to and d > date_to:
                            continue
                rec = dict(data)
                rec["_folder"] = label
                rec["_folder_path"] = path
                rec["_key"] = key
                dk = self._obs_dedupe_key(rec, label)
                if dk in seen and dk[0]:  # ignore empty species keys less aggressively
                    # si GPS+espèce+heure identiques → doublon
                    continue
                seen.add(dk)
                merged.append(rec)
        return merged

    def open_project_hub(self):
        """Projet multi-sorties (ex. Meuse 2026) : sorties, filtres, stats, carte cumulée."""
        proj = self._load_project()
        self.project = proj

        win = ctk.CTkToplevel(self)
        win.title(f"Projet — {proj.get('name', 'Meuse')}")
        win.geometry("720x640")
        win.configure(fg_color=UI.get("bg", "#0c1210"))
        self._prepare_tool_window(win)

        head = ctk.CTkFrame(win, fg_color=UI.get("card", "#1a2420"), corner_radius=12)
        head.pack(fill="x", padx=14, pady=12)
        ctk.CTkLabel(head, text="Projet naturaliste", font=ctk.CTkFont(size=14, weight="bold")).pack(
            anchor="w", padx=12, pady=(10, 2)
        )
        name_var = tk.StringVar(value=proj.get("name", f"Meuse {datetime.now().year}"))
        year_var = tk.StringVar(value=str(proj.get("year", datetime.now().year)))
        row = ctk.CTkFrame(head, fg_color="transparent")
        row.pack(fill="x", padx=12, pady=6)
        ctk.CTkEntry(row, textvariable=name_var, width=220).pack(side="left", padx=(0, 8))
        ctk.CTkEntry(row, textvariable=year_var, width=80).pack(side="left")
        ctk.CTkLabel(row, text="(année)", text_color=UI.get("text_muted", "#888")).pack(side="left", padx=6)

        def save_meta():
            proj["name"] = name_var.get().strip() or proj.get("name")
            try:
                proj["year"] = int(year_var.get().strip())
            except Exception:
                pass
            self._save_project(proj)
            win.title(f"Projet — {proj['name']}")
            self.log(f"Projet enregistré : {proj['name']}")

        ctk.CTkButton(head, text="Enregistrer le nom", width=140, command=save_meta).pack(
            anchor="w", padx=12, pady=(0, 10)
        )

        ctk.CTkLabel(win, text="Sorties du projet", font=ctk.CTkFont(weight="bold")).pack(
            anchor="w", padx=16, pady=(4, 2)
        )
        list_frame = ctk.CTkScrollableFrame(win, height=180, fg_color=UI.get("card", "#1a2420"))
        list_frame.pack(fill="x", padx=14, pady=4)

        def refresh_list():
            for w in list_frame.winfo_children():
                w.destroy()
            folders = proj.get("folders") or []
            if not folders:
                ctk.CTkLabel(
                    list_frame, text="Aucune sortie — ouvrez un dossier puis « Ajouter la sortie courante ».",
                    text_color=UI.get("text_dim", "#9db0a6")
                ).pack(anchor="w", padx=8, pady=8)
                return
            for entry in folders:
                line = ctk.CTkFrame(list_frame, fg_color="transparent")
                line.pack(fill="x", pady=2)
                ctk.CTkLabel(line, text=entry.get("label") or entry.get("path"), anchor="w").pack(
                    side="left", fill="x", expand=True, padx=6
                )
                path = entry.get("path")

                def remove(p=path):
                    proj["folders"] = [f for f in proj.get("folders", []) if os.path.normpath(f.get("path", "")) != os.path.normpath(p)]
                    self._save_project(proj)
                    refresh_list()

                ctk.CTkButton(line, text="Retirer", width=70, height=24, command=remove).pack(side="right", padx=4)

        refresh_list()

        filt = ctk.CTkFrame(win, fg_color="transparent")
        filt.pack(fill="x", padx=14, pady=8)
        ctk.CTkLabel(filt, text="Filtre espèce :").pack(side="left")
        species_var = tk.StringVar(value="")
        ctk.CTkEntry(filt, textvariable=species_var, width=180, placeholder_text="ex. Chevreuil").pack(
            side="left", padx=6
        )

        stats_lbl = ctk.CTkLabel(win, text="", justify="left", font=ctk.CTkFont(size=12))
        stats_lbl.pack(anchor="w", padx=16, pady=4)

        def run_stats():
            obs = self.gather_project_observations(proj, species_filter=species_var.get().strip() or None)
            notes_like = {f"{o.get('_folder')}:{o.get('_key')}": o for o in obs}
            st = self._stats_from_notes(notes_like)
            stats_lbl.configure(
                text=(
                    f"{st['n_obs']} observations fusionnées · {st['n_especes']} espèces\n"
                    f"Photos {st['avec_photo']} · sans photo {st['sans_photo']} · indices ~{st['indices']}\n"
                    + "Top : " + ", ".join(f"{k} ({v})" for k, v in st["especes"].most_common(8))
                )
            )

        def open_cumul_map():
            # Réutilise known folders = sorties projet temporairement pour la carte existante
            # en injectant les paths projet dans un chargement dédié
            self._open_project_cumulative_map(proj, species_filter=species_var.get().strip() or None)

        def brief_projet():
            obs = self.gather_project_observations(proj, species_filter=species_var.get().strip() or None)
            if not obs:
                messagebox.showinfo("Vide", "Aucune observation dans le projet (ou filtre trop strict).")
                return
            notes_like = {f"{o.get('_folder')}:{o.get('_key')}": o for o in obs}
            st = self._stats_from_notes(notes_like)
            lines = [
                f"Projet {proj.get('name')} ({proj.get('year')})",
                f"{st['n_obs']} obs fusionnées, {st['n_especes']} espèces, {len(proj.get('folders') or [])} sorties.",
                "Espèces : " + ", ".join(f"{k} ({v})" for k, v in st["especes"].most_common()),
                "",
                "Rédige une synthèse de saison naturaliste Meuse, sans inventer d'espèces.",
                "",
                "=== DONNÉES ===",
            ]
            for o in obs[:200]:
                lines.append(
                    f"- [{o.get('_folder')}] {o.get('heure','')} | {o.get('espece')} ×{o.get('nombre','?')} | {o.get('lieu','')}"
                )
            prompt = "\n".join(lines)
            tw = ctk.CTkToplevel(win)
            tw.title("Brief projet")
            tw.geometry("640x520")
            box = ctk.CTkTextbox(tw)
            box.pack(fill="both", expand=True, padx=10, pady=10)
            box.insert("0.0", prompt)
            ctk.CTkButton(
                tw, text="Copier",
                command=lambda: (self.clipboard_clear(), self.clipboard_append(prompt), messagebox.showinfo("OK", "Copié"))
            ).pack(pady=8)

        actions = ctk.CTkFrame(win, fg_color="transparent")
        actions.pack(fill="x", padx=14, pady=10)
        ctk.CTkButton(actions, text="Ajouter la sortie courante", command=lambda: (self.project_add_current_folder(), proj.update(self._load_project()), refresh_list())).pack(side="left", padx=2)
        ctk.CTkButton(actions, text="Stats / fusion", command=run_stats, fg_color=UI.get("accent", "#3eb4a0")).pack(side="left", padx=2)
        ctk.CTkButton(actions, text="Carte cumulée", command=open_cumul_map, fg_color=UI.get("success", "#5ecf8a")).pack(side="left", padx=2)
        ctk.CTkButton(actions, text="Brief IA projet", command=brief_projet, fg_color=UI.get("purple", "#a78bfa")).pack(side="left", padx=2)

        run_stats()

    def _open_project_cumulative_map(self, project, species_filter=None):
        """Carte projet + liste des sorties cliquables."""
        folders = project.get("folders") or []
        obs = self.gather_project_observations(project, species_filter=species_filter)
        pts = []
        for o in obs:
            try:
                lat, lon = float(o.get("lat")), float(o.get("lon"))
            except Exception:
                continue
            pts.append((
                lat, lon,
                o.get("espece") or "?",
                o.get("categorie") or "Autre",
                o.get("_folder") or "",
                o.get("_folder_path") or "",
            ))

        win = ctk.CTkToplevel(self)
        win.title(f"Carte projet — {project.get('name')}")
        win.geometry("1100x720")

        top = ctk.CTkFrame(win, fg_color="transparent")
        top.pack(fill="x", padx=10, pady=8)
        ctk.CTkLabel(
            top,
            text=f"{len(pts)} points · {len(folders)} sorties"
            + (f" · filtre « {species_filter} »" if species_filter else ""),
            font=ctk.CTkFont(weight="bold"),
        ).pack(side="left")

        body = ctk.CTkFrame(win, fg_color=UI.get("bg", "#101612"))
        body.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        body.grid_columnconfigure(1, weight=1)
        body.grid_rowconfigure(0, weight=1)

        side = ctk.CTkScrollableFrame(body, width=260, fg_color=UI.get("card", "#1c2620"))
        side.grid(row=0, column=0, sticky="nsew", padx=(0, 8))
        ctk.CTkLabel(side, text="Sorties du projet", font=ctk.CTkFont(weight="bold")).pack(
            anchor="w", padx=8, pady=8
        )

        def open_one(path):
            self.open_existing_sortie(path, switch_tab=True, close_window=win)

        for entry in folders:
            path = entry.get("path") or ""
            label = entry.get("label") or os.path.basename(path)
            ctk.CTkButton(
                side, text=f"📂 {label}", anchor="w", height=30,
                fg_color=UI.get("card_alt", "#24302a"),
                hover_color=UI.get("accent", "#4ec4b0"),
                command=lambda p=path: open_one(p),
            ).pack(fill="x", padx=6, pady=2)

        map_host = tk.Frame(body, bg=UI.get("card", "#1c2620"), highlightthickness=0)
        map_host.grid(row=0, column=1, sticky="nsew")
        map_w = TkinterMapView(map_host, corner_radius=0)
        map_w.pack(fill="both", expand=True)
        try:
            map_w.set_tile_server("https://tile.openstreetmap.org/{z}/{x}/{y}.png")
        except Exception:
            pass
        if pts:
            map_w.set_position(pts[0][0], pts[0][1])
            map_w.set_zoom(11)
            for lat, lon, esp, cat, folder, fpath in pts:
                color = CATEGORY_COLORS.get(cat, "#7f7f7f")

                def mk_cmd(m, fp=fpath, e=esp, f=folder):
                    if fp and messagebox.askyesno(
                        "Ouvrir la sortie ?",
                        f"{e} ({f})\n\nOuvrir cette sortie dans le carnet ?"
                    ):
                        open_one(fp)

                map_w.set_marker(
                    lat, lon, text=f"{esp}\n({folder})",
                    marker_color_circle=color, marker_color_outside=color,
                    command=mk_cmd,
                )
        else:
            map_w.set_position(49.16, 5.38)
            map_w.set_zoom(9)
            messagebox.showinfo("Carte", "Aucun point GPS dans les observations fusionnées.")

    def open_calendrier_chasse(self):
        """Affiche le calendrier de chasse indicatif Meuse (12 mois / especes)."""
        win = ctk.CTkToplevel(self)
        win.title("Calendrier de chasse — Meuse (indicatif)")
        win.geometry("640x560")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Calendrier de chasse indicatif — Meuse (55)",
            font=ctk.CTkFont(size=15, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=14, pady=(12, 4))
        ctk.CTkLabel(
            win,
            text=MEUSE_CALENDRIER_CHASSE.get("disclaimer", ""),
            font=ctk.CTkFont(size=11),
            text_color=UI.get("warning", "#c4841a"),
            wraplength=600, justify="left",
        ).pack(anchor="w", padx=14, pady=(0, 6))

        src = MEUSE_CALENDRIER_CHASSE.get("sources_suggerees") or []
        if src:
            ctk.CTkLabel(
                win, text="Sources a consulter : " + " · ".join(src),
                font=ctk.CTkFont(size=10),
                text_color=UI.get("text_dim"),
                wraplength=600, justify="left",
            ).pack(anchor="w", padx=14, pady=(0, 8))

        box = ctk.CTkTextbox(win, height=380)
        box.pack(fill="both", expand=True, padx=14, pady=6)
        mois_noms = ["", "Janvier", "Fevrier", "Mars", "Avril", "Mai", "Juin",
                     "Juillet", "Aout", "Septembre", "Octobre", "Novembre", "Decembre"]
        lines = []
        for mois in range(1, 13):
            periodes = meuse_chasse_pour_mois(mois)
            lines.append("=== %s ===" % mois_noms[mois])
            if not periodes:
                lines.append("  (pas d'entree indicative majeure)")
            for p in periodes:
                lines.append("  • %s — %s" % (p.get("espece"), p.get("mode") or ""))
                if p.get("note"):
                    lines.append("      %s" % p.get("note"))
            lines.append("")
        box.insert("0.0", "\n".join(lines))
        box.configure(state="disabled")

        ctk.CTkButton(
            win, text="Fermer", height=32, fg_color=UI.get("card_alt"),
            command=win.destroy,
        ).pack(pady=10)

    def open_season_suggestions(self):
        """Suggestions d'espèces Meuse par mois — navigation libre sur les 12 mois."""
        mois_noms = ["", "Janvier", "Février", "Mars", "Avril", "Mai", "Juin",
                     "Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre"]
        state = {"month": datetime.now().month}

        win = ctk.CTkToplevel(self)
        win.title("Cibles de saison — Meuse")
        win.geometry("560x700")
        try:
            win.configure(fg_color=UI.get("bg", "#0c1210"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        title_lbl = ctk.CTkLabel(win, text="", font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"))
        title_lbl.pack(anchor="w", padx=16, pady=(14, 4))

        notes_lbl = ctk.CTkLabel(
            win, text="", font=ctk.CTkFont(size=12),
            text_color=UI.get("text_dim", "#9db0a6"),
            wraplength=500, justify="left",
        )
        notes_lbl.pack(anchor="w", padx=16, pady=4)

        ctk.CTkLabel(win, text="Habitats à privilégier", font=ctk.CTkFont(weight="bold")).pack(
            anchor="w", padx=16, pady=(10, 2)
        )
        habitats_lbl = ctk.CTkLabel(
            win, text="", text_color=UI.get("text_accent", "#8eebd8"),
            wraplength=500, justify="left",
        )
        habitats_lbl.pack(anchor="w", padx=16)

        ctk.CTkLabel(
            win, text="Especes / indices a chercher",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(12, 4))
        box = ctk.CTkTextbox(win, height=150)
        box.pack(fill="both", expand=True, padx=16, pady=4)

        ctk.CTkLabel(
            win, text="Calendrier de chasse (indicatif Meuse)",
            font=ctk.CTkFont(weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(10, 2))
        chasse_lbl = ctk.CTkLabel(
            win,
            text="",
            font=ctk.CTkFont(size=11),
            text_color=UI.get("text_dim"),
            wraplength=500, justify="left",
        )
        chasse_lbl.pack(anchor="w", padx=16, pady=(0, 2))
        box_chasse = ctk.CTkTextbox(win, height=120)
        box_chasse.pack(fill="both", expand=True, padx=16, pady=(0, 4))

        def show(m):
            m = ((m - 1) % 12) + 1  # garde 1..12
            state["month"] = m
            data = MEUSE_SAISON_ESPECES.get(m) or {"notes": "", "habitats": [], "cibles": []}
            title_lbl.configure(text=f"Meuse 55 — {mois_noms[m]}")
            notes_lbl.configure(text=data.get("notes", "") or "")
            habitats_lbl.configure(text=" · ".join(data.get("habitats") or []) or "—")
            box.configure(state="normal")
            box.delete("0.0", "end")
            box.insert("0.0", chr(10).join("• %s" % x for x in (data.get("cibles") or [])))
            box.configure(state="disabled")
            periodes = meuse_chasse_pour_mois(m)
            try:
                chasse_lbl.configure(
                    text=(MEUSE_CALENDRIER_CHASSE.get("disclaimer") or "")[:200] + "…"
                )
            except Exception:
                pass
            box_chasse.configure(state="normal")
            box_chasse.delete("0.0", "end")
            if periodes:
                lines = []
                for p in periodes:
                    lines.append("• %s — %s" % (p.get("espece"), p.get("mode") or ""))
                    if p.get("note"):
                        lines.append("    %s" % p.get("note"))
                    if p.get("photo"):
                        lines.append("    Photo/obs. : %s" % p.get("photo"))
                box_chasse.insert("0.0", chr(10).join(lines))
            else:
                box_chasse.insert(
                    "0.0",
                    "Aucune periode indicative majeure pour ce mois." + chr(10)
                    + "Verifiez l'arrete prefectorel Meuse.",
                )
            box_chasse.configure(state="disabled")
            win.title(f"Cibles de saison — {mois_noms[m]}")
            # Maj libellés boutons
            prev_m = 12 if m == 1 else m - 1
            next_m = 1 if m == 12 else m + 1
            btn_prev.configure(text=f"← {mois_noms[prev_m]}")
            btn_next.configure(text=f"{mois_noms[next_m]} →")
            month_menu.set(mois_noms[m])

        nav = ctk.CTkFrame(win, fg_color="transparent")
        nav.pack(fill="x", padx=16, pady=10)

        btn_prev = ctk.CTkButton(
            nav, text="←", width=130, height=32,
            command=lambda: show(state["month"] - 1),
        )
        btn_prev.pack(side="left", padx=(0, 4))

        month_menu = ctk.CTkOptionMenu(
            nav,
            values=mois_noms[1:],
            width=140,
            command=lambda name: show(mois_noms.index(name)),
        )
        month_menu.pack(side="left", padx=4)

        btn_next = ctk.CTkButton(
            nav, text="→", width=130, height=32,
            command=lambda: show(state["month"] + 1),
        )
        btn_next.pack(side="left", padx=4)

        ctk.CTkButton(nav, text="Fermer", width=90, command=win.destroy).pack(side="right")

        show(state["month"])

    def open_inpn_search(self):
        """Recherche nom scientifique : INPN (souvent HS) + secours GBIF / Wikipedia."""
        q = ""
        if hasattr(self, "entry_nom_sci"):
            q = self.entry_nom_sci.get().strip()
        if not q and hasattr(self, "choice_species"):
            try:
                q = self.choice_species.get().strip()
            except Exception:
                q = ""
        if not q:
            messagebox.showinfo(
                "Nom scientifique",
                "Saisissez un nom d'espece ou un nom scientifique,\npuis recliquez sur INPN.",
            )
            return

        q_enc = urllib.parse.quote(q)
        urls = [
            ("INPN (MNHN)", f"https://inpn.mnhn.fr/recherche?q={q_enc}"),
            ("GBIF", f"https://www.gbif.org/species/search?q={q_enc}"),
            ("Wikipedia FR", f"https://fr.wikipedia.org/w/index.php?search={q_enc}"),
        ]

        win = ctk.CTkToplevel(self)
        win.title("Recherche nom scientifique")
        win.geometry("520x440")
        win.configure(fg_color=UI.get("bg", "#101612"))
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text=f"Recherche : {q}",
            font=ctk.CTkFont(size=14, weight="bold"),
        ).pack(anchor="w", padx=14, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text=(
                "Le site INPN (MNHN) peut renvoyer une erreur 500\n"
                "ou etre hors service (incident informatique MNHN).\n"
                "Utilisez GBIF / Wikipedia, ou la suggestion automatique."
            ),
            font=ctk.CTkFont(size=11),
            text_color=UI.get("text_dim", "#888"),
            justify="left",
        ).pack(anchor="w", padx=14, pady=(0, 10))

        for label, url in urls:
            row = ctk.CTkFrame(win, fg_color="transparent")
            row.pack(fill="x", padx=14, pady=3)
            ctk.CTkButton(
                row, text=label, width=160, height=30,
                fg_color=UI.get("card_alt", "#24302a"),
                hover_color=UI.get("accent", "#4ec4b0"),
                command=lambda u=url, lab=label: (
                    webbrowser.open(u),
                    self.log(f"Ouverture {lab} pour {q}"),
                ),
            ).pack(side="left")

        status = ctk.CTkLabel(win, text="", text_color=UI.get("text_accent", "#8eebd8"))
        status.pack(anchor="w", padx=14, pady=(12, 4))

        def suggest_gbif():
            status.configure(text="Interrogation GBIF...")
            win.update_idletasks()

            def work():
                result = None
                err = None
                try:
                    import requests
                    r = requests.get(
                        "https://api.gbif.org/v1/species/match",
                        params={"name": q, "strict": "false"},
                        timeout=12,
                    )
                    if r.status_code != 200:
                        err = f"GBIF HTTP {r.status_code}"
                    else:
                        data = r.json()
                        if data.get("matchType") in ("EXACT", "FUZZY", "HIGHERRANK") and data.get("scientificName"):
                            result = data
                        else:
                            r2 = requests.get(
                                "https://api.gbif.org/v1/species/search",
                                params={"q": q, "limit": 5, "status": "ACCEPTED"},
                                timeout=12,
                            )
                            if r2.status_code == 200:
                                res = (r2.json() or {}).get("results") or []
                                if res:
                                    result = res[0]
                                else:
                                    err = "Aucun resultat GBIF"
                            else:
                                err = f"GBIF search HTTP {r2.status_code}"
                except Exception as e:
                    err = str(e)

                def apply():
                    if not win.winfo_exists():
                        return
                    if err and not result:
                        status.configure(text=f"Echec : {err}")
                        return
                    sci = (
                        (result or {}).get("scientificName")
                        or (result or {}).get("canonicalName")
                        or ""
                    )
                    if sci and hasattr(self, "entry_nom_sci"):
                        self.entry_nom_sci.delete(0, "end")
                        self.entry_nom_sci.insert(0, sci)
                        status.configure(text=f"Nom propose : {sci}")
                        self.log(f"GBIF -> nom scientifique : {sci}")
                    else:
                        status.configure(text="Pas de nom scientifique exploitable.")

                try:
                    win.after(0, apply)
                except Exception:
                    pass

            threading.Thread(target=work, daemon=True).start()

        ctk.CTkButton(
            win,
            text="Suggérer le nom scientifique (GBIF)",
            height=36,
            fg_color=UI.get("accent", "#4ec4b0"),
            hover_color=UI.get("accent_hover", "#3aab98"),
            command=suggest_gbif,
        ).pack(fill="x", padx=14, pady=(8, 6))

        ctk.CTkButton(
            win, text="Fermer", height=30,
            fg_color=UI.get("card_alt", "#24302a"),
            command=win.destroy,
        ).pack(fill="x", padx=14, pady=(0, 14))


    def _gpx_track_length_km(self, gpx_path):
        """Distance approximative d'une trace GPX (km)."""
        if not gpx_path or not os.path.isfile(gpx_path):
            return None
        try:
            with open(gpx_path, "r", encoding="utf-8", errors="ignore") as f:
                raw = f.read()
            pts = re.findall(
                r'lat="([+-]?\\d+(?:\\.\\d+)?)"[^>]*lon="([+-]?\\d+(?:\\.\\d+)?)"|'
                r'lon="([+-]?\\d+(?:\\.\\d+)?)"[^>]*lat="([+-]?\\d+(?:\\.\\d+)?)"',
                raw,
            )
            coords = []
            for a, b, c, d in pts:
                if a and b:
                    coords.append((float(a), float(b)))
                elif c and d:
                    coords.append((float(d), float(c)))
            if len(coords) < 2:
                return 0.0
            def haversine(lat1, lon1, lat2, lon2):
                r = 6371.0
                p1, p2 = math.radians(lat1), math.radians(lat2)
                dp = math.radians(lat2 - lat1)
                dl = math.radians(lon2 - lon1)
                x = math.sin(dp / 2) ** 2 + math.cos(p1) * math.cos(p2) * math.sin(dl / 2) ** 2
                return 2 * r * math.asin(math.sqrt(x))
            dist = 0.0
            for i in range(1, len(coords)):
                dist += haversine(coords[i - 1][0], coords[i - 1][1], coords[i][0], coords[i][1])
            return round(dist, 2)
        except Exception:
            return None

    def _effort_from_notes(self, notes, gpx_path=None):
        """Estime l'effort : plage horaire des obs + km GPX si dispo."""
        heures = []
        for data in (notes or {}).values():
            h = (data.get("heure") or "").strip()
            m = re.match(r"^(\\d{1,2}):(\\d{2})", h)
            if m:
                heures.append(int(m.group(1)) * 60 + int(m.group(2)))
        duree_h = None
        if len(heures) >= 2:
            duree_h = round((max(heures) - min(heures)) / 60.0, 2)
            if duree_h < 0:
                duree_h = None
        km = self._gpx_track_length_km(gpx_path or self.gpx_file_path)
        return {"duree_heures": duree_h, "km_trace": km, "n_obs": len(notes or {})}

    def open_analysis_panel(self):
        """Effort, richesse spécifique, comparaison années / secteurs."""
        win = ctk.CTkToplevel(self)
        win.title("Analyse naturaliste")
        win.geometry("700x620")
        win.configure(fg_color=UI.get("bg", "#0c1210"))
        self._prepare_tool_window(win)

        box = ctk.CTkTextbox(win, font=ctk.CTkFont(size=13))
        box.pack(fill="both", expand=True, padx=14, pady=14)

        lines = ["=== ANALYSE ===", ""]

        # Sortie courante
        if self.photo_folder_path:
            notes_path = os.path.join(self.photo_folder_path, NOTES_FILE)
            notes = {}
            if os.path.isfile(notes_path):
                try:
                    with open(notes_path, "r", encoding="utf-8") as f:
                        notes = json.load(f) or {}
                except Exception:
                    pass
            st = self._stats_from_notes(notes)
            eff = self._effort_from_notes(notes, self.gpx_file_path)
            lines += [
                f"Sortie : {os.path.basename(self.photo_folder_path)}",
                f"Observations : {st['n_obs']} · Espèces : {st['n_especes']}",
                f"Effort estimé : {eff['duree_heures'] if eff['duree_heures'] is not None else '—'} h "
                f"(plage horaires obs) · Trace GPX : {eff['km_trace'] if eff['km_trace'] is not None else '—'} km",
                "Espèces : " + ", ".join(f"{k} ({v})" for k, v in st["especes"].most_common()),
                "",
            ]
        else:
            lines.append("Aucune sortie ouverte.\n")

        # Projet / known folders — richesse par mois et comparaison années
        try:
            proj = self._load_project()
            obs = self.gather_project_observations(proj)
        except Exception:
            obs = []
            for e in self._load_known_folders():
                path = e.get("path")
                np = os.path.join(path or "", NOTES_FILE)
                if not os.path.isfile(np):
                    continue
                try:
                    with open(np, "r", encoding="utf-8") as f:
                        for k, d in (json.load(f) or {}).items():
                            if isinstance(d, dict):
                                d = dict(d)
                                d["_folder"] = e.get("label") or os.path.basename(path)
                                obs.append(d)
                except Exception:
                    pass

        by_month = collections.defaultdict(set)
        by_year = collections.defaultdict(set)
        by_sector = collections.defaultdict(set)
        for o in obs:
            folder = o.get("_folder") or ""
            d = None
            if hasattr(self, "_parse_sortie_date"):
                try:
                    d = self._parse_sortie_date(folder)
                except Exception:
                    d = None
            esp = (o.get("espece") or "").strip()
            if not esp:
                continue
            if d:
                by_month[d.month].add(esp)
                by_year[d.year].add(esp)
            lieu = (o.get("lieu") or folder or "—").strip()
            # secteur = premier segment du lieu
            sector = lieu.split("–")[0].split("-")[0].strip() or "—"
            by_sector[sector].add(esp)

        lines.append("--- Richesse spécifique par mois (projet / sorties connues) ---")
        mois = ["", "Jan", "Fév", "Mar", "Avr", "Mai", "Juin", "Juil", "Aoû", "Sep", "Oct", "Nov", "Déc"]
        if by_month:
            for m in range(1, 13):
                if m in by_month:
                    lines.append(f"  {mois[m]} : {len(by_month[m])} espèces")
        else:
            lines.append("  (pas assez de dates dans les noms de dossiers)")

        lines.append("")
        lines.append("--- Comparaison par année ---")
        if by_year:
            for y in sorted(by_year.keys()):
                lines.append(f"  {y} : {len(by_year[y])} espèces — " + ", ".join(sorted(by_year[y])[:12]))
            years = sorted(by_year.keys())
            if len(years) >= 2:
                a, b = years[-2], years[-1]
                only_b = by_year[b] - by_year[a]
                only_a = by_year[a] - by_year[b]
                lines.append(f"  Nouveautés {b} vs {a} : " + (", ".join(sorted(only_b)[:15]) or "—"))
                lines.append(f"  Vues en {a} absentes en {b} : " + (", ".join(sorted(only_a)[:15]) or "—"))
        else:
            lines.append("  (ajoutez des sorties au projet avec une date dans le nom de dossier)")

        lines.append("")
        lines.append("--- Richesse par secteur (lieu) ---")
        for sec, espset in sorted(by_sector.items(), key=lambda x: -len(x[1]))[:15]:
            lines.append(f"  {sec} : {len(espset)} espèces")

        box.insert("0.0", "\n".join(lines))
        box.configure(state="disabled")
        ctk.CTkButton(win, text="Fermer", command=win.destroy).pack(pady=8)

    def export_field_package(self):
        """Exporte un paquet terrain (JSON + points) pour OsmAnd / Locus / archive."""
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier", "Ouvrez une sortie ou un carnet.")
            return
        notes_path = os.path.join(self.photo_folder_path, NOTES_FILE)
        notes = {}
        if os.path.isfile(notes_path):
            try:
                with open(notes_path, "r", encoding="utf-8") as f:
                    notes = json.load(f) or {}
            except Exception as e:
                messagebox.showerror("Erreur", str(e))
                return
        dest_dir = filedialog.askdirectory(title="Dossier de destination du paquet terrain")
        if not dest_dir:
            return
        stamp = datetime.now().strftime("%Y%m%d_%H%M")
        name = os.path.basename(self.photo_folder_path.rstrip("\\/")) or "sortie"
        pack_dir = os.path.join(dest_dir, f"GeoExif_terrain_{name}_{stamp}")
        os.makedirs(pack_dir, exist_ok=True)

        # Package JSON
        package = {
            "format": "geoexif_field_package",
            "version": APP_VERSION,
            "sortie": name,
            "exported": datetime.now().isoformat(),
            "observations": notes,
        }
        with open(os.path.join(pack_dir, "observations.json"), "w", encoding="utf-8") as f:
            json.dump(notes, f, ensure_ascii=False, indent=2)
        with open(os.path.join(pack_dir, "package.json"), "w", encoding="utf-8") as f:
            json.dump(package, f, ensure_ascii=False, indent=2)

        # GeoJSON points
        features = []
        for key, data in notes.items():
            try:
                lat = float(data.get("lat"))
                lon = float(data.get("lon"))
            except Exception:
                info = self.photos_data.get(key) or {}
                try:
                    lat, lon = float(info.get("lat")), float(info.get("lon"))
                except Exception:
                    continue
            features.append({
                "type": "Feature",
                "geometry": {"type": "Point", "coordinates": [lon, lat]},
                "properties": {
                    "name": data.get("espece") or key,
                    "categorie": data.get("categorie"),
                    "heure": data.get("heure"),
                    "lieu": data.get("lieu"),
                    "comportement": data.get("comportement"),
                    "certitude": data.get("certitude"),
                    "type_indice": data.get("type_indice"),
                    "nom_scientifique": data.get("nom_scientifique"),
                    "notes": data.get("notes_libres"),
                    "id": key,
                },
            })
        geojson = {"type": "FeatureCollection", "features": features}
        with open(os.path.join(pack_dir, "points.geojson"), "w", encoding="utf-8") as f:
            json.dump(geojson, f, ensure_ascii=False, indent=2)

        # GPX waypoints (OsmAnd / Locus)
        gpx_parts = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<gpx version="1.1" creator="GeoExif" xmlns="http://www.topografix.com/GPX/1/1">',
        ]
        for feat in features:
            lon, lat = feat["geometry"]["coordinates"]
            prop = feat["properties"]
            nm = (prop.get("name") or "obs").replace("&", "&amp;").replace("<", "&lt;")
            desc = " | ".join(
                filter(None, [
                    prop.get("categorie"), prop.get("type_indice"),
                    prop.get("certitude"), prop.get("notes"),
                ])
            ).replace("&", "&amp;").replace("<", "&lt;")
            gpx_parts.append(
                f'  <wpt lat="{lat:.6f}" lon="{lon:.6f}"><name>{nm}</name><desc>{desc}</desc></wpt>'
            )
        gpx_parts.append("</gpx>")
        with open(os.path.join(pack_dir, "waypoints.gpx"), "w", encoding="utf-8") as f:
            f.write("\n".join(gpx_parts))

        if self.gpx_file_path and os.path.isfile(self.gpx_file_path):
            try:
                shutil.copy2(self.gpx_file_path, os.path.join(pack_dir, "trace.gpx"))
            except Exception:
                pass

        self.log(f"Paquet terrain : {pack_dir} ({len(features)} points)")
        messagebox.showinfo(
            "Paquet terrain",
            f"Exporté dans :\n{pack_dir}\n\n"
            f"• observations.json / package.json\n"
            f"• points.geojson\n"
            f"• waypoints.gpx (OsmAnd / Locus)\n"
            f"{len(features)} point(s) GPS"
        )

    def import_field_package(self):
        """Importe observations depuis JSON, GeoJSON ou GPX (waypoints Locus/OsmAnd)."""
        if not self.photo_folder_path:
            if not self.create_or_open_carnet_folder():
                return
        path = filedialog.askopenfilename(
            title="Importer paquet / points terrain",
            filetypes=[
                ("JSON / GeoJSON / GPX", "*.json *.geojson *.gpx"),
                ("Tous", "*.*"),
            ],
        )
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        imported = 0
        notes_path = os.path.join(self.photo_folder_path, NOTES_FILE)
        all_notes = {}
        if os.path.isfile(notes_path):
            try:
                with open(notes_path, "r", encoding="utf-8") as f:
                    all_notes = json.load(f) or {}
            except Exception:
                pass

        try:
            if ext in (".json", ".geojson"):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict) and data.get("type") == "FeatureCollection":
                    for feat in data.get("features") or []:
                        geom = feat.get("geometry") or {}
                        prop = feat.get("properties") or {}
                        coords = geom.get("coordinates") or []
                        if len(coords) < 2:
                            continue
                        lon, lat = float(coords[0]), float(coords[1])
                        cle = prop.get("id") or f"_import_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
                        if cle in all_notes:
                            cle = f"_import_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
                        all_notes[cle] = {
                            "departement": "55 - Meuse",
                            "categorie": prop.get("categorie") or "Autre",
                            "espece": prop.get("name") or prop.get("espece") or "Import",
                            "nombre": prop.get("nombre") or "1",
                            "heure": prop.get("heure") or "",
                            "lieu": prop.get("lieu") or "",
                            "comportement": prop.get("comportement") or "",
                            "certitude": prop.get("certitude") or "",
                            "type_indice": prop.get("type_indice") or "",
                            "nom_scientifique": prop.get("nom_scientifique") or "",
                            "notes_libres": prop.get("notes") or prop.get("desc") or "Import terrain",
                            "sans_photo": True,
                            "lat": lat,
                            "lon": lon,
                        }
                        imported += 1
                elif isinstance(data, dict) and "observations" in data:
                    for k, v in (data.get("observations") or {}).items():
                        if k not in all_notes:
                            all_notes[k] = v
                            imported += 1
                elif isinstance(data, dict):
                    # observations.json brut
                    for k, v in data.items():
                        if isinstance(v, dict) and k not in all_notes:
                            all_notes[k] = v
                            imported += 1
            elif ext == ".gpx":
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    raw = f.read()
                for m in re.finditer(
                    r'<wpt[^>]*lat="([^"]+)"[^>]*lon="([^"]+)"[^>]*>(.*?)</wpt>',
                    raw, re.S | re.I,
                ):
                    lat, lon, body = float(m.group(1)), float(m.group(2)), m.group(3)
                    nm = re.search(r"<name>(.*?)</name>", body, re.S | re.I)
                    desc = re.search(r"<desc>(.*?)</desc>", body, re.S | re.I)
                    name = (nm.group(1).strip() if nm else "Waypoint")
                    cle = f"_gpx_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
                    all_notes[cle] = {
                        "departement": "55 - Meuse",
                        "categorie": "Autre",
                        "espece": name,
                        "nombre": "1",
                        "notes_libres": (desc.group(1).strip() if desc else "Import GPX"),
                        "sans_photo": True,
                        "lat": lat,
                        "lon": lon,
                        "type_indice": "",
                    }
                    imported += 1
        except Exception as e:
            messagebox.showerror("Import", str(e))
            return

        if imported:
            try:
                self._save_notes_dict(all_notes, force_backup=True)
            except Exception:
                with open(notes_path, "w", encoding="utf-8") as f:
                    json.dump(all_notes, f, ensure_ascii=False, indent=2)
            # refresh map data
            for k, d in all_notes.items():
                if d.get("lat") is not None and d.get("lon") is not None:
                    self.photos_data[k] = {
                        "path": "", "lat": d["lat"], "lon": d["lon"],
                        "date": d.get("heure") or "", "sans_photo": True,
                    }
            try:
                self.refresh_map_markers()
                self.refresh_daily_counter()
            except Exception:
                pass
            self.log(f"Import terrain : {imported} observation(s)")
            messagebox.showinfo("Import", f"{imported} observation(s) importée(s) dans le carnet.")
        else:
            messagebox.showinfo("Import", "Aucune observation nouvelle trouvée dans le fichier.")

    def _clear_map_measure_graphics(self):
        """Efface marqueurs et corde de mesure sur la carte."""
        for mk in getattr(self, "_measure_markers", []) or []:
            try:
                mk.delete()
            except Exception:
                pass
        self._measure_markers = []
        path = getattr(self, "_measure_path", None)
        if path is not None:
            try:
                path.delete()
            except Exception:
                pass
        self._measure_path = None
        self._measure_pts = []

    def toggle_map_measure(self):
        """Mode mesure : 2 clics → corde + distance (km). Recliquer pour quitter et effacer."""
        if not hasattr(self, "_measure_mode"):
            self._measure_mode = False
        self._measure_mode = not self._measure_mode
        self._clear_map_measure_graphics()

        if self._measure_mode:
            self.log("Mesure : cliquez 2 points — une corde s'affiche. Recliquez « Mesure » pour effacer.")
            try:
                self.map_widget.add_left_click_map_command(self._on_map_measure_click)
            except Exception:
                pass
            messagebox.showinfo(
                "Mesure de distance",
                "1) Cliquez le premier point\n"
                "2) Cliquez le second → corde + distance\n"
                "3) Recliquez « Mesure carte » pour effacer et quitter\n"
                "   (ou continuez : un nouveau couple de points remplace le précédent)."
            )
        else:
            self.log("Mesure carte : désactivée — corde et points effacés.")
            try:
                if getattr(self, "_place_obs_mode", False):
                    self.map_widget.add_left_click_map_command(self._on_map_left_click)
                else:
                    self.map_widget.add_left_click_map_command(None)
            except Exception:
                pass

    def _on_map_measure_click(self, coords):
        if not getattr(self, "_measure_mode", False):
            return
        lat, lon = float(coords[0]), float(coords[1])
        if not hasattr(self, "_measure_pts"):
            self._measure_pts = []
        if not hasattr(self, "_measure_markers"):
            self._measure_markers = []

        if len(self._measure_pts) >= 2:
            self._clear_map_measure_graphics()

        self._measure_pts.append((lat, lon))
        n = len(self._measure_pts)
        try:
            mk = self.map_widget.set_marker(
                lat, lon,
                text=("A" if n == 1 else "B"),
                marker_color_circle="#f0b45c",
                marker_color_outside="#d99a3e",
            )
            self._measure_markers.append(mk)
        except Exception:
            try:
                mk = self.map_widget.set_marker(lat, lon, text=f"M{n}")
                self._measure_markers.append(mk)
            except Exception:
                pass

        if len(self._measure_pts) >= 2:
            (lat1, lon1), (lat2, lon2) = self._measure_pts[0], self._measure_pts[1]
            try:
                self._measure_path = self.map_widget.set_path(
                    [(lat1, lon1), (lat2, lon2)],
                    color="#f0b45c",
                    width=3,
                )
            except TypeError:
                try:
                    self._measure_path = self.map_widget.set_path([(lat1, lon1), (lat2, lon2)])
                except Exception:
                    self._measure_path = None
            except Exception:
                self._measure_path = None

            r = 6371.0
            p1, p2 = math.radians(lat1), math.radians(lat2)
            dp = math.radians(lat2 - lat1)
            dl = math.radians(lon2 - lon1)
            x = math.sin(dp / 2) ** 2 + math.cos(p1) * math.cos(p2) * math.sin(dl / 2) ** 2
            dist = 2 * r * math.asin(min(1.0, math.sqrt(x)))
            self.log(f"Mesure : {dist:.2f} km ({dist * 1000:.0f} m) — corde affichée")
            messagebox.showinfo(
                "Distance",
                f"{dist:.2f} km\n({dist * 1000:.0f} m)\n\n"
                "La corde reste sur la carte.\n"
                "• Nouveau clic = nouvelle mesure\n"
                "• Bouton « Mesure carte » = tout effacer"
            )

    def _toggle_gps_advanced(self):
        self._gps_advanced_visible = not getattr(self, "_gps_advanced_visible", False)
        if self._gps_advanced_visible:
            # Re-pack juste avant le bouton sync pour rester dans le flux
            self.gps_advanced_frame.pack(fill="x", padx=10, pady=(0, 4), before=self.btn_sync)
            self.btn_gps_advanced.configure(text="▾ Options avancées")
        else:
            self.gps_advanced_frame.pack_forget()
            self.btn_gps_advanced.configure(text="▸ Options avancées")

    def _apply_ctk_theme_defaults(self):
        """Force les couleurs par défaut CTk (nouveaux widgets + cohérence globale)."""
        try:
            tm = ctk.ThemeManager.theme
        except Exception:
            return
        text = UI.get("text", "#f3f7f4")
        text_dim = UI.get("text_dim", text)
        muted = UI.get("text_muted", text_dim)
        card = UI.get("card", "#1c2620")
        card_alt = UI.get("card_alt", "#24302a")
        border = UI.get("border", "#354840")
        accent = UI.get("accent", "#4ec4b0")
        accent_h = UI.get("accent_hover", accent)
        bg = UI.get("bg", "#101612")

        def set_keys(section, mapping):
            if section not in tm:
                return
            for k, v in mapping.items():
                try:
                    tm[section][k] = v
                except Exception:
                    pass

        set_keys("CTkLabel", {"text_color": [text, text]})
        set_keys("CTkButton", {
            "text_color": [text, text],
            "text_color_disabled": [muted, muted],
            "fg_color": [card_alt, card_alt],
            "hover_color": [accent, accent],
            "border_color": [border, border],
        })
        set_keys("CTkEntry", {
            "text_color": [text, text],
            "fg_color": [card_alt, card_alt],
            "border_color": [border, border],
            "placeholder_text_color": [muted, muted],
        })
        set_keys("CTkTextbox", {
            "text_color": [text_dim, text_dim],
            "fg_color": [card, card],
            "border_color": [border, border],
        })
        set_keys("CTkComboBox", {
            "text_color": [text, text],
            "fg_color": [card, card],
            "border_color": [border, border],
            "button_color": [accent, accent],
            "button_hover_color": [accent_h, accent_h],
            "dropdown_fg_color": [card, card],
            "dropdown_text_color": [text, text],
            "dropdown_hover_color": [card_alt, card_alt],
        })
        set_keys("CTkOptionMenu", {
            "text_color": ["#ffffff", "#ffffff"],
            "fg_color": [accent, accent],
            "button_color": [accent_h, accent_h],
            "button_hover_color": [accent_h, accent_h],
            "dropdown_fg_color": [card, card],
            "dropdown_text_color": [text, text],
            "dropdown_hover_color": [card_alt, card_alt],
        })
        set_keys("CTkCheckBox", {
            "text_color": [text, text],
            "fg_color": [accent, accent],
            "hover_color": [accent_h, accent_h],
            "border_color": [border, border],
        })
        set_keys("CTkRadioButton", {
            "text_color": [text, text],
            "fg_color": [accent, accent],
            "hover_color": [accent_h, accent_h],
            "border_color": [border, border],
        })
        set_keys("CTkSwitch", {
            "text_color": [text, text],
            "progress_color": [accent, accent],
            "button_color": [card_alt, card_alt],
            "button_hover_color": [accent, accent],
        })
        set_keys("CTkSlider", {
            "progress_color": [accent, accent],
            "button_color": [accent, accent],
            "button_hover_color": [accent_h, accent_h],
            "fg_color": [card_alt, card_alt],
        })
        set_keys("CTkProgressBar", {
            "progress_color": [accent, accent],
            "fg_color": [card_alt, card_alt],
        })
        set_keys("CTkFrame", {
            "fg_color": [card, card],
            "border_color": [border, border],
            "top_fg_color": [card, card],
        })
        set_keys("CTkScrollableFrame", {
            "fg_color": [card, card],
            "label_fg_color": [card, card],
            "label_text_color": [text, text],
            "border_color": [border, border],
        })
        set_keys("CTkTabview", {
            "fg_color": [card, card],
            "segmented_button_fg_color": [card_alt, card_alt],
            "segmented_button_selected_color": [accent, accent],
            "segmented_button_selected_hover_color": [accent_h, accent_h],
            "segmented_button_unselected_color": [card_alt, card_alt],
            "segmented_button_unselected_hover_color": [border, border],
            "text_color": [text, text],
        })
        set_keys("CTkSegmentedButton", {
            "fg_color": [card_alt, card_alt],
            "selected_color": [accent, accent],
            "selected_hover_color": [accent_h, accent_h],
            "unselected_color": [card_alt, card_alt],
            "unselected_hover_color": [border, border],
            "text_color": [text, text],
            "text_color_disabled": [muted, muted],
        })
        set_keys("CTkToplevel", {"fg_color": [bg, bg]})
        set_keys("CTk", {"fg_color": [bg, bg]})
        set_keys("DropdownMenu", {
            "fg_color": [card, card],
            "hover_color": [card_alt, card_alt],
            "text_color": [text, text],
        })

    def _style_dialog(self, win):
        """Applique le thème courant a une fenetre dialogue (fond + textes + champs)."""
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        try:
            self._restyle_widget_tree(win, depth=0)
        except Exception:
            pass
        # Force lisibilite des labels / entrees de ce dialogue
        text = UI.get("text", "#111111")
        dim = UI.get("text_dim", "#555555")
        card = UI.get("card", "#ffffff")
        card_alt = UI.get("card_alt", "#f0f0f0")
        accent = UI.get("accent", "#2f9e5f")
        border = UI.get("border", "#cccccc")

        def walk(w, depth=0):
            if depth > 14:
                return
            cls = w.__class__.__name__
            try:
                if cls == "CTkLabel":
                    # ne pas ecraser les labels sur boutons colores si texte deja blanc volontaire
                    cur = None
                    try:
                        cur = w.cget("text_color")
                    except Exception:
                        pass
                    if cur in (None, "transparent", "", "#999999", "#888888", "#888", "#999"):
                        w.configure(text_color=text)
                    elif isinstance(cur, (list, tuple)) and cur[0] in ("#999999", "#888888"):
                        w.configure(text_color=text)
                    else:
                        # forcer quand meme si contraste douteux (gris fixes)
                        try:
                            w.configure(text_color=text)
                        except Exception:
                            pass
                elif cls == "CTkEntry":
                    w.configure(
                        text_color=text,
                        fg_color=card_alt,
                        border_color=border,
                        placeholder_text_color=dim,
                    )
                elif cls == "CTkTextbox":
                    w.configure(text_color=text, fg_color=card, border_color=border)
                elif cls == "CTkComboBox":
                    w.configure(
                        text_color=text,
                        fg_color=card,
                        border_color=border,
                        button_color=accent,
                        button_hover_color=UI.get("accent_hover", accent),
                        dropdown_fg_color=card,
                        dropdown_text_color=text,
                    )
                elif cls == "CTkOptionMenu":
                    w.configure(
                        text_color="#ffffff",
                        fg_color=accent,
                        button_color=UI.get("accent_hover", accent),
                        dropdown_fg_color=card,
                        dropdown_text_color=text,
                    )
                elif cls in ("CTkFrame", "CTkScrollableFrame"):
                    try:
                        fc = w.cget("fg_color")
                        if fc in ("transparent", "Transparent", None, ""):
                            pass
                        elif fc in ("#2b2b2b", "#3a3a3a", "#1a1a1a"):
                            w.configure(fg_color=card)
                    except Exception:
                        pass
                elif cls == "CTkButton":
                    try:
                        fg = w.cget("fg_color")
                        # boutons gris sombres illisibles en theme clair
                        if fg in ("#2b2b2b", "#3a3a3a", "#333333", "#1a1a1a", "#242424"):
                            w.configure(
                                fg_color=card_alt,
                                hover_color=UI.get("accent_hover", accent),
                                text_color=text,
                            )
                    except Exception:
                        pass
            except Exception:
                pass
            try:
                for ch in w.winfo_children():
                    walk(ch, depth + 1)
            except Exception:
                pass

        try:
            walk(win)
        except Exception:
            pass

    def apply_theme(self, name="sombre", *, save=True, silent=False):
        """Applique un thème partout : fenêtre principale + dialogues ouverts."""
        name = (name or "sombre").lower().strip()
        if name not in THEMES:
            name = "sombre"
        palette = THEMES[name]
        UI.clear()
        UI.update(palette)
        self._theme_name = name

        try:
            ctk.set_appearance_mode(palette.get("ctk_mode", "Dark"))
        except Exception:
            pass

        # Défauts CTk pour tous les prochains widgets + cohérence
        try:
            self._apply_ctk_theme_defaults()
        except Exception:
            pass

        # Fenêtre racine
        try:
            self.configure(fg_color=UI["bg"])
        except Exception:
            pass

        # Zones principales (références connues)
        for attr, fg in (
            ("sidebar", UI["sidebar"]),
            ("main_frame", UI["bg"]),
            ("icon_toolbar", UI["card"]),
            ("card_gps", UI["card"]),
            ("card_saison", UI["card"]),
            ("log_box", UI["card"]),
            ("media_explorer", UI["card"]),
            ("preview_container", UI["card"]),
            ("upper_panel", UI["card"]),
            ("compteur_frame", UI["card"]),
            ("form_frame", UI["card"]),
            ("thumb_scroll", UI["card_alt"]),
        ):
            w = getattr(self, attr, None)
            if w is None:
                continue
            try:
                w.configure(fg_color=fg, border_color=UI.get("border", fg))
            except Exception:
                try:
                    w.configure(fg_color=fg)
                except Exception:
                    pass

        try:
            self.log_box.configure(text_color=UI["text_dim"])
        except Exception:
            pass

        tv = getattr(self, "tab_view", None)
        if tv is not None:
            try:
                tv.configure(
                    fg_color=UI["card"],
                    segmented_button_fg_color=UI["card_alt"],
                    segmented_button_selected_color=UI["accent"],
                    segmented_button_selected_hover_color=UI["accent_hover"],
                    segmented_button_unselected_color=UI["card_alt"],
                    segmented_button_unselected_hover_color=UI["border"],
                    text_color=UI["text"],
                )
            except Exception:
                pass

        try:
            self.mode_switch.configure(
                selected_color=UI["accent"],
                selected_hover_color=UI["accent_hover"],
                unselected_color=UI["card_alt"],
                unselected_hover_color=UI["border"],
                text_color=UI["text"],
            )
        except Exception:
            pass

        for attr, fg, hover in (
            ("btn_browse_photos", UI["card_alt"], UI["border"]),
            ("btn_browse_gpx", UI["card_alt"], UI["border"]),
            ("btn_outils_hub", UI["card_alt"], UI["border"]),
            ("btn_undo_sync", UI["card_alt"], UI["border"]),
            ("btn_toggle_sidebar", UI["card_elevated"], UI["accent"]),
        ):
            w = getattr(self, attr, None)
            if w is None:
                continue
            try:
                w.configure(
                    fg_color=fg, hover_color=hover,
                    text_color=UI["text"], border_color=UI["border"],
                )
            except Exception:
                try:
                    w.configure(fg_color=fg, hover_color=hover, text_color=UI["text"])
                except Exception:
                    pass

        try:
            self.title_label.configure(text_color=UI["text"])
        except Exception:
            pass
        try:
            self.lbl_progress.configure(text_color=UI["text_muted"])
        except Exception:
            pass
        try:
            self.toolbar_mode_lbl.configure(text_color=UI["text_accent"])
        except Exception:
            pass
        try:
            self.selecteur_carte.configure(
                fg_color=UI["card_alt"], button_color=UI["accent"],
                button_hover_color=UI["accent_hover"], text_color=UI["text"],
                dropdown_fg_color=UI["card"], dropdown_text_color=UI["text"],
            )
        except Exception:
            pass
        try:
            self.entry_geosync.configure(
                fg_color=UI["card_alt"], text_color=UI["text"],
                border_color=UI["border"], placeholder_text_color=UI["text_muted"],
            )
        except Exception:
            pass
        try:
            self.progress_bar.configure(progress_color=UI["accent"], fg_color=UI["card_alt"])
        except Exception:
            pass

        # Barre d'icônes
        try:
            for b in getattr(self, "_toolbar_buttons", []) or []:
                b.configure(
                    fg_color=UI["card_alt"], hover_color=UI["accent"],
                    text_color=UI["text"], border_color=UI.get("border"),
                )
        except Exception:
            pass
        try:
            for s in getattr(self, "_toolbar_seps", []) or []:
                s.configure(fg_color=UI["border"])
        except Exception:
            pass

        # Listbox carnet
        lb = getattr(self, "photo_listbox", None)
        if lb is not None:
            try:
                lb.configure(
                    bg=UI["card_alt"], fg=UI.get("list_pending", UI["text"]),
                    selectbackground=UI["accent"], selectforeground=UI["bg"],
                    highlightbackground=UI["border"], highlightcolor=UI["border"],
                )
            except Exception:
                pass
            try:
                self._refresh_listbox_annotation_status()
            except Exception:
                pass

        try:
            self.lbl_preview.configure(fg_color=UI["card_alt"], text_color=UI["text_dim"])
        except Exception:
            pass
        try:
            self.weather_box.configure(
                fg_color=UI["card_alt"], text_color=UI["text_dim"], border_color=UI["border"],
            )
        except Exception:
            pass
        try:
            self.note_text.configure(
                fg_color=UI["card_alt"], text_color=UI["text"], border_color=UI["border"],
            )
        except Exception:
            pass

        # Arbre principal + toutes les fenêtres secondaires ouvertes
        roots = [self]
        try:
            for child in self.winfo_children():
                try:
                    if child.winfo_class() in ("Toplevel", "CTkToplevel") or isinstance(
                        child, (ctk.CTkToplevel, tk.Toplevel)
                    ):
                        roots.append(child)
                except Exception:
                    pass
        except Exception:
            pass
        # Toplevels CTk parfois hors winfo_children selon version
        try:
            for w in self._winfo_toplevels_safe():
                if w not in roots:
                    roots.append(w)
        except Exception:
            pass

        for root in roots:
            try:
                if root is not self:
                    try:
                        root.configure(fg_color=UI["bg"])
                    except Exception:
                        pass
                self._restyle_widget_tree(root)
            except Exception:
                pass

        if save:
            cfg = dict(self.app_config or {})
            cfg["theme"] = name
            try:
                self._save_app_config(cfg)
            except Exception:
                pass

        if not silent:
            labels = {"sombre": "Sombre (forêt)", "clair": "Clair", "papier": "Papier"}
            self.log(f"Thème : {labels.get(name, name)} (appliqué partout)")
            try:
                if hasattr(self, "toolbar_mode_lbl"):
                    mode = "Saison" if getattr(self, "_ui_mode", "") == "saison" else "Sortie"
                    self.toolbar_mode_lbl.configure(
                        text=f"  Mode {mode}  ·  {labels.get(name, name)}  ",
                        text_color=UI.get("text_accent", UI["text"]),
                    )
            except Exception:
                pass

    def _winfo_toplevels_safe(self):
        """Liste les Toplevel connus rattachés à l'app."""
        found = []
        try:
            for attr in dir(self):
                if attr.startswith("_"):
                    continue
                try:
                    obj = getattr(self, attr)
                except Exception:
                    continue
                try:
                    if isinstance(obj, (ctk.CTkToplevel, tk.Toplevel)) and obj.winfo_exists():
                        found.append(obj)
                except Exception:
                    pass
        except Exception:
            pass
        return found

    def _restyle_widget_tree(self, widget, depth=0):
        """Parcourt récursivement et applique couleurs texte/fond du thème actif."""
        if depth > 16:
            return
        try:
            cls = widget.__class__.__name__
        except Exception:
            return

        action_btns = {
            getattr(self, n, None)
            for n in (
                "btn_sync", "btn_carnet_only", "btn_ai_summary", "btn_manual_obs",
                "btn_edit_gps", "btn_place_obs", "btn_save_note",
            )
        }

        # Mots-clés de boutons métier à ne pas griser
        keep_colored_txt = (
            "Été", "Hiver", "Synchroniser", "Sans photo", "Brief", "Carnet sans",
            "Placer observation", "Enregistrer", "Ouvrir", "INPN", "GBIF",
        )

        try:
            if widget in action_btns:
                pass
            elif cls in ("CTkFrame", "CTkScrollableFrame"):
                try:
                    fc = widget.cget("fg_color")
                except Exception:
                    fc = None
                if fc not in (None, "transparent", "Transparent", ""):
                    if widget in (
                        getattr(self, "sidebar", None),
                        getattr(self, "main_frame", None),
                        getattr(self, "icon_toolbar", None),
                    ):
                        pass
                    else:
                        try:
                            widget.configure(fg_color=UI["card"], border_color=UI["border"])
                        except Exception:
                            try:
                                widget.configure(fg_color=UI["card"])
                            except Exception:
                                pass
            elif cls == "CTkLabel":
                try:
                    widget.configure(text_color=UI["text"])
                except Exception:
                    pass
            elif cls == "CTkTextbox":
                try:
                    widget.configure(
                        fg_color=UI["card"], text_color=UI["text_dim"],
                        border_color=UI["border"],
                    )
                except Exception:
                    pass
            elif cls == "CTkEntry":
                try:
                    widget.configure(
                        fg_color=UI["card_alt"], text_color=UI["text"],
                        border_color=UI["border"],
                        placeholder_text_color=UI["text_muted"],
                    )
                except Exception:
                    pass
            elif cls in ("CTkComboBox", "CTkOptionMenu"):
                try:
                    widget.configure(
                        fg_color=UI["card_alt"], text_color=UI["text"],
                        button_color=UI["accent"],
                        button_hover_color=UI["accent_hover"],
                        dropdown_fg_color=UI["card"],
                        dropdown_text_color=UI["text"],
                        dropdown_hover_color=UI["card_alt"],
                    )
                except Exception:
                    try:
                        widget.configure(fg_color=UI["card_alt"], text_color=UI["text"])
                    except Exception:
                        pass
            elif cls == "CTkCheckBox":
                try:
                    widget.configure(text_color=UI["text"], border_color=UI["border"])
                except Exception:
                    pass
            elif cls == "CTkRadioButton":
                try:
                    widget.configure(text_color=UI["text"])
                except Exception:
                    pass
            elif cls == "CTkSwitch":
                try:
                    widget.configure(text_color=UI["text"], progress_color=UI["accent"])
                except Exception:
                    pass
            elif cls == "CTkSlider":
                try:
                    widget.configure(
                        progress_color=UI["accent"], button_color=UI["accent"],
                        fg_color=UI["card_alt"],
                    )
                except Exception:
                    pass
            elif cls == "CTkProgressBar":
                try:
                    widget.configure(progress_color=UI["accent"], fg_color=UI["card_alt"])
                except Exception:
                    pass
            elif cls == "CTkSegmentedButton":
                try:
                    widget.configure(
                        selected_color=UI["accent"],
                        selected_hover_color=UI["accent_hover"],
                        unselected_color=UI["card_alt"],
                        unselected_hover_color=UI["border"],
                        text_color=UI["text"],
                    )
                except Exception:
                    pass
            elif cls == "CTkTabview":
                try:
                    widget.configure(
                        fg_color=UI["card"],
                        segmented_button_fg_color=UI["card_alt"],
                        segmented_button_selected_color=UI["accent"],
                        segmented_button_selected_hover_color=UI["accent_hover"],
                        segmented_button_unselected_color=UI["card_alt"],
                        segmented_button_unselected_hover_color=UI["border"],
                        text_color=UI["text"],
                    )
                except Exception:
                    pass
            elif cls == "CTkButton" and widget not in action_btns:
                try:
                    txt = str(widget.cget("text") or "")
                except Exception:
                    txt = ""
                if any(k in txt for k in keep_colored_txt):
                    try:
                        widget.configure(text_color=UI["text"] if "INPN" in txt or "Ouvrir" in txt else None)
                    except Exception:
                        pass
                    # garder couleurs métier des boutons action textuels
                    if any(k in txt for k in ("Été", "Hiver", "Synchroniser", "Sans photo", "Brief", "Carnet sans", "Placer", "Enregistrer")):
                        pass
                    else:
                        try:
                            widget.configure(text_color=UI["text"])
                        except Exception:
                            pass
                else:
                    try:
                        w = int(widget.cget("width") or 0)
                    except Exception:
                        w = 0
                    try:
                        if w and w <= 48:
                            widget.configure(
                                fg_color=UI["card_alt"], hover_color=UI["accent"],
                                text_color=UI["text"],
                            )
                        else:
                            widget.configure(
                                fg_color=UI["card_alt"], hover_color=UI["border"],
                                text_color=UI["text"], border_color=UI["border"],
                            )
                    except Exception:
                        try:
                            widget.configure(text_color=UI["text"])
                        except Exception:
                            pass
            elif cls == "Listbox":
                try:
                    widget.configure(
                        bg=UI["card_alt"], fg=UI.get("list_pending", UI["text"]),
                        selectbackground=UI["accent"], selectforeground=UI["bg"],
                        highlightbackground=UI["border"],
                    )
                except Exception:
                    pass
            elif cls == "Text":
                try:
                    widget.configure(
                        bg=UI["card_alt"], fg=UI["text"],
                        insertbackground=UI["text"],
                        highlightbackground=UI["border"],
                    )
                except Exception:
                    pass
        except Exception:
            pass

        try:
            for child in widget.winfo_children():
                self._restyle_widget_tree(child, depth + 1)
        except Exception:
            pass

    def open_theme_picker(self):
        win = ctk.CTkToplevel(self)
        win.title("Thème d'interface")
        win.geometry("420x280")
        win.configure(fg_color=UI.get("bg", "#101612"))
        self._prepare_tool_window(win)
        ctk.CTkLabel(
            win, text="Choisir un thème",
            font=ctk.CTkFont(size=16, weight="bold"),
        ).pack(pady=(16, 8))
        ctk.CTkLabel(
            win,
            text="Le thème est mémorisé pour les prochains démarrages.\n"
                 "Certains boutons gardent leur couleur d'action.",
            font=ctk.CTkFont(size=12), text_color=UI.get("text_dim"),
            justify="center",
        ).pack(pady=(0, 12))

        def pick(name):
            self.apply_theme(name, save=True)
            win.destroy()

        for name, label, color in (
            ("sombre", "🌑  Sombre (forêt)", THEMES["sombre"]["accent"]),
            ("clair", "☀  Clair", THEMES["clair"]["accent"]),
            ("papier", "📜  Papier (carnet)", THEMES["papier"]["accent"]),
        ):
            ctk.CTkButton(
                win, text=label, height=36, corner_radius=10,
                fg_color=color, hover_color=THEMES[name]["accent_hover"],
                command=lambda n=name: pick(n),
            ).pack(fill="x", padx=40, pady=5)

    def _attach_tooltip(self, widget, text):
        """Infobulle au survol (barre d'icônes, boutons)."""
        if not text:
            return

        def show(_event=None):
            self._hide_tooltip()
            try:
                tip = tk.Toplevel(self)
                tip.wm_overrideredirect(True)
                tip.attributes("-topmost", True)
                lbl = tk.Label(
                    tip, text=text, justify="left",
                    background=UI.get("card_elevated", "#2c3a32"),
                    foreground=UI.get("text", "#f3f7f4"),
                    relief="solid", borderwidth=1,
                    font=("Segoe UI", 9), padx=8, pady=4,
                )
                lbl.pack()
                x = widget.winfo_rootx() + 8
                y = widget.winfo_rooty() + widget.winfo_height() + 6
                tip.geometry(f"+{x}+{y}")
                self._tooltip_win = tip
            except Exception:
                self._tooltip_win = None

        def schedule(e=None):
            self._hide_tooltip()
            self._tooltip_after = self.after(450, show)

        def cancel(e=None):
            aid = getattr(self, "_tooltip_after", None)
            if aid:
                try:
                    self.after_cancel(aid)
                except Exception:
                    pass
                self._tooltip_after = None
            self._hide_tooltip()

        widget.bind("<Enter>", schedule, add="+")
        widget.bind("<Leave>", cancel, add="+")
        widget.bind("<ButtonPress>", cancel, add="+")

    def _hide_tooltip(self):
        tip = getattr(self, "_tooltip_win", None)
        if tip is not None:
            try:
                tip.destroy()
            except Exception:
                pass
        self._tooltip_win = None

    def _build_icon_toolbar(self):
        """Barre d'icônes horizontale sous le menu (esprit QGIS / FastStone)."""
        bar = ctk.CTkFrame(
            self, fg_color=UI.get("card", "#1c2620"), height=46,
            corner_radius=0, border_width=0,
        )
        bar.grid(row=0, column=0, columnspan=2, sticky="ew")
        bar.grid_columnconfigure(20, weight=1)
        self.icon_toolbar = bar

        def grp():
            f = ctk.CTkFrame(bar, fg_color="transparent")
            f.pack(side="left", padx=(6, 2), pady=6)
            return f

        self._toolbar_buttons = []
        self._toolbar_seps = []
        self._toolbar_groups = []

        def ibtn(parent, text, cmd, tip=""):
            b = ctk.CTkButton(
                parent, text=text, width=36, height=32,
                font=ctk.CTkFont(size=14),
                fg_color=UI.get("card_alt", "#24302a"),
                hover_color=UI.get("accent", "#4ec4b0"),
                text_color=UI.get("text", "#f3f7f4"),
                corner_radius=8, command=cmd,
            )
            b.pack(side="left", padx=2)
            self._toolbar_buttons.append(b)
            if tip:
                self._attach_tooltip(b, tip)
            return b

        def sep(parent):
            s = ctk.CTkFrame(parent, width=1, height=28, fg_color=UI.get("border", "#354840"))
            s.pack(side="left", padx=6, pady=2)
            self._toolbar_seps.append(s)
            return s

        # Fichier / dossiers
        g = grp()
        ibtn(g, "📂", self.select_photo_folder, "Ouvrir un dossier photos / sortie")
        ibtn(g, "👁️", self.create_or_open_carnet_folder, "Carnet sans photos")
        ibtn(g, "🛰️", self.select_gpx_file, "Charger une trace GPX")
        sep(g)

        # GPS
        g2 = grp()
        ibtn(g2, "⚡", self.start_sync_thread, "Synchroniser photos & GPX")
        ibtn(g2, "☀", lambda: self.apply_geosync_preset("-1:00:00", "été"), "Décalage été −1 h")
        ibtn(g2, "❄", lambda: self.apply_geosync_preset("-2:00:00", "hiver"), "Décalage hiver −2 h")
        ibtn(g2, "↩️", self.undo_last_sync, "Annuler la dernière synchro GPS")
        sep(g2)

        # Carnet / carte
        g3 = grp()
        ibtn(g3, "💾", lambda: self.save_current_note(silent=False), "Enregistrer l'observation")
        ibtn(g3, "📝", self.open_manual_observation_dialog, "Observation sans photo")
        ibtn(g3, "🎧", self.open_import_son_menu, "Import son (BirdNET Live / Chirpity / Birda)")
        ibtn(g3, "🤖", self.generate_ai_summary, "Brief IA de la sortie")
        ibtn(g3, "🗺️", self.open_multi_sorties_map, "Carte cumulée de toutes les sorties")
        ibtn(g3, "📏", self.toggle_map_measure, "Mesurer une distance sur la carte")
        ibtn(g3, "📸", self.capture_map_screenshot, "Capture de la carte")
        ibtn(g3, "📋", self.open_gps_metadata_panel, "Métadonnées GPS de la sélection")
        ibtn(g3, "📷", self.open_devices_manager, "Caméras de chasse & enregistreurs")
        sep(g3)

        # Saison / outils
        g4 = grp()
        ibtn(g4, "📁", self.open_project_hub, "Projet Meuse / saison")
        ibtn(g4, "📊", self.open_analysis_panel, "Analyse effort et richesse")
        ibtn(g4, "🌿", self.open_season_suggestions, "Cibles d'espèces selon la saison")
        ibtn(g4, "📦", self.export_field_package, "Exporter le paquet terrain")
        ibtn(g4, "✨", self.open_outils_hub, "Tous les outils et rapports")
        sep(g4)

        # Afficher / masquer panneau + thème
        g5 = grp()
        ibtn(g5, "☰", self.toggle_sidebar, "Afficher / masquer la barre latérale")
        ibtn(g5, "🎨", self.open_theme_picker, "Changer le thème (sombre / clair / papier)")

        # Indicateur mode à droite
        self.toolbar_mode_lbl = ctk.CTkLabel(
            bar, text="  Mode Sortie  ",
            font=ctk.CTkFont(size=11, weight="bold"),
            text_color=UI.get("text_accent", "#8eebd8"),
        )
        self.toolbar_mode_lbl.pack(side="right", padx=12)

    def detach_map_window(self):
        """Ouvre la carte dans une fenêtre séparée (idéal 2ᵉ écran)."""
        if not hasattr(self, "_detached_wins"):
            self._detached_wins = {}
        existing = self._detached_wins.get("map")
        if existing is not None:
            try:
                if existing.winfo_exists():
                    existing.lift()
                    existing.focus_force()
                    return
            except Exception:
                pass

        # Masquer la carte intégrée
        try:
            self.map_widget.grid_forget()
        except Exception:
            pass
        if not hasattr(self, "_map_detached_placeholder") or self._map_detached_placeholder is None:
            self._map_detached_placeholder = ctk.CTkLabel(
                self.tab_map,
                text="Carte détachée dans une fenêtre séparée.\n"
                     "Affichage → Réattacher les fenêtres\npour la ramener ici.",
                font=ctk.CTkFont(size=14),
                text_color=UI.get("text_dim"),
            )
        try:
            self._map_detached_placeholder.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
        except Exception:
            pass

        win = ctk.CTkToplevel(self)
        win.title("GeoExif — Carte")
        cfg = self._load_app_config() or {}
        geo = (cfg.get("workspace") or {}).get("map_geometry") or "1100x750+100+80"
        try:
            win.geometry(geo)
        except Exception:
            win.geometry("1100x750")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass

        # barre mini
        bar = ctk.CTkFrame(win, fg_color=UI.get("card"), height=36)
        bar.pack(fill="x", padx=6, pady=4)
        ctk.CTkButton(
            bar, text="Réattacher", width=100, height=28,
            command=lambda: self._reattach_map(win),
        ).pack(side="left", padx=4)
        ctk.CTkButton(
            bar, text="Rafraîchir marqueurs", width=140, height=28,
            command=self.refresh_map_markers,
        ).pack(side="left", padx=4)
        ctk.CTkLabel(
            bar, text="Glissez cette fenêtre sur le 2ᵉ écran",
            text_color=UI.get("text_dim"), font=ctk.CTkFont(size=11),
        ).pack(side="right", padx=8)

        host = ctk.CTkFrame(win, fg_color=UI.get("card"))
        host.pack(fill="both", expand=True, padx=6, pady=(0, 6))
        host.grid_rowconfigure(0, weight=1)
        host.grid_columnconfigure(0, weight=1)

        try:
            self.map_widget.grid_forget()
        except Exception:
            pass
        # Reparent impossible en Tk → on garde le même widget en le packant dans win via un conteneur tk
        # Astuce : placer map_widget dans host (même App, master peut rester tab_map pour certains widgets)
        # TkinterMapView exige souvent le même parent ; on recrée si besoin.
        try:
            self.map_widget.master
            self.map_widget.grid(in_=host, row=0, column=0, sticky="nsew")
        except Exception:
            try:
                # recreation
                pos = (49.16, 5.38)
                z = 10
                try:
                    pos = self.map_widget.get_position()
                    z = self.map_widget.zoom
                except Exception:
                    pass
                self.map_widget = TkinterMapView(host, corner_radius=10)
                self.map_widget.grid(row=0, column=0, sticky="nsew")
                try:
                    self.map_widget.set_position(pos[0], pos[1])
                    self.map_widget.set_zoom(int(z) if z else 10)
                except Exception:
                    self.map_widget.set_position(49.1627, 5.3854)
                    self.map_widget.set_zoom(9)
                try:
                    self.changer_fond_carte("Plan IGN (Moderne)")
                except Exception:
                    pass
                try:
                    self.map_widget.add_left_click_map_command(self._on_map_left_click)
                except Exception:
                    pass
                try:
                    self.refresh_map_markers()
                except Exception:
                    pass
            except Exception as e:
                ctk.CTkLabel(host, text="Impossible d'afficher la carte : %s" % e).pack(pady=20)

        self._detached_wins["map"] = win
        self._map_detached = True

        def on_close():
            self._reattach_map(win)

        try:
            win.protocol("WM_DELETE_WINDOW", on_close)
        except Exception:
            pass
        self.log("Carte détachée — placez la fenêtre sur le 2ᵉ écran si besoin.")

    def _reattach_map(self, win=None):
        """Ramène la carte dans l'onglet principal."""
        try:
            if hasattr(self, "_map_detached_placeholder") and self._map_detached_placeholder:
                self._map_detached_placeholder.grid_forget()
        except Exception:
            pass
        try:
            self.map_widget.grid_forget()
        except Exception:
            pass
        try:
            self.map_widget.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
        except Exception:
            # recreer dans tab_map
            try:
                self.map_widget = TkinterMapView(self.tab_map, corner_radius=10)
                self.map_widget.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
                self.map_widget.set_position(49.1627, 5.3854)
                self.map_widget.set_zoom(9)
                self.changer_fond_carte("Plan IGN (Moderne)")
                self.map_widget.add_left_click_map_command(self._on_map_left_click)
                self.refresh_map_markers()
            except Exception:
                pass
        if win is not None:
            try:
                # sauver geometrie
                cfg = dict(self._load_app_config() or {})
                ws = dict(cfg.get("workspace") or {})
                try:
                    ws["map_geometry"] = win.geometry()
                except Exception:
                    pass
                cfg["workspace"] = ws
                self._save_app_config(cfg)
            except Exception:
                pass
            try:
                win.destroy()
            except Exception:
                pass
        if hasattr(self, "_detached_wins"):
            self._detached_wins.pop("map", None)
        self._map_detached = False
        try:
            self.refresh_map_markers()
        except Exception:
            pass
        self.log("Carte réattachée à l'onglet principal.")

    def detach_log_window(self):
        """Journal dans une fenêtre séparée."""
        if not hasattr(self, "_detached_wins"):
            self._detached_wins = {}
        existing = self._detached_wins.get("log")
        if existing is not None:
            try:
                if existing.winfo_exists():
                    existing.lift()
                    return
            except Exception:
                pass

        win = ctk.CTkToplevel(self)
        win.title("GeoExif — Journal")
        cfg = self._load_app_config() or {}
        geo = (cfg.get("workspace") or {}).get("log_geometry") or "520x400+80+120"
        try:
            win.geometry(geo)
        except Exception:
            win.geometry("520x400")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass

        bar = ctk.CTkFrame(win, fg_color="transparent")
        bar.pack(fill="x", padx=8, pady=6)
        ctk.CTkButton(
            bar, text="Réattacher", width=100,
            command=lambda: self._reattach_log(win),
        ).pack(side="left")

        # Miroir du journal : zone texte synchro a chaque log() si possible
        box = ctk.CTkTextbox(win, font=ctk.CTkFont(size=12))
        box.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        try:
            if hasattr(self, "log_box"):
                txt = self.log_box.get("0.0", "end")
                box.insert("0.0", txt)
        except Exception:
            pass
        box.configure(state="disabled")
        self._detached_log_box = box
        self._detached_wins["log"] = win

        def on_close():
            self._reattach_log(win)

        try:
            win.protocol("WM_DELETE_WINDOW", on_close)
        except Exception:
            pass
        self.log("Journal détaché.")

    def _reattach_log(self, win=None):
        if win is not None:
            try:
                cfg = dict(self._load_app_config() or {})
                ws = dict(cfg.get("workspace") or {})
                try:
                    ws["log_geometry"] = win.geometry()
                except Exception:
                    pass
                cfg["workspace"] = ws
                self._save_app_config(cfg)
            except Exception:
                pass
            try:
                win.destroy()
            except Exception:
                pass
        if hasattr(self, "_detached_wins"):
            self._detached_wins.pop("log", None)
        self._detached_log_box = None
        self.log("Journal réattaché.")

    def reattach_all_windows(self):
        """Ferme les fenêtres détachées et ramène carte / journal."""
        wins = dict(getattr(self, "_detached_wins", {}) or {})
        if "map" in wins:
            self._reattach_map(wins.get("map"))
        if "log" in wins:
            self._reattach_log(wins.get("log"))
        # autres fenetres outils : ne pas fermer (brief, etc.)
        self.log("Fenêtres détachées réattachées.")

    def save_workspace_layout(self):
        """Mémorise géométrie fenêtre principale + fenêtres détachées + sidebar."""
        cfg = dict(self._load_app_config() or {})
        ws = dict(cfg.get("workspace") or {})
        try:
            ws["main_geometry"] = self.geometry()
        except Exception:
            pass
        try:
            ws["sidebar_visible"] = bool(getattr(self, "_sidebar_visible", True))
        except Exception:
            pass
        try:
            ws["ui_mode"] = getattr(self, "_ui_mode", "sortie")
        except Exception:
            pass
        for key in ("map", "log"):
            w = (getattr(self, "_detached_wins", {}) or {}).get(key)
            if w is not None:
                try:
                    if w.winfo_exists():
                        ws["%s_geometry" % key] = w.geometry()
                        ws["%s_detached" % key] = True
                except Exception:
                    pass
            else:
                ws["%s_detached" % key] = False
        cfg["workspace"] = ws
        try:
            self._save_app_config(cfg)
            self.app_config = cfg
        except Exception as e:
            messagebox.showerror("Espace de travail", str(e))
            return
        messagebox.showinfo(
            "Espace de travail",
            "Disposition enregistrée.\n"
            "(fenêtre principale, barre latérale, carte/journal détachés)\n\n"
            "Restauration au prochain démarrage ou via Affichage → Restaurer.",
        )
        self.log("Espace de travail enregistré.")

    def restore_workspace_layout(self, *, silent=False):
        """Restaure la disposition sauvegardée."""
        cfg = self._load_app_config() or {}
        ws = cfg.get("workspace") or {}
        if not ws:
            if not silent:
                messagebox.showinfo("Espace de travail", "Aucune disposition enregistrée.")
            return
        try:
            g = ws.get("main_geometry")
            if g:
                self.geometry(g)
        except Exception:
            pass
        try:
            want_side = ws.get("sidebar_visible", True)
            if bool(getattr(self, "_sidebar_visible", True)) != bool(want_side):
                self.toggle_sidebar()
        except Exception:
            pass
        try:
            mode = ws.get("ui_mode")
            if mode and hasattr(self, "_on_ui_mode_change"):
                self._on_ui_mode_change("Saison" if mode == "saison" else "Sortie")
        except Exception:
            pass
        # détachements
        try:
            if ws.get("map_detached") and not getattr(self, "_map_detached", False):
                self.detach_map_window()
            elif not ws.get("map_detached") and getattr(self, "_map_detached", False):
                self._reattach_map((getattr(self, "_detached_wins", {}) or {}).get("map"))
        except Exception:
            pass
        try:
            if ws.get("log_detached"):
                self.detach_log_window()
        except Exception:
            pass
        if not silent:
            messagebox.showinfo("Espace de travail", "Disposition restaurée.")
            self.log("Espace de travail restauré.")

    def toggle_fullscreen(self, event=None):
        """Plein écran (F11 / F12). Rappuyer pour quitter."""
        try:
            current = bool(self.attributes("-fullscreen"))
        except Exception:
            current = getattr(self, "_fullscreen", False)
        new_state = not current
        self._fullscreen = new_state
        try:
            self.attributes("-fullscreen", new_state)
        except Exception:
            # Repli : maximiser / restaurer
            try:
                if new_state:
                    self._geometry_before_fs = self.geometry()
                    self.state("zoomed")
                else:
                    self.state("normal")
                    if getattr(self, "_geometry_before_fs", None):
                        self.geometry(self._geometry_before_fs)
            except Exception:
                pass
        try:
            if new_state:
                self.log("Plein écran — F11 ou F12 ou Échap pour quitter")
            else:
                self.log("Plein écran désactivé")
        except Exception:
            pass
        return "break"

    def _quit_fullscreen(self, event=None):
        """Échap quitte le plein écran uniquement (ne ferme pas l'app)."""
        if getattr(self, "_fullscreen", False):
            self.toggle_fullscreen()
            return "break"

    def toggle_sidebar(self):
        """Affiche ou masque la barre d'outils gauche."""
        self._sidebar_visible = not getattr(self, "_sidebar_visible", True)
        if self._sidebar_visible:
            self.sidebar.grid(row=1, column=0, sticky="nsew", padx=(10, 6), pady=(4, 10))
            self.grid_columnconfigure(0, weight=1, minsize=280)
            self.main_frame.grid_configure(padx=(6, 12))
            try:
                self.btn_toggle_sidebar.configure(text="◀", width=26)
                self.btn_toggle_sidebar.place(x=6, rely=0.42, anchor="w")
                self.btn_toggle_sidebar.lift()
            except Exception:
                pass
            self.log("Barre d'outils affichée")
        else:
            self.sidebar.grid_remove()
            self.grid_columnconfigure(0, weight=0, minsize=0)
            self.main_frame.grid_configure(padx=(12, 12))
            try:
                self.btn_toggle_sidebar.configure(text="▶ Outils", width=78)
                self.btn_toggle_sidebar.place(x=6, rely=0.42, anchor="w")
                self.btn_toggle_sidebar.lift()
            except Exception:
                pass
            self.log("Barre d'outils masquée — ▶ Outils, menu Fichier, ou Ctrl+B")

    def _on_ui_mode_change(self, value):
        """Bascule Sortie (terrain) / Saison (analyse)."""
        mode = "saison" if str(value).lower().startswith("saison") else "sortie"
        self._ui_mode = mode
        if mode == "saison":
            self.card_gps.grid_remove()
            self.card_saison.grid(row=2, column=0, padx=12, pady=6, sticky="ew")
            self.log("Mode Saison — projet, carte cumulée, analyses.")
            try:
                self.toolbar_mode_lbl.configure(text="  Mode Saison  ")
            except Exception:
                pass
            try:
                self.tab_view.set("🗺️   Carte")
            except Exception:
                pass
        else:
            self.card_saison.grid_remove()
            self.card_gps.grid(row=2, column=0, padx=12, pady=6, sticky="ew")
            self.log("Mode Sortie — dossier, GPS, carnet.")
            try:
                self.toolbar_mode_lbl.configure(text="  Mode Sortie  ")
            except Exception:
                pass

    def _build_menubar(self):
        """Menu allégé sans doublons : Fichier, GPS, Carnet, Carte, Outils, Sauvegardes, Affichage, Aide."""
        try:
            self.option_add("*Menu.font", "{Segoe UI} 10")
        except Exception:
            pass

        menubar = tk.Menu(self, tearoff=0)

        # —— Fichier ——
        m_file = tk.Menu(menubar, tearoff=0)
        m_file.add_command(label="Ouvrir dossier photos / sortie", command=self.select_photo_folder)
        m_file.add_command(label="Carnet sans photos…", command=self.create_or_open_carnet_folder)
        m_file.add_command(label="Charger une trace GPX…", command=self.select_gpx_file)
        m_file.add_separator()
        m_file.add_command(label="Exporter paquet terrain…", command=self.export_field_package)
        m_file.add_command(label="Importer paquet / GPX…", command=self.import_field_package)
        m_file.add_separator()
        m_file.add_command(label="Quitter", command=self.destroy)
        menubar.add_cascade(label="Fichier", menu=m_file)

        # —— GPS ——
        m_gps = tk.Menu(menubar, tearoff=0)
        m_gps.add_command(label="Synchroniser photos & GPX", command=self.start_sync_thread)
        m_gps.add_command(label="Décalage été (−1 h)", command=lambda: self.apply_geosync_preset("-1:00:00", "été"))
        m_gps.add_command(label="Décalage hiver (−2 h)", command=lambda: self.apply_geosync_preset("-2:00:00", "hiver"))
        m_gps.add_separator()
        m_gps.add_command(label="Annuler dernière synchro", command=self.undo_last_sync)
        m_gps.add_command(label="Éditer GPS photo…", command=self.open_edit_gps_dialog)
        m_gps.add_command(label="Traitement par lot (GPS)…", command=self.open_batch_processor)
        m_gps.add_command(label="Métadonnées GPS…", command=self.open_gps_metadata_panel)
        m_gps.add_command(label="Affût / série (GPS + horaires)…", command=self.open_affut_series_dialog)
        menubar.add_cascade(label="GPS", menu=m_gps)

        # —— Carnet ——
        m_carnet = tk.Menu(menubar, tearoff=0)
        m_carnet.add_command(label="Enregistrer l'observation", command=lambda: self.save_current_note(silent=False))
        m_carnet.add_command(label="Observation sans photo…", command=self.open_manual_observation_dialog)
        m_carnet.add_command(label="Brief IA de la sortie", command=self.generate_ai_summary)
        m_carnet.add_command(label="Indices — dossier + import…", command=self.open_indices_folder_workflow)
        m_carnet.add_command(label="Rappels indices / suivi…", command=self.open_rappels_indices)
        m_carnet.add_separator()
        m_carnet.add_command(label="Import son…", command=self.open_import_son_menu)
        m_carnet.add_command(label="Birda CLI — analyser un audio…", command=self.open_birda_cli_analyzer)
        m_carnet.add_command(label="Session d'écoute (guidée)…", command=self.open_ecoute_session_wizard)
        m_carnet.add_command(label="Importer session BirdNET Live…", command=self.import_birdnet_live_session)
        m_carnet.add_command(label="Birda — prise de son & import…", command=self.open_birda_tools)
        m_carnet.add_command(label="Importer détections Birda (JSON/CSV)…", command=self.import_birda_detections)
        m_carnet.add_command(label="Importer CSV Chirpity…", command=self.import_chirpity_csv)
        m_carnet.add_separator()
        m_carnet.add_command(label="Supprimer l'observation sélectionnée", command=self.delete_current_observation)
        m_carnet.add_command(label="Point d'écoute GPS (carte)…", command=self.toggle_point_ecoute_gps)
        menubar.add_cascade(label="Carnet", menu=m_carnet)

        # —— Carte ——
        m_map = tk.Menu(menubar, tearoff=0)
        m_map.add_command(label="Carte cumulée…", command=self.open_multi_sorties_map)
        m_map.add_command(label="Placer observation", command=self.toggle_place_observation_mode)
        m_map.add_command(label="Mesure distance (km)", command=self.toggle_map_measure)
        m_map.add_command(label="Rafraîchir marqueurs", command=self.refresh_map_markers)
        m_map.add_command(label="Capture carte…", command=self.capture_map_screenshot)
        menubar.add_cascade(label="Carte", menu=m_map)

        # —— Outils (saison, exports, terrain) ——
        m_outils = tk.Menu(menubar, tearoff=0)
        m_outils.add_command(label="Projet Meuse / saison…", command=self.open_project_hub)
        m_outils.add_command(label="Analyse effort / richesse", command=self.open_analysis_panel)
        m_outils.add_command(label="Cibles de saison", command=self.open_season_suggestions)
        m_outils.add_command(label="Calendrier de chasse (indicatif)", command=self.open_calendrier_chasse)
        m_outils.add_command(label="Brief multi-sorties…", command=self.open_aggregated_report)
        m_outils.add_command(label="Débrief texte (archive, sans photos)…", command=self.open_offline_text_debrief)
        m_outils.add_separator()
        if hasattr(self, "export_csv"):
            m_outils.add_command(label="Export CSV", command=self.export_csv)
        if hasattr(self, "export_kml"):
            m_outils.add_command(label="Export KML / KMZ", command=self.export_kml)
        if hasattr(self, "export_geojson"):
            m_outils.add_command(label="Export GeoJSON", command=self.export_geojson)
        if hasattr(self, "export_daily_pdf"):
            m_outils.add_command(label="Export PDF du jour", command=self.export_daily_pdf)
        m_outils.add_separator()
        m_outils.add_command(label="Birda — prise de son & import…", command=self.open_birda_tools)
        m_outils.add_command(label="Caméras & enregistreurs…", command=self.open_devices_manager)
        m_outils.add_command(label="Poser un dispositif sur la carte", command=self.toggle_device_place_mode)
        m_outils.add_command(label="Cloud (Drive / Dropbox)…", command=self.open_cloud_settings)
        m_outils.add_command(label="Capture fenêtre…", command=self.capture_window_screenshot)
        m_outils.add_command(label="Tous les outils…", command=self.open_outils_hub)
        menubar.add_cascade(label="Outils", menu=m_outils)

        # —— Sauvegardes / recovery ——
        m_bak = tk.Menu(menubar, tearoff=0)
        m_bak.add_command(label="Configuration…", command=lambda: self.open_setup_wizard(first_run=False))
        m_bak.add_command(label="Sauvegardes automatiques…", command=self.open_backup_settings)
        m_bak.add_separator()
        m_bak.add_command(label="Rescanner mes sorties…", command=self.scan_parent_for_sorties)
        m_bak.add_command(label="Gérer / retirer des sorties…", command=self.open_manage_sorties)
        m_bak.add_command(label="Archive locale (recovery)…", command=self.open_local_archive_browser)
        m_bak.add_command(
            label="Copier sorties accessibles → archive PC",
            command=self.mirror_all_known_sorties_to_archive,
        )
        m_bak.add_command(label="Pack recovery ZIP…", command=self.create_recovery_pack_zip)
        m_bak.add_separator()
        m_bak.add_command(label="Vider le cache aperçus…", command=self.clear_preview_cache)
        menubar.add_cascade(label="Sauvegardes", menu=m_bak)

        # —— Affichage ——
        m_view = tk.Menu(menubar, tearoff=0)
        m_view.add_command(label="Thème sombre", command=lambda: self.apply_theme("sombre"))
        m_view.add_command(label="Thème clair", command=lambda: self.apply_theme("clair"))
        m_view.add_command(label="Thème papier", command=lambda: self.apply_theme("papier"))
        m_view.add_separator()
        m_view.add_command(label="Choisir un thème…", command=self.open_theme_picker)
        m_view.add_command(label="Afficher / masquer barre latérale", command=self.toggle_sidebar)
        m_view.add_command(label="Plein écran (F11 / F12)", command=self.toggle_fullscreen)
        m_view.add_separator()
        m_view.add_command(label="Détacher la carte (2ᵉ écran)…", command=self.detach_map_window)
        m_view.add_command(label="Détacher le journal…", command=self.detach_log_window)
        m_view.add_command(label="Réattacher les fenêtres", command=self.reattach_all_windows)
        m_view.add_separator()
        m_view.add_command(label="Enregistrer l'espace de travail", command=self.save_workspace_layout)
        m_view.add_command(label="Restaurer l'espace de travail", command=self.restore_workspace_layout)
        m_view.add_separator()
        m_view.add_command(label="Raccourcis clavier…", command=self.open_shortcuts_settings)
        menubar.add_cascade(label="Affichage", menu=m_view)

        # —— Aide ——
        m_help = tk.Menu(menubar, tearoff=0)
        if hasattr(self, "open_faq"):
            m_help.add_command(label="FAQ", command=self.open_faq)
        m_help.add_command(label="Raccourcis clavier (F1)…", command=self.open_shortcuts_settings)
        m_help.add_command(
            label="À propos",
            command=lambda: messagebox.showinfo(
                "À propos",
                "GeoExif Meuse 55\nVersion %s\n%s\n\n"
                "Mode Sortie = terrain du jour\n"
                "Mode Saison = projet & analyses\n\n"
                "Photos : Amazon (jour le jour)\n"
                "Carnets : archive PC + pack recovery ZIP\n\n"
                "F1 = raccourcis clavier"
                % (APP_VERSION, APP_AUTHOR),
            ),
        )
        menubar.add_cascade(label="Aide", menu=m_help)

        try:
            self.config(menu=menubar)
            self._menubar = menubar
            # Les raccourcis sont gérés par _setup_keyboard_shortcuts (configurables)
            self._setup_keyboard_shortcuts()
        except Exception as e:
            try:
                self.log(f"Barre de menu indisponible : {e}")
            except Exception:
                pass


    # ========== Dispositifs terrain (cameras / enregistreurs) ==========

    def _devices_path(self, folder=None):
        folder = folder or self.photo_folder_path
        if not folder:
            return None
        return os.path.join(folder, DEVICES_FILE)

    def _load_devices(self, folder=None):
        path = self._devices_path(folder)
        if not path or not os.path.isfile(path):
            return []
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                return data
            return data.get("devices") or []
        except Exception:
            return []

    def _save_devices(self, devices=None, folder=None):
        folder = folder or self.photo_folder_path
        path = self._devices_path(folder)
        if not path:
            messagebox.showwarning("Dispositifs", "Ouvrez d'abord un dossier de sortie.")
            return False
        devices = devices if devices is not None else getattr(self, "devices_data", [])
        try:
            payload = {
                "devices": devices,
                "updated": datetime.now().isoformat(timespec="seconds"),
            }
            with open(path, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
            self.devices_data = devices
            return True
        except Exception as e:
            messagebox.showerror("Dispositifs", "Sauvegarde impossible :\n%s" % e)
            return False

    def _device_due_date(self, dev):
        if dev.get("date_releve_prevue"):
            try:
                return datetime.strptime(str(dev["date_releve_prevue"])[:10], "%Y-%m-%d").date()
            except Exception:
                pass
        try:
            pose = datetime.fromisoformat(str(dev.get("date_pose", ""))[:19])
            jours = int(dev.get("delai_jours") or 4)
            return (pose + timedelta(days=jours)).date()
        except Exception:
            return None

    def _device_status_label(self, dev):
        if (dev.get("statut") or "") == "recupere":
            return "Recupere"
        due = self._device_due_date(dev)
        if not due:
            return "En place"
        today = datetime.now().date()
        delta = (due - today).days
        if delta < 0:
            return "En retard (%d j)" % abs(delta)
        if delta == 0:
            return "A relever aujourd'hui"
        if delta <= 1:
            return "Bientot (%d j)" % delta
        return "En place (J-%d)" % delta

    def reload_devices_for_folder(self):
        self.devices_data = self._load_devices()
        try:
            self.refresh_device_map_markers()
        except Exception:
            pass
        try:
            key = getattr(self, "photo_folder_path", "") or ""
            if key and key != getattr(self, "_device_reminded_folder", None):
                self.check_device_reminders(silent=False)
                try:
                    self.check_indice_reminders(silent=True, days_ahead=30)
                except Exception:
                    pass
                try:
                    # e-mail seulement si active dans la config (pas de popup)
                    cfg = self._load_email_alert_config()
                    if cfg.get("enabled"):
                        threading.Thread(target=lambda: self.send_reminder_email_digest(force=False), daemon=True).start()
                except Exception:
                    pass
                self._device_reminded_folder = key
            else:
                self.check_device_reminders(silent=True)
        except Exception:
            pass

    def toggle_device_place_mode(self):
        if not self.photo_folder_path:
            messagebox.showinfo(
                "Dispositifs",
                "Ouvrez d'abord un dossier de sortie\n(fichier dispositifs.json)."
            )
            return
        self.device_place_mode = not getattr(self, "device_place_mode", False)
        if self.device_place_mode and getattr(self, "_place_obs_mode", False):
            try:
                self.toggle_place_observation_mode()
            except Exception:
                self._place_obs_mode = False
        if self.device_place_mode:
            self.log("Mode pose dispositif : cliquez sur la carte.")
            messagebox.showinfo(
                "Pose dispositif",
                "Cliquez sur la carte pour placer\nune camera, un enregistreur ou autre."
            )
        else:
            self.log("Mode pose dispositif desactive.")

    def open_device_dialog(self, lat=None, lon=None, device=None):
        if not self.photo_folder_path and device is None:
            messagebox.showinfo("Dispositifs", "Ouvrez un dossier de sortie.")
            return
        edit = device is not None
        win = ctk.CTkToplevel(self)
        win.title("Dispositif terrain")
        win.geometry("440x560")
        try:
            win.configure(fg_color=UI.get("bg", "#101612"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Camera / Enregistreur / Autre",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 6))

        form = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=12)
        form.pack(fill="both", expand=True, padx=14, pady=8)
        form.grid_columnconfigure(1, weight=1)

        def add_row(r, label, widget):
            ctk.CTkLabel(form, text=label, text_color=UI.get("text")).grid(
                row=r, column=0, sticky="w", padx=10, pady=6
            )
            widget.grid(row=r, column=1, sticky="ew", padx=10, pady=6)

        type_var = tk.StringVar(value=(device or {}).get("type") or "camera")
        name_var = tk.StringVar(value=(device or {}).get("name") or "")
        lieu_var = tk.StringVar(value=(device or {}).get("lieu") or "")
        delai_var = tk.StringVar(value=str((device or {}).get("delai_jours") or 4))
        bat_raw = (device or {}).get("batterie_pct")
        bat_var = tk.StringVar(value="" if bat_raw is None else str(bat_raw))
        bat_type_var = tk.StringVar(value=(device or {}).get("batterie_type") or "lithium")
        note_var = tk.StringVar(value=(device or {}).get("note") or "")
        lat_var = tk.StringVar(value="")
        lon_var = tk.StringVar(value="")
        if edit and device.get("lat") is not None:
            lat_var.set("%.6f" % float(device["lat"]))
            lon_var.set("%.6f" % float(device["lon"]))
        elif lat is not None:
            lat_var.set("%.6f" % float(lat))
            lon_var.set("%.6f" % float(lon))

        add_row(0, "Type", ctk.CTkOptionMenu(
            form, variable=type_var, values=["camera", "enregistreur", "autre"],
            fg_color=UI.get("card_alt"), text_color=UI.get("text"),
        ))
        add_row(1, "Nom / code", ctk.CTkEntry(form, textvariable=name_var, placeholder_text="ex. Cam Bois Nord"))
        add_row(2, "Latitude", ctk.CTkEntry(form, textvariable=lat_var))
        add_row(3, "Longitude", ctk.CTkEntry(form, textvariable=lon_var))
        add_row(4, "Lieu", ctk.CTkEntry(form, textvariable=lieu_var, placeholder_text="lieu-dit..."))
        add_row(5, "Releve (jours)", ctk.CTkOptionMenu(
            form, variable=delai_var,
            values=["1", "2", "3", "4", "5", "7", "10", "14", "21", "30"],
            fg_color=UI.get("card_alt"), text_color=UI.get("text"),
        ))
        add_row(6, "Batterie %", ctk.CTkEntry(form, textvariable=bat_var, placeholder_text="ex. 85"))
        add_row(7, "Type batterie", ctk.CTkOptionMenu(
            form, variable=bat_type_var,
            values=["lithium", "alcaline", "NiMH", "externe", "inconnu"],
            fg_color=UI.get("card_alt"), text_color=UI.get("text"),
        ))
        add_row(8, "Note", ctk.CTkEntry(form, textvariable=note_var, placeholder_text="orientation, SD..."))

        def save():
            try:
                la = float(lat_var.get().replace(",", "."))
                lo = float(lon_var.get().replace(",", "."))
            except Exception:
                messagebox.showwarning("GPS", "Latitude / longitude invalides.")
                return
            try:
                delai = int(delai_var.get())
            except Exception:
                delai = 4
            bat = None
            if bat_var.get().strip():
                try:
                    bat = max(0, min(100, int(float(bat_var.get().replace(",", ".")))))
                except Exception:
                    messagebox.showwarning("Batterie", "Pourcentage invalide (0-100).")
                    return
            pose = datetime.now()
            if edit and device.get("date_pose"):
                try:
                    pose = datetime.fromisoformat(str(device["date_pose"])[:19])
                except Exception:
                    pass
            due = (pose + timedelta(days=delai)).date().isoformat()
            nom = name_var.get().strip()
            if not nom:
                tlabel = {"camera": "Camera", "enregistreur": "Enregistreur", "autre": "Dispositif"}.get(type_var.get(), "Dispositif")
                nom = "%s %s" % (tlabel, pose.strftime("%d/%m %H:%M"))
            entry = {
                "id": (device or {}).get("id") or datetime.now().strftime("%Y%m%d%H%M%S%f"),
                "type": type_var.get(),
                "name": nom,
                "lat": la,
                "lon": lo,
                "lieu": lieu_var.get().strip(),
                "date_pose": pose.isoformat(timespec="seconds"),
                "delai_jours": delai,
                "date_releve_prevue": due,
                "batterie_pct": bat,
                "batterie_type": bat_type_var.get(),
                "note": note_var.get().strip(),
                "statut": (device or {}).get("statut") or "en_place",
                "date_recup": (device or {}).get("date_recup"),
            }
            devices = list(self._load_devices())
            if edit:
                devices = [entry if d.get("id") == entry["id"] else d for d in devices]
            else:
                devices.append(entry)
            if self._save_devices(devices):
                self.device_place_mode = False
                self.refresh_device_map_markers()
                self.log("Dispositif enregistre : %s — releve %s" % (nom, due))
                win.destroy()
                if messagebox.askyesno("Rappel calendrier", "Relève prevue le %s.\n\nCreer un fichier .ics ?" % due):
                    self.export_device_ics(entry)

        btns = ctk.CTkFrame(win, fg_color="transparent")
        btns.pack(fill="x", padx=14, pady=(4, 14))
        ctk.CTkButton(btns, text="Enregistrer", height=36, fg_color=UI.get("success", "#2f9e5f"), command=save).pack(side="left", fill="x", expand=True, padx=(0, 6))
        ctk.CTkButton(btns, text="Annuler", height=36, fg_color=UI.get("card_alt"), command=win.destroy).pack(side="left", fill="x", expand=True)

    def open_devices_manager(self):
        if not self.photo_folder_path:
            messagebox.showinfo("Dispositifs", "Ouvrez un dossier de sortie.\nFichier : dispositifs.json")
            return
        self.devices_data = self._load_devices()
        win = ctk.CTkToplevel(self)
        win.title("Dispositifs terrain")
        win.geometry("720x520")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        head = ctk.CTkFrame(win, fg_color="transparent")
        head.pack(fill="x", padx=12, pady=10)
        ctk.CTkLabel(head, text="Cameras · Enregistreurs · Releves", font=ctk.CTkFont(size=15, weight="bold"), text_color=UI.get("text")).pack(side="left")
        ctk.CTkButton(head, text="+ Poser sur la carte", width=150, height=30, fg_color=UI.get("accent"), command=lambda: (win.destroy(), self.toggle_device_place_mode())).pack(side="right", padx=4)
        ctk.CTkButton(head, text="+ Saisie manuelle", width=130, height=30, fg_color=UI.get("card_alt"), command=lambda: (self.open_device_dialog(), refresh())).pack(side="right", padx=4)

        scroll = ctk.CTkScrollableFrame(win, fg_color=UI.get("card"))
        scroll.pack(fill="both", expand=True, padx=12, pady=(0, 8))

        def mark_recovered(dev):
            devices = self._load_devices()
            for d in devices:
                if d.get("id") == dev.get("id"):
                    d["statut"] = "recupere"
                    d["date_recup"] = datetime.now().isoformat(timespec="seconds")
            self._save_devices(devices)
            self.refresh_device_map_markers()
            refresh()

        def delete_dev(dev):
            if not messagebox.askyesno("Supprimer", "Supprimer « %s » ?" % dev.get("name")):
                return
            devices = [d for d in self._load_devices() if d.get("id") != dev.get("id")]
            self._save_devices(devices)
            self.refresh_device_map_markers()
            refresh()

        def refresh():
            for w in scroll.winfo_children():
                w.destroy()
            devices = self._load_devices()
            self.devices_data = devices
            if not devices:
                ctk.CTkLabel(scroll, text="Aucun dispositif.\nPosez-en un sur la carte ou en saisie manuelle.", text_color=UI.get("text_dim")).pack(pady=20)
                return

            def sort_key(d):
                due = self._device_due_date(d) or datetime.now().date()
                rec = 1 if d.get("statut") == "recupere" else 0
                return (rec, due)

            for dev in sorted(devices, key=sort_key):
                row = ctk.CTkFrame(scroll, fg_color=UI.get("card_alt"), corner_radius=10)
                row.pack(fill="x", pady=4, padx=4)
                typ = {"camera": "Camera", "enregistreur": "Enregistreur", "autre": "Autre"}.get(dev.get("type"), "Dispositif")
                st = self._device_status_label(dev)
                bat = dev.get("batterie_pct")
                bat_s = " · Batterie %s%%" % bat if bat is not None else ""
                if bat is not None and dev.get("batterie_type"):
                    bat_s += " (%s)" % dev.get("batterie_type")
                due = self._device_due_date(dev)
                due_s = due.isoformat() if due else "?"
                lieu = dev.get("lieu") or ("%.4f, %.4f" % (float(dev.get("lat") or 0), float(dev.get("lon") or 0)))
                ctk.CTkLabel(row, text="%s — %s" % (typ, dev.get("name", "")), font=ctk.CTkFont(size=13, weight="bold"), text_color=UI.get("text"), anchor="w").pack(fill="x", padx=10, pady=(8, 0))
                ctk.CTkLabel(row, text="%s · releve %s · %s%s" % (st, due_s, lieu, bat_s), font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"), anchor="w").pack(fill="x", padx=10, pady=(2, 4))
                if dev.get("note"):
                    ctk.CTkLabel(row, text=dev.get("note"), font=ctk.CTkFont(size=10), text_color=UI.get("text_muted"), anchor="w").pack(fill="x", padx=10, pady=(0, 4))
                actions = ctk.CTkFrame(row, fg_color="transparent")
                actions.pack(fill="x", padx=8, pady=(0, 8))
                if dev.get("statut") != "recupere":
                    ctk.CTkButton(actions, text="Marquer recupere", width=130, height=26, fg_color=UI.get("success"), command=lambda d=dev: mark_recovered(d)).pack(side="left", padx=2)
                    ctk.CTkButton(actions, text="Rappel .ics", width=90, height=26, fg_color=UI.get("card"), command=lambda d=dev: self.export_device_ics(d)).pack(side="left", padx=2)
                ctk.CTkButton(actions, text="Editer", width=70, height=26, fg_color=UI.get("accent"), command=lambda d=dev: (self.open_device_dialog(device=d), refresh())).pack(side="left", padx=2)
                ctk.CTkButton(actions, text="Suppr.", width=60, height=26, fg_color=UI.get("danger", "#c44"), command=lambda d=dev: delete_dev(d)).pack(side="right", padx=2)

        refresh()
        foot = ctk.CTkFrame(win, fg_color="transparent")
        foot.pack(fill="x", padx=12, pady=(0, 12))
        ctk.CTkButton(foot, text="Rafraichir marqueurs carte", height=30, fg_color=UI.get("card_alt"), command=self.refresh_device_map_markers).pack(side="left")
        ctk.CTkButton(foot, text="Fermer", height=30, width=100, fg_color=UI.get("card_alt"), command=win.destroy).pack(side="right")

    def export_device_ics(self, device):
        due = self._device_due_date(device)
        if not due:
            messagebox.showwarning("ICS", "Date de releve inconnue.")
            return
        name = device.get("name") or "Dispositif"
        lieu = device.get("lieu") or ("%s %s" % (device.get("lat"), device.get("lon")))
        typ = device.get("type") or "dispositif"
        bat = device.get("batterie_pct")
        bat_line = ""
        if bat is not None:
            bat_line = "Batterie a la pose : %s%% (%s). " % (bat, device.get("batterie_type") or "?")
        desc = "Releve %s : %s. Lieu : %s. Pose : %s. %sNote : %s. GPS : %s %s" % (
            typ, name, lieu, str(device.get("date_pose") or "")[:16], bat_line,
            device.get("note") or "", device.get("lat"), device.get("lon"),
        )
        uid = "%s@geoexif" % device.get("id")
        dt_start = due.strftime("%Y%m%d")
        crlf = "\r\n"
        parts = [
            "BEGIN:VCALENDAR",
            "VERSION:2.0",
            "PRODID:-//GeoExif//Dispositifs//FR",
            "CALSCALE:GREGORIAN",
            "BEGIN:VEVENT",
            "UID:" + uid,
            "DTSTAMP:" + datetime.utcnow().strftime("%Y%m%dT%H%M%SZ"),
            "DTSTART;VALUE=DATE:" + dt_start,
            "DTEND;VALUE=DATE:" + dt_start,
            "SUMMARY:Releve %s - %s" % (typ, name),
            "DESCRIPTION:" + desc,
            "LOCATION:" + str(lieu),
            "BEGIN:VALARM",
            "TRIGGER:-P1D",
            "ACTION:DISPLAY",
            "DESCRIPTION:Releve %s demain" % name,
            "END:VALARM",
            "END:VEVENT",
            "END:VCALENDAR",
            "",
        ]
        ics = crlf.join(parts)
        path = filedialog.asksaveasfilename(
            title="Enregistrer le rappel calendrier",
            defaultextension=".ics",
            initialfile="releve_%s_%s.ics" % (name.replace(" ", "_"), due.isoformat()),
            filetypes=[("Calendrier iCalendar", "*.ics"), ("Tous", "*.*")],
        )
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8", newline="") as f:
                f.write(ics)
            self.log("Rappel calendrier : %s" % path)
            messagebox.showinfo(
                "Rappel .ics",
                "Fichier cree :\n%s\n\nDouble-clic pour l'ajouter a l'agenda." % path,
            )
        except Exception as e:
            messagebox.showerror("ICS", str(e))

    def refresh_device_map_markers(self):
        for m in getattr(self, "_device_map_markers", []) or []:
            try:
                m.delete()
            except Exception:
                pass
        self._device_map_markers = []
        if not hasattr(self, "map_widget") or self.map_widget is None:
            return
        devices = self._load_devices() if self.photo_folder_path else []
        self.devices_data = devices
        for dev in devices:
            if dev.get("statut") == "recupere":
                continue
            try:
                lat, lon = float(dev["lat"]), float(dev["lon"])
            except Exception:
                continue
            typ = dev.get("type") or "autre"
            color = {"camera": "#e67e22", "enregistreur": "#9b59b6", "autre": "#3498db"}.get(typ, "#3498db")
            label = dev.get("name") or typ
            st = self._device_status_label(dev)
            try:
                mk = self.map_widget.set_marker(
                    lat, lon, text="%s | %s" % (label, st),
                    marker_color_circle=color, marker_color_outside=color,
                    command=lambda _m=None, d=dev: self.open_device_dialog(device=d),
                )
                self._device_map_markers.append(mk)
            except Exception:
                try:
                    mk = self.map_widget.set_marker(lat, lon, text=label)
                    self._device_map_markers.append(mk)
                except Exception:
                    pass


    def _parse_rappel_date(self, s):
        s = (s or "").strip()[:10]
        if not s:
            return None
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                return datetime.strptime(s, fmt).date()
            except Exception:
                pass
        return None

    def _collect_indice_rappels(self, folders=None):
        """Parcourt des dossiers (sortie courante + projet) et retourne les rappels d'indices."""
        out = []
        seen = set()
        if folders is None:
            folders = []
            if self.photo_folder_path:
                folders.append(self.photo_folder_path)
            # projet saison
            try:
                proj = getattr(self, "projet_data", None) or {}
                for s in proj.get("sorties") or proj.get("folders") or []:
                    if isinstance(s, dict):
                        fp = s.get("path") or s.get("dossier") or s.get("folder")
                    else:
                        fp = s
                    if fp and os.path.isdir(fp):
                        folders.append(fp)
            except Exception:
                pass
            # known folders
            try:
                for fp in getattr(self, "known_folders", []) or []:
                    if fp and os.path.isdir(fp):
                        folders.append(fp)
            except Exception:
                pass
        today = datetime.now().date()
        for folder in folders:
            folder = os.path.normpath(folder)
            if folder in seen:
                continue
            seen.add(folder)
            notes_path = os.path.join(folder, NOTES_FILE)
            if not os.path.isfile(notes_path):
                continue
            try:
                with open(notes_path, "r", encoding="utf-8") as f:
                    notes = json.load(f)
            except Exception:
                continue
            if not isinstance(notes, dict):
                continue
            for key, data in notes.items():
                if not isinstance(data, dict):
                    continue
                rd = self._parse_rappel_date(data.get("rappel_date"))
                if not rd:
                    continue
                if data.get("rappel_done"):
                    continue
                delta = (rd - today).days
                out.append({
                    "folder": folder,
                    "key": key,
                    "date": rd.isoformat(),
                    "delta": delta,
                    "espece": data.get("espece") or "?",
                    "type_indice": data.get("type_indice") or data.get("type_observation") or "",
                    "lieu": data.get("lieu") or "",
                    "note": data.get("rappel_note") or data.get("notes_libres") or "",
                    "lat": data.get("lat"),
                    "lon": data.get("lon"),
                })
        out.sort(key=lambda x: x.get("date") or "9999")
        return out

    def check_indice_reminders(self, silent=False, days_ahead=45):
        """Alerte si un rappel d'indice est du / proche (ex. retour terrier en mars)."""
        rappels = self._collect_indice_rappels()
        due = [r for r in rappels if r["delta"] <= days_ahead]
        if not due:
            if not silent:
                messagebox.showinfo("Rappels indices", "Aucun rappel a venir (dans %d jours)." % days_ahead)
            return
        lines = []
        for r in due[:25]:
            when = (
                "en retard (%d j)" % abs(r["delta"]) if r["delta"] < 0
                else ("aujourd'hui" if r["delta"] == 0 else "dans %d j" % r["delta"])
            )
            lines.append(
                "• %s — %s (%s)\n  %s · %s\n  %s\n  Dossier : %s" % (
                    r["date"], when,
                    r.get("type_indice") or "indice",
                    r.get("espece"),
                    r.get("lieu") or "lieu ?",
                    (r.get("note") or "")[:120],
                    os.path.basename(r["folder"]),
                )
            )
        msg = (
            "%d rappel(s) d'indice a traiter (echus ou dans %d j) :\n\n%s" % (
                len(due), days_ahead, "\n\n".join(lines)
            )
        )
        if len(due) > 25:
            msg += "\n\n… et %d autres." % (len(due) - 25)
        if silent:
            self.log("Rappels indices : %d" % len(due))
            try:
                messagebox.showwarning("Rappels indices", msg)
            except Exception:
                pass
        else:
            messagebox.showinfo("Rappels indices", msg)

    def _ics_escape(self, s):
        s = str(s or "").replace("\\", "\\\\").replace(";", "\\;").replace(",", "\\,")
        s = s.replace("\n", "\\n").replace("\r", "")
        return s

    def _build_ics_event(self, uid, summary, description, date_obj, days_span=0):
        """Evenement ICS journee entiere (date Python)."""
        if hasattr(date_obj, "strftime"):
            ds = date_obj.strftime("%Y%m%d")
        else:
            ds = str(date_obj).replace("-", "")[:8]
        # fin = lendemain (DTEND exclusif en all-day)
        try:
            from datetime import timedelta
            if hasattr(date_obj, "year"):
                end_d = date_obj + timedelta(days=max(1, days_span + 1))
            else:
                end_d = datetime.strptime(ds, "%Y%m%d").date() + timedelta(days=1)
            de = end_d.strftime("%Y%m%d")
        except Exception:
            de = ds
        stamp = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
        return "\r\n".join([
            "BEGIN:VEVENT",
            "UID:%s" % uid,
            "DTSTAMP:%s" % stamp,
            "DTSTART;VALUE=DATE:%s" % ds,
            "DTEND;VALUE=DATE:%s" % de,
            "SUMMARY:%s" % self._ics_escape(summary),
            "DESCRIPTION:%s" % self._ics_escape(description),
            "STATUS:CONFIRMED",
            "END:VEVENT",
        ])

    def export_indice_rappels_ics(self):
        """Exporte les rappels d'indices (+ option dispositifs) vers un fichier .ics agenda."""
        rappels = self._collect_indice_rappels()
        events = []
        for r in rappels:
            try:
                d = datetime.strptime(r["date"][:10], "%Y-%m-%d").date()
            except Exception:
                continue
            summary = "GeoExif suivi : %s" % (r.get("type_indice") or r.get("espece") or "indice")
            desc = "Espece : %s\\nLieu : %s\\nNote : %s\\nDossier : %s\\nGPS : %s %s" % (
                r.get("espece"), r.get("lieu"), r.get("note"),
                r.get("folder"), r.get("lat"), r.get("lon"),
            )
            uid = "indice-%s-%s@geoexif" % (
                abs(hash(r.get("folder") or "")) % 10**8,
                abs(hash(r.get("key") or "")) % 10**8,
            )
            events.append(self._build_ics_event(uid, summary, desc, d))

        # Dispositifs en cours
        try:
            for folder in {r.get("folder") for r in rappels} | (
                {self.photo_folder_path} if self.photo_folder_path else set()
            ):
                if not folder:
                    continue
                # temporairement lire dispositifs du dossier
                prev = self.photo_folder_path
                try:
                    self.photo_folder_path = folder
                    devices = self._load_devices()
                finally:
                    self.photo_folder_path = prev
                for dev in devices or []:
                    if dev.get("statut") == "recupere":
                        continue
                    due = self._device_due_date(dev)
                    if not due:
                        continue
                    summary = "GeoExif releve : %s" % (dev.get("name") or "dispositif")
                    desc = "Type : %s\\nLieu : %s\\nNote : %s" % (
                        dev.get("type"), dev.get("lieu"), dev.get("note"),
                    )
                    uid = "%s@geoexif" % (dev.get("id") or abs(hash(str(dev))) % 10**8)
                    events.append(self._build_ics_event(uid, summary, desc, due))
        except Exception:
            pass

        if not events:
            messagebox.showinfo("Agenda", "Aucun rappel a exporter.")
            return
        path = filedialog.asksaveasfilename(
            title="Exporter les rappels vers l'agenda (.ics)",
            defaultextension=".ics",
            initialfile="geoexif_rappels_%s.ics" % datetime.now().strftime("%Y%m%d"),
            filetypes=[("iCalendar", "*.ics"), ("Tous", "*.*")],
        )
        if not path:
            return
        body = "BEGIN:VCALENDAR\r\nVERSION:2.0\r\nPRODID:-//GeoExif//Rappels//FR\r\nCALSCALE:GREGORIAN\r\nMETHOD:PUBLISH\r\n"
        body += "\r\n".join(events)
        body += "\r\nEND:VCALENDAR\r\n"
        try:
            with open(path, "w", encoding="utf-8", newline="") as f:
                f.write(body)
            messagebox.showinfo(
                "Agenda",
                "Fichier ICS cree (%d evenement(s)) :\n%s\n\n"
                "Importez-le dans Google Agenda, Outlook ou l'app Calendrier Windows."
                % (len(events), path),
            )
            self.log("Export ICS rappels : %s (%d)" % (path, len(events)))
        except Exception as e:
            messagebox.showerror("Agenda", str(e))

    def _load_email_alert_config(self):
        cfg = {}
        try:
            # reutilise app_config si present
            path = getattr(self, "app_config_path", None)
            if not path:
                base = os.path.dirname(os.path.abspath(sys.argv[0]))
                path = os.path.join(base, "app_config.json")
            if os.path.isfile(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                cfg = data.get("email_alerts") or {}
        except Exception:
            cfg = {}
        return cfg

    def _save_email_alert_config(self, email_cfg):
        try:
            path = getattr(self, "app_config_path", None)
            if not path:
                base = os.path.dirname(os.path.abspath(sys.argv[0]))
                path = os.path.join(base, "app_config.json")
            data = {}
            if os.path.isfile(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
            data["email_alerts"] = email_cfg
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            messagebox.showerror("Config e-mail", str(e))

    def open_email_alert_settings(self):
        """Configure SMTP pour les alertes de rappels (indices + dispositifs)."""
        cfg = self._load_email_alert_config()
        win = ctk.CTkToplevel(self)
        win.title("Alertes e-mail — configuration")
        win.geometry("520x480")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)
        ctk.CTkLabel(
            win, text="Alertes e-mail (SMTP)",
            font=ctk.CTkFont(size=15, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=14, pady=(12, 4))
        ctk.CTkLabel(
            win,
            text="Optionnel. GeoExif envoie un resume des rappels proches via votre serveur SMTP "
                 "(ex. smtp.gmail.com). Mot de passe d'application recommande pour Gmail.",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            wraplength=480, justify="left",
        ).pack(anchor="w", padx=14, pady=(0, 8))

        form = ctk.CTkFrame(win, fg_color=UI.get("card"))
        form.pack(fill="both", expand=True, padx=14, pady=6)
        fields = {}
        rows = [
            ("enabled", "Activer (oui/non)", "oui" if cfg.get("enabled") else "non"),
            ("smtp_host", "Serveur SMTP", cfg.get("smtp_host") or "smtp.gmail.com"),
            ("smtp_port", "Port", str(cfg.get("smtp_port") or "587")),
            ("smtp_user", "Utilisateur", cfg.get("smtp_user") or ""),
            ("smtp_password", "Mot de passe", cfg.get("smtp_password") or ""),
            ("mail_from", "Expediteur", cfg.get("mail_from") or cfg.get("smtp_user") or ""),
            ("mail_to", "Destinataire", cfg.get("mail_to") or ""),
            ("days_ahead", "Jours a l'avance", str(cfg.get("days_ahead") or "14")),
        ]
        for i, (key, label, val) in enumerate(rows):
            ctk.CTkLabel(form, text=label, text_color=UI.get("text")).grid(
                row=i, column=0, sticky="w", padx=10, pady=4
            )
            show = "*" if key == "smtp_password" else None
            e = ctk.CTkEntry(form, width=280, show=show)
            e.grid(row=i, column=1, sticky="ew", padx=10, pady=4)
            e.insert(0, val)
            fields[key] = e
        form.grid_columnconfigure(1, weight=1)

        def save():
            email_cfg = {
                "enabled": fields["enabled"].get().strip().lower() in ("oui", "yes", "1", "true"),
                "smtp_host": fields["smtp_host"].get().strip(),
                "smtp_port": int(fields["smtp_port"].get().strip() or "587"),
                "smtp_user": fields["smtp_user"].get().strip(),
                "smtp_password": fields["smtp_password"].get().strip(),
                "mail_from": fields["mail_from"].get().strip(),
                "mail_to": fields["mail_to"].get().strip(),
                "days_ahead": int(fields["days_ahead"].get().strip() or "14"),
            }
            self._save_email_alert_config(email_cfg)
            messagebox.showinfo("E-mail", "Configuration enregistree.")
            win.destroy()

        def test_send():
            save()
            self.send_reminder_email_digest(force=True)

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=14, pady=10)
        ctk.CTkButton(bf, text="Enregistrer", command=save).pack(side="left", padx=4)
        ctk.CTkButton(bf, text="Tester l'envoi", fg_color=UI.get("accent"), command=test_send).pack(side="left", padx=4)
        ctk.CTkButton(bf, text="Fermer", fg_color=UI.get("card_alt"), command=win.destroy).pack(side="right")

    def send_reminder_email_digest(self, force=False):
        """Envoie un e-mail resume des rappels indices + dispositifs proches."""
        cfg = self._load_email_alert_config()
        if not force and not cfg.get("enabled"):
            return False
        if not cfg.get("smtp_host") or not cfg.get("mail_to"):
            if force:
                messagebox.showwarning("E-mail", "Configurez d'abord le serveur SMTP et le destinataire.")
            return False
        days = int(cfg.get("days_ahead") or 14)
        rappels = [r for r in self._collect_indice_rappels() if r["delta"] <= days]
        # dispositifs dossier courant
        dev_lines = []
        try:
            for dev in self._load_devices() or []:
                if dev.get("statut") == "recupere":
                    continue
                due = self._device_due_date(dev)
                if not due:
                    continue
                delta = (due - datetime.now().date()).days
                if delta <= days:
                    dev_lines.append(
                        "• Releve %s (%s) — %s (dans %d j) — %s" % (
                            dev.get("name"), dev.get("type"), due.isoformat(), delta,
                            dev.get("lieu") or "",
                        )
                    )
        except Exception:
            pass

        if not rappels and not dev_lines:
            if force:
                messagebox.showinfo("E-mail", "Aucun rappel proche a signaler.")
            return False

        lines = ["Rappels GeoExif — resume automatique", ""]
        if rappels:
            lines.append("=== Indices / sites ===")
            for r in rappels:
                lines.append(
                    "• %s | %s | %s | %s | %s" % (
                        r["date"], r.get("type_indice"), r.get("espece"),
                        r.get("lieu"), (r.get("note") or "")[:80],
                    )
                )
            lines.append("")
        if dev_lines:
            lines.append("=== Dispositifs ===")
            lines.extend(dev_lines)
        body = "\n".join(lines)

        try:
            msg = MIMEMultipart()
            msg["From"] = cfg.get("mail_from") or cfg.get("smtp_user")
            msg["To"] = cfg.get("mail_to")
            msg["Subject"] = "GeoExif — %d rappel(s) a traiter" % (len(rappels) + len(dev_lines))
            msg.attach(MIMEText(body, "plain", "utf-8"))
            port = int(cfg.get("smtp_port") or 587)
            with smtplib.SMTP(cfg["smtp_host"], port, timeout=30) as server:
                server.ehlo()
                try:
                    server.starttls()
                    server.ehlo()
                except Exception:
                    pass
                if cfg.get("smtp_user"):
                    server.login(cfg.get("smtp_user"), cfg.get("smtp_password") or "")
                server.sendmail(msg["From"], [cfg["mail_to"]], msg.as_string())
            self.log("E-mail rappels envoye a %s" % cfg.get("mail_to"))
            if force:
                messagebox.showinfo("E-mail", "Message envoye a %s" % cfg.get("mail_to"))
            return True
        except Exception as e:
            if force:
                messagebox.showerror("E-mail", "Echec envoi :\n%s" % e)
            self.log("E-mail rappels echec : %s" % e)
            return False

    def open_rappels_indices(self):
        """Liste tous les rappels d'indices (projet + sortie courante)."""
        rappels = self._collect_indice_rappels()
        win = ctk.CTkToplevel(self)
        win.title("Rappels de suivi (indices)")
        win.geometry("680x560")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)
        ctk.CTkLabel(
            win, text="Rappels de suivi — indices / sites",
            font=ctk.CTkFont(size=15, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=14, pady=(12, 4))
        ctk.CTkLabel(
            win,
            text="Ex. terrier decouvert en aout → rappel mars suivant. "
                 "Definir la date dans le carnet (champ Rappel suivi). "
                 "Export ICS = agenda · E-mail = alerte optionnelle SMTP.",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            wraplength=640, justify="left",
        ).pack(anchor="w", padx=14, pady=(0, 8))
        box = ctk.CTkTextbox(win, height=320)
        box.pack(fill="both", expand=True, padx=14, pady=6)
        if not rappels:
            box.insert("0.0", "Aucun rappel programme.\n\nDans le carnet, renseignez « Rappel suivi » (date AAAA-MM-JJ) sur une observation.")
        else:
            lines = []
            for r in rappels:
                flag = "⚠ " if r["delta"] <= 14 else "· "
                lines.append(
                    "%s%s  |  %s  |  %s  |  %s\n   %s\n   %s" % (
                        flag, r["date"],
                        r.get("type_indice") or "indice",
                        r.get("espece"),
                        r.get("lieu") or "—",
                        (r.get("note") or "")[:150],
                        r["folder"],
                    )
                )
            box.insert("0.0", "\n\n".join(lines))
        box.configure(state="disabled")
        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=14, pady=10)
        ctk.CTkButton(
            bf, text="Verifier proches",
            command=lambda: self.check_indice_reminders(silent=False, days_ahead=60),
        ).pack(side="left", padx=2)
        ctk.CTkButton(
            bf, text="Export agenda (.ics)",
            fg_color=UI.get("accent"),
            command=self.export_indice_rappels_ics,
        ).pack(side="left", padx=2)
        ctk.CTkButton(
            bf, text="Config e-mail",
            command=self.open_email_alert_settings,
        ).pack(side="left", padx=2)
        ctk.CTkButton(
            bf, text="Envoyer e-mail",
            command=lambda: self.send_reminder_email_digest(force=True),
        ).pack(side="left", padx=2)
        ctk.CTkButton(bf, text="Fermer", fg_color=UI.get("card_alt"), command=win.destroy).pack(side="right")

    def check_device_reminders(self, silent=False):
        if not self.photo_folder_path:
            return
        devices = self._load_devices()
        alerts = []
        for dev in devices:
            if dev.get("statut") == "recupere":
                continue
            due = self._device_due_date(dev)
            if not due:
                continue
            delta = (due - datetime.now().date()).days
            if delta <= 1:
                lieu = dev.get("lieu") or ("%s %s" % (dev.get("lat"), dev.get("lon")))
                when = "en retard" if delta < 0 else ("aujourd'hui" if delta == 0 else "demain")
                bat = dev.get("batterie_pct")
                bat_s = " - batterie pose %s%%" % bat if bat is not None else ""
                alerts.append(
                    "• %s (%s) - releve %s\n  Lieu : %s%s" % (
                        dev.get("name"), dev.get("type"), when, lieu, bat_s
                    )
                )
        if not alerts:
            return
        msg = "Dispositifs a relever :\n\n" + "\n\n".join(alerts) + "\n\nOuvrir le gestionnaire ?"
        if silent:
            self.log("Rappel dispositifs :\n" + "\n".join(alerts))
            return
        if messagebox.askyesno("Releves dispositifs", msg):
            self.open_devices_manager()


    # ========== Traitement par lot · métadonnées GPS · captures ==========

    def open_batch_processor(self):
        """Enfile plusieurs dossiers pour synchronisation GPS automatique (GPX + decalage)."""
        win = ctk.CTkToplevel(self)
        win.title("Traitement par lot — GPS")
        win.geometry("720x560")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Traitement par lot (geolocalisation)",
            font=ctk.CTkFont(size=15, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=14, pady=(12, 4))
        ctk.CTkLabel(
            win,
            text="Ajoutez des dossiers photos + un GPX par ligne, puis lancez. Chaque dossier est traite l un apres l autre.",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            wraplength=680, justify="left",
        ).pack(anchor="w", padx=14, pady=(0, 8))

        jobs = []

        mid = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=10)
        mid.pack(fill="both", expand=True, padx=12, pady=6)
        list_box = ctk.CTkTextbox(mid, height=220)
        list_box.pack(fill="both", expand=True, padx=8, pady=8)
        list_box.insert("0.0", "(aucun lot)\n")
        list_box.configure(state="disabled")

        form = ctk.CTkFrame(win, fg_color="transparent")
        form.pack(fill="x", padx=12, pady=4)
        folder_var = tk.StringVar(value=self.photo_folder_path or "")
        gpx_var = tk.StringVar(value=self.gpx_file_path or "")
        offset_var = tk.StringVar(value="")
        try:
            if hasattr(self, "entry_geosync"):
                offset_var.set(self.entry_geosync.get().strip() or "")
        except Exception:
            pass
        force_var = tk.BooleanVar(value=False)

        def row_entry(parent, label, var, browse=None):
            r = ctk.CTkFrame(parent, fg_color="transparent")
            r.pack(fill="x", pady=2)
            ctk.CTkLabel(r, text=label, width=80, anchor="w", text_color=UI.get("text")).pack(side="left")
            e = ctk.CTkEntry(r, textvariable=var)
            e.pack(side="left", fill="x", expand=True, padx=4)
            if browse:
                ctk.CTkButton(r, text="...", width=36, command=browse).pack(side="left")

        def browse_folder():
            d = filedialog.askdirectory(title="Dossier photos")
            if d:
                folder_var.set(d)

        def browse_gpx():
            f = filedialog.askopenfilename(
                title="Trace GPX",
                filetypes=[("GPX", "*.gpx"), ("Tous", "*.*")],
            )
            if f:
                gpx_var.set(f)

        row_entry(form, "Dossier", folder_var, browse_folder)
        row_entry(form, "GPX", gpx_var, browse_gpx)
        row_entry(form, "Decalage", offset_var)
        ctk.CTkCheckBox(
            form, text="Forcer la reecriture GPS", variable=force_var,
            text_color=UI.get("text"),
        ).pack(anchor="w", pady=4)

        def refresh_list():
            list_box.configure(state="normal")
            list_box.delete("0.0", "end")
            if not jobs:
                list_box.insert("0.0", "(aucun lot)\n")
            else:
                lines = []
                for i, j in enumerate(jobs, 1):
                    lines.append(
                        "%d. %s\n   GPX: %s | decalage: %s | force: %s" % (
                            i,
                            j.get("folder"),
                            os.path.basename(j.get("gpx") or "") or "-",
                            j.get("offset") or "0",
                            "oui" if j.get("force") else "non",
                        )
                    )
                list_box.insert("0.0", "\n".join(lines))
            list_box.configure(state="disabled")

        def add_job():
            folder = folder_var.get().strip()
            gpx = gpx_var.get().strip()
            if not folder or not os.path.isdir(folder):
                messagebox.showwarning("Lot", "Dossier photos invalide.")
                return
            if not gpx or not os.path.isfile(gpx):
                messagebox.showwarning("Lot", "Fichier GPX invalide.")
                return
            jobs.append({
                "folder": folder,
                "gpx": gpx,
                "offset": offset_var.get().strip(),
                "force": bool(force_var.get()),
            })
            refresh_list()

        def add_current():
            if self.photo_folder_path:
                folder_var.set(self.photo_folder_path)
            if self.gpx_file_path:
                gpx_var.set(self.gpx_file_path)
            add_job()

        def clear_jobs():
            jobs.clear()
            refresh_list()

        btns = ctk.CTkFrame(win, fg_color="transparent")
        btns.pack(fill="x", padx=12, pady=6)
        ctk.CTkButton(btns, text="Ajouter la ligne", width=130, command=add_job).pack(side="left", padx=3)
        ctk.CTkButton(btns, text="Dossier courant", width=120, command=add_current).pack(side="left", padx=3)
        ctk.CTkButton(btns, text="Vider la liste", width=100, fg_color=UI.get("card_alt"), command=clear_jobs).pack(side="left", padx=3)

        status = ctk.CTkLabel(win, text="", text_color=UI.get("text_accent"))
        status.pack(anchor="w", padx=14)

        def run_batch():
            if not jobs:
                messagebox.showinfo("Lot", "Ajoutez au moins un dossier.")
                return
            exe = self.get_exiftool_path()
            if not exe:
                messagebox.showerror("ExifTool", "exiftool.exe introuvable.")
                return
            if not messagebox.askyesno("Lancer le lot", "%d dossier(s) vont etre geolocalises.\nContinuer ?" % len(jobs)):
                return

            queue = list(jobs)
            prev_folder = self.photo_folder_path
            prev_gpx = self.gpx_file_path
            try:
                prev_offset = self.entry_geosync.get() if hasattr(self, "entry_geosync") else ""
            except Exception:
                prev_offset = ""
            try:
                prev_force = self.force_resync.get() if hasattr(self, "force_resync") else False
            except Exception:
                prev_force = False

            def work():
                ok_n, err_n = 0, 0
                for idx, job in enumerate(queue):
                    def st(msg):
                        try:
                            win.after(0, lambda m=msg: status.configure(text=m))
                        except Exception:
                            pass
                    st("Lot %d/%d — %s" % (idx + 1, len(queue), os.path.basename(job["folder"])))
                    try:
                        self.photo_folder_path = job["folder"]
                        self.gpx_file_path = job["gpx"]
                        try:
                            if hasattr(self, "entry_geosync"):
                                self.entry_geosync.delete(0, "end")
                                self.entry_geosync.insert(0, job.get("offset") or "")
                        except Exception:
                            pass
                        try:
                            if hasattr(self, "force_resync"):
                                self.force_resync.set(bool(job.get("force")))
                        except Exception:
                            pass
                        self.execute_geotagging(exe)
                        ok_n += 1
                    except Exception as e:
                        err_n += 1
                        self.log("Lot erreur (%s): %s" % (job["folder"], e))
                self.photo_folder_path = prev_folder
                self.gpx_file_path = prev_gpx
                try:
                    if hasattr(self, "entry_geosync"):
                        self.entry_geosync.delete(0, "end")
                        self.entry_geosync.insert(0, prev_offset or "")
                    if hasattr(self, "force_resync"):
                        self.force_resync.set(prev_force)
                except Exception:
                    pass

                def done():
                    status.configure(text="Termine — OK: %d · erreurs: %d" % (ok_n, err_n))
                    messagebox.showinfo("Lot termine", "Reussis : %d\nErreurs : %d" % (ok_n, err_n))
                    try:
                        if prev_folder and os.path.isdir(prev_folder):
                            self.recharger_donnees_gps()
                    except Exception:
                        pass
                try:
                    win.after(0, done)
                except Exception:
                    pass

            threading.Thread(target=work, daemon=True).start()

        ctk.CTkButton(
            win, text="Lancer le traitement par lot", height=36,
            fg_color=UI.get("success", "#2f9e5f"), command=run_batch,
        ).pack(fill="x", padx=14, pady=(4, 6))
        ctk.CTkButton(
            win, text="Fermer", height=30, fg_color=UI.get("card_alt"), command=win.destroy,
        ).pack(fill="x", padx=14, pady=(0, 12))

    def open_gps_metadata_panel(self):
        """Voir / effacer / modifier les metadonnees GPS des photos selectionnees."""
        if not self.photo_folder_path:
            messagebox.showwarning("GPS", "Ouvrez d abord un dossier de photos.")
            return
        selection = []
        try:
            selection = [self.photo_listbox.get(i) for i in self.photo_listbox.curselection()]
        except Exception:
            pass
        if not selection:
            selection = self.get_supported_images()[:1]
            if not selection:
                messagebox.showinfo("GPS", "Aucune photo dans le dossier.")
                return
            messagebox.showinfo(
                "GPS",
                "Aucune selection : affichage de la premiere photo.\n"
                "Selectionnez des fichiers dans la liste pour un traitement groupe.",
            )

        win = ctk.CTkToplevel(self)
        win.title("Metadonnees GPS")
        win.geometry("560x520")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Metadonnees GPS",
            font=ctk.CTkFont(size=15, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=14, pady=(12, 4))
        ctk.CTkLabel(
            win, text="%d fichier(s) concerne(s)" % len(selection),
            text_color=UI.get("text_dim"),
        ).pack(anchor="w", padx=14)

        box = ctk.CTkTextbox(win, height=280)
        box.pack(fill="both", expand=True, padx=14, pady=8)

        def load_info():
            exe = self.get_exiftool_path()
            lines = []
            if not exe:
                lines.append("ExifTool introuvable.")
            else:
                for fn in selection[:40]:
                    path = os.path.join(self.photo_folder_path, fn)
                    if not os.path.isfile(path):
                        lines.append("%s — fichier absent" % fn)
                        continue
                    try:
                        startupinfo = self._exiftool_startupinfo()
                        cmd = [
                            exe, "-n", "-json", "-s",
                            "-GPSLatitude", "-GPSLongitude", "-GPSAltitude",
                            "-CreateDate", "-DateTimeOriginal", "-OffsetTimeOriginal",
                            path,
                        ]
                        r = subprocess.run(
                            cmd, capture_output=True, text=True, timeout=30,
                            startupinfo=startupinfo, encoding="utf-8", errors="replace",
                        )
                        data = json.loads(r.stdout or "[]")
                        item = data[0] if data else {}
                        lat = item.get("GPSLatitude")
                        lon = item.get("GPSLongitude")
                        alt = item.get("GPSAltitude")
                        dt = item.get("DateTimeOriginal") or item.get("CreateDate") or "?"
                        if lat is not None and lon is not None:
                            gps_s = "%.6f, %.6f" % (float(lat), float(lon))
                        else:
                            gps_s = "(pas de GPS)"
                        alt_s = (" alt %.0fm" % float(alt)) if alt is not None else ""
                        lines.append("%s\n  GPS: %s%s\n  Date: %s" % (fn, gps_s, alt_s, dt))
                    except Exception as e:
                        lines.append("%s — erreur: %s" % (fn, e))
                if len(selection) > 40:
                    lines.append("... (%d autres non listes)" % (len(selection) - 40))
            def apply():
                box.configure(state="normal")
                box.delete("0.0", "end")
                box.insert("0.0", "\n\n".join(lines))
                box.configure(state="disabled")
            try:
                win.after(0, apply)
            except Exception:
                pass

        def remove_gps():
            if not messagebox.askyesno(
                "Effacer GPS",
                "Supprimer les balises GPS de %d fichier(s) ?" % len(selection),
            ):
                return
            exe = self.get_exiftool_path()
            if not exe:
                messagebox.showerror("ExifTool", "exiftool.exe introuvable.")
                return
            paths = [
                os.path.join(self.photo_folder_path, fn)
                for fn in selection
                if os.path.isfile(os.path.join(self.photo_folder_path, fn))
            ]
            if not paths:
                return
            try:
                startupinfo = self._exiftool_startupinfo()
                cmd = [exe, "-gps:all=", "-overwrite_original"] + paths
                subprocess.run(
                    cmd, capture_output=True, text=True, timeout=max(60, 5 * len(paths)),
                    startupinfo=startupinfo, encoding="utf-8", errors="replace",
                )
                for fn in selection:
                    if fn in getattr(self, "photos_data", {}):
                        self.photos_data[fn]["lat"] = None
                        self.photos_data[fn]["lon"] = None
                self.log("GPS efface sur %d fichier(s)" % len(paths))
                messagebox.showinfo("GPS", "Metadonnees GPS supprimees.")
                try:
                    self.recharger_donnees_gps()
                except Exception:
                    pass
                load_info()
            except Exception as e:
                messagebox.showerror("GPS", str(e))

        def open_edit():
            win.destroy()
            self.open_edit_gps_dialog()

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=14, pady=(0, 12))
        ctk.CTkButton(bf, text="Rafraichir", width=100, command=lambda: threading.Thread(target=load_info, daemon=True).start()).pack(side="left", padx=3)
        ctk.CTkButton(bf, text="Modifier GPS...", width=120, fg_color=UI.get("accent"), command=open_edit).pack(side="left", padx=3)
        ctk.CTkButton(bf, text="Effacer GPS", width=110, fg_color=UI.get("danger", "#c44"), command=remove_gps).pack(side="left", padx=3)
        ctk.CTkButton(bf, text="Fermer", width=90, fg_color=UI.get("card_alt"), command=win.destroy).pack(side="right", padx=3)

        threading.Thread(target=load_info, daemon=True).start()

    def _grab_widget_screenshot(self, widget, path):
        """Capture un widget Tk/CTk via bbox ecran."""
        try:
            from PIL import ImageGrab
        except Exception as e:
            raise RuntimeError("PIL.ImageGrab indisponible: %s" % e)
        widget.update_idletasks()
        x = widget.winfo_rootx()
        y = widget.winfo_rooty()
        w = widget.winfo_width()
        h = widget.winfo_height()
        if w < 10 or h < 10:
            raise RuntimeError("Zone trop petite a capturer")
        img = ImageGrab.grab(bbox=(x, y, x + w, y + h))
        img.save(path)
        return path

    def capture_map_screenshot(self):
        """Enregistre une capture de la carte."""
        if not hasattr(self, "map_widget") or self.map_widget is None:
            messagebox.showwarning("Capture", "Carte indisponible.")
            return
        initial = "carte_%s.png" % datetime.now().strftime("%Y%m%d_%H%M%S")
        path = filedialog.asksaveasfilename(
            title="Enregistrer la capture de carte",
            defaultextension=".png",
            initialfile=initial,
            filetypes=[("PNG", "*.png"), ("JPEG", "*.jpg"), ("Tous", "*.*")],
        )
        if not path:
            return
        try:
            target = getattr(self, "map_widget", None)
            try:
                parent = target.master
                if parent is not None:
                    target = parent
            except Exception:
                pass
            self._grab_widget_screenshot(target, path)
            self.log("Capture carte : %s" % path)
            messagebox.showinfo("Capture", "Carte enregistree :\n%s" % path)
        except Exception as e:
            messagebox.showerror("Capture", "Impossible de capturer la carte :\n%s" % e)

    def capture_window_screenshot(self):
        """Capture la fenetre principale GeoExif."""
        initial = "geoexif_%s.png" % datetime.now().strftime("%Y%m%d_%H%M%S")
        path = filedialog.asksaveasfilename(
            title="Enregistrer la capture d ecran",
            defaultextension=".png",
            initialfile=initial,
            filetypes=[("PNG", "*.png"), ("JPEG", "*.jpg"), ("Tous", "*.*")],
        )
        if not path:
            return
        try:
            self.update_idletasks()
            self._grab_widget_screenshot(self, path)
            self.log("Capture fenetre : %s" % path)
            messagebox.showinfo("Capture", "Fenetre enregistree :\n%s" % path)
        except Exception as e:
            messagebox.showerror("Capture", "Impossible de capturer :\n%s" % e)

    def capture_quick_to_folder(self):
        """Capture carte dans le dossier de sortie (ou Images)."""
        if self.photo_folder_path and os.path.isdir(self.photo_folder_path):
            dest_dir = self.photo_folder_path
        else:
            dest_dir = os.path.join(os.path.expanduser("~"), "Pictures")
            os.makedirs(dest_dir, exist_ok=True)
        path = os.path.join(
            dest_dir, "capture_carte_%s.png" % datetime.now().strftime("%Y%m%d_%H%M%S")
        )
        try:
            target = getattr(self, "map_widget", None)
            if target is None:
                messagebox.showwarning("Capture", "Carte indisponible.")
                return
            try:
                if target.master is not None:
                    target = target.master
            except Exception:
                pass
            self._grab_widget_screenshot(target, path)
            self.log("Capture rapide : %s" % path)
            messagebox.showinfo("Capture", "Enregistre dans :\n%s" % path)
        except Exception as e:
            messagebox.showerror("Capture", str(e))


    def open_outils_hub(self):
        """Panneau central regroupant tous les outils secondaires par catégorie, pour désencombrer la sidebar."""

        win = ctk.CTkToplevel(self)
        win.title("Outils, rapports & réglages")
        win.geometry("480x680")
        try:
            win.configure(fg_color=UI.get("bg", "#101612"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        scroll = ctk.CTkScrollableFrame(
            win, fg_color="transparent",
            label_text_color=UI.get("text", "#111"),
        )
        scroll.pack(fill="both", expand=True, padx=15, pady=15)

        def lancer(fonction):
            win.destroy()
            fonction()

        def section(titre, description=None, couleur_titre=None):
            ctk.CTkLabel(
                scroll, text=titre,
                font=ctk.CTkFont(size=13, weight="bold"),
                text_color=couleur_titre or UI.get("text", "#111"),
            ).pack(anchor="w", pady=(14, 0))
            if description:
                ctk.CTkLabel(
                    scroll, text=description,
                    font=ctk.CTkFont(size=11),
                    text_color=UI.get("text_dim", "#444"),
                    anchor="w", justify="left", wraplength=430,
                ).pack(anchor="w", pady=(0, 6))
            else:
                ctk.CTkLabel(scroll, text="", height=2).pack()
            cadre = ctk.CTkFrame(scroll, fg_color="transparent")
            cadre.pack(fill="x")
            cadre.grid_columnconfigure((0, 1), weight=1)
            return cadre

        def bouton(cadre, texte, commande, ligne, col, span=1, couleur=None, survol=None):
            couleur = couleur or UI.get("accent", "#2b6cb0")
            survol = survol or UI.get("accent_hover", "#1f4f80")
            # Texte toujours lisible sur fond coloré
            b = ctk.CTkButton(
                cadre, text=texte,
                fg_color=couleur, hover_color=survol,
                text_color="#ffffff",
                font=ctk.CTkFont(size=11, weight="bold"),
                height=34,
                command=lambda: lancer(commande),
            )
            b.grid(row=ligne, column=col, columnspan=span, sticky="ew", padx=3, pady=3)
            return b

        c = section("📤  Exports du jour", "Exporte les observations du dossier actuellement ouvert (dossier de rapports configurable).")
        bouton(c, "📄 CSV", self.export_csv, 0, 0)
        bouton(c, "🌍 GPX", self.export_gpx, 0, 1)
        bouton(c, "🌍 KML / KMZ", self.export_kml, 1, 0, couleur="#1f7d5a", survol="#175f45")
        bouton(c, "🗺️ GeoJSON", self.export_geojson, 1, 1, couleur="#1f7d5a", survol="#175f45")
        bouton(c, "📕 PDF", self.export_daily_pdf, 2, 0, couleur="#b03a2e", survol="#8a2d24")
        bouton(c, "📘 Word", self.export_daily_docx, 2, 1)
        bouton(c, "Lot GPS multi-dossiers", self.open_batch_processor, 3, 0, couleur="#e67e22", survol="#d35400")
        bouton(c, "Metadonnees GPS", self.open_gps_metadata_panel, 3, 1, couleur="#1f5d8a", survol="#174868")
        bouton(c, "Affut / serie GPS+heure", self.open_affut_series_dialog, 5, 0, span=2, couleur="#8e44ad", survol="#6c3483")
        bouton(c, "Vider cache apercus", self.clear_preview_cache, 6, 0, span=2, couleur="#5a6570", survol="#3d4650")
        bouton(c, "Rappels indices / suivi", self.open_rappels_indices, 7, 0, span=2, couleur="#c4841a", survol="#9a6510")
        bouton(c, "Configuration", lambda: self.open_setup_wizard(False), 8, 0, span=2, couleur="#3d7ea6", survol="#2c5f7c")
        bouton(c, "Sauvegardes automatiques", self.open_backup_settings, 9, 0, span=2, couleur="#6b7280", survol="#4b5563")
        bouton(c, "Rescanner mes sorties", self.scan_parent_for_sorties, 10, 0, span=2, couleur="#0d9488", survol="#0f766e")
        bouton(c, "Gérer / retirer des sorties", self.open_manage_sorties, 11, 0, span=2, couleur="#c0392b", survol="#a93226")
        bouton(c, "Archive locale (recovery)", self.open_local_archive_browser, 12, 0, span=2, couleur="#0369a1", survol="#075985")
        bouton(c, "Débrief texte (archive, sans photos)", self.open_offline_text_debrief, 13, 0, span=2, couleur="#15803d", survol="#166534")
        bouton(c, "Pack recovery ZIP", self.create_recovery_pack_zip, 14, 0, span=2, couleur="#7c3aed", survol="#6d28d9")
        bouton(c, "Indices — dossier + import photos", self.open_indices_folder_workflow, 15, 0, span=2, couleur="#b45309", survol="#92400e")
        bouton(c, "Capture carte", self.capture_map_screenshot, 4, 0, couleur="#5a6570", survol="#3d4650")
        bouton(c, "Capture fenetre", self.capture_window_screenshot, 4, 1, couleur="#5a6570", survol="#3d4650")

        c = section(
            "🗺️  Vue d'ensemble — toutes sorties",
            "Combine automatiquement TOUS les dossiers déjà ouverts dans l'app (pas seulement celui du jour) : "
            "carte cumulée = tous les points sur une carte ; rapport agrégé = synthèse écrite/PDF de plusieurs sorties."
        )
        bouton(c, "🗂️ Carte cumulée", self.open_multi_sorties_map, 0, 0, couleur="#1f7d5a", survol="#175f45")
        bouton(c, "📅 Brief IA multi-sorties", self.open_aggregated_report, 0, 1, couleur="#6f42c1", survol="#59339d")

        c = section(
            "📊  Analyse",
            "Statistiques et historique par espèce, calculés sur toutes les sorties connues."
        )
        bouton(c, "📊 Tableau de bord", self.open_dashboard, 0, 0, couleur="#1f7d5a", survol="#175f45")
        bouton(c, "🔎 Fiche espèce", self.open_species_sheet, 0, 1)
        bouton(c, "📷 Analyse EXIF", self.open_exif_analysis, 1, 0, couleur="#e67e22", survol="#d35400")
        bouton(c, "🗺️ GeoJSON toutes sorties", self.export_geojson_all_sorties, 1, 1, couleur="#1f7d5a", survol="#175f45")
        bouton(c, "❓ FAQ", self.open_faq, 2, 0, couleur="#5a6570", survol="#3d4650")
        bouton(c, "☁ Cloud (Drive/Dropbox)", self.open_cloud_settings, 3, 0, couleur="#3d9cf0", survol="#2b7fc4")
        bouton(c, "📁 Projet Meuse / saison", self.open_project_hub, 4, 0, couleur="#3eb4a0", survol="#2f9a88")
        bouton(c, "🌿 Cibles de saison", self.open_season_suggestions, 4, 1, couleur="#1f7d5a", survol="#175f45")
        bouton(c, "📊 Analyse effort / richesse", self.open_analysis_panel, 5, 0, couleur="#e67e22", survol="#d35400")
        bouton(c, "📦 Export paquet terrain", self.export_field_package, 5, 1, couleur="#1f7d5a", survol="#175f45")
        bouton(c, "📥 Import paquet / GPX", self.import_field_package, 6, 0, couleur="#3d9cf0", survol="#2b7fc4")
        bouton(c, "📏 Mesure carte (km)", self.toggle_map_measure, 6, 1, couleur="#3a3a3a", survol="#4a4a4a")
        bouton(c, "☁ Sync carnet maintenant", self.sync_carnet_to_cloud, 3, 1, couleur="#1f7d5a", survol="#175f45")
        bouton(c, "🐦 BirdNET (chants)", self.open_birdnet_panel, 2, 1, couleur="#1f7d5a", survol="#175f45")

        c = section(
            "🔮  Terrain",
            "Prévisions météo, phases lunaires, et mode compagnon smartphone (saisie GPS + note)."
        )
        bouton(c, "🔮 Prévisions & Phases lunaires", self.open_forecast_view, 0, 0, couleur="#1f5d8a", survol="#174868")
        bouton(c, "📱 Compagnon smartphone", self.open_companion_panel, 0, 1, couleur="#1f7d5a", survol="#175f45")
        bouton(c, "📷 Cameras & enregistreurs", self.open_devices_manager, 1, 0, span=2, couleur="#e67e22", survol="#d35400")
        bouton(c, "📷 Caméras & enregistreurs", self.open_devices_manager, 1, 0, span=2, couleur="#e67e22", survol="#d35400")

        c = section(
            "⚙️  Configuration",
            "Réglages : dossier de sauvegarde des rapports, clé IA, dictionnaire d'espèces."
        )
        bouton(c, "📁 Dossier des rapports", self.open_reports_dir_settings, 0, 0, couleur="#1f5d8a", survol="#174868")
        bouton(c, "💾 Sauvegardes carnet", self.open_notes_backup_manager, 0, 1, couleur="#1f5d8a", survol="#174868")
        bouton(c, "🤖 Identification IA", self.open_ai_settings, 1, 0, couleur="#3a3a3a", survol="#4a4a4a")
        bouton(c, "🗂️ Gérer les espèces", self.open_species_manager, 1, 1, couleur="#3a3a3a", survol="#4a4a4a")
        bouton(c, "🏷️ Mots-clés Lightroom (XMP)", self.toggle_xmp_keywords_pref, 2, 0, span=2, couleur="#e67e22", survol="#d35400")

    # --- Mots-clés Lightroom (XMP/IPTC) écrits automatiquement à l'enregistrement ---

    def _xmp_keywords_enabled(self):
        """Renvoie True si l'écriture automatique des mots-clés XMP/IPTC est activée."""
        return bool((getattr(self, "app_config", {}) or {}).get("xmp_keywords_enabled", False))

    def toggle_xmp_keywords_pref(self):
        """Fenêtre de réglage : active/désactive l'écriture auto des mots-clés Lightroom (XMP/IPTC)."""
        win = ctk.CTkToplevel(self)
        win.title("🏷️ Mots-clés Lightroom (XMP)")
        win.geometry("520x260")
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win,
            text="Quand cette option est active, chaque enregistrement d'observation écrit\n"
                 "automatiquement des mots-clés (espèce, catégorie, lieu, département) dans les\n"
                 "champs XMP:Subject / IPTC:Keywords des photos — reconnus par Lightroom, Bridge, etc.",
            font=ctk.CTkFont(size=12), wraplength=480, justify="left"
        ).pack(anchor="w", padx=15, pady=(15, 10))

        var_actif = ctk.BooleanVar(value=self._xmp_keywords_enabled())

        def on_toggle():
            cfg = dict(self.app_config or {})
            cfg["xmp_keywords_enabled"] = bool(var_actif.get())
            self._save_app_config(cfg)
            etat = "activée" if var_actif.get() else "désactivée"
            self.log(f"🏷️ Écriture automatique des mots-clés XMP/IPTC {etat}.")

        ctk.CTkSwitch(
            win, text="Écrire automatiquement les mots-clés XMP/IPTC",
            variable=var_actif, command=on_toggle,
            font=ctk.CTkFont(size=12),
        ).pack(anchor="w", padx=15, pady=(0, 15))

        ctk.CTkLabel(
            win,
            text="Astuce : les mots-clés sont ajoutés au fil des enregistrements (pas de doublons\n"
                 "supprimés automatiquement). Rouvrez le fichier dans Lightroom (F5) après écriture.",
            font=ctk.CTkFont(size=11), text_color="#8ab4ff", wraplength=480, justify="left"
        ).pack(anchor="w", padx=15, pady=(0, 10))

        ctk.CTkButton(win, text="Fermer", width=100, command=win.destroy).pack(anchor="e", padx=15, pady=(5, 15))
        win.bind("<Escape>", lambda e: win.destroy())

    def _write_xmp_keywords_to_files(self, paths, data, silent=False):
        """Écrit les mots-clés (espèce, catégorie, lieu, département) en XMP:Subject / IPTC:Keywords via ExifTool."""
        exe_path = self.get_exiftool_path()
        if not exe_path:
            if not silent:
                self.after(0, lambda: messagebox.showerror(
                    "ExifTool introuvable", "Posez exiftool.exe à côté du script ou installez-le dans le PATH."
                ))
            return

        chemins = [p for p in paths if p and os.path.isfile(p)]
        if not chemins:
            return

        # Construit la liste de mots-clés (dédupliqués, ordre préservé), à partir des champs de l'observation
        mots_cles = []
        for champ in ("espece", "categorie", "lieu", "departement"):
            val = data.get(champ)
            if isinstance(val, str):
                val = val.strip()
            if val:
                val = str(val).strip()
                if val not in mots_cles:
                    mots_cles.append(val)
        if "Affût" not in mots_cles:
            mots_cles.append("Affût")

        if not mots_cles:
            return

        try:
            startupinfo = self._exiftool_startupinfo()
            cmd = [exe_path, "-P", "-overwrite_original", "-codedcharacterset=utf8"]
            for kw in mots_cles:
                cmd.append(f"-XMP-dc:Subject+={kw}")
                cmd.append(f"-IPTC:Keywords+={kw}")
            cmd += chemins
            proc = subprocess.run(
                cmd, startupinfo=startupinfo, capture_output=True, text=True,
                encoding="utf-8", errors="ignore", timeout=180
            )
            if proc.returncode == 0:
                self.log(f"🏷️ Mots-clés XMP/IPTC écrits sur {len(chemins)} photo(s) : {', '.join(mots_cles)}")
            else:
                err = (proc.stderr or proc.stdout or "code retour non nul").strip()[:300]
                self.log(f"❌ Échec écriture mots-clés XMP : {err}")
                if not silent:
                    self.after(0, lambda: messagebox.showerror("Échec", f"Impossible d'écrire les mots-clés :\n{err}"))
        except Exception as e:
            msg = str(e)
            self.log(f"❌ Échec écriture mots-clés XMP : {msg}")
            if not silent:
                self.after(0, lambda: messagebox.showerror("Échec", msg))

    def open_dashboard(self):
        """Tableau de bord : stats multi-sorties + raccourcis carte / fiche espèce."""
        try:
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
            from matplotlib.figure import Figure
        except ImportError:
            messagebox.showerror(
                "Bibliothèque manquante",
                "Le tableau de bord nécessite matplotlib.\n\npip install matplotlib"
            )
            return

        folders = self._load_known_folders()
        if not folders:
            messagebox.showinfo("Aucune donnée", "Aucune sortie connue pour établir des statistiques.")
            return

        win = ctk.CTkToplevel(self)
        win.title("Tableau de bord — Toutes les sorties")
        win.geometry("1100x820")
        win.configure(fg_color=UI.get("bg", "#0f1419"))
        self._prepare_tool_window(win)

        header = ctk.CTkFrame(win, fg_color=UI.get("card", "#1a2332"), corner_radius=12)
        header.pack(fill="x", padx=14, pady=(14, 8))
        ctk.CTkLabel(
            header, text="Vue d'ensemble de vos sorties",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(side="left", padx=14, pady=12)
        lbl_status = ctk.CTkLabel(header, text="Calcul…", text_color=UI.get("text_dim", "#888"))
        lbl_status.pack(side="right", padx=14)

        def worker():
            dataset = self._build_report_dataset(
                [f["path"] for f in folders if os.path.isdir(f.get("path", ""))],
                read_exif=False,
            )
            self.after(0, lambda: self._render_dashboard(win, lbl_status, dataset, len(folders)))

        threading.Thread(target=worker, daemon=True).start()

    def _render_dashboard(self, win, lbl_status, dataset, n_folders):
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        from matplotlib.figure import Figure

        if not win.winfo_exists():
            return
        try:
            lbl_status.configure(text="Prêt")
        except Exception:
            pass

        if not dataset:
            ctk.CTkLabel(win, text="Aucune observation dans vos sorties connues.").pack(pady=20)
            return

        especes_count = collections.Counter()
        categories_count = collections.Counter()
        mois_count = collections.Counter()
        n_gps = 0
        for obs in dataset:
            especes_count[obs.get("espece") or "Inconnu"] += 1
            categories_count[obs.get("categorie") or "Non classé"] += 1
            if obs.get("lat") and obs.get("lon"):
                n_gps += 1
            date_str = None
            met = obs.get("meteo")
            if isinstance(met, dict):
                date_str = met.get("date")
            if not date_str:
                # essayer heure / dossier
                for cand in (
                    obs.get("date"),
                    (obs.get("heure") or "")[:10],
                    obs.get("dossier"),
                ):
                    if not cand:
                        continue
                    s = str(cand)
                    m = re.search(r"(20\d{2}-\d{2}-\d{2})", s)
                    if m:
                        date_str = m.group(1)
                        break
                    m = re.search(r"(20\d{2}-\d{2})", s)
                    if m and len(m.group(1)) == 7:
                        date_str = m.group(1) + "-01"
                        break
            if date_str:
                try:
                    if len(date_str) >= 7:
                        mois = date_str[:7]
                        mois_count[mois] += 1
                except Exception:
                    pass

        # Cartes résumé
        cards = ctk.CTkFrame(win, fg_color="transparent")
        cards.pack(fill="x", padx=14, pady=6)
        for titre, val, color in (
            ("Observations", str(len(dataset)), UI.get("accent", "#3d9cf0")),
            ("Sorties", str(n_folders), UI.get("success", "#2ecc71")),
            ("Espèces", str(len(especes_count)), UI.get("purple", "#9b59b6")),
            ("Avec GPS", str(n_gps), UI.get("warning", "#e67e22")),
        ):
            card = ctk.CTkFrame(cards, fg_color=UI.get("card", "#1a2332"), corner_radius=12, border_width=1, border_color=UI.get("border", "#2a3544"))
            card.pack(side="left", expand=True, fill="x", padx=4)
            ctk.CTkLabel(card, text=val, font=ctk.CTkFont(size=22, weight="bold"), text_color=color).pack(pady=(12, 0))
            ctk.CTkLabel(card, text=titre, font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#888")).pack(pady=(0, 12))

        actions = ctk.CTkFrame(win, fg_color="transparent")
        actions.pack(fill="x", padx=14, pady=(4, 8))
        ctk.CTkButton(
            actions, text="Carte cumulée", command=self.open_multi_sorties_map,
            fg_color=UI.get("success", "#2ecc71"), hover_color=UI.get("success_hover", "#27ae60"), height=32, width=140
        ).pack(side="left", padx=(0, 6))
        ctk.CTkButton(
            actions, text="Fiche espèce + carte", command=self.open_species_sheet,
            fg_color=UI.get("purple", "#9b59b6"), hover_color=UI.get("purple_hover", "#8e44ad"), height=32, width=160
        ).pack(side="left", padx=6)
        ctk.CTkButton(
            actions, text="Brief IA multi-sorties", command=self.open_aggregated_report,
            fg_color=UI.get("accent", "#3d9cf0"), hover_color=UI.get("accent_hover", "#2b7fc4"), height=32, width=160
        ).pack(side="left", padx=6)

        fig_bg = UI.get("bg", "#0f1419")
        fig = Figure(figsize=(10.2, 6.4), dpi=100, facecolor=fig_bg)

        ax1 = fig.add_subplot(221)
        top_especes = especes_count.most_common(10)
        st = self._style_chart_axes(ax1)
        if top_especes:
            noms = [n for n, _ in top_especes][::-1]
            vals = [v for _, v in top_especes][::-1]
            ax1.barh(noms, vals, color=UI.get("accent", "#3d9cf0"))
            ax1.set_title("Especes les plus observees", color=st["title"], fontsize=10)
            try:
                ax1.tick_params(axis="y", labelsize=8, colors=st["tick"])
            except Exception:
                pass

        ax2 = fig.add_subplot(222)
        st2 = self._style_chart_axes(ax2)
        if categories_count:
            labels = list(categories_count.keys())
            vals = list(categories_count.values())
            couleurs = [CATEGORY_COLORS.get(l, "#7f7f7f") for l in labels]
            text_c = st2["tick"]
            ax2.pie(
                vals, labels=labels, autopct="%1.0f%%", colors=couleurs,
                textprops={"color": text_c, "fontsize": 8},
            )
            ax2.set_title("Repartition par categorie", color=st2["title"], fontsize=10)

        ax3 = fig.add_subplot(212)
        st3 = self._style_chart_axes(ax3)
        if mois_count:
            mois_tries = sorted(mois_count.keys())
            vals = [mois_count[m] for m in mois_tries]
            ax3.bar(mois_tries, vals, color=UI.get("success", "#2ecc71"))
            ax3.set_title("Observations par mois", color=st3["title"], fontsize=10)
            ax3.tick_params(axis="x", rotation=45, colors=st3["tick"])
            ax3.tick_params(axis="y", colors=st3["tick"])
        else:
            ax3.text(
                0.5, 0.5, "Pas assez de dates pour la courbe mensuelle",
                ha="center", va="center", color=st3["tick"],
            )
            ax3.set_axis_off()

        fig.tight_layout()
        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True, padx=12, pady=(0, 12))

    def _enrich_dataset_gps(self, dataset):
        """Complete lat/lon manquants via ExifTool (par dossier), pour fiche espece / cartes cumulees."""
        if not dataset:
            return dataset
        exe = self.get_exiftool_path()
        if not exe:
            self.log("Fiche espece : ExifTool introuvable — GPS fichiers non relus.")
            return dataset
        # Grouper les fichiers sans GPS par dossier
        by_folder = {}
        for o in dataset:
            if o.get("lat") is not None and o.get("lon") is not None:
                continue
            fp = o.get("dossier_path") or ""
            fn = o.get("fichier") or ""
            if not fp or not fn or str(fn).startswith("_"):
                continue
            full = o.get("chemin_complet") or os.path.join(fp, fn)
            if not os.path.isfile(full):
                continue
            by_folder.setdefault(fp, []).append((fn, full, o))
        if not by_folder:
            return dataset
        startupinfo = self._exiftool_startupinfo()
        found = 0
        for folder, items in by_folder.items():
            # lots de 40
            for i in range(0, len(items), 40):
                chunk = items[i:i+40]
                paths = [c[1] for c in chunk]
                try:
                    cmd = [exe, "-n", "-json", "-s", "-GPSLatitude", "-GPSLongitude"] + paths
                    r = subprocess.run(
                        cmd, capture_output=True, text=True, timeout=max(45, 8 * len(paths)),
                        startupinfo=startupinfo, encoding="utf-8", errors="replace",
                    )
                    if not r.stdout:
                        continue
                    meta = json.loads(r.stdout)
                    by_name = {}
                    for it in meta:
                        src = it.get("SourceFile") or ""
                        by_name[os.path.basename(src.replace("/", os.sep))] = it
                    for fn, full, o in chunk:
                        it = by_name.get(fn) or by_name.get(os.path.basename(full))
                        if not it:
                            continue
                        lat = self.parse_coord(it.get("GPSLatitude"))
                        lon = self.parse_coord(it.get("GPSLongitude"))
                        if lat is None or lon is None:
                            continue
                        o["lat"], o["lon"] = lat, lon
                        found += 1
                        # persister dans le carnet du dossier pour les prochaines fois
                        try:
                            notes = self._load_notes_dict(folder)
                            if isinstance(notes.get(fn), dict):
                                notes[fn]["lat"] = lat
                                notes[fn]["lon"] = lon
                                self._save_notes_dict(notes, folder_path=folder, create_backup=False, silent=True)
                        except Exception:
                            pass
                except Exception as e:
                    self.log("Enrich GPS dossier %s : %s" % (os.path.basename(folder), e))
        if found:
            self.log("Fiche espece : %d GPS recuperes depuis les fichiers photo." % found)
        return dataset

    def open_species_sheet(self):
        """Fiche espèce : stats + carte cumulée de tous les points GPS de l'espèce."""
        folders = self._load_known_folders()
        if not folders:
            messagebox.showinfo("Aucune donnée", "Aucune sortie connue pour le moment.")
            return

        self.configure(cursor="watch")
        self.update_idletasks()
        try:
            dataset_complet = self._build_report_dataset(
                [f["path"] for f in folders if os.path.isdir(f.get("path", ""))],
                read_exif=False,
            )
            # Relire les GPS dans les fichiers si absents du carnet (sinon carte vide)
            dataset_complet = self._enrich_dataset_gps(dataset_complet)
        finally:
            self.configure(cursor="")

        especes_dispo = sorted({o.get("espece") for o in dataset_complet if o.get("espece")})
        if not especes_dispo:
            messagebox.showinfo("Aucune donnée", "Aucune espèce enregistrée pour le moment.")
            return

        win = ctk.CTkToplevel(self)
        win.title("Fiche espèce — cumul multi-sorties")
        win.geometry("1100x780")
        win.configure(fg_color=UI.get("bg", "#0f1419"))
        self._prepare_tool_window(win)

        top = ctk.CTkFrame(win, fg_color=UI.get("card", "#1a2332"), corner_radius=12)
        top.pack(fill="x", padx=14, pady=(14, 8))
        inner = ctk.CTkFrame(top, fg_color="transparent")
        inner.pack(fill="x", padx=12, pady=10)
        ctk.CTkLabel(inner, text="Espèce", font=ctk.CTkFont(size=12), text_color=UI.get("text_dim", "#888")).pack(side="left", padx=(0, 8))
        choix_espece = ctk.CTkOptionMenu(
            inner, values=especes_dispo, width=300,
            fg_color=UI.get("card_alt", "#1e2a3a"), button_color=UI.get("accent", "#3d9cf0"),
        )
        choix_espece.pack(side="left")
        choix_espece.set(especes_dispo[0])
        lbl_resume = ctk.CTkLabel(inner, text="", font=ctk.CTkFont(size=12), text_color=UI.get("text_dim", "#888"))
        lbl_resume.pack(side="right", padx=8)

        body = ctk.CTkFrame(win, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=14, pady=(0, 14))
        body.grid_columnconfigure(0, weight=1)
        body.grid_columnconfigure(1, weight=1)
        body.grid_rowconfigure(0, weight=1)

        left = ctk.CTkScrollableFrame(body, fg_color=UI.get("card", "#1a2332"), corner_radius=12)
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 6))
        right = ctk.CTkFrame(body, fg_color=UI.get("card", "#1a2332"), corner_radius=12)
        right.grid(row=0, column=1, sticky="nsew", padx=(6, 0))
        right.grid_rowconfigure(1, weight=1)
        right.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(right, text="Carte — tous les points de l'espèce", font=ctk.CTkFont(size=13, weight="bold")).grid(
            row=0, column=0, sticky="w", padx=12, pady=(10, 4)
        )
        map_w = TkinterMapView(right, corner_radius=10)
        map_w.grid(row=1, column=0, sticky="nsew", padx=10, pady=(0, 10))
        try:
            map_w.set_tile_server(
                "https://data.geopf.fr/wmts?SERVICE=WMTS&REQUEST=GetTile&VERSION=1.0.0"
                "&LAYER=GEOGRAPHICALGRIDSYSTEMS.PLANIGNV2&STYLE=normal&FORMAT=image/png"
                "&TILEMATRIXSET=PM&TILEMATRIX={z}&TILEROW={y}&TILECOL={x}",
                max_zoom=19,
            )
        except Exception:
            pass
        map_w.set_position(49.1627, 5.3854)
        map_w.set_zoom(9)
        map_markers = []

        def render(*_):
            for w in left.winfo_children():
                w.destroy()
            for m in map_markers:
                try:
                    m.delete()
                except Exception:
                    pass
            map_markers.clear()

            espece = choix_espece.get()
            def _same_esp(a, b):
                return (a or "").strip().casefold() == (b or "").strip().casefold()
            obs_espece = [o for o in dataset_complet if _same_esp(o.get("espece"), espece)]
            total_individus = 0
            for o in obs_espece:
                try:
                    total_individus += int(str(o.get("nombre", "1")).split("-")[0].replace("+", "").strip() or "1")
                except Exception:
                    total_individus += 1

            lieux = collections.Counter(o.get("lieu") for o in obs_espece if o.get("lieu"))
            dossiers = collections.Counter(o.get("dossier") for o in obs_espece if o.get("dossier"))
            cat = next((o.get("categorie") for o in obs_espece if o.get("categorie")), "Non classé")

            lbl_resume.configure(
                text=f"{len(obs_espece)} obs.  ·  ~{total_individus} ind.  ·  {len(dossiers)} sortie(s)"
            )

            ctk.CTkLabel(left, text=espece, font=ctk.CTkFont(size=20, weight="bold"), text_color=UI.get("text", "#eee")).pack(
                anchor="w", padx=12, pady=(12, 2)
            )
            ctk.CTkLabel(
                left,
                text=f"Catégorie : {cat}",
                font=ctk.CTkFont(size=12), text_color=CATEGORY_COLORS.get(cat, "#95a5a6"),
            ).pack(anchor="w", padx=12, pady=(0, 8))

            # --- Résumé Wikipédia (chargé en arrière-plan, mis en cache localement) ---
            wiki_frame = ctk.CTkFrame(left, fg_color=UI.get("card_alt", "#1e2a3a"), corner_radius=10)
            wiki_frame.pack(fill="x", padx=12, pady=(0, 10))

            lbl_wiki_photo = ctk.CTkLabel(wiki_frame, text="", width=70)
            lbl_wiki_photo.pack(side="left", padx=(10, 6), pady=10)

            lbl_wiki_extract = ctk.CTkLabel(
                wiki_frame, text="🔎 Recherche sur Wikipédia…", font=ctk.CTkFont(size=11),
                text_color=UI.get("text_dim", "#888"), justify="left", anchor="w", wraplength=460,
            )
            lbl_wiki_extract.pack(side="left", padx=(0, 10), pady=10, fill="x", expand=True)

            btn_wiki = ctk.CTkButton(
                wiki_frame, text="🔗 Wikipédia", width=110, height=26,
                fg_color=UI.get("accent", "#3d9cf0"), state="disabled",
            )
            btn_wiki.pack(side="right", padx=10, pady=10)

            def charger_wiki(espece_visee=espece):
                info = self._fetch_wikipedia_info(espece_visee)

                def appliquer():
                    # L'utilisateur a changé d'espèce pendant la requête réseau : résultat obsolète, on ignore.
                    if choix_espece.get() != espece_visee:
                        return
                    if info is None:
                        lbl_wiki_extract.configure(text="⚠️ Wikipédia inaccessible (pas de connexion internet ?).")
                        return
                    if info.get("not_found"):
                        lbl_wiki_extract.configure(text="Aucune fiche Wikipédia trouvée pour cette espèce.")
                        return
                    extrait = (info.get("extract") or "").strip()
                    if len(extrait) > 400:
                        extrait = extrait[:400].rsplit(" ", 1)[0] + "…"
                    lbl_wiki_extract.configure(text=extrait or "(résumé indisponible)")
                    url = info.get("url")
                    if url:
                        btn_wiki.configure(state="normal", command=lambda u=url: webbrowser.open(u))
                    thumb = info.get("thumbnail")
                    if thumb:
                        try:
                            img_data = requests.get(thumb, timeout=6).content
                            pil_img = Image.open(io.BytesIO(img_data)).convert("RGB")
                            pil_img.thumbnail((70, 70))
                            ctk_img = ctk.CTkImage(light_image=pil_img, dark_image=pil_img, size=pil_img.size)
                            lbl_wiki_photo.configure(image=ctk_img, text="")
                            lbl_wiki_photo.image = ctk_img  # référence conservée (sinon Tk libère l'image)
                        except Exception:
                            pass

                self.after(0, appliquer)

            threading.Thread(target=charger_wiki, daemon=True).start()

            stats = ctk.CTkFrame(left, fg_color=UI.get("card_alt", "#1e2a3a"), corner_radius=10)
            stats.pack(fill="x", padx=12, pady=6)
            for titre, val in (
                ("Observations", str(len(obs_espece))),
                ("Individus ~", str(total_individus)),
                ("Sorties", str(len(dossiers))),
                ("Avec GPS", str(sum(1 for o in obs_espece if o.get("lat") and o.get("lon")))),
            ):
                cell = ctk.CTkFrame(stats, fg_color="transparent")
                cell.pack(side="left", expand=True, padx=6, pady=10)
                ctk.CTkLabel(cell, text=val, font=ctk.CTkFont(size=16, weight="bold")).pack()
                ctk.CTkLabel(cell, text=titre, font=ctk.CTkFont(size=10), text_color=UI.get("text_dim", "#888")).pack()

            if lieux:
                ctk.CTkLabel(left, text="Lieux fréquentés", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=12, pady=(12, 4))
                for lieu, n in lieux.most_common(8):
                    ctk.CTkLabel(left, text=f"• {lieu}  ({n})", font=ctk.CTkFont(size=12), text_color=UI.get("text_dim", "#aaa")).pack(
                        anchor="w", padx=18, pady=1
                    )

            ctk.CTkLabel(left, text="Journal des observations", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=12, pady=(14, 4))
            for o in sorted(obs_espece, key=lambda x: (str(x.get("dossier", "")), x.get("heure") or "")):
                card = ctk.CTkFrame(left, fg_color=UI.get("card_alt", "#1e2a3a"), corner_radius=8)
                card.pack(fill="x", padx=12, pady=3)
                doss = o.get("dossier") or "?"
                ligne1 = f"{doss}  ·  {o.get('heure') or '--:--'}  ·  ×{o.get('nombre') or '?'}"
                if o.get("sans_photo"):
                    ligne1 += "  ·  sans photo"
                ctk.CTkLabel(card, text=ligne1, font=ctk.CTkFont(size=11), anchor="w").pack(anchor="w", padx=10, pady=(6, 0), fill="x")
                if o.get("lieu"):
                    ctk.CTkLabel(card, text=f"📍 {o.get('lieu')}", font=ctk.CTkFont(size=10), text_color=UI.get("text_dim", "#888"), anchor="w").pack(
                        anchor="w", padx=10, pady=(0, 6), fill="x"
                    )
                elif o.get("lat") and o.get("lon"):
                    ctk.CTkLabel(
                        card, text=f"📍 {float(o['lat']):.5f}, {float(o['lon']):.5f}",
                        font=ctk.CTkFont(size=10), text_color=UI.get("text_dim", "#888"), anchor="w"
                    ).pack(anchor="w", padx=10, pady=(0, 6), fill="x")

            # Carte : tous les points GPS de l'espèce
            coords = []
            for o in obs_espece:
                lat, lon = o.get("lat"), o.get("lon")
                if not (lat and lon):
                    continue
                try:
                    lat, lon = float(lat), float(lon)
                except Exception:
                    continue
                forme = self._detect_marker_shape(o.get("espece") or "")
                type_obs = (o.get("type_observation") or "").lower()
                if "empreinte" in type_obs or "trace" in type_obs:
                    forme = "triangle"
                elif "terrier" in type_obs:
                    forme = "square"
                elif "coulée" in type_obs or "coulee" in type_obs or "passage" in type_obs:
                    forme = "diamond"
                icon = self._get_marker_icon(
                    o.get("categorie") or cat, 1, shape=forme,
                    sans_photo=bool(o.get("sans_photo")),
                )
                mk = map_w.set_marker(lat, lon, text=None, icon=icon)
                map_markers.append(mk)
                coords.append((lat, lon))

            if coords:
                lats = [c[0] for c in coords]
                lons = [c[1] for c in coords]
                if len(coords) == 1:
                    map_w.set_position(*coords[0])
                    map_w.set_zoom(14)
                else:
                    try:
                        map_w.fit_bounding_box((max(lats), min(lons)), (min(lats), max(lons)))
                    except Exception:
                        map_w.set_position(sum(lats) / len(lats), sum(lons) / len(lons))
                        map_w.set_zoom(12)
            else:
                ctk.CTkLabel(
                    left,
                    text="Aucun point GPS pour cette espèce (géotaggez ou enregistrez les fiches).",
                    font=ctk.CTkFont(size=11), text_color=UI.get("warning", "#e67e22"), wraplength=360, justify="left",
                ).pack(anchor="w", padx=12, pady=10)

        choix_espece.configure(command=render)
        render()

    def select_photo_by_filename(self, filename):
        """Sélectionne une photo dans la liste principale (utilisé par les clics sur la carte)."""
        items = list(self.photo_listbox.get(0, tk.END))
        if filename not in items:
            self.entry_search_photo.delete(0, tk.END)
            self.filter_photo_list()
            items = list(self.photo_listbox.get(0, tk.END))
        if filename in items:
            idx = items.index(filename)
            self.photo_listbox.selection_clear(0, tk.END)
            self.photo_listbox.selection_set(idx)
            self.photo_listbox.activate(idx)
            self.photo_listbox.see(idx)
            self.on_photo_select(None)

    def _persist_gps_into_notes(self):
        """Écrit lat/lon du cache photos_data dans observations.json (pour carte cumulée multi-sorties)."""
        if not self.photo_folder_path:
            return
        notes = self._load_notes_dict()
        changed = False
        for fn, info in (self.photos_data or {}).items():
            if str(fn).startswith("_"):
                continue
            lat, lon = info.get("lat"), info.get("lon")
            if lat is None or lon is None:
                continue
            entry = notes.get(fn)
            if not isinstance(entry, dict):
                # Cree une fiche minimale pour conserver le GPS entre versions / dossiers
                entry = {
                    "espece": "",
                    "categorie": "",
                    "nombre": "1",
                    "notes_libres": "",
                }
            if entry.get("lat") != lat or entry.get("lon") != lon:
                entry["lat"] = lat
                entry["lon"] = lon
                notes[fn] = entry
                changed = True
        if changed:
            try:
                self._save_notes_dict(notes, silent=True)
            except Exception:
                pass

    def _check_existing_gps_worker(self, exe_path):
        """Lit les GPS existants en arrière-plan (ExifTool optimisé + progression)."""
        try:
            self._set_progress(0.05, "Lecture GPS existants…")
            images = self.get_supported_images()
            if not images:
                self._set_progress(0, "En attente…")
                return
            dossier = os.path.normpath(self.photo_folder_path)
            chemins = [os.path.join(dossier, img) for img in images]
            tags = ["GPSLatitude", "GPSLongitude", "DateTimeOriginal"]
            try:
                metadata = self._exiftool_read_json_chunked(
                    exe_path, chemins, tags,
                    chunk_size=50, timeout_per_chunk=75,
                    progress_label="Lecture GPS"
                )
            except subprocess.TimeoutExpired:
                self.log("⚠️ Lecture GPS interrompue (délai dépassé). Utilisez « Rafraîchir » sur la carte.")
                self._set_progress(0, "Prêt")
                return

            if not metadata:
                self.log("⚠️ Vérification GPS : ExifTool n'a rien renvoyé.")
                self._set_progress(0, "Prêt")
                return

            temp_photos_data = {}
            with_gps = 0
            for item in metadata:
                full_path = item.get("SourceFile", "").replace("/", os.sep)
                filename = os.path.basename(full_path)
                f_lat = self.parse_coord(item.get("GPSLatitude"))
                f_lon = self.parse_coord(item.get("GPSLongitude"))
                temp_photos_data[filename] = {
                    "path": full_path if full_path else os.path.join(dossier, filename),
                    "lat": f_lat, "lon": f_lon,
                    "date": item.get("DateTimeOriginal", "") or "",
                }
                if f_lat and f_lon:
                    with_gps += 1

            total = len(metadata)
            self._set_progress(1.0, f"GPS lus : {with_gps}/{total}")
            self.after(0, lambda: self._apply_existing_gps(temp_photos_data, with_gps, total))
            self.after(1500, lambda: self._set_progress(0, "Prêt"))
        except Exception as e:
            self.log(f"⚠️ Vérification GPS impossible : {e}")
            self._set_progress(0, "Prêt")

    def _apply_existing_gps(self, photos_data, with_gps, total):
        if with_gps == 0:
            self.log("ℹ️ Aucune balise GPS dans les fichiers. Les points déjà en carnet (si présents) restent affichés. Synchro GPX si besoin.")
            try:
                self.refresh_map_markers()
            except Exception:
                pass
            return

        # Fusion : ne pas ecraser un GPS carnet par un None ExifTool
        for fn, info in (photos_data or {}).items():
            if not isinstance(info, dict):
                continue
            lat, lon = info.get("lat"), info.get("lon")
            if lat is None or lon is None:
                # conserve l'entree existante (notes) si elle a deja un GPS
                continue
            if fn not in self.photos_data:
                self.photos_data[fn] = {}
            self.photos_data[fn].update({
                "path": info.get("path") or self.photos_data[fn].get("path"),
                "lat": lat,
                "lon": lon,
                "date": info.get("date") or self.photos_data[fn].get("date") or "",
            })
        self._persist_gps_into_notes()
        self.refresh_map_markers()
        self.refresh_daily_counter()
        self._refresh_listbox_annotation_status()

        self.is_synced = True  # débloque la carte et le carnet sans passer par l'écriture ExifTool

        if with_gps == total:
            self.log(f"✅ {with_gps}/{total} photos possèdent déjà des coordonnées GPS. Synchronisation GPX non nécessaire — carte et carnet disponibles immédiatement.")
            self.btn_sync.configure(text="⚡ Re-synchroniser (optionnel)")
            self.lbl_progress.configure(text="Prêt (GPS déjà présent)")
        else:
            self.log(f"📍 {with_gps}/{total} photos ont déjà des coordonnées GPS. Les {total - with_gps} restantes nécessitent une synchronisation GPX.")

    def start_sync_thread(self):
        if not self.photo_folder_path:
            messagebox.showwarning("Dossier manquant", "Sélectionnez un dossier de photos d'abord.")
            return

        exe_path = self.get_exiftool_path()
        if not exe_path:
            self.log("❌ ERREUR : 'exiftool.exe' introuvable.")
            messagebox.showerror("ExifTool Introuvable", "Posez 'exiftool.exe' dans le dossier du script.")
            return

        images = self.get_supported_images()
        forcer = self.force_resync.get()
        all_already_geotagged = bool(images) and all(
            self.photos_data.get(img, {}).get("lat") and self.photos_data.get(img, {}).get("lon")
            for img in images
        )

        if not self.gpx_file_path and (forcer or not all_already_geotagged):
            messagebox.showwarning(
                "Fichiers manquants",
                "Sélectionnez un fichier GPX, ou utilisez un dossier dont les photos ont déjà des coordonnées GPS."
            )
            return

        self.is_processing = True
        self.progress_bar.set(0.0)

        if not forcer and not self.gpx_file_path and all_already_geotagged:
            # Toutes les photos ont déjà du GPS : pas besoin d'écrire quoi que ce soit, on relit juste les données.
            self.btn_sync.configure(state="disabled", text="⚡ Lecture...")
            self.lbl_progress.configure(text="Toutes les photos ont déjà des coordonnées GPS, lecture en cours...")
            threading.Thread(target=self.extract_coordinates_after_sync, args=(exe_path,), daemon=True).start()
        else:
            self.btn_sync.configure(state="disabled", text="⚡ Écriture GPS...")
            self.lbl_progress.configure(text="ExifTool modifie les fichiers..." + (" (réécriture forcée)" if forcer else ""))
            threading.Thread(target=self.execute_geotagging, args=(exe_path,), daemon=True).start()

    def apply_geosync_preset(self, offset, label=""):
        """Remplit le décalage photo→GPX (ex. été -1h, hiver -2h)."""
        self.entry_geosync.delete(0, tk.END)
        self.entry_geosync.insert(0, offset)
        msg = f"Décalage {label}: {offset}" if label else f"Décalage: {offset}"
        self.log(msg + " — lancez la synchronisation GPS.")
        try:
            self.lbl_progress.configure(text=f"Décalage {offset} prêt")
        except Exception:
            pass

    def show_geosync_help(self):
        messagebox.showinfo(
            "Décalage horaire photo → GPX",
            "Le GPX est en UTC, l'appareil photo en heure locale (France).\n\n"
            "Boutons rapides (votre calage) :\n"
            "  ☀ Été   −1:00:00\n"
            "  ❄ Hiver −2:00:00\n\n"
            "Vous pouvez aussi saisir un décalage libre (ex. -0:30:00).\n"
            "Puis lancez « Synchroniser GPS ».\n\n"
            "Cochez « Forcer la réécriture » si les photos ont déjà un GPS incorrect."
        )

    def _executer_lot_exiftool(self, commande, startupinfo, timeout_secondes, callback_progression=None):
        """Exécute une commande ExifTool en suivant la progression fichier par fichier en temps réel
        (option -progress d'ExifTool, format confirmé : '======== fichier [N/M]' sur stdout), avec un
        délai maximum de sécurité (watchdog) qui tue le processus s'il ne répond plus, plutôt que de
        bloquer indéfiniment le thread de synchronisation."""
        processus = subprocess.Popen(
            commande, startupinfo=startupinfo, stdout=subprocess.PIPE, stderr=subprocess.DEVNULL,
            text=True, encoding="utf-8", errors="ignore", bufsize=1
        )

        etat = {"tue": False}

        def tuer():
            etat["tue"] = True
            try:
                processus.kill()
            except Exception:
                pass

        minuteur = threading.Timer(timeout_secondes, tuer)
        minuteur.start()

        lignes = []
        try:
            while True:
                ligne = processus.stdout.readline()
                if ligne == "" and processus.poll() is not None:
                    break
                if not ligne:
                    continue
                lignes.append(ligne)
                if callback_progression:
                    m = re.search(r"\[(\d+)/(\d+)\]", ligne)
                    if m:
                        callback_progression(int(m.group(1)), int(m.group(2)))
        finally:
            minuteur.cancel()

        code_retour = processus.wait()
        if etat["tue"]:
            raise subprocess.TimeoutExpired(cmd=commande, timeout=timeout_secondes)
        return "".join(lignes), code_retour

    def execute_geotagging(self, exe_path):
        try:
            dossier_propre = os.path.normpath(self.photo_folder_path)
            gpx_propre = os.path.normpath(self.gpx_file_path)

            for file in os.listdir(dossier_propre):
                if file.endswith("_exiftool_tmp"):
                    try: os.remove(os.path.join(dossier_propre, file))
                    except: pass

            images = self.get_supported_images()
            total_images = len(images)
            startupinfo = self._exiftool_startupinfo()

            forcer = self.force_resync.get()
            decalage = self.entry_geosync.get().strip()

            a_ecrire, deja_ok = [], []
            for img in images:
                existing = self.photos_data.get(img)
                if not forcer and existing and existing.get("lat") and existing.get("lon"):
                    deja_ok.append(img)
                else:
                    a_ecrire.append(img)

            skipped_count = len(deja_ok)
            if skipped_count:
                self.log(f"⏭️ {skipped_count} photo(s) ignorée(s) car déjà géolocalisée(s).")

            self.log(f"Écriture GPS par lots : {len(a_ecrire)} photo(s) à traiter sur {total_images}…")
            self._set_progress(0.02, f"Préparation : {len(a_ecrire)} fichier(s)…")

            sans_backup = self.no_backup.get() if hasattr(self, "no_backup") else False

            backup_root = os.path.join(dossier_propre, ".geoexif_backups")
            backup_dir = None
            backed_up = 0
            if a_ecrire:
                shutil.rmtree(backup_root, ignore_errors=True)
                if not sans_backup:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    backup_dir = os.path.join(backup_root, timestamp)
                    os.makedirs(backup_dir, exist_ok=True)
                    self._set_progress(0.05, "Sauvegarde des originaux…")

                    def _copy_one(img):
                        try:
                            shutil.copy2(
                                os.path.join(dossier_propre, img),
                                os.path.join(backup_dir, img),
                            )
                            return 1
                        except Exception:
                            return 0

                    # Copies parallèles (I/O bound)
                    with ThreadPoolExecutor(max_workers=min(CPU_WORKERS, 6)) as pool:
                        backed_up = sum(pool.map(_copy_one, a_ecrire))
                    self._set_progress(0.12, f"Sauvegarde : {backed_up} fichier(s)")

            # Écriture GPS par LOTS : un seul appel ExifTool pour plusieurs fichiers, au lieu d'un appel par photo.
            # C'est le principal gain de vitesse : le coût de démarrage d'ExifTool ne se paie plus qu'une fois par lot.
            # -P (preserve) : empêche ExifTool de mettre à jour la date de modification du fichier lors de l'écriture,
            # pour que la date de prise de vue affichée dans l'explorateur/Lightroom ne change pas.
            # Lot volontairement plus petit (15) avec un délai généreux et proportionnel : sur de gros fichiers RAW
            # (Z5 II, ~45 Mo/photo) ou un disque lent, chaque lot reste gérable, et si UN lot échoue ou dépasse le
            # délai, les autres lots continuent quand même — on ne perd plus toute une sortie pour un seul incident.
            # Lots plus grands sur machines multi-cœurs (ExifTool démarre moins souvent)
            # Lots plus grands = moins de démarrages ExifTool ; timeout proportionnel à la taille du lot
            TAILLE_LOT = 25 if CPU_WORKERS >= 4 else 15
            TIMEOUT_PAR_LOT = 90 + TAILLE_LOT * 8  # ~8 s/photo RAW, min 90 s
            success_count = skipped_count
            lots_en_echec = []
            n_lots = max(1, (len(a_ecrire) + TAILLE_LOT - 1) // TAILLE_LOT)

            for i in range(0, len(a_ecrire), TAILLE_LOT):
                lot = a_ecrire[i:i + TAILLE_LOT]
                chemins = [os.path.join(dossier_propre, img) for img in lot]
                lot_num = i // TAILLE_LOT + 1

                # -P conserve la date fichier ; -progress pour la barre ; pas de -fast en écriture géotag
                commande = [
                    exe_path, f"-geotag={gpx_propre}", "-P", "-progress",
                    "-api", "largefilesupport=1",
                ]
                if decalage:
                    commande.append(f"-geosync={decalage}")
                commande += ["-overwrite_original"] + chemins

                def maj_progression(n_dans_lot, total_dans_lot, deja_avant=i, deja_saute=skipped_count, ln=lot_num, nl=n_lots):
                    fait_global = deja_avant + n_dans_lot + deja_saute
                    progression = 0.12 + 0.75 * (fait_global / total_images) if total_images else 1.0
                    self._set_progress(
                        progression,
                        f"Écriture GPS lot {ln}/{nl} — photo {fait_global}/{total_images}"
                    )

                try:
                    sortie, code_retour = self._executer_lot_exiftool(commande, startupinfo, TIMEOUT_PAR_LOT, maj_progression)
                    maj = 0
                    for ligne in sortie.splitlines():
                        if "image files updated" in ligne:
                            try:
                                maj = int(ligne.strip().split()[0])
                            except Exception:
                                maj = 0
                    if maj == 0 and code_retour != 0:
                        lots_en_echec.append((i // TAILLE_LOT + 1, lot))
                        self.log(f"⚠️ Lot {i // TAILLE_LOT + 1} : ExifTool a retourné une erreur ({len(lot)} photo(s) concernée(s), les autres lots continuent).")
                    success_count += maj if maj else (len(lot) if code_retour == 0 else 0)
                except subprocess.TimeoutExpired:
                    lots_en_echec.append((i // TAILLE_LOT + 1, lot))
                    self.log(f"⚠️ Lot {i // TAILLE_LOT + 1} : délai dépassé ({TIMEOUT_PAR_LOT}s, {len(lot)} photo(s) concernée(s)) — les autres lots continuent. Certaines de ces photos ont peut-être quand même été écrites avant l'interruption ; une resynchronisation les détectera.")
                except Exception as e:
                    lots_en_echec.append((i // TAILLE_LOT + 1, lot))
                    self.log(f"⚠️ Lot {i // TAILLE_LOT + 1} : erreur inattendue ({e}) — les autres lots continuent.")

                fait = min(i + TAILLE_LOT, len(a_ecrire))
                progression = (fait + skipped_count) / total_images if total_images else 1.0
                self.after(0, lambda p=progression: self.progress_bar.set(p))
                self.after(0, lambda f=fait, t=len(a_ecrire): self.lbl_progress.configure(text=f"Écriture GPS par lots : {f}/{t} photos"))

            if lots_en_echec:
                total_en_echec = sum(len(lot) for _, lot in lots_en_echec)
                self.log(f"⚠️ {len(lots_en_echec)} lot(s) sur {total_en_echec} photo(s) au total ont rencontré un problème. Relancez simplement la synchronisation : les photos déjà écrites seront ignorées, seules celles en échec seront retentées.")

            if backed_up:
                self.last_backup_dir = backup_dir
                self.after(0, lambda: self.btn_undo_sync.configure(state="normal"))
                self.log(f"💾 {backed_up} photo(s) sauvegardée(s) avant écriture (annulation possible, ancienne sauvegarde remplacée).")
            elif sans_backup and a_ecrire:
                self.last_backup_dir = None
                self.after(0, lambda: self.btn_undo_sync.configure(state="disabled"))
                self.log("⚠️ Écriture directe sans sauvegarde (option activée) : annulation impossible pour cette synchro.")

            self.log(f"ExifTool : Écriture terminée avec succès pour {success_count}/{total_images} fichiers.")
            
            self.is_synced = True
            self.after(0, lambda: self.lbl_progress.configure(text="Lecture des coordonnées injectées..."))
            self.extract_coordinates_after_sync(exe_path)

        except Exception as e:
            self.is_processing = False
            self.log(f"❌ Erreur système lors de la synchronisation : {str(e)}")
            self.after(0, lambda: self.progress_bar.set(0))
            self.after(0, lambda: self.lbl_progress.configure(text="Prêt"))
            self.after(0, lambda: self.btn_sync.configure(state="normal", text="⚡ Synchroniser Photos & GPX"))

    @staticmethod
    def parse_coord(val):
        if not val: return None
        if isinstance(val, (int, float)): return float(val)
        try:
            if "deg" in str(val):
                parts = str(val).split("deg")
                deg = float(parts[0])
                min_sec = parts[1].strip().split("'")
                minutes = float(min_sec[0])
                seconds = float(min_sec[1].replace('"', '').split()[0])
                sign = -1 if any(x in str(val) for x in ['S', 'W', 'O']) else 1
                return sign * (deg + minutes/60.0 + seconds/3600.0)
            return float(str(val).split()[0])
        except: return None

    def extract_coordinates_after_sync(self, exe_path):
        """Relit les GPS après écriture (lots parallèles + barre de progression)."""
        try:
            self._set_progress(0.05, "Relecture des coordonnées…")
            images = self.get_supported_images()
            dossier = os.path.normpath(self.photo_folder_path)
            chemins = [os.path.join(dossier, img) for img in images]
            tags = ["GPSLatitude", "GPSLongitude", "DateTimeOriginal"]
            metadata = self._exiftool_read_json_chunked(
                exe_path, chemins, tags,
                chunk_size=50, timeout_per_chunk=75,
                progress_label="Relecture GPS"
            )

            coordinates_found = []
            if metadata:
                for item in metadata:
                    full_path = item.get("SourceFile", "").replace("/", os.sep)
                    filename = os.path.basename(full_path)
                    f_lat = self.parse_coord(item.get("GPSLatitude"))
                    f_lon = self.parse_coord(item.get("GPSLongitude"))
                    self.photos_data[filename] = {
                        "path": full_path if full_path else os.path.join(dossier, filename),
                        "lat": f_lat, "lon": f_lon,
                        "date": item.get("DateTimeOriginal", "") or "",
                    }
                    if f_lat and f_lon:
                        coordinates_found.append((f_lat, f_lon))

            # Réhydrate observations manuelles
            notes = self._load_notes_dict()
            for cle, data in notes.items():
                if isinstance(data, dict) and data.get("sans_photo") and data.get("lat") and data.get("lon"):
                    self.photos_data[cle] = {
                        "lat": data["lat"], "lon": data["lon"], "date": "",
                    }

            def finish():
                self._persist_gps_into_notes()
                self.refresh_map_markers()
                self.refresh_daily_counter()
                self._refresh_listbox_annotation_status()
                self.is_processing = False
                self.btn_sync.configure(state="normal", text="⚡  Synchroniser photos & GPX")
                n = len(coordinates_found)
                self.lbl_progress.configure(text=f"Terminé — {n} point(s) GPS")
                self.progress_bar.set(1.0)
                self.log(f"✅ {n} photo(s) avec coordonnées GPS chargée(s).")
                if n:
                    self.is_synced = True
                self.after(2000, lambda: self._set_progress(0, "Prêt"))

            self.after(0, finish)
        except subprocess.TimeoutExpired:
            self.is_processing = False
            self.log("⚠️ Relecture GPS : délai dépassé.")
            self.after(0, lambda: self.btn_sync.configure(state="normal", text="⚡  Synchroniser photos & GPX"))
            self._set_progress(0, "Prêt")
        except Exception as e:
            self.is_processing = False
            self.log(f"⚠️ Relecture GPS impossible : {e}")
            self.after(0, lambda: self.btn_sync.configure(state="normal", text="⚡  Synchroniser photos & GPX"))
            self._set_progress(0, "Prêt")

    def set_carnet_view_mode(self, mode):
        """Bascule Liste (léger) / Miniatures (lazy, sans saturater le CPU)."""
        mode = "thumbs" if mode == "thumbs" else "list"
        self.carnet_view_mode = mode
        if mode == "list":
            self.btn_view_list.configure(fg_color=UI["accent"], hover_color=UI["accent_hover"])
            self.btn_view_thumbs.configure(fg_color=UI["card_alt"], hover_color=UI["border"])
            try:
                self.thumb_scroll.grid_forget()
            except Exception:
                pass
            self.photo_listbox.grid(row=1, column=0, sticky="nsew", padx=6, pady=(2, 6))
            self._populate_listbox_from_filtered()
        else:
            self.btn_view_thumbs.configure(fg_color=UI["accent"], hover_color=UI["accent_hover"])
            self.btn_view_list.configure(fg_color=UI["card_alt"], hover_color=UI["border"])
            self.photo_listbox.grid_forget()
            self.thumb_scroll.grid(row=1, column=0, sticky="nsew", padx=6, pady=(2, 6))
            self._rebuild_thumb_explorer(lazy=True)

    def _is_virtual_observation(self, key, data=None):
        """Observation sans fichier média : Birda, Chirpity, manuel, terrain, indices…"""
        if data is None:
            data = {}
        if not isinstance(data, dict):
            data = {}
        k = str(key or "")
        if k.startswith("_"):
            return True
        if data.get("sans_photo"):
            return True
        src = (data.get("source") or "").lower()
        if src in ("birda", "chirpity", "field_notebook", "manuel", "audio"):
            return True
        typ = (data.get("type_observation") or data.get("type_indice") or "").lower()
        if "prise de son" in typ or "birda" in typ or "chirpity" in typ:
            return True
        return False

    def _virtual_note_keys(self):
        """Clés d'observations sans fichier média (Birda, Chirpity, manuel…)."""
        if not self.photo_folder_path:
            return []
        try:
            notes = self._load_notes_dict(self.photo_folder_path) or {}
        except Exception:
            notes = {}
        keys = []
        for k, data in notes.items():
            if not isinstance(data, dict):
                continue
            if not self._is_virtual_observation(k, data):
                continue
            if self.photo_folder_path and os.path.isfile(os.path.join(self.photo_folder_path, k)):
                continue
            keys.append(k)
        keys.sort()
        return keys

    def _display_name_for_list(self, key):
        """Libellé liste carnet : média réel ou obs. virtuelle (Birda / sans photo)."""
        if self._is_video_file(key):
            return "🎬 " + key
        if str(key).startswith("_birda") or (isinstance(key, str) and "birda" in key.lower()):
            notes = {}
            try:
                notes = self._load_notes_dict(self.photo_folder_path) or {}
            except Exception:
                pass
            data = notes.get(key) or {}
            esp = (data.get("espece") or "Prise de son").strip()
            conf = data.get("birda_confidence")
            conf_s = " %.0f%%" % (float(conf) * 100) if conf is not None else ""
            try:
                if conf is not None and float(conf) > 1:
                    conf_s = " %.0f%%" % float(conf)
            except Exception:
                pass
            return "🎧 %s%s" % (esp, conf_s)
        if str(key).startswith("_") or not os.path.isfile(os.path.join(self.photo_folder_path or "", key)):
            notes = {}
            try:
                notes = self._load_notes_dict(self.photo_folder_path) or {}
            except Exception:
                pass
            data = notes.get(key) or {}
            esp = (data.get("espece") or "Sans photo").strip()
            typ = (data.get("type_indice") or data.get("type_observation") or "").strip()
            if typ:
                return "📍 %s (%s)" % (esp, typ)
            return "📍 %s" % esp
        return key

    def _current_media_filenames(self):
        """Liste filtrée : photos + vidéos + observations virtuelles (Birda, sans photo)."""
        if getattr(self, "_filtered_media_names", None):
            base = list(self._filtered_media_names)
        else:
            base = self.get_supported_media() if self.photo_folder_path else []
        virtual = self._virtual_note_keys()
        # virtual at end, avoid duplicates
        seen = set(base)
        for v in virtual:
            if v not in seen:
                base.append(v)
                seen.add(v)
        return base

    def _populate_listbox_from_filtered(self):
        names = self._current_media_filenames()
        self.photo_listbox.delete(0, tk.END)
        for name in names:
            self.photo_listbox.insert(tk.END, self._display_name_for_list(name))
        if hasattr(self, "lbl_explorer_count"):
            n_vid = sum(1 for n in names if self._is_video_file(n))
            n_virt = sum(1 for n in names if str(n).startswith("_") or (
                self.photo_folder_path and not os.path.isfile(os.path.join(self.photo_folder_path, n))
                and not self._is_video_file(n)
            ))
            n_ph = max(0, len(names) - n_vid - n_virt)
            parts = []
            if n_ph:
                parts.append("%d photo(s)" % n_ph)
            if n_vid:
                parts.append("%d vidéo(s)" % n_vid)
            if n_virt:
                parts.append("%d sans média (Birda / manuel)" % n_virt)
            self.lbl_explorer_count.configure(text=" · ".join(parts) if parts else "0 élément")
        self._refresh_listbox_annotation_status()

    def _rebuild_thumb_explorer(self, lazy=True, batch_size=36):
        """Grille de miniatures : placeholders immédiats, chargement par lots en arrière-plan."""
        for w in self.thumb_scroll.winfo_children():
            w.destroy()
        self._thumb_widgets.clear()
        # ne pas vider tout le cache disque/RAM des thumbs déjà calculés
        names = self._current_media_filenames()
        if hasattr(self, "lbl_explorer_count"):
            self.lbl_explorer_count.configure(text=f"{len(names)} fichier(s) — miniatures")

        if not names:
            ctk.CTkLabel(
                self.thumb_scroll, text="Aucun média",
                text_color=UI["text_muted"]
            ).pack(pady=20)
            return

        cols = 3
        self._thumb_load_token += 1
        token = self._thumb_load_token

        # Zone grille
        grid = ctk.CTkFrame(self.thumb_scroll, fg_color="transparent")
        grid.pack(fill="both", expand=True)
        for c in range(cols):
            grid.grid_columnconfigure(c, weight=1)

        to_load = []
        for i, name in enumerate(names):
            cell = ctk.CTkFrame(grid, fg_color=UI["card"], corner_radius=8, width=100, height=110)
            cell.grid(row=i // cols, column=i % cols, padx=4, pady=4, sticky="nsew")
            cell.grid_propagate(False)

            is_vid = self._is_video_file(name)
            placeholder = "🎬" if is_vid else "…"
            btn = ctk.CTkButton(
                cell, text=placeholder, width=90, height=72,
                fg_color=UI["card_alt"], hover_color=UI["border"],
                font=ctk.CTkFont(size=18),
                command=lambda n=name: self._on_thumb_click(n),
            )
            btn.pack(padx=4, pady=(4, 0))
            # Double-clic (widget tk sous-jacent) → ouvrir en grand
            try:
                btn.bind("<Double-Button-1>", lambda e, n=name: self._on_thumb_double_click(n))
                cell.bind("<Double-Button-1>", lambda e, n=name: self._on_thumb_double_click(n))
            except Exception:
                pass
            short = name if len(name) <= 14 else name[:11] + "…"
            lbl = ctk.CTkLabel(cell, text=short, font=ctk.CTkFont(size=9), text_color=UI["text_dim"])
            lbl.pack(pady=(2, 4))
            self._thumb_widgets[name] = btn
            if not is_vid:
                to_load.append(name)

        # Bouton charger plus si beaucoup de fichiers (affiche déjà tous les placeholders)
        if len(names) > batch_size:
            ctk.CTkLabel(
                self.thumb_scroll,
                text=f"Miniatures chargées progressivement ({min(batch_size, len(to_load))} d'abord)…",
                font=ctk.CTkFont(size=10), text_color=UI["text_muted"],
            ).pack(pady=(4, 8))

        # Charge seulement un premier lot pour ne pas saturer
        first_batch = to_load[:batch_size] if lazy else to_load
        rest = to_load[batch_size:] if lazy else []

        def load_batch(batch, then_rest=None):
            exe = self.get_exiftool_path()

            def work():
                for name in batch:
                    if token != self._thumb_load_token:
                        return
                    path = os.path.join(self.photo_folder_path, name)
                    try:
                        data = self._get_thumbnail_bytes(path, exe, max_dim=120)
                        if not data:
                            continue
                        img = Image.open(io.BytesIO(data)).convert("RGB")
                        img.thumbnail((88, 66))
                        # CTkImage thread-safe via after
                        self.after(0, lambda n=name, im=img: self._apply_thumb_image(n, im, token))
                    except Exception:
                        continue
                if then_rest and token == self._thumb_load_token:
                    # enchaîne le reste par paquets de 24, pause légère
                    nxt = then_rest[:24]
                    more = then_rest[24:]
                    if nxt:
                        self.after(200, lambda: threading.Thread(
                            target=lambda: load_batch(nxt, more), daemon=True
                        ).start())

            threading.Thread(target=work, daemon=True).start()

        load_batch(first_batch, rest)

    def _apply_thumb_image(self, filename, pil_img, token):
        if token != getattr(self, "_thumb_load_token", 0):
            return
        btn = self._thumb_widgets.get(filename)
        if not btn:
            return
        try:
            ctk_img = ctk.CTkImage(light_image=pil_img, dark_image=pil_img, size=(pil_img.width, pil_img.height))
            self._thumb_photo_refs[filename] = ctk_img  # évite GC
            btn.configure(image=ctk_img, text="")
        except Exception:
            pass

    def _on_thumb_click(self, filename):
        """Clic miniature : sélectionne le média ; si vidéo → lecture immédiate."""
        if not filename or not self.photo_folder_path:
            return
        path = os.path.join(self.photo_folder_path, filename)
        info = self.photos_data.get(filename)
        if not info:
            info = {
                "path": path, "lat": None, "lon": None, "date": None,
                "is_video": self._is_video_file(filename),
            }
            self.photos_data[filename] = info
        self.selected_photo_path = path

        # Alimente la listbox si elle contient encore les entrées (mode mixte / outils)
        try:
            display = f"🎬 {filename}" if self._is_video_file(filename) else filename
            self.photo_listbox.selection_clear(0, tk.END)
            for i in range(self.photo_listbox.size()):
                if self._listbox_filename(self.photo_listbox.get(i)) == filename:
                    self.photo_listbox.selection_set(i)
                    self.photo_listbox.see(i)
                    break
            else:
                # listbox vide (mode miniatures) : injecte une entrée pour les autres outils
                if self.photo_listbox.size() == 0:
                    self.photo_listbox.insert(tk.END, display)
                    self.photo_listbox.selection_set(0)
        except Exception:
            pass

        # Charge la fiche (aperçu / formulaire) sans dépendre de la listbox
        self.on_photo_select(filename=filename)

        # Vidéo : un clic = lecture (lecteur système)
        if self._is_video_file(filename):
            self.open_video_external(path)

    def _on_thumb_double_click(self, filename):
        """Double-clic miniature : aperçu interne (photo) ou lecteur (vidéo)."""
        if not filename or not self.photo_folder_path:
            return
        path = os.path.join(self.photo_folder_path, filename)
        if os.path.isfile(path):
            self.open_media_large(path)

    def _listbox_filename(self, display_name):
        name = (display_name or "").strip()
        if name.startswith("🎬 "):
            name = name[2:].strip()
        return name

    def _clear_preview_frame(self, message="Chargement…"):
        """Réinitialise le cadre aperçu de façon fiable (CTkImage ne supporte pas bien image=None)."""
        try:
            blank = Image.new("RGB", (12, 12), (34, 46, 40))
            blank_ctk = ctk.CTkImage(light_image=blank, dark_image=blank, size=(12, 12))
            self._preview_blank_ref = blank_ctk
            self.lbl_preview.configure(
                image=blank_ctk, text=message,
                width=260, height=195,
                text_color=UI.get("text_dim", "#9db0a6"),
            )
            self.lbl_preview.image = blank_ctk
        except Exception:
            try:
                self.lbl_preview.configure(text=message)
            except Exception:
                pass

    def _update_preview_display(self):
        """Affiche _current_full_preview dans le petit cadre (chaque sélection doit refaire un CTkImage neuf)."""
        if self._current_full_preview is None:
            self._clear_preview_frame("Aperçu indisponible")
            return
        try:
            size = int(self.slider_thumb_size.get())
        except Exception:
            size = 260
        w, h = max(80, size), max(60, int(size * 0.75))
        try:
            src = self._current_full_preview
            if src.mode not in ("RGB", "RGBA"):
                src = src.convert("RGB")
            else:
                src = src.copy()
            src.thumbnail((w, h))
            # Nouvelle image à chaque fois (sinon 2e sélection reste figée)
            img_ctk = ctk.CTkImage(
                light_image=src, dark_image=src.copy(),
                size=(max(1, src.width), max(1, src.height)),
            )
            self._preview_ctk_ref = img_ctk
            self.lbl_preview.configure(
                image=img_ctk, text="",
                width=src.width, height=src.height,
            )
            self.lbl_preview.image = img_ctk
            try:
                self.lbl_preview.update_idletasks()
            except Exception:
                pass
        except Exception:
            self._clear_preview_frame("Aperçu erreur")

    def fetch_historical_weather(self, lat, lon, datetime_str, filename):
        weather_data = None
        try:
            parts = datetime_str.split(" ")
            date_iso = parts[0].replace(":", "-") 
            heure_xml = parts[1].split(":")[0]     
            
            url = f"https://archive-api.open-meteo.com/v1/archive?latitude={lat}&longitude={lon}&start_date={date_iso}&end_date={date_iso}&hourly=temperature_2m,relative_humidity_2m,weather_code"
            response = requests.get(url, timeout=5).json()
            
            if "hourly" in response:
                heure_index = int(heure_xml)
                temp = response["hourly"]["temperature_2m"][heure_index]
                hum = response["hourly"]["relative_humidity_2m"][heure_index]
                code = response["hourly"]["weather_code"][heure_index]
                
                c_meteo = "Dégagé"
                if code in [1, 2, 3]: c_meteo = "Nuageux"
                elif code in [45, 48]: c_meteo = "Brouillard"
                elif code in [51, 53, 55, 61, 63, 65]: c_meteo = "Pluvieux"
                elif code in [71, 73, 75]: c_meteo = "Neige"
                elif code in [95, 96, 99]: c_meteo = "Orageux"

                # Lot 1 : phase lunaire (calcul local, aucun appel réseau)
                date_obj = datetime.strptime(date_iso, "%Y-%m-%d")
                phase_lune = self.moon_phase(date_obj)

                # Lot 1 : cumul de pluie des 3 jours précédant l'observation
                pluie_3j = None
                try:
                    debut_prev = (date_obj - timedelta(days=3)).strftime("%Y-%m-%d")
                    fin_prev = (date_obj - timedelta(days=1)).strftime("%Y-%m-%d")
                    url_prev = (
                        f"https://archive-api.open-meteo.com/v1/archive?latitude={lat}&longitude={lon}"
                        f"&start_date={debut_prev}&end_date={fin_prev}&daily=precipitation_sum&timezone=auto"
                    )
                    resp_prev = requests.get(url_prev, timeout=5).json()
                    valeurs = [v for v in resp_prev.get("daily", {}).get("precipitation_sum", []) if v is not None]
                    if valeurs:
                        pluie_3j = round(sum(valeurs), 1)
                except Exception:
                    pluie_3j = None

                weather_data = {
                    "date": date_iso, "heure": f"{heure_xml}h",
                    "temperature": temp, "humidite": hum, "ciel": c_meteo,
                    "phase_lunaire": phase_lune, "pluie_3j_precedents": pluie_3j,
                }
                weather_report = (
                    f"🌤️ MÉTÉO OPEN-METEO\n----------------------------\n"
                    f"📅 Le : {date_iso} à {heure_xml}h\n🌡️ Température : {temp}°C\n"
                    f"💧 Humidité : {hum}%\n☁️ Ciel : {c_meteo}\n{phase_lune}"
                )
                if pluie_3j is not None:
                    weather_report += f"\n🌧️ Cumul pluie (3j avant) : {pluie_3j} mm"
            else:
                weather_report = "❌ Données indisponibles."
        except Exception:
            weather_report = "⚠️ Serveur OpenMeteo injoignable."

        self.after(0, self.update_weather_ui, weather_report, weather_data, filename)

    def fetch_location(self, lat, lon, filename):
        lieu = self.reverse_geocode(lat, lon)
        self.after(0, lambda: self.update_location_ui(lieu, filename))

    def on_category_change(self, choix_categorie):
        liste_especes = self.faune_meuse.get(choix_categorie, ["Autre"])
        self.choice_species.configure(values=liste_especes)
        self.choice_species.set(liste_especes[0])

    def _resolve_list_selection_to_key(self, display_or_key):
        """Convertit le libellé affiché dans la listbox vers la clé réelle (fichier ou _birda_…)."""
        if not display_or_key:
            return display_or_key
        s = str(display_or_key)
        if s.startswith("🎬 "):
            s = s[2:].strip()
        # Si c'est déjà une clé connue
        names = self._current_media_filenames()
        if s in names:
            return s
        if display_or_key in names:
            return display_or_key
        # Index sélection → clé
        try:
            sel = self.photo_listbox.curselection()
            if sel:
                idx = sel[0]
                if 0 <= idx < len(names):
                    return names[idx]
        except Exception:
            pass
        # Fallback préfixes
        for prefix in ("🎧 ", "📍 ", "🎬 "):
            if s.startswith(prefix):
                s = s[len(prefix):].strip()
        return s

    def on_photo_select(self, event=None, filename=None):
        if not filename:
            selection = self.photo_listbox.curselection()
            if not selection:
                return
            filename = self.photo_listbox.get(selection[0])
        filename = self._resolve_list_selection_to_key(filename)
        # Ferme la loupe survol si ouverte (évite conflit avec le cadre aperçu)
        try:
            self._hide_hover_loupe()
        except Exception:
            pass
        is_virtual = (
            str(filename).startswith("_")
            or not os.path.isfile(os.path.join(self.photo_folder_path or "", filename))
        )
        info = self.photos_data.get(filename)
        if not info:
            notes = {}
            try:
                notes = self._load_notes_dict(self.photo_folder_path) or {}
            except Exception:
                pass
            data = notes.get(filename) or {}
            info = {
                "path": "" if is_virtual else os.path.join(self.photo_folder_path, filename),
                "lat": data.get("lat"),
                "lon": data.get("lon"),
                "date": data.get("heure") or "",
                "is_video": False if is_virtual else self._is_video_file(filename),
                "sans_photo": bool(is_virtual or data.get("sans_photo")),
            }
            self.photos_data[filename] = info
        # mémoriser la clé active pour save_current_note
        self._current_note_key = filename
        # path vide pour virtuel, mais selected_photo_path non-None pour activer Enregistrer
        self.selected_photo_path = info.get("path") or ("virtual:" + str(filename))

        is_video = bool(info.get("is_video") or self._is_video_file(filename))

        if is_video:
            self.current_burst_files = []
            if hasattr(self, "burst_frame"):
                try:
                    self.burst_frame.pack_forget()
                except Exception:
                    pass
        else:
            self.current_burst_files = self._detect_burst(filename)
            if self.current_burst_files:
                self.lbl_burst_info.configure(text=f"📸 Rafale détectée : {len(self.current_burst_files)} photos")
                self.burst_frame.pack(fill="x", padx=5, pady=(10, 0), before=self.lbl_details_header)
            else:
                if hasattr(self, "burst_frame"):
                    self.burst_frame.pack_forget()

        if not getattr(self, "selected_photo_path", None):
            self.selected_photo_path = info.get("path") or ("virtual:" + str(filename))

        img_loaded = False
        ext = os.path.splitext(filename)[1].lower()

        if is_video:
            self._current_full_preview = None
            try:
                self.lbl_preview.configure(
                    image=None,
                    text=f"🎬 Vidéo\n{filename}\n\nLecture auto\nou bouton ▶",
                    width=260, height=195,
                )
                self.lbl_preview.image = None
            except Exception:
                pass
            if hasattr(self, "btn_play_video"):
                try:
                    self.btn_play_video.pack(side="left", fill="x", expand=True, padx=(2, 0))
                except Exception:
                    pass
            img_loaded = True
        else:
            if hasattr(self, "btn_play_video"):
                try:
                    self.btn_play_video.pack_forget()
                except Exception:
                    pass
            self._current_full_preview = None
            self._clear_preview_frame("Chargement…")

        if is_video and hasattr(self, "btn_play_video"):
            try:
                self.btn_play_video.pack(side="left", fill="x", expand=True, padx=(2, 0))
            except Exception:
                pass

        # Observation virtuelle (Birda / sans photo) : pas d'aperçu fichier
        if is_virtual:
            self._current_full_preview = None
            notes_v = {}
            try:
                notes_v = self._load_notes_dict(self.photo_folder_path) or {}
            except Exception:
                pass
            data_v = notes_v.get(filename) or {}
            esp_v = data_v.get("espece") or "Observation"
            typ_v = data_v.get("type_indice") or data_v.get("type_observation") or ""
            gps_v = ""
            if data_v.get("lat") is not None and data_v.get("lon") is not None:
                try:
                    gps_v = "\nGPS %.5f, %.5f" % (float(data_v["lat"]), float(data_v["lon"]))
                except Exception:
                    pass
            try:
                self.lbl_preview.configure(
                    image=None,
                    text="🎧 / 📍 Sans média\n%s\n%s%s\n\nModifiez la fiche\nà droite puis Enregistrer"
                    % (esp_v, typ_v, gps_v),
                    width=260, height=195,
                )
                self.lbl_preview.image = None
            except Exception:
                pass
            img_loaded = True

        # Aperçu photo : charge + affiche (jeton pour ignorer une réponse obsolète si clic rapide)
        if not is_video and not is_virtual:
            path = info.get("path") or os.path.join(self.photo_folder_path, filename)
            self._preview_load_token = getattr(self, "_preview_load_token", 0) + 1
            token = self._preview_load_token
            exe = self.get_exiftool_path()

            def _load_and_show(p=path, tok=token, e=ext):
                img = None
                try:
                    data = self._get_thumbnail_bytes(p, exe, max_dim=900)
                    if data and len(data) > 100:
                        img = Image.open(io.BytesIO(data)).convert("RGB")
                        img.load()
                    if img is None and e in (".jpg", ".jpeg", ".png", ".tif", ".tiff"):
                        img = Image.open(p).convert("RGB")
                        img.load()
                except Exception as err:
                    msg = str(err)
                    self.after(0, lambda: self.log(f"Aperçu : {msg}"))
                    img = None

                def apply():
                    if tok != getattr(self, "_preview_load_token", 0):
                        return  # une sélection plus récente a pris le relais
                    if img is not None:
                        self._current_full_preview = img
                        self._update_preview_display()
                    else:
                        self._current_full_preview = None
                        self._clear_preview_frame("Aperçu indisponible")

                self.after(0, apply)

            # petit délai 0 : laisse l'UI afficher « Chargement… » puis charge
            threading.Thread(target=_load_and_show, daemon=True).start()
            img_loaded = True  # géré en async

        if info["date"] and " " in info["date"]:
            heure_brute = info["date"].split(" ")[1]
            heure_propre = ":".join(heure_brute.split(":")[:2])
            self.entry_time.delete(0, tk.END)
            self.entry_time.insert(0, heure_propre)
        else:
            self.entry_time.delete(0, tk.END)

        self.choice_category.set("Mammifère")
        self.on_category_change("Mammifère")
        self.choice_count.set("1")
        self.note_text.delete("0.0", "end")
        self.entry_lieu.delete(0, tk.END)
        if hasattr(self, "choice_comportement"):
            self.choice_comportement.set("")
        if hasattr(self, "choice_certitude"):
            self.choice_certitude.set("")
        if hasattr(self, "choice_type_indice"):
            self.choice_type_indice.set("")
        if hasattr(self, "entry_nom_sci"):
            self.entry_nom_sci.delete(0, tk.END)
        if hasattr(self, "entry_amazon_photo"):
            self.entry_amazon_photo.delete(0, tk.END)
        if hasattr(self, "entry_amazon_album"):
            self.entry_amazon_album.delete(0, tk.END)
            try:
                alb = self._load_sortie_amazon_album(self.photo_folder_path)
                if alb:
                    self.entry_amazon_album.insert(0, alb)
            except Exception:
                pass

        lieu_deja_connu = False
        chemin_notes = os.path.join(self.photo_folder_path, NOTES_FILE)
        if os.path.exists(chemin_notes):
            try:
                with open(chemin_notes, "r", encoding="utf-8") as f:
                    all_notes = json.load(f)
                    data_photo = all_notes.get(filename)
                    if data_photo:
                        cat = data_photo.get("categorie") or "Oiseau"
                        self.choice_category.set(cat)
                        self.on_category_change(cat)
                        esp = data_photo.get("espece") or ""
                        # Autoriser un nom hors liste (latin Birda, etc.)
                        try:
                            vals = list(self.choice_species.cget("values") or [])
                            if esp and esp not in vals:
                                vals = [esp] + vals
                                self.choice_species.configure(values=vals)
                        except Exception:
                            pass
                        if esp:
                            self.choice_species.set(esp)
                        self.choice_count.set(data_photo.get("nombre", "1"))
                        if data_photo.get("heure"):
                            self.entry_time.delete(0, tk.END)
                            self.entry_time.insert(0, data_photo.get("heure"))
                        if hasattr(self, "choice_comportement"):
                            self.choice_comportement.set(data_photo.get("comportement") or "")
                        if hasattr(self, "choice_certitude"):
                            self.choice_certitude.set(data_photo.get("certitude") or "")
                        if hasattr(self, "choice_type_indice"):
                            self.choice_type_indice.set(data_photo.get("type_indice") or data_photo.get("type_observation") or "")
                        if hasattr(self, "entry_rappel_date"):
                            self.entry_rappel_date.delete(0, tk.END)
                            if data_photo.get("rappel_date"):
                                self.entry_rappel_date.insert(0, str(data_photo.get("rappel_date")))
                        if hasattr(self, "entry_rappel_note"):
                            self.entry_rappel_note.delete(0, tk.END)
                            if data_photo.get("rappel_note"):
                                self.entry_rappel_note.insert(0, str(data_photo.get("rappel_note")))
                        if hasattr(self, "entry_nom_sci"):
                            self.entry_nom_sci.delete(0, tk.END)
                            if data_photo.get("nom_scientifique"):
                                self.entry_nom_sci.insert(0, data_photo.get("nom_scientifique"))
                        if hasattr(self, "entry_amazon_photo"):
                            self.entry_amazon_photo.delete(0, tk.END)
                            if data_photo.get("lien_amazon") or data_photo.get("amazon_url"):
                                self.entry_amazon_photo.insert(
                                    0, data_photo.get("lien_amazon") or data_photo.get("amazon_url") or ""
                                )
                        self.note_text.insert("0.0", data_photo.get("notes_libres", ""))
                        if data_photo.get("lieu"):
                            self.entry_lieu.insert(0, data_photo.get("lieu"))
                            self.location_cache[filename] = data_photo.get("lieu")
                            lieu_deja_connu = True
            except Exception:
                pass

        self.weather_box.configure(state="normal")
        self.weather_box.delete("0.0", "end")

        meteo_en_cache = self.weather_cache.get(filename)
        if meteo_en_cache:
            texte_cache = (
                f"🌤️ MÉTÉO OPEN-METEO (mémoire)\n----------------------------\n"
                f"📅 Le : {meteo_en_cache.get('date','?')} à {meteo_en_cache.get('heure','?')}\n"
                f"🌡️ Température : {meteo_en_cache.get('temperature','?')}°C\n"
                f"💧 Humidité : {meteo_en_cache.get('humidite','?')}%\n☁️ Ciel : {meteo_en_cache.get('ciel','?')}\n"
                f"{meteo_en_cache.get('phase_lunaire','')}"
            )
            if meteo_en_cache.get("pluie_3j_precedents") is not None:
                texte_cache += f"\n🌧️ Cumul pluie (3j avant) : {meteo_en_cache.get('pluie_3j_precedents')} mm"
            self.weather_box.insert("0.0", texte_cache)
            self.weather_box.configure(state="disabled")
        elif info["lat"] and info["lon"] and info["date"]:
            self.weather_box.insert("0.0", "⏳ Récupération de la météo Open-Meteo...")
            self.weather_box.configure(state="disabled")
            threading.Thread(target=self.fetch_historical_weather, args=(info["lat"], info["lon"], info["date"], filename), daemon=True).start()
        else:
            self.weather_box.insert("0.0", "ℹ️ Pas encore de GPS sur cette photo (synchronisez pour activer météo/carte).\nL'annotation reste possible dès maintenant.")
            self.weather_box.configure(state="disabled")

        if info["lat"] and info["lon"]:
            if lieu_deja_connu:
                pass  # déjà rempli ci-dessus, on évite un appel réseau inutile
            else:
                self.entry_lieu.configure(placeholder_text="Recherche du lieu...")
                threading.Thread(target=self.fetch_location, args=(info["lat"], info["lon"], filename), daemon=True).start()
        else:
            self.entry_lieu.configure(placeholder_text="Pas de GPS")

    def on_thumb_size_change(self, value):
        size = int(float(value))
        self.lbl_thumb_size_value.configure(text=f"Taille de l'aperçu : {size} px")
        self._update_preview_display()

    def create_or_open_carnet_folder(self):
        """Crée ou ouvre un dossier de carnet pour observations sans photo (ni GPX obligatoires)."""
        win = ctk.CTkToplevel(self)
        win.title("Carnet sans photos")
        win.geometry("480x280")
        win.configure(fg_color=UI.get("bg", "#0c1210"))
        self._prepare_tool_window(win)
        result = {"ok": False}

        card = ctk.CTkFrame(win, fg_color=UI.get("card", "#1a2420"), corner_radius=12)
        card.pack(fill="both", expand=True, padx=14, pady=14)
        ctk.CTkLabel(
            card, text="Sortie / carnet sans photos",
            font=ctk.CTkFont(size=15, weight="bold")
        ).pack(anchor="w", padx=14, pady=(12, 4))
        ctk.CTkLabel(
            card,
            text="Idéal pour billebaude, affût, indices, chants…\n"
                 "Un dossier local stocke observations.json (et le cloud si configuré).",
            font=ctk.CTkFont(size=12), text_color=UI.get("text_dim", "#9db0a6"),
            justify="left",
        ).pack(anchor="w", padx=14, pady=(0, 10))

        def ouvrir_existant():
            d = filedialog.askdirectory(title="Ouvrir un dossier de carnet existant")
            if not d:
                return
            self._activate_carnet_folder(d)
            result["ok"] = True
            win.destroy()

        def creer_nouveau():
            parent = filedialog.askdirectory(title="Emplacement du nouveau carnet")
            if not parent:
                return
            nom = simpledialog.askstring(
                "Nom de la sortie",
                "Nom du dossier carnet :",
                initialvalue=datetime.now().strftime("%Y-%m-%d_sortie"),
            )
            if not nom:
                return
            nom_safe = re.sub(r'[<>:"/\\|?*]', "_", nom.strip()) or datetime.now().strftime("%Y-%m-%d_sortie")
            dest = os.path.join(parent, nom_safe)
            try:
                os.makedirs(dest, exist_ok=True)
                notes_path = os.path.join(dest, NOTES_FILE)
                if not os.path.exists(notes_path):
                    with open(notes_path, "w", encoding="utf-8") as f:
                        json.dump({}, f, ensure_ascii=False, indent=2)
            except Exception as e:
                messagebox.showerror("Erreur", str(e))
                return
            self._activate_carnet_folder(dest)
            result["ok"] = True
            win.destroy()

        ctk.CTkButton(
            card, text="📂 Ouvrir un carnet existant", command=ouvrir_existant,
            height=36, fg_color=UI.get("accent", "#3eb4a0"),
        ).pack(fill="x", padx=14, pady=6)
        ctk.CTkButton(
            card, text="✨ Créer un nouveau carnet", command=creer_nouveau,
            height=36, fg_color=UI.get("success", "#5ecf8a"),
        ).pack(fill="x", padx=14, pady=6)
        ctk.CTkButton(card, text="Annuler", command=win.destroy, height=32).pack(fill="x", padx=14, pady=(6, 12))

        win.grab_set()
        self.wait_window(win)
        return result["ok"]

    def _activate_carnet_folder(self, path):
        """Active un dossier comme sortie courante (photos optionnelles)."""
        self.photo_folder_path = path
        try:
            self.btn_browse_photos.configure(text=f"📁 {os.path.basename(path)}")
        except Exception:
            pass
        self.photos_data = {}
        self.selected_photo_path = ""
        self.is_synced = False
        # Charge les observations déjà notées (y compris sans photo)
        notes_path = os.path.join(path, NOTES_FILE)
        if os.path.exists(notes_path):
            try:
                with open(notes_path, "r", encoding="utf-8") as f:
                    all_notes = json.load(f)
                for key, data in (all_notes or {}).items():
                    if data.get("lat") is not None and data.get("lon") is not None:
                        self.photos_data[key] = {
                            "path": "",
                            "lat": data.get("lat"),
                            "lon": data.get("lon"),
                            "date": data.get("heure") or "",
                            "is_video": False,
                            "sans_photo": bool(data.get("sans_photo") or str(key).startswith("_manuel_")),
                        }
            except Exception:
                pass
        # Liste média vide ou photos s'il y en a
        try:
            if hasattr(self, "entry_search_photo"):
                self.entry_search_photo.delete(0, tk.END)
            images = self.get_supported_images() if hasattr(self, "get_supported_images") else []
            videos = self.get_supported_videos() if hasattr(self, "get_supported_videos") else []
            self._filtered_media_names = images + videos
            if getattr(self, "carnet_view_mode", "list") == "thumbs":
                self._rebuild_thumb_explorer(lazy=True)
            else:
                self._populate_listbox_from_filtered()
        except Exception:
            pass
        try:
            self.refresh_map_markers()
            self.refresh_daily_counter()
        except Exception:
            pass
        n_obs = 0
        try:
            with open(os.path.join(path, NOTES_FILE), "r", encoding="utf-8") as f:
                n_obs = len(json.load(f) or {})
        except Exception:
            pass
        self.log(
            f"Carnet : {path}\n"
            f"{len(getattr(self, '_filtered_media_names', []) or [])} média(s) · "
            f"{n_obs} observation(s) dans le carnet.\n"
            f"Utilisez « Sans photo » ou « Placer observation » sur la carte."
        )

        try:
            self.reload_devices_for_folder()
        except Exception:
            pass

    def open_indices_folder_workflow(self):
        """Créer un dossier dédié aux indices (terriers, empreintes…) et y importer des photos."""
        win = ctk.CTkToplevel(self)
        win.title("Indices — dossier photo")
        win.geometry("520x420")
        try:
            win.configure(fg_color=UI.get("bg"))
        except Exception:
            pass
        self._prepare_tool_window(win)

        ctk.CTkLabel(
            win, text="Dossier photo pour indices",
            font=ctk.CTkFont(size=16, weight="bold"), text_color=UI.get("text"),
        ).pack(anchor="w", padx=16, pady=(14, 4))
        ctk.CTkLabel(
            win,
            text="Créez un dossier de sortie pour vos photos d'indices (empreintes, terriers, coulées…),\n"
                 "puis importez-y des fichiers depuis la carte SD ou un autre dossier.\n"
                 "Vous pourrez ensuite annoter chaque photo (type d'indice, espèce, GPS).",
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim"),
            wraplength=480, justify="left",
        ).pack(anchor="w", padx=16, pady=(0, 10))

        card = ctk.CTkFrame(win, fg_color=UI.get("card"), corner_radius=10)
        card.pack(fill="both", expand=True, padx=16, pady=6)

        parent_var = tk.StringVar(value="")
        try:
            cfg = self._load_app_config() or {}
            if cfg.get("default_photo_folder") and os.path.isdir(cfg["default_photo_folder"]):
                parent_var.set(cfg["default_photo_folder"])
            elif self.photo_folder_path and os.path.isdir(os.path.dirname(self.photo_folder_path)):
                parent_var.set(os.path.dirname(self.photo_folder_path))
        except Exception:
            pass

        ctk.CTkLabel(card, text="Dossier parent", text_color=UI.get("text")).pack(anchor="w", padx=12, pady=(12, 2))
        row_p = ctk.CTkFrame(card, fg_color="transparent")
        row_p.pack(fill="x", padx=12, pady=2)
        ctk.CTkEntry(row_p, textvariable=parent_var).pack(side="left", fill="x", expand=True)
        ctk.CTkButton(
            row_p, text="…", width=36,
            command=lambda: parent_var.set(
                filedialog.askdirectory(title="Dossier parent") or parent_var.get()
            ),
        ).pack(side="left", padx=4)

        name_var = tk.StringVar(value=datetime.now().strftime("%Y-%m-%d_indices"))
        ctk.CTkLabel(card, text="Nom du dossier", text_color=UI.get("text")).pack(anchor="w", padx=12, pady=(10, 2))
        ctk.CTkEntry(card, textvariable=name_var).pack(fill="x", padx=12, pady=2)

        import_var = tk.BooleanVar(value=True)
        ctk.CTkCheckBox(
            card, text="Importer des photos juste après la création",
            variable=import_var, text_color=UI.get("text"),
        ).pack(anchor="w", padx=12, pady=12)

        def run():
            parent = parent_var.get().strip()
            if not parent or not os.path.isdir(parent):
                messagebox.showwarning("Indices", "Choisissez un dossier parent valide.")
                return
            nom = re.sub(r'[<>:"/\\|?*]', "_", (name_var.get() or "").strip())
            if not nom:
                nom = datetime.now().strftime("%Y-%m-%d_indices")
            dest = os.path.join(parent, nom)
            try:
                os.makedirs(dest, exist_ok=True)
                notes_path = os.path.join(dest, NOTES_FILE)
                if not os.path.exists(notes_path):
                    with open(notes_path, "w", encoding="utf-8") as f:
                        json.dump({}, f, ensure_ascii=False, indent=2)
            except Exception as e:
                messagebox.showerror("Indices", str(e))
                return

            n_copy = 0
            if import_var.get():
                files = filedialog.askopenfilenames(
                    title="Photos d'indices à importer",
                    filetypes=[
                        ("Images", "*.jpg *.jpeg *.png *.tif *.tiff *.nef *.cr2 *.cr3 *.arw *.dng *.raf *.orf *.rw2"),
                        ("Tous", "*.*"),
                    ],
                )
                for src in files or []:
                    try:
                        base = os.path.basename(src)
                        target = os.path.join(dest, base)
                        if os.path.abspath(src) == os.path.abspath(target):
                            continue
                        if os.path.exists(target):
                            stem, ext = os.path.splitext(base)
                            target = os.path.join(dest, "%s_import%s" % (stem, ext))
                        shutil.copy2(src, target)
                        n_copy += 1
                    except Exception as e:
                        self.log("Import indice %s : %s" % (src, e))

            self._activate_carnet_folder(dest)
            self._register_known_folder(dest)
            # Relire GPS éventuels des photos importées
            try:
                exe = self.get_exiftool_path()
                if exe:
                    threading.Thread(
                        target=self._check_existing_gps_worker, args=(exe,), daemon=True
                    ).start()
            except Exception:
                pass

            win.destroy()
            messagebox.showinfo(
                "Indices",
                "Dossier prêt :\n%s\n\n%d photo(s) importée(s).\n\n"
                "Annotez chaque photo dans le carnet (type d'indice, espèce)\n"
                "ou placez des observations sans photo sur la carte."
                % (dest, n_copy),
            )
            self.log("Dossier indices : %s (%d import)" % (dest, n_copy))

        bf = ctk.CTkFrame(win, fg_color="transparent")
        bf.pack(fill="x", padx=16, pady=12)
        ctk.CTkButton(
            bf, text="Créer / importer", fg_color=UI.get("success", "#2f9e5f"), command=run,
        ).pack(side="right", padx=4)
        ctk.CTkButton(
            bf, text="Annuler", fg_color=UI.get("card_alt"), command=win.destroy,
        ).pack(side="right", padx=4)
        win.after(200, lambda: self._style_dialog(win) if hasattr(self, "_style_dialog") else None)

    def open_manual_observation_dialog(self, lat=None, lon=None):
        """Formulaire pour enregistrer une observation sans photo : animal vu ou entendu, sans déclenchement.
        lat/lon optionnels : préremplis si l'utilisateur a cliqué sur la carte (mode « Placer observation »)."""
        if not self.photo_folder_path:
            if messagebox.askyesno(
                "Carnet sans photos",
                "Aucun dossier de sortie n'est ouvert.\n\n"
                "Créer ou choisir un dossier de carnet\n"
                "(observations sans photo, pas besoin de photos) ?"
            ):
                if not self.create_or_open_carnet_folder():
                    return
            else:
                return

        win = ctk.CTkToplevel(self)
        win.title("👁️ Observation sans photo")
        win.geometry("500x700")
        self._prepare_tool_window(win)

        if lat is not None and lon is not None:
            intro = f"Point placé sur la carte : {lat:.5f}, {lon:.5f}\nComplétez l'espèce et les détails ci-dessous."
        else:
            intro = ("Animal vu ou entendu sans photo.\n"
                     "Astuce : sur l'onglet Carte, utilisez « ➕ Placer observation » puis cliquez "
                     "à l'endroit voulu pour préremplir le GPS.")
        ctk.CTkLabel(
            win, text=intro,
            font=ctk.CTkFont(size=11), text_color=UI.get("text_dim", "#666666"), wraplength=460, justify="left"
        ).pack(anchor="w", padx=15, pady=(15, 10))

        form = ctk.CTkScrollableFrame(win, fg_color="transparent")
        form.pack(fill="both", expand=True, padx=15, pady=(0, 10))
        form.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(form, text="Type :", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, sticky="w", pady=8)
        choix_type = ctk.CTkOptionMenu(form, values=[
            "👁️ Vu (sans photo)", "👂 Entendu",
            "🐾 Empreinte / Trace", "🕳️ Terrier", "➡️ Coulée / Passage", "💩 Latrines / Indice"
        ])
        choix_type.grid(row=0, column=1, sticky="ew", pady=8)

        ctk.CTkLabel(form, text="Catégorie :", font=ctk.CTkFont(weight="bold")).grid(row=1, column=0, sticky="w", pady=8)
        choix_cat = ctk.CTkOptionMenu(
            form, values=list(self.faune_meuse.keys()) or ["Autre"],
            command=lambda c: (
                choix_esp.configure(values=self.faune_meuse.get(c, ["Autre"])),
                choix_esp.set(self.faune_meuse.get(c, ["Autre"])[0])
            )
        )
        choix_cat.grid(row=1, column=1, sticky="ew", pady=8)

        ctk.CTkLabel(form, text="Espèce :", font=ctk.CTkFont(weight="bold")).grid(row=2, column=0, sticky="w", pady=8)
        premiere_cat = choix_cat.get()
        choix_esp = ctk.CTkComboBox(form, values=self.faune_meuse.get(premiere_cat, ["Autre"]))
        choix_esp.grid(row=2, column=1, sticky="ew", pady=8)
        if self.faune_meuse.get(premiere_cat):
            choix_esp.set(self.faune_meuse[premiere_cat][0])

        # Si le type est un indice, bascule auto vers catégorie Autre + espèce correspondante
        def on_type_change(val):
            mapping = {
                "🐾 Empreinte / Trace": ("Autre", "Traces / Empreintes"),
                "🕳️ Terrier": ("Autre", "Terrier"),
                "➡️ Coulée / Passage": ("Autre", "Coulée / Passage"),
                "💩 Latrines / Indice": ("Autre", "Latrines"),
            }
            if val in mapping:
                cat, esp = mapping[val]
                choix_cat.set(cat)
                choix_esp.configure(values=self.faune_meuse.get(cat, ["Autre"]))
                if esp in self.faune_meuse.get(cat, []):
                    choix_esp.set(esp)
                else:
                    choix_esp.set(self.faune_meuse.get(cat, ["Autre"])[0])
        choix_type.configure(command=on_type_change)

        ctk.CTkLabel(form, text="Nombre :", font=ctk.CTkFont(weight="bold")).grid(row=3, column=0, sticky="w", pady=8)
        entry_nb = ctk.CTkComboBox(form, values=["1", "2", "3", "4", "5", "6-10", "10+", "Nombreux", "Vol important", "Indéterminé"])
        entry_nb.grid(row=3, column=1, sticky="ew", pady=8)
        entry_nb.set("1")

        ctk.CTkLabel(form, text="Heure :", font=ctk.CTkFont(weight="bold")).grid(row=4, column=0, sticky="w", pady=8)
        entry_heure = ctk.CTkEntry(form, placeholder_text="HH:MM")
        entry_heure.grid(row=4, column=1, sticky="ew", pady=8)
        entry_heure.insert(0, datetime.now().strftime("%H:%M"))

        ctk.CTkLabel(form, text="Lieu :", font=ctk.CTkFont(weight="bold")).grid(row=5, column=0, sticky="w", pady=8)
        entry_lieu = ctk.CTkEntry(form, placeholder_text="ex : Lisière bois de Verdun")
        entry_lieu.grid(row=5, column=1, sticky="ew", pady=8)

        if lat is not None and lon is not None:
            lat_ref, lon_ref = lat, lon
        else:
            lat_ref, lon_ref, _ = self._get_reference_location()
        ctk.CTkLabel(form, text="GPS :", font=ctk.CTkFont(weight="bold")).grid(row=6, column=0, sticky="w", pady=8)
        gps_frame = ctk.CTkFrame(form, fg_color="transparent")
        gps_frame.grid(row=6, column=1, sticky="ew", pady=8)
        entry_lat = ctk.CTkEntry(gps_frame, placeholder_text="latitude", width=110)
        entry_lat.pack(side="left", padx=(0, 5))
        entry_lat.insert(0, f"{lat_ref:.5f}")
        entry_lon = ctk.CTkEntry(gps_frame, placeholder_text="longitude", width=110)
        entry_lon.pack(side="left")
        entry_lon.insert(0, f"{lon_ref:.5f}")
        ctk.CTkButton(
            gps_frame, text="✕", width=26, fg_color="#3a3a3a", hover_color="#4a4a4a",
            command=lambda: (entry_lat.delete(0, tk.END), entry_lon.delete(0, tk.END))
        ).pack(side="left", padx=(5, 0))

        # Géocodage inverse optionnel pour préremplir le lieu
        if lat is not None and lon is not None:
            def remplir_lieu():
                lieu = self.reverse_geocode(lat, lon)
                if lieu:
                    entry_lieu.delete(0, tk.END)
                    entry_lieu.insert(0, lieu)
            threading.Thread(target=remplir_lieu, daemon=True).start()

        ctk.CTkLabel(form, text="Comportement :", font=ctk.CTkFont(weight="bold")).grid(row=7, column=0, sticky="w", pady=6)
        choix_comp = ctk.CTkComboBox(form, values=COMPORTEMENTS_OBS)
        choix_comp.grid(row=7, column=1, sticky="ew", pady=6)
        choix_comp.set("")

        ctk.CTkLabel(form, text="Certitude :", font=ctk.CTkFont(weight="bold")).grid(row=8, column=0, sticky="w", pady=6)
        choix_cert = ctk.CTkComboBox(form, values=CERTITUDES_OBS)
        choix_cert.grid(row=8, column=1, sticky="ew", pady=6)
        choix_cert.set("Sûr")

        ctk.CTkLabel(form, text="Nom scientifique :", font=ctk.CTkFont(weight="bold")).grid(row=9, column=0, sticky="w", pady=6)
        entry_sci = ctk.CTkEntry(form, placeholder_text="optionnel")
        entry_sci.grid(row=9, column=1, sticky="ew", pady=6)

        ctk.CTkLabel(form, text="Notes :", font=ctk.CTkFont(weight="bold")).grid(row=10, column=0, sticky="nw", pady=8)
        note_manuelle = ctk.CTkTextbox(form, height=80)
        note_manuelle.grid(row=10, column=1, sticky="ew", pady=8)

        def enregistrer():
            espece = choix_esp.get().strip()
            if not espece:
                messagebox.showwarning("Espèce manquante", "Indiquez au moins une espèce.")
                return

            cle = f"_manuel_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
            chemin_notes = os.path.join(self.photo_folder_path, NOTES_FILE)
            all_notes = {}
            if os.path.exists(chemin_notes):
                try:
                    with open(chemin_notes, "r", encoding="utf-8") as f:
                        all_notes = json.load(f)
                except Exception:
                    pass

            all_notes[cle] = {
                "departement": "55 - Meuse",
                "categorie": choix_cat.get(),
                "espece": espece,
                "nombre": entry_nb.get().strip() or "1",
                "heure": entry_heure.get().strip(),
                "lieu": entry_lieu.get().strip(),
                "notes_libres": note_manuelle.get("0.0", "end-1c"),
                "type_observation": choix_type.get(),
                "type_indice": choix_type.get() if "Empreinte" in choix_type.get() or "Terrier" in choix_type.get() or "Coulée" in choix_type.get() or "Latrine" in choix_type.get() else "",
                "comportement": choix_comp.get().strip(),
                "certitude": choix_cert.get().strip(),
                "nom_scientifique": entry_sci.get().strip(),
                "sans_photo": True,
                "meteo": None,
            }

            lat_txt, lon_txt = entry_lat.get().strip().replace(",", "."), entry_lon.get().strip().replace(",", ".")
            if lat_txt and lon_txt:
                try:
                    all_notes[cle]["lat"] = float(lat_txt)
                    all_notes[cle]["lon"] = float(lon_txt)
                except ValueError:
                    messagebox.showwarning("GPS invalide", "Latitude/longitude non numériques — observation enregistrée sans GPS.")

            try:
                self._save_notes_dict(all_notes, force_backup=True)
            except Exception as e:
                messagebox.showerror("Erreur de sauvegarde", str(e))
                return

            if all_notes[cle].get("lat") is not None and all_notes[cle].get("lon") is not None:
                self.photos_data[cle] = {
                    "lat": all_notes[cle]["lat"],
                    "lon": all_notes[cle]["lon"],
                    "date": "",
                }

            self.log(f"👁️ Observation sans photo enregistrée : {choix_type.get()} → {espece} ({entry_nb.get().strip()})")
            self.refresh_daily_counter()
            self.refresh_map_markers()
            win.destroy()
            messagebox.showinfo("Enregistré", f"Observation « {espece} » ajoutée sur la carte.")

        btn_row = ctk.CTkFrame(win, fg_color="transparent")
        btn_row.pack(fill="x", padx=15, pady=15)
        btn_row.grid_columnconfigure((0, 1), weight=1)
        ctk.CTkButton(
            btn_row, text="Annuler", command=win.destroy,
            fg_color="#3a3a3a", hover_color="#4a4a4a"
        ).grid(row=0, column=0, sticky="ew", padx=(0, 4))
        ctk.CTkButton(
            btn_row, text="💾 Enregistrer l'observation", command=enregistrer,
            fg_color="#2ba14b", hover_color="#1f7d37", font=ctk.CTkFont(weight="bold")
        ).grid(row=0, column=1, sticky="ew", padx=(4, 0))

    def relocaliser_photo(self):
        """Relance manuellement la recherche du lieu pour la photo actuellement sélectionnée."""
        if not self.selected_photo_path:
            return
        filename = os.path.basename(self.selected_photo_path)
        info = self.photos_data.get(filename)
        if not info or not info.get("lat") or not info.get("lon"):
            messagebox.showinfo("Pas de GPS", "Cette photo n'a pas de coordonnées GPS valides.")
            return
        self.entry_lieu.configure(placeholder_text="Recherche du lieu...")
        self.entry_lieu.delete(0, tk.END)
        threading.Thread(target=self.fetch_location, args=(info["lat"], info["lon"], filename), daemon=True).start()

    def reverse_geocode(self, lat, lon):
        """Traduit des coordonnées GPS en nom de lieu lisible via Nominatim (OpenStreetMap)."""
        try:
            url = (
                f"https://nominatim.openstreetmap.org/reverse?lat={lat}&lon={lon}"
                f"&format=json&accept-language=fr&zoom=15"
            )
            headers = {"User-Agent": f"GeoExifMeuse55/{APP_VERSION} (usage naturaliste local)"}
            response = requests.get(url, headers=headers, timeout=6).json()
            addr = response.get("address", {})

            lieu_precis = (
                addr.get("hamlet") or addr.get("village") or addr.get("locality")
                or addr.get("suburb") or addr.get("town") or addr.get("city")
            )
            commune = addr.get("municipality") or addr.get("county") or addr.get("town") or addr.get("city")
            foret = addr.get("forest") or addr.get("natural")

            morceaux = []
            if foret and foret != lieu_precis:
                morceaux.append(foret)
            if lieu_precis:
                morceaux.append(lieu_precis)
            if commune and commune not in morceaux:
                morceaux.append(commune)

            if morceaux:
                return " – ".join(morceaux)

            display_name = response.get("display_name", "")
            if display_name:
                return ", ".join(display_name.split(",")[:2]).strip()

            return f"{lat:.4f}, {lon:.4f}"
        except Exception:
            return None

    def delete_current_observation(self):
        """Supprime l'observation sélectionnée du carnet (surtout Birda / sans photo)."""
        if not self.photo_folder_path:
            messagebox.showwarning("Carnet", "Aucune sortie ouverte.")
            return
        key = getattr(self, "_current_note_key", None)
        if not key:
            try:
                sel = self.photo_listbox.curselection()
                if sel:
                    key = self._resolve_list_selection_to_key(self.photo_listbox.get(sel[0]))
            except Exception:
                key = None
        if not key:
            messagebox.showinfo(
                "Supprimer",
                "Sélectionnez d'abord une observation dans la liste du carnet\n"
                "(ligne 🎧 Birda ou 📍 sans photo).",
            )
            return

        notes = self._load_notes_dict(self.photo_folder_path) or {}
        data = notes.get(key) if isinstance(notes.get(key), dict) else {}
        label = (data.get("espece") or key) if data else key
        is_file = os.path.isfile(os.path.join(self.photo_folder_path, key))

        if is_file:
            msg = (
                "« %s » est aussi un fichier photo/vidéo du dossier.\n\n"
                "Supprimer uniquement la fiche carnet (le fichier reste sur le disque) ?"
                % label
            )
        else:
            msg = (
                "Supprimer l'observation « %s » du carnet ?\n\n"
                "(Birda / sans photo — aucune photo n'est effacée du disque)"
                % label
            )
        if not messagebox.askyesno("Supprimer l'observation", msg):
            return

        if key in notes:
            del notes[key]
        try:
            self._save_notes_dict(notes, force_backup=True)
        except Exception as e:
            messagebox.showerror("Supprimer", str(e))
            return
        self.photos_data.pop(key, None)
        if getattr(self, "_current_note_key", None) == key:
            self._current_note_key = None
            self.selected_photo_path = None
        try:
            self._filtered_media_names = None
            self._populate_listbox_from_filtered()
            self.refresh_map_markers()
            self.refresh_daily_counter()
        except Exception:
            pass
        self.log("Observation supprimée du carnet : %s" % label)
        messagebox.showinfo("Supprimé", "« %s » a été retiré du carnet." % label)

    def save_current_note(self, silent=False):
        # Clé = observation virtuelle (Birda / manuel) ou fichier média
        filename = getattr(self, "_current_note_key", None)
        if not filename:
            if not self.selected_photo_path:
                return
            filename = os.path.basename(self.selected_photo_path)
        is_virtual = (
            str(filename).startswith("_")
            or not os.path.isfile(os.path.join(self.photo_folder_path or "", filename))
        )
        chemin_notes = os.path.join(self.photo_folder_path, NOTES_FILE)
        
        all_notes = {}
        if os.path.exists(chemin_notes):
            try:
                with open(chemin_notes, "r", encoding="utf-8") as f:
                    all_notes = json.load(f)
            except Exception:
                pass

        prev = all_notes.get(filename) if isinstance(all_notes.get(filename), dict) else {}
        all_notes[filename] = {
            "departement": "55 - Meuse",
            "categorie": self.choice_category.get(),
            "espece": self.choice_species.get(),
            "nombre": self.choice_count.get(),
            "heure": self.entry_time.get(),
            "lieu": self.entry_lieu.get().strip(),
            "comportement": self.choice_comportement.get().strip() if hasattr(self, "choice_comportement") else "",
            "certitude": self.choice_certitude.get().strip() if hasattr(self, "choice_certitude") else "",
            "type_indice": self.choice_type_indice.get().strip() if hasattr(self, "choice_type_indice") else "",
            "type_observation": prev.get("type_observation") or (
                "Prise de son (Birda)" if "birda" in str(filename).lower() else ""
            ),
            "nom_scientifique": self.entry_nom_sci.get().strip() if hasattr(self, "entry_nom_sci") else "",
            "lien_amazon": self.entry_amazon_photo.get().strip() if hasattr(self, "entry_amazon_photo") else "",
            "notes_libres": self.note_text.get("0.0", "end-1c"),
            "meteo": self.weather_cache.get(filename) or prev.get("meteo"),
            "sans_photo": True if is_virtual else bool(prev.get("sans_photo")),
            "source": prev.get("source") or ("birda" if "birda" in str(filename).lower() else ""),
            "birda_confidence": prev.get("birda_confidence"),
        }
        # Album sortie (si saisi)
        try:
            if hasattr(self, "entry_amazon_album"):
                alb = self.entry_amazon_album.get().strip()
                if alb:
                    self._save_sortie_amazon_album(self.photo_folder_path, alb)
        except Exception:
            pass
        # GPS si connu
        info = self.photos_data.get(filename) or {}
        if info.get("lat") is not None and info.get("lon") is not None:
            all_notes[filename]["lat"] = info.get("lat")
            all_notes[filename]["lon"] = info.get("lon")
        
        try:
            # force_backup sur sauvegarde manuelle ; autosave respecte l'intervalle mini
            self._save_notes_dict(all_notes, force_backup=not silent, silent=silent)
            if silent:
                self.log(f"💾 Sauvegarde auto : {filename}")
            else:
                self.log(f"📝 Enregistré : {filename} ➔ {self.choice_species.get()}")
            self.refresh_map_markers()
            self.refresh_daily_counter()
            self._refresh_listbox_annotation_status()
            # Sync cloud optionnel (après enregistrement manuel ou auto)
            if (self.app_config or {}).get("cloud_auto_sync") and (self.app_config or {}).get("cloud_folder"):
                try:
                    self.sync_carnet_to_cloud(silent=True)
                except Exception:
                    pass
        except Exception as e:
            if not silent:
                messagebox.showerror("Erreur de sauvegarde", str(e))

    def update_location_ui(self, lieu, filename):
        if lieu:
            self.location_cache[filename] = lieu
        # N'affiche le résultat que si la photo sélectionnée n'a pas changé entre-temps
        if self.selected_photo_path and os.path.basename(self.selected_photo_path) == filename:
            self.entry_lieu.delete(0, tk.END)
            if lieu:
                self.entry_lieu.insert(0, lieu)
                self.entry_lieu.configure(placeholder_text="")
            else:
                self.entry_lieu.configure(placeholder_text="Lieu introuvable (saisie manuelle possible)")

    def update_weather_ui(self, text, weather_data, filename):
        self.weather_box.configure(state="normal")
        self.weather_box.delete("0.0", "end")
        self.weather_box.insert("0.0", text)
        self.weather_box.configure(state="disabled")
        if weather_data:
            self.weather_cache[filename] = weather_data
        else:
            self.weather_cache.pop(filename, None)



if __name__ == "__main__":
    # Obligatoire pour les EXE Windows (PyInstaller / cx_Freeze) :
    # sans freeze_support, chaque ProcessPoolExecutor relance tout le script
    # → plusieurs fenêtres GeoExif à l'ouverture d'un dossier photos.
    multiprocessing.freeze_support()

    def _find_splash_image():
        """Cherche une illustration de boot (PNG/JPG) a cote de l'app ou dans les donnees utilisateur."""
        names = (
            "geoexif_splash.png", "splash.png", "boot.png",
            "geoexif_splash.jpg", "splash.jpg", "boot.jpg",
            "geoexif_boot.png",
        )
        bases = []
        try:
            bases.append(os.path.dirname(os.path.abspath(sys.argv[0])))
        except Exception:
            pass
        try:
            bases.append(os.getcwd())
        except Exception:
            pass
        if os.name == "nt":
            la = os.environ.get("LOCALAPPDATA")
            if la:
                bases.append(os.path.join(la, "GeoExif"))
        try:
            bases.append(os.path.join(os.path.expanduser("~"), ".geoexif"))
        except Exception:
            pass
        for base in bases:
            for name in names:
                path = os.path.join(base, name)
                if os.path.isfile(path):
                    return path
        return None

    class SplashScreen(ctk.CTk):
        """Ecran de demarrage : illustration optionnelle + auteur + version."""
        def __init__(self, on_finish):
            super().__init__()
            self.on_finish = on_finish
            self.overrideredirect(True)
            self._splash_photo = None  # garde ref PIL/CTkImage

            splash_path = _find_splash_image()
            img_disp = None
            if splash_path:
                try:
                    im = Image.open(splash_path).convert("RGB")
                    # Taille max splash (portrait ou paysage)
                    max_w, max_h = 520, 620
                    im.thumbnail((max_w, max_h), Image.Resampling.LANCZOS if hasattr(Image, "Resampling") else Image.LANCZOS)
                    img_disp = im
                except Exception:
                    img_disp = None

            if img_disp is not None:
                width = max(420, img_disp.width + 40)
                height = img_disp.height + 120
            else:
                width, height = 480, 300

            self.update_idletasks()
            sw = self.winfo_screenwidth()
            sh = self.winfo_screenheight()
            x = (sw - width) // 2
            y = (sh - height) // 2
            self.geometry("%dx%d+%d+%d" % (width, height, x, y))
            try:
                self.configure(fg_color=UI.get("bg", "#101612"))
            except Exception:
                pass

            content = ctk.CTkFrame(self, fg_color="transparent")
            content.pack(expand=True, fill="both", padx=12, pady=10)

            if img_disp is not None:
                try:
                    ctk_img = ctk.CTkImage(
                        light_image=img_disp, dark_image=img_disp,
                        size=(img_disp.width, img_disp.height),
                    )
                    self._splash_photo = ctk_img
                    ctk.CTkLabel(content, text="", image=ctk_img).pack(pady=(4, 8))
                except Exception:
                    ctk.CTkLabel(content, text="GeoExif", font=ctk.CTkFont(size=28, weight="bold")).pack(pady=20)
            else:
                ctk.CTkLabel(content, text="📸", font=ctk.CTkFont(size=48)).pack(pady=(28, 6))
                ctk.CTkLabel(
                    content, text="GeoExif Sync & Naturalist Notebook",
                    font=ctk.CTkFont(size=17, weight="bold"), text_color=UI.get("text", "#eee"),
                ).pack()
                ctk.CTkLabel(
                    content, text="Meuse (55)", font=ctk.CTkFont(size=12),
                    text_color=UI.get("text_dim", "#888"),
                ).pack(pady=(0, 12))

            self.progress = ctk.CTkProgressBar(
                content, width=min(320, width - 60),
                progress_color=UI.get("accent", "#3eb4a0"),
                fg_color=UI.get("card_alt", "#333"),
                height=8,
            )
            self.progress.pack(pady=(4, 10))
            self.progress.set(0)

            ctk.CTkLabel(
                content, text=APP_AUTHOR,
                font=ctk.CTkFont(size=12, slant="italic"),
                text_color=UI.get("text", "#eee"),
            ).pack()
            ctk.CTkLabel(
                content, text="Version %s" % APP_VERSION,
                font=ctk.CTkFont(size=10),
                text_color=UI.get("text_muted", "#888"),
            ).pack(pady=(2, 4))

            # Duree un peu plus longue si illustration (laisser apprecier)
            self._animate(0, total_steps=36 if img_disp is not None else 24)

        def _animate(self, step, total_steps=24):
            try:
                self.progress.set(step / max(1, total_steps))
            except Exception:
                pass
            if step < total_steps:
                delay = 50 if total_steps > 30 else 45
                self.after(delay, lambda: self._animate(step + 1, total_steps))
            else:
                try:
                    self.destroy()
                except Exception:
                    pass
                self.on_finish()

    def _launch_main_app():
        app = GeoExifIgnApp()
        app.mainloop()

    splash = SplashScreen(on_finish=_launch_main_app)
    splash.mainloop()
